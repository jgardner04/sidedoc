"""Reconstruct Word documents from sidedoc format."""

import hashlib
import re
import warnings
from pathlib import Path
from typing import Any, Optional
from urllib.parse import unquote
from docx import Document
from docx.shared import Pt, Inches
from docx.document import Document as DocumentType
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.opc.part import Part
from docx.opc.packuri import PackURI
from docx.oxml import OxmlElement
from lxml import etree  # type: ignore[import-untyped]
import click
from sidedoc.models import Block, ColumnDefinition, SectionProperties, Style, TrackChange, deserialize_sections
from sidedoc.constants import (
    DEFAULT_IMAGE_WIDTH_INCHES,
    ALIGNMENT_STRING_TO_ENUM,
    GFM_SEPARATOR_PATTERNS,
    DEFAULT_ALIGNMENT,
    MAX_TABLE_ROWS,
    MAX_TABLE_COLS,
    MAX_TABLE_LINES,
    INSERTION_PATTERN,
    DELETION_PATTERN,
    SUBSTITUTION_PATTERN,
    WORDPROCESSINGML_NS,
    XML_SPACE_NS,
    VALID_BORDER_STYLES,
    VALID_PATTERN_FILLS,
    HEX_COLOR_PATTERN,
    MAX_BORDER_WIDTH,
    FOOTNOTES_RT,
    ENDNOTES_RT,
    FOOTNOTES_CT,
    ENDNOTES_CT,
)
from sidedoc.store import SidedocStore

import mistune

# Known limitation: multi-line footnote definitions not supported.
# Only single-line [^N]: text definitions are captured; indented continuation
# lines (per Markdown spec) are silently dropped.
FOOTNOTE_DEF_PATTERN = re.compile(r'^\[\^(\d+)\]:\s*(.+)$')
# Note: only numeric markers are supported. Alphanumeric labels like [^label]
# (valid in standard Markdown) will be left as literal text. This is consistent
# with the extraction side, which always generates numeric markers.
FOOTNOTE_REF_PATTERN = re.compile(r'\[\^(\d+)\]')

# Pattern for parsing inline markdown formatting (bold/italic) in footnote text.
# Matches ***bold italic***, **bold**, or *italic* spans.
_INLINE_FORMAT_PATTERN = re.compile(r'(\*{1,3})(.*?)\1')

# Cache the mistune parser at module level to avoid recreating per paragraph
_MARKDOWN_PARSER = mistune.create_markdown(renderer=None)


def _extract_textbox_inner_content(content: str) -> str:
    """Extract inner text from textbox markdown markers.

    Args:
        content: Block content like "<!-- textbox -->\\nText here.\\n<!-- /textbox -->"

    Returns:
        Inner text with markers stripped.
    """
    lines = content.split("\n")
    inner = []
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("<!-- textbox") or stripped == "<!-- /textbox -->":
            continue
        inner.append(line)
    return "\n".join(inner)


def apply_inline_formatting(paragraph: Any, content: str) -> None:
    """Apply inline formatting from markdown to a paragraph.

    Uses mistune for robust markdown parsing to handle:
    - Nested formatting (**bold *italic* text**)
    - Escaped asterisks (\\*literal\\*)
    - Malformed markdown (graceful degradation)

    Args:
        paragraph: python-docx Paragraph object
        content: Text content with markdown formatting
    """
    runs = _parse_inline_markdown(content)

    if not runs:
        paragraph.add_run(content)
    else:
        for text, bold, italic in runs:
            run = paragraph.add_run(text)
            if bold:
                run.bold = True
            if italic:
                run.italic = True


def _parse_inline_markdown(content: str) -> list[tuple[str, bool, bool]]:
    """Parse inline markdown formatting using mistune.

    Handles nested formatting, escaped markers, and malformed markdown.
    Returns plain text on parse error.
    """
    try:
        tokens, _ = _MARKDOWN_PARSER.parse(content)
    except Exception:
        return [(content, False, False)]

    runs: list[tuple[str, bool, bool]] = []
    token_list: list[dict[str, Any]] = list(tokens) if isinstance(tokens, list) else []
    for block_token in token_list:
        if block_token.get("type") == "paragraph":
            children = block_token.get("children", [])
            _process_tokens(children, runs, bold=False, italic=False)
        else:
            raw = block_token.get("raw", "")
            if raw:
                runs.append((raw, False, False))

    return runs if runs else [(content, False, False)]


def _process_tokens(
    tokens: list[dict[str, Any]],
    runs: list[tuple[str, bool, bool]],
    bold: bool,
    italic: bool,
) -> None:
    """Recursively process mistune tokens into text runs.

    Args:
        tokens: List of mistune inline tokens
        runs: Output list to append runs to
        bold: Whether current context is bold
        italic: Whether current context is italic
    """
    for token in tokens:
        token_type = token.get("type", "")

        if token_type == "text":
            raw = token.get("raw", "")
            if raw:
                runs.append((raw, bold, italic))

        elif token_type == "strong":
            children = token.get("children", [])
            _process_tokens(children, runs, bold=True, italic=italic)

        elif token_type == "emphasis":
            children = token.get("children", [])
            _process_tokens(children, runs, bold=bold, italic=True)

        elif token_type == "codespan":
            raw = token.get("raw", "")
            if raw:
                runs.append((raw, bold, italic))

        elif token_type == "softbreak" or token_type == "linebreak":
            runs.append((" ", bold, italic))

        else:
            raw = token.get("raw", "")
            children = token.get("children", [])

            if children:
                _process_tokens(children, runs, bold, italic)
            elif raw:
                runs.append((raw, bold, italic))


# Regex to match markdown hyperlinks: [text](url)
# The link text pattern handles escaped brackets (e.g., \[ and \]) by matching either:
# - any character except ] or \
# - OR a backslash followed by any character (which handles \[ and \])
HYPERLINK_PATTERN = re.compile(r'\[((?:[^\]\\]|\\.)*)\]\(([^)]+)\)')

# Pre-compiled CriticMarkup patterns
_INS_RE = re.compile(INSERTION_PATTERN)
_DEL_RE = re.compile(DELETION_PATTERN)
_SUB_RE = re.compile(SUBSTITUTION_PATTERN)

# Standard hyperlink color (blue)
HYPERLINK_COLOR = "0563C1"


def parse_criticmarkup(text: str) -> list[tuple[str, str]]:
    """Parse CriticMarkup syntax into segments.

    Identifies insertions {++text++}, deletions {--text--}, and substitutions {~~old~>new~~}
    in the text and returns a list of segments with their types.

    Args:
        text: Text that may contain CriticMarkup syntax

    Returns:
        List of (type, content) tuples where type is "text", "insertion", "deletion", or "substitution"
    """
    segments: list[tuple[str, str]] = []
    last_end = 0

    # Find all matches with their positions
    all_matches: list[tuple[int, int, str, str, Optional[str]]] = []

    for match in _INS_RE.finditer(text):
        all_matches.append((match.start(), match.end(), "insertion", match.group(1), None))

    for match in _DEL_RE.finditer(text):
        all_matches.append((match.start(), match.end(), "deletion", match.group(1), None))

    for match in _SUB_RE.finditer(text):
        # Substitution has two parts: old text and new text
        all_matches.append((match.start(), match.end(), "substitution", match.group(1), match.group(2)))

    # Sort by position
    all_matches.sort(key=lambda x: x[0])

    for start, end, match_type, content, new_content in all_matches:
        # Add any text before this match
        if start > last_end:
            segments.append(("text", text[last_end:start]))

        if match_type == "substitution":
            # Substitution expands to deletion followed by insertion
            segments.append(("deletion", content))
            if new_content:
                segments.append(("insertion", new_content))
        else:
            segments.append((match_type, content))

        last_end = end

    # Add remaining text after last match
    if last_end < len(text):
        segments.append(("text", text[last_end:]))

    # If no CriticMarkup was found, return the whole text as a single segment
    if not segments:
        segments.append(("text", text))

    return segments


def create_ins_element(text: str, author: str, date: str, revision_id: str) -> Any:
    """Create a w:ins XML element for track change insertion.

    Args:
        text: The inserted text
        author: Author who made the change
        date: ISO 8601 timestamp
        revision_id: Unique revision ID

    Returns:
        lxml Element representing the w:ins element
    """
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), revision_id)
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), date)

    # Create run inside insertion
    run = OxmlElement("w:r")
    ins.append(run)

    # Create text element
    t = OxmlElement("w:t")
    t.text = text
    t.set(XML_SPACE_NS, "preserve")
    run.append(t)

    return ins


def create_del_element(text: str, author: str, date: str, revision_id: str) -> Any:
    """Create a w:del XML element for track change deletion.

    Args:
        text: The deleted text
        author: Author who made the change
        date: ISO 8601 timestamp
        revision_id: Unique revision ID

    Returns:
        lxml Element representing the w:del element
    """
    del_elem = OxmlElement("w:del")
    del_elem.set(qn("w:id"), revision_id)
    del_elem.set(qn("w:author"), author)
    del_elem.set(qn("w:date"), date)

    # Create run inside deletion
    run = OxmlElement("w:r")
    del_elem.append(run)

    # Create delText element (not regular w:t)
    del_text = OxmlElement("w:delText")
    del_text.text = text
    del_text.set(XML_SPACE_NS, "preserve")
    run.append(del_text)

    return del_elem


def add_text_with_track_changes(
    paragraph: Any,
    content: str,
    track_changes: Optional[list[TrackChange]] = None,
    default_author: str = "Sidedoc AI",
    default_date: str = "",
) -> None:
    """Add text content with track changes to a paragraph.

    Parses CriticMarkup syntax and creates appropriate w:ins and w:del elements.
    If track_changes metadata is provided, uses those author/date values.

    Args:
        paragraph: python-docx Paragraph object
        content: Text content with CriticMarkup syntax
        track_changes: Optional list of TrackChange metadata
        default_author: Default author for new track changes
        default_date: Default date for new track changes
    """
    from sidedoc.utils import get_iso_timestamp

    if not default_date:
        default_date = get_iso_timestamp()

    segments = parse_criticmarkup(content)

    # Create a mapping from track change content to metadata
    tc_lookup: dict[tuple[str, str], TrackChange] = {}
    if track_changes:
        for tc in track_changes:
            if tc.type == "deletion" and tc.deleted_text:
                tc_lookup[("deletion", tc.deleted_text)] = tc
            elif tc.type == "insertion":
                # For insertions, we need to figure out the text from position
                # This is approximate - we'll use the content between positions
                pass

    revision_counter = [1]  # Use list to allow modification in nested function

    def get_next_revision_id() -> str:
        rid = str(revision_counter[0])
        revision_counter[0] += 1
        return rid

    para_elem = paragraph._p

    for seg_type, seg_content in segments:
        if seg_type == "text":
            # Regular text - add as run
            if seg_content:
                run = OxmlElement("w:r")
                t = OxmlElement("w:t")
                t.text = seg_content
                t.set(XML_SPACE_NS, "preserve")
                run.append(t)
                para_elem.append(run)

        elif seg_type == "insertion":
            # Look up metadata or use defaults
            author = default_author
            date = default_date
            revision_id = get_next_revision_id()

            # Try to find matching track change
            if track_changes:
                for tc in track_changes:
                    if tc.type == "insertion":
                        author = tc.author
                        date = tc.date
                        revision_id = tc.revision_id or get_next_revision_id()
                        break

            ins_elem = create_ins_element(seg_content, author, date, revision_id)
            para_elem.append(ins_elem)

        elif seg_type == "deletion":
            # Look up metadata or use defaults
            author = default_author
            date = default_date
            revision_id = get_next_revision_id()

            # Try to find matching track change
            key = ("deletion", seg_content)
            if key in tc_lookup:
                tc = tc_lookup[key]
                author = tc.author
                date = tc.date
                revision_id = tc.revision_id or get_next_revision_id()
            elif track_changes:
                for tc in track_changes:
                    if tc.type == "deletion" and tc.deleted_text == seg_content:
                        author = tc.author
                        date = tc.date
                        revision_id = tc.revision_id or get_next_revision_id()
                        break

            del_elem = create_del_element(seg_content, author, date, revision_id)
            para_elem.append(del_elem)


def has_criticmarkup(content: str) -> bool:
    """Check if content contains CriticMarkup syntax.

    Args:
        content: Text content to check

    Returns:
        True if CriticMarkup is present
    """
    return bool(
        _INS_RE.search(content)
        or _DEL_RE.search(content)
        or _SUB_RE.search(content)
    )


def validate_criticmarkup(content: str) -> list[str]:
    """Validate CriticMarkup syntax and return any errors.

    Checks for:
    - Unclosed {++ (missing ++})
    - Unclosed {-- (missing --})
    - Malformed {~~ (missing ~> or ~~})

    Args:
        content: Text content to validate

    Returns:
        List of error messages (empty if no errors)
    """
    errors: list[str] = []
    lines = content.split('\n')

    for line_num, line in enumerate(lines, 1):
        # Check for unclosed insertions
        ins_opens = line.count('{++')
        ins_closes = len(re.findall(r'\+\+\}', line))
        if ins_opens > ins_closes:
            errors.append(f"Line {line_num}: Unclosed insertion {{++ (missing ++}})")

        # Check for unclosed deletions
        del_opens = line.count('{--')
        del_closes = len(re.findall(r'--\}', line))
        if del_opens > del_closes:
            errors.append(f"Line {line_num}: Unclosed deletion {{-- (missing --}})")

        # Check for malformed substitutions
        sub_opens = line.count('{~~')
        sub_closes = len(re.findall(r'~~\}', line))
        if sub_opens > sub_closes:
            errors.append(f"Line {line_num}: Unclosed substitution {{~~ (missing ~~}})")

        # Check for substitution missing ~> separator
        for match in re.finditer(r'\{~~([^~]*)~~\}', line):
            if '~>' not in match.group(1):
                errors.append(f"Line {line_num}: Malformed substitution (missing ~> separator)")

    return errors


def add_hyperlink_to_paragraph(
    paragraph: Any, text: str, url: str,
    bold: bool = False, italic: bool = False
) -> None:
    """Add a hyperlink to a paragraph.

    Creates a proper w:hyperlink element with the text and URL.

    Args:
        paragraph: python-docx Paragraph object
        text: Display text for the hyperlink
        url: URL the hyperlink points to
        bold: Whether the hyperlink text should be bold
        italic: Whether the hyperlink text should be italic
    """
    # Decode percent-encoded URLs for storage in relationship
    decoded_url = unquote(url)

    # Get the document part to create relationships
    part = paragraph.part

    # Create the relationship for the hyperlink
    r_id = part.relate_to(
        decoded_url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )

    # Create the hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a run for the text
    run = OxmlElement('w:r')

    # Add run properties (hyperlink styling: blue color and underline)
    rPr = OxmlElement('w:rPr')

    color = OxmlElement('w:color')
    color.set(qn('w:val'), HYPERLINK_COLOR)
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    # Add bold formatting if requested
    if bold:
        bold_elem = OxmlElement('w:b')
        rPr.append(bold_elem)

    # Add italic formatting if requested
    if italic:
        italic_elem = OxmlElement('w:i')
        rPr.append(italic_elem)

    run.append(rPr)

    # Add the text
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    # Preserve spaces
    text_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    run.append(text_elem)

    # Add run to hyperlink
    hyperlink.append(run)

    # Add hyperlink to paragraph
    paragraph._p.append(hyperlink)


def parse_link_text_formatting(link_text: str) -> tuple[str, bool, bool]:
    """Parse formatting markers from link text.

    Handles markdown formatting markers like **bold**, *italic*, ***bold italic***.

    Args:
        link_text: The raw link text that may contain markdown formatting markers

    Returns:
        Tuple of (plain_text, is_bold, is_italic)
    """
    text = link_text
    is_bold = False
    is_italic = False

    # Check for bold+italic: ***text*** or ___text___
    if (text.startswith("***") and text.endswith("***") and len(text) > 6):
        text = text[3:-3]
        is_bold = True
        is_italic = True
    elif (text.startswith("___") and text.endswith("___") and len(text) > 6):
        text = text[3:-3]
        is_bold = True
        is_italic = True
    # Check for bold: **text** or __text__
    elif (text.startswith("**") and text.endswith("**") and len(text) > 4):
        text = text[2:-2]
        is_bold = True
    elif (text.startswith("__") and text.endswith("__") and len(text) > 4):
        text = text[2:-2]
        is_bold = True
    # Check for italic: *text* or _text_
    elif (text.startswith("*") and text.endswith("*") and len(text) > 2):
        text = text[1:-1]
        is_italic = True
    elif (text.startswith("_") and text.endswith("_") and len(text) > 2):
        text = text[1:-1]
        is_italic = True

    # Unescape brackets that were escaped during extraction
    text = text.replace("\\[", "[").replace("\\]", "]")

    return text, is_bold, is_italic


def add_text_with_hyperlinks(paragraph: Any, content: str) -> None:
    """Add text content to a paragraph, converting markdown hyperlinks.

    Parses [text](url) patterns and creates proper hyperlinks.
    Also handles formatting markers inside link text like [**bold**](url).

    Args:
        paragraph: python-docx Paragraph object
        content: Text content that may contain markdown hyperlinks
    """
    last_end = 0

    for match in HYPERLINK_PATTERN.finditer(content):
        # Add text before the hyperlink
        before_text = content[last_end:match.start()]
        if before_text:
            paragraph.add_run(before_text)

        # Get the link text and URL
        link_text = match.group(1)
        link_url = match.group(2)

        # Parse formatting from link text
        plain_text, is_bold, is_italic = parse_link_text_formatting(link_text)

        # Add the hyperlink with formatting
        add_hyperlink_to_paragraph(paragraph, plain_text, link_url, bold=is_bold, italic=is_italic)

        last_end = match.end()

    # Add any remaining text after the last hyperlink
    if last_end < len(content):
        paragraph.add_run(content[last_end:])


def parse_gfm_alignments(separator_line: str) -> list[str]:
    """Parse GFM alignment indicators from separator line.

    Args:
        separator_line: The separator line like "| --- | :---: | ---: |"

    Returns:
        List of alignments: 'left', 'center', or 'right' for each column
    """
    stripped = separator_line.strip()
    if not stripped:
        return []

    # Remove outer pipes if present
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]

    cells = [c.strip() for c in stripped.split("|")]
    alignments = []

    for cell in cells:
        if not cell:
            continue

        # Detect alignment using GFM_SEPARATOR_PATTERNS
        starts_colon = cell.startswith(":")
        ends_colon = cell.endswith(":")

        detected = GFM_SEPARATOR_PATTERNS.get(
            (starts_colon, ends_colon), DEFAULT_ALIGNMENT
        )
        alignments.append(detected)

    return alignments


def is_table_separator_line(line: str) -> bool:
    """Check if a line is a GFM table separator line.

    Separator lines have the format: | --- | --- | or |:---:|---:|
    """
    stripped = line.strip()
    if not stripped.startswith("|") or not stripped.endswith("|"):
        return False

    # Remove outer pipes and split by pipe
    inner = stripped[1:-1]
    cells = [c.strip() for c in inner.split("|")]

    if not cells:
        return False

    for cell in cells:
        # Each cell should be only dashes with optional colons for alignment
        if not cell:
            continue
        # Remove alignment colons
        cell = cell.strip(":")
        # Should be only dashes
        if not cell or not all(c == "-" for c in cell):
            return False

    return True


def is_table_row(line: str) -> bool:
    """Check if a line could be a GFM table row (starts and ends with pipe)."""
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def parse_markdown_to_blocks(markdown_content: str) -> list[Block]:
    """Parse markdown content into Block objects.

    Args:
        markdown_content: Markdown text

    Returns:
        List of Block objects
    """
    blocks: list[Block] = []
    lines = markdown_content.split("\n")
    block_id = 0
    content_position = 0
    table_index = 0
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped_line = line.strip()

        if not stripped_line:
            i += 1
            continue

        # Skip footnote/endnote definition lines (they don't become blocks)
        if FOOTNOTE_DEF_PATTERN.match(stripped_line):
            i += 1
            continue

        # Check if this is a textbox block
        if stripped_line == "<!-- textbox -->" or stripped_line.startswith("<!-- textbox"):
            textbox_lines = [stripped_line]
            j = i + 1
            while j < len(lines):
                tl = lines[j].strip()
                textbox_lines.append(tl)
                if tl == "<!-- /textbox -->":
                    break
                j += 1

            textbox_content = "\n".join(textbox_lines)

            block = Block(
                id=f"block-{block_id}",
                type="textbox",
                content=textbox_content,
                docx_paragraph_index=block_id,
                content_start=content_position,
                content_end=content_position + len(textbox_content),
                content_hash=hashlib.sha256(textbox_content.encode()).hexdigest(),
            )
            blocks.append(block)
            block_id += 1
            content_position += len(textbox_content) + 1
            i = j + 1
            continue

        # Check if this is the start of a GFM table
        # A table starts with a row that starts/ends with pipes
        # and is followed by a separator line
        if is_table_row(stripped_line) and i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            if is_table_separator_line(next_line):
                # This is a table! Collect all table lines
                table_lines = [stripped_line, next_line]
                j = i + 2
                while j < len(lines) and is_table_row(lines[j].strip()):
                    table_lines.append(lines[j].strip())
                    j += 1

                table_content = "\n".join(table_lines)

                # Parse table dimensions
                num_rows = len(table_lines) - 1  # Subtract separator row
                num_cols = len(split_gfm_row(stripped_line))

                # Validate table dimensions
                if num_rows > MAX_TABLE_ROWS:
                    raise ValueError(f"Table has too many rows ({num_rows}), maximum is {MAX_TABLE_ROWS}")
                if num_cols > MAX_TABLE_COLS:
                    raise ValueError(f"Table has too many columns ({num_cols}), maximum is {MAX_TABLE_COLS}")

                # Parse column alignments from separator line
                column_alignments = parse_gfm_alignments(next_line)

                block = Block(
                    id=f"block-{block_id}",
                    type="table",
                    content=table_content,
                    docx_paragraph_index=-1,
                    content_start=content_position,
                    content_end=content_position + len(table_content),
                    content_hash=hashlib.sha256(table_content.encode()).hexdigest(),
                    table_metadata={
                        "rows": num_rows,
                        "cols": num_cols,
                        "cells": [],
                        "column_alignments": column_alignments,
                        "docx_table_index": table_index
                    }
                )
                blocks.append(block)
                block_id += 1
                table_index += 1
                content_position += len(table_content) + 1
                i = j
                continue

        # Handle other block types
        if stripped_line.startswith("![") and "](" in stripped_line and stripped_line.endswith(")"):
            # Extract alt text and image path from markdown
            alt_start = 2  # after "!["
            alt_end = stripped_line.find("](")
            alt_text = stripped_line[alt_start:alt_end]
            start_idx = alt_end + 2
            end_idx = stripped_line.rfind(")")
            image_path = stripped_line[start_idx:end_idx]

            # Convention: alt text starting with "Chart" identifies chart blocks.
            # This enables round-trip type preservation (extract → build → extract).
            # Edge case: user-authored images with alt text starting with "Chart"
            # (e.g., "Chart of accounts") will be misclassified as chart blocks.
            block_type = "chart" if alt_text.startswith("Chart") else "image"

            block = Block(
                id=f"block-{block_id}",
                type=block_type,
                content=stripped_line,
                docx_paragraph_index=block_id,
                content_start=content_position,
                content_end=content_position + len(stripped_line),
                content_hash=hashlib.sha256(stripped_line.encode()).hexdigest(),
                image_path=image_path,
            )
        elif stripped_line.startswith("#"):
            level = 0
            while level < len(stripped_line) and stripped_line[level] == "#":
                level += 1

            block = Block(
                id=f"block-{block_id}",
                type="heading",
                content=stripped_line,
                docx_paragraph_index=block_id,
                content_start=content_position,
                content_end=content_position + len(stripped_line),
                content_hash=hashlib.sha256(stripped_line.encode()).hexdigest(),
                level=level,
            )
        else:
            block = Block(
                id=f"block-{block_id}",
                type="paragraph",
                content=stripped_line,
                docx_paragraph_index=block_id,
                content_start=content_position,
                content_end=content_position + len(stripped_line),
                content_hash=hashlib.sha256(stripped_line.encode()).hexdigest(),
            )

        blocks.append(block)
        block_id += 1
        content_position += len(stripped_line) + 1
        i += 1

    return blocks


def has_hyperlinks(content: str) -> bool:
    """Check if content contains markdown hyperlinks.

    Args:
        content: Text content to check

    Returns:
        True if hyperlinks are present
    """
    return bool(HYPERLINK_PATTERN.search(content))


def split_gfm_row(line: str) -> list[str]:
    """Split a GFM table row respecting escaped pipes.

    A naive split on '|' would incorrectly split cells containing escaped
    pipes (\\|). This function respects escape sequences.

    Args:
        line: A table row line like "| Name | Test\\|Pipe |"

    Returns:
        List of cell values with escaped characters unescaped
    """
    # Remove outer pipes if present
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]

    # Split respecting escaped pipes
    cells = []
    current: list[str] = []
    i = 0

    while i < len(stripped):
        # Check for escaped pipe
        if i < len(stripped) - 1 and stripped[i] == '\\' and stripped[i + 1] == '|':
            current.append('|')  # Unescape: \| becomes literal |
            i += 2
        elif stripped[i] == '|':
            # Unescaped pipe - cell boundary
            cells.append(''.join(current).strip())
            current = []
            i += 1
        else:
            current.append(stripped[i])
            i += 1

    # Don't forget the last cell
    cells.append(''.join(current).strip())

    return cells


def parse_gfm_table(table_content: str) -> tuple[list[list[str]], list[str]]:
    """Parse GFM table content into a 2D array of cell values and alignments.

    Args:
        table_content: GFM pipe table markdown

    Returns:
        Tuple of (rows, alignments) where rows is a list of rows and alignments
        is a list of column alignments
    """
    lines = table_content.strip().split("\n")
    rows: list[list[str]] = []
    alignments: list[str] = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Check for separator row and extract alignments
        if is_table_separator_line(stripped):
            alignments = parse_gfm_alignments(stripped)
            continue

        # Parse row cells using proper escape-aware splitting
        cells = split_gfm_row(stripped)
        rows.append(cells)

    # Normalize column counts: pad shorter rows / truncate longer rows to match header
    if rows:
        header_col_count = len(rows[0])
        for i in range(1, len(rows)):
            if len(rows[i]) < header_col_count:
                rows[i].extend([""] * (header_col_count - len(rows[i])))
            elif len(rows[i]) > header_col_count:
                rows[i] = rows[i][:header_col_count]

    return rows, alignments


def apply_cell_shading(cell: Any, color: str, pattern: str | None = None) -> None:
    """Apply background shading to a table cell.

    Args:
        cell: python-docx table cell object
        color: Hex color string like 'D9E2F3'
        pattern: Optional pattern fill name like 'diagStripe'
    """
    if not re.match(HEX_COLOR_PATTERN, color):
        return
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing_shd = tcPr.find(qn('w:shd'))
    if existing_shd is not None:
        tcPr.remove(existing_shd)
    shd = OxmlElement('w:shd')
    if pattern and pattern in VALID_PATTERN_FILLS:
        shd.set(qn('w:val'), pattern)
    else:
        shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def apply_cell_borders(cell: Any, borders: dict[str, dict[str, Any]]) -> None:
    """Apply border styles to a table cell.

    Args:
        cell: python-docx table cell object
        borders: Dictionary with border data per side (top/bottom/left/right)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'bottom', 'left', 'right']:
        if side not in borders:
            continue
        border_data = borders[side]
        style = border_data.get('style', 'single')
        if style not in VALID_BORDER_STYLES:
            continue
        width = border_data.get('width', 4)
        if not isinstance(width, int) or width < 0 or width > MAX_BORDER_WIDTH:
            continue
        color = border_data.get('color', 'auto')
        if color != 'auto' and not re.match(HEX_COLOR_PATTERN, color):
            continue
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), style)
        border.set(qn('w:sz'), str(width))
        border.set(qn('w:color'), color)
        border.set(qn('w:space'), '0')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _apply_cell_styles(table: Any, cell_styles: dict[str, dict[str, Any]]) -> None:
    """Apply cell-level styles (shading, borders) to a table.

    Args:
        table: python-docx Table object
        cell_styles: Dictionary keyed by 'row,col' with formatting data
    """
    for key, style_data in cell_styles.items():
        parts = key.split(',')
        if len(parts) != 2:
            continue
        try:
            row_idx, col_idx = int(parts[0]), int(parts[1])
        except ValueError:
            continue
        if row_idx < 0 or row_idx >= len(table.rows):
            continue
        if col_idx < 0 or col_idx >= len(table.columns):
            continue
        cell = table.cell(row_idx, col_idx)
        if 'background_color' in style_data:
            apply_cell_shading(cell, style_data['background_color'],
                               pattern=style_data.get('pattern_fill'))
        if 'borders' in style_data:
            apply_cell_borders(cell, style_data['borders'])


def validate_gfm_table_dimensions(table_content: str) -> None:
    """Pre-validate GFM table dimensions before parsing into memory.

    Counts rows and columns from raw text to reject oversized tables
    before parse_gfm_table() builds a full 2D array in memory.

    Args:
        table_content: GFM pipe table markdown

    Raises:
        ValueError: If table dimensions exceed MAX_TABLE_ROWS or MAX_TABLE_COLS
    """
    lines = table_content.strip().split("\n")

    if len(lines) > MAX_TABLE_LINES:
        raise ValueError(
            f"Table has too many lines ({len(lines)}), maximum is {MAX_TABLE_LINES}"
        )

    # Count data rows (non-empty, non-separator lines)
    data_rows = 0
    first_data_line = None
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if is_table_separator_line(stripped):
            continue
        data_rows += 1
        if first_data_line is None:
            first_data_line = stripped

    if data_rows > MAX_TABLE_ROWS:
        raise ValueError(
            f"Table has too many rows ({data_rows}), maximum is {MAX_TABLE_ROWS}"
        )

    # Count columns from first data line
    if first_data_line:
        num_cols = len(split_gfm_row(first_data_line))
        if num_cols > MAX_TABLE_COLS:
            raise ValueError(
                f"Table has too many columns ({num_cols}), maximum is {MAX_TABLE_COLS}"
            )


def create_table_from_gfm(
    doc: DocumentType,
    table_content: str,
    styles: dict[str, Any],
    block_id: str,
    table_metadata: Optional[dict[str, Any]] = None,
) -> None:
    """Create a table in the document from GFM table content.

    Args:
        doc: Document object
        table_content: GFM pipe table markdown
        styles: Style information dictionary
        block_id: Block ID for looking up table styles
        table_metadata: Optional table metadata with merged_cells, header_rows etc.

    Raises:
        ValueError: If table dimensions exceed reasonable limits
    """
    # Pre-validate dimensions before building 2D array in memory
    validate_gfm_table_dimensions(table_content)

    rows, alignments = parse_gfm_table(table_content)

    if not rows:
        return

    num_rows = len(rows)
    num_cols = max(len(row) for row in rows)

    if num_cols == 0:
        return

    # Create the table
    table = doc.add_table(rows=num_rows, cols=num_cols)

    # Populate cells
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < len(table.columns):
                cell = table.cell(row_idx, col_idx)

                # Apply content with the same priority chain as non-table blocks
                if has_criticmarkup(cell_text):
                    cell.text = ""
                    add_text_with_track_changes(cell.paragraphs[0], cell_text)
                elif has_hyperlinks(cell_text):
                    cell.text = ""
                    add_text_with_hyperlinks(cell.paragraphs[0], cell_text)
                elif re.search(r'\*\*.+?\*\*|\*[^*]+?\*', cell_text):
                    cell.text = ""
                    apply_inline_formatting(cell.paragraphs[0], cell_text)
                else:
                    cell.text = cell_text

                # Apply alignment to the cell's paragraph
                if alignments and col_idx < len(alignments):
                    alignment = alignments[col_idx]
                    if cell.paragraphs and alignment in ALIGNMENT_STRING_TO_ENUM:
                        cell.paragraphs[0].alignment = ALIGNMENT_STRING_TO_ENUM[alignment]

    # Apply header rows from table_metadata
    if table_metadata:
        header_row_count = table_metadata.get("header_rows", 1)
        for row_idx in range(min(header_row_count, num_rows)):
            row = table.rows[row_idx]
            trPr = row._tr.get_or_add_trPr()
            tblHeader = OxmlElement('w:tblHeader')
            trPr.append(tblHeader)

    # Apply merged cells from table_metadata if available
    if table_metadata:
        merged_cells = table_metadata.get("merged_cells", [])
        for merge in merged_cells:
            start_row = merge.get("start_row", 0)
            start_col = merge.get("start_col", 0)
            row_span = merge.get("row_span", 1)
            col_span = merge.get("col_span", 1)
            end_row = start_row + row_span - 1
            end_col = start_col + col_span - 1
            # Validate bounds
            if (0 <= start_row <= end_row < num_rows and
                    0 <= start_col <= end_col < num_cols):
                table.cell(start_row, start_col).merge(
                    table.cell(end_row, end_col)
                )

    # Apply column widths from styles if available
    block_style = styles.get("block_styles", {}).get(block_id, {})
    table_formatting = block_style.get("table_formatting", {})

    if table_formatting:
        column_widths = table_formatting.get("column_widths", [])
        for col_idx, width in enumerate(column_widths):
            if col_idx < len(table.columns):
                # Width is in inches, convert to EMUs
                table.columns[col_idx].width = Inches(width)

        # Apply cell-level styles (shading, borders)
        cell_styles = table_formatting.get("cell_styles", {})
        if cell_styles:
            _apply_cell_styles(table, cell_styles)


def _apply_block_formatting(para: Any, block_style: dict[str, Any]) -> None:
    """Apply formatting from block_style to a paragraph.

    Args:
        para: python-docx Paragraph object
        block_style: Style dictionary with font_name, font_size, alignment
    """
    if not block_style:
        return

    if "font_name" in block_style and para.style:
        para.style.font.name = block_style["font_name"]
    if "font_size" in block_style and para.style:
        para.style.font.size = Pt(block_style["font_size"])

    alignment = block_style.get("alignment", DEFAULT_ALIGNMENT)
    if alignment in ALIGNMENT_STRING_TO_ENUM:
        para.alignment = ALIGNMENT_STRING_TO_ENUM[alignment]


def _parse_footnote_definitions(content_md: str) -> dict[int, str]:
    """Parse [^N]: text definitions from markdown content.

    Returns:
        Dict mapping marker number to definition text.
    """
    defs = {}
    for line in content_md.split("\n"):
        m = FOOTNOTE_DEF_PATTERN.match(line.strip())
        if m:
            defs[int(m.group(1))] = m.group(2)
    return defs


def _has_footnote_refs(text: str) -> bool:
    """Check if text contains [^N] footnote references."""
    return bool(FOOTNOTE_REF_PATTERN.search(text))


def _get_or_create_notes_part(doc, note_tag, rel_type, content_type, part_path):
    """Get or create a footnotes/endnotes XML part in the document.

    Args:
        doc: Document object
        note_tag: XML tag for notes (e.g. "w:footnotes" or "w:endnotes")
        rel_type: Relationship type constant
        content_type: Content type constant
        part_path: Part URI path (e.g. "/word/footnotes.xml")
    """
    for rel in doc.part.rels.values():
        if rel.reltype == rel_type:
            return rel.target_part

    _NOTE_SINGULAR = {"w:footnotes": "w:footnote", "w:endnotes": "w:endnote"}
    note_singular = _NOTE_SINGULAR[note_tag]
    nsmap = {"w": WORDPROCESSINGML_NS, "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships"}
    root = etree.Element(qn(note_tag), nsmap=nsmap)
    for sep_id, sep_type in [("-1", "separator"), ("0", "continuationSeparator")]:
        note = etree.SubElement(root, qn(note_singular))
        note.set(qn("w:type"), sep_type)
        note.set(qn("w:id"), sep_id)
        p = etree.SubElement(note, qn("w:p"))
        r = etree.SubElement(p, qn("w:r"))
        if sep_type == "separator":
            etree.SubElement(r, qn("w:separator"))
        else:
            etree.SubElement(r, qn("w:continuationSeparator"))

    xml_bytes = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    part_name = PackURI(part_path)
    notes_part = Part(part_name, content_type, xml_bytes, doc.part.package)
    # Cache parsed element so _add_note_to_part can accumulate mutations
    # without re-parsing from blob on each call.
    notes_part._element = root
    doc.part.relate_to(notes_part, rel_type)
    return notes_part


def _get_or_create_footnotes_part(doc):
    """Get or create the footnotes.xml part in the document."""
    return _get_or_create_notes_part(doc, "w:footnotes", FOOTNOTES_RT, FOOTNOTES_CT, "/word/footnotes.xml")


def _get_or_create_endnotes_part(doc):
    """Get or create the endnotes.xml part in the document."""
    return _get_or_create_notes_part(doc, "w:endnotes", ENDNOTES_RT, ENDNOTES_CT, "/word/endnotes.xml")


def _parse_formatted_segments(text: str) -> list[tuple[str, bool, bool]]:
    """Parse markdown inline formatting into (text, bold, italic) segments.

    Returns a list of tuples, each containing the plain text and its
    bold/italic state. Handles **bold**, *italic*, and ***bold italic***.
    """
    segments: list[tuple[str, bool, bool]] = []
    last_end = 0
    for m in _INLINE_FORMAT_PATTERN.finditer(text):
        # Add any plain text before this match
        if m.start() > last_end:
            segments.append((text[last_end:m.start()], False, False))
        stars = len(m.group(1))
        content = m.group(2)
        bold = stars >= 2
        italic = stars % 2 == 1
        segments.append((content, bold, italic))
        last_end = m.end()
    # Add any remaining plain text
    if last_end < len(text):
        segments.append((text[last_end:], False, False))
    # If no formatting found, return the whole text as plain
    if not segments:
        segments.append((text, False, False))
    return segments


def _add_formatted_runs_to_note(p, text):
    """Add text runs to a footnote/endnote paragraph, preserving bold/italic formatting.

    The leading space (OOXML convention after the ref mark) is prepended to the first segment.
    """
    segments = _parse_formatted_segments(text)
    for idx, (seg_text, bold, italic) in enumerate(segments):
        if idx == 0:
            seg_text = f" {seg_text}"
        text_run = etree.SubElement(p, qn("w:r"))
        if bold or italic:
            rPr = etree.SubElement(text_run, qn("w:rPr"))
            if bold:
                etree.SubElement(rPr, qn("w:b"))
            if italic:
                etree.SubElement(rPr, qn("w:i"))
        t = etree.SubElement(text_run, qn("w:t"))
        t.set(qn("xml:space"), "preserve")
        t.text = seg_text


def _add_note_to_part(part, note_id, text, note_tag, ref_tag, style_name):
    """Add a footnote or endnote element to a notes part.

    Uses a lazy _element cache: on first call, parses blob into an element tree
    and caches it on the part. Subsequent calls mutate the cached tree directly.
    The blob is updated after each mutation so Part serializes correctly.
    """
    if hasattr(part, '_element'):
        root = part._element
    else:
        root = etree.fromstring(part.blob)
        part._element = root

    note = etree.SubElement(root, qn(note_tag))
    note.set(qn("w:id"), str(note_id))
    p = etree.SubElement(note, qn("w:p"))
    ref_run = etree.SubElement(p, qn("w:r"))
    rPr = etree.SubElement(ref_run, qn("w:rPr"))
    style = etree.SubElement(rPr, qn("w:rStyle"))
    style.set(qn("w:val"), style_name)
    etree.SubElement(ref_run, qn(ref_tag))
    _add_formatted_runs_to_note(p, text)

    part._blob = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    return note_id


def _add_footnote_to_part(part, note_id, text):
    """Add a footnote element to the footnotes part."""
    return _add_note_to_part(part, note_id, text, "w:footnote", "w:footnoteRef", "FootnoteReference")


def _add_endnote_to_part(part, note_id, text):
    """Add an endnote element to the endnotes part."""
    return _add_note_to_part(part, note_id, text, "w:endnote", "w:endnoteRef", "EndnoteReference")


def _insert_footnote_reference(run, note_id):
    """Insert a w:footnoteReference element into a run."""
    rPr = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "FootnoteReference")
    rPr.append(style)
    run._element.insert(0, rPr)
    ref = OxmlElement("w:footnoteReference")
    ref.set(qn("w:id"), str(note_id))
    run._element.append(ref)


def _insert_endnote_reference(run, note_id):
    """Insert a w:endnoteReference element into a run."""
    rPr = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "EndnoteReference")
    rPr.append(style)
    run._element.insert(0, rPr)
    ref = OxmlElement("w:endnoteReference")
    ref.set(qn("w:id"), str(note_id))
    run._element.append(ref)


def _add_text_with_footnotes(paragraph, text, doc, footnote_defs, footnote_meta):
    """Add text to a paragraph, replacing [^N] with proper footnote references.

    Args:
        paragraph: python-docx Paragraph object
        text: Text content potentially containing [^N] markers
        doc: Document object for creating footnote parts
        footnote_defs: Dict mapping marker number to definition text
        footnote_meta: Dict mapping marker number to note type info
    """
    parts = FOOTNOTE_REF_PATTERN.split(text)
    # parts alternates: text, number, text, number, ...
    for idx, part in enumerate(parts):
        if idx % 2 == 0:
            # Regular text
            if part:
                paragraph.add_run(part)
        else:
            # Footnote reference number
            marker_num = int(part)
            def_text = footnote_defs.get(marker_num, "")
            note_info = footnote_meta.get(marker_num, {})
            note_type = note_info.get("note_type", "footnote")

            if note_type == "endnote":
                en_part = _get_or_create_endnotes_part(doc)
                note_id = _add_endnote_to_part(en_part, marker_num, def_text)
                run = paragraph.add_run()
                _insert_endnote_reference(run, note_id)
            else:
                fn_part = _get_or_create_footnotes_part(doc)
                note_id = _add_footnote_to_part(fn_part, marker_num, def_text)
                run = paragraph.add_run()
                _insert_footnote_reference(run, note_id)


COLUMN_BREAK_MARKER = "<!-- column-break -->"


def _add_paragraph_with_column_breaks(doc: DocumentType, content: str) -> Any:
    """Add a paragraph, inserting column breaks where markers appear.

    Applies inline formatting (bold/italic/links) to text segments between
    column break markers.

    Args:
        doc: Document to add paragraph to
        content: Text content possibly containing <!-- column-break --> markers

    Returns:
        The created paragraph
    """
    parts = content.split(COLUMN_BREAK_MARKER)
    para = doc.add_paragraph()
    first_text = parts[0].strip()
    if first_text:
        apply_inline_formatting(para, first_text)
    for part in parts[1:]:
        # Add column break
        run = para.add_run()
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'column')
        run._element.append(br)
        # Add text after break with inline formatting
        text = part.strip()
        if text:
            apply_inline_formatting(para, text)
    return para


def _build_cols_element(section: SectionProperties) -> Any:
    """Build a w:cols XML element from SectionProperties.

    Args:
        section: SectionProperties with column configuration

    Returns:
        OxmlElement for w:cols
    """
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), str(section.column_count))
    if section.column_spacing is not None:
        cols.set(qn('w:space'), str(section.column_spacing))
    cols.set(qn('w:equalWidth'), '1' if section.equal_width else '0')
    if not section.equal_width:
        if section.columns:
            for col_def in section.columns:
                col_elem = OxmlElement('w:col')
                col_elem.set(qn('w:w'), str(col_def.width))
                if col_def.space is not None:
                    col_elem.set(qn('w:space'), str(col_def.space))
                cols.append(col_elem)
    return cols


def _apply_sections_to_doc(doc: DocumentType, sections: list[SectionProperties]) -> None:
    """Apply section properties (columns) to a document.

    Mid-document sections get a sectPr in the last paragraph's pPr.
    The final section gets a sectPr as direct child of w:body.

    Args:
        doc: Document to modify
        sections: List of SectionProperties in document order
    """
    if not sections:
        return

    body = doc.element.body
    # Build combined list of block elements (paragraphs AND tables) in document
    # order. end_block_index is a block index covering both types, so indexing
    # into only w:p elements would be wrong when tables are present.
    block_elements = [
        child for child in body
        if child.tag.endswith('}p') or child.tag.endswith('}tbl')
    ]

    for i, section in enumerate(sections):
        is_last = (i == len(sections) - 1)

        if is_last:
            # Final section: apply to body-level sectPr
            sect_pr = body.find(qn('w:sectPr'))
            if sect_pr is None:
                sect_pr = OxmlElement('w:sectPr')
                body.append(sect_pr)
            # Remove existing cols if any
            existing_cols = sect_pr.find(qn('w:cols'))
            if existing_cols is not None:
                sect_pr.remove(existing_cols)
            if section.column_count > 1 or not section.equal_width:
                sect_pr.append(_build_cols_element(section))
        else:
            # Mid-document section: insert sectPr in last block's pPr.
            # If the block is a table, insert sectPr in the last paragraph
            # of that table (OOXML requires sectPr inside a w:pPr).
            end_idx = section.end_block_index
            if end_idx is not None and end_idx < len(block_elements):
                block_elem = block_elements[end_idx]
                # If block is a table, find the last paragraph inside it
                if block_elem.tag.endswith('}tbl'):
                    paragraphs_in_table = block_elem.findall('.//' + qn('w:p'))
                    if paragraphs_in_table:
                        p_elem = paragraphs_in_table[-1]
                    else:
                        continue  # No paragraph to attach sectPr to
                else:
                    p_elem = block_elem
                pPr = p_elem.find(qn('w:pPr'))
                if pPr is None:
                    pPr = OxmlElement('w:pPr')
                    p_elem.insert(0, pPr)
                sect_pr = OxmlElement('w:sectPr')
                if section.column_count > 1 or not section.equal_width:
                    sect_pr.append(_build_cols_element(section))
                pPr.append(sect_pr)


def create_docx_from_blocks(
    blocks: list[Block],
    styles: dict[str, Any],
    assets_dir: Optional[Path] = None,
    style_id_remap: Optional[dict[str, str]] = None,
    default_tc_author: Optional[str] = None,
    default_tc_date: Optional[str] = None,
    footnote_defs: Optional[dict[int, str]] = None,
    footnote_meta: Optional[dict[int, dict]] = None,
    content_md: Optional[str] = None,
    sections: Optional[list[SectionProperties]] = None,
) -> DocumentType:
    """Create a Word document from Block objects.

    Args:
        blocks: List of Block objects
        styles: Style information dictionary
        assets_dir: Optional path to assets directory for image files
        style_id_remap: Optional mapping from new block IDs to old block IDs for style lookup
        default_tc_author: Optional default author for CriticMarkup track changes (sync use case)
        default_tc_date: Optional default date for CriticMarkup track changes (sync use case)
        sections: Optional list of SectionProperties for column layouts

    Returns:
        Document object
    """
    doc = Document()
    para = None  # Track current paragraph for styling

    # Auto-parse footnote definitions if not provided
    if footnote_defs is None:
        footnote_defs = {}
        if content_md:
            footnote_defs = _parse_footnote_definitions(content_md)
    if footnote_meta is None:
        footnote_meta = {}

    for block in blocks:
        if block.type == "heading" and block.level:
            style_name = f"Heading {block.level}"
            text = block.content.lstrip("#").strip()

            if has_criticmarkup(text):
                para = doc.add_paragraph(style=style_name)
                if default_tc_author is not None:
                    add_text_with_track_changes(para, text, default_author=default_tc_author, default_date=default_tc_date or "")
                else:
                    add_text_with_track_changes(para, text, block.track_changes)
            elif _has_footnote_refs(text) and footnote_defs:
                para = doc.add_paragraph(style=style_name)
                _add_text_with_footnotes(para, text, doc, footnote_defs, footnote_meta)
            elif has_hyperlinks(text):
                para = doc.add_paragraph(style=style_name)
                add_text_with_hyperlinks(para, text)
            else:
                para = doc.add_paragraph(text, style=style_name)
        elif block.type == "table":
            # For tables, use remapped ID for style lookup if available
            table_style_id = block.id
            if style_id_remap and block.id in style_id_remap:
                table_style_id = style_id_remap[block.id]
            create_table_from_gfm(doc, block.content, styles, table_style_id, block.table_metadata)
            para = None
        elif block.type == "textbox":
            if block.text_box_metadata and "drawing_xml" in block.text_box_metadata:
                para = doc.add_paragraph()
                run = para.add_run()
                drawing_elem = etree.fromstring(block.text_box_metadata["drawing_xml"])
                run._element.append(drawing_elem)
            else:
                inner_text = _extract_textbox_inner_content(block.content)
                para = doc.add_paragraph(inner_text)
        elif block.type == "image":
            if block.image_path and assets_dir:
                image_filename = block.image_path.split("/")[-1]
                image_file_path = assets_dir / image_filename

                if image_file_path.exists():
                    para = doc.add_paragraph()
                    run = para.add_run()
                    run.add_picture(str(image_file_path), width=Inches(DEFAULT_IMAGE_WIDTH_INCHES))
                else:
                    para = doc.add_paragraph(f"[Missing image: {block.image_path}]")
            else:
                para = doc.add_paragraph("[Image]")
        elif block.type == "chart":
            # Charts reconstruct as images for now (full-fidelity in JON-108)
            if block.image_path and assets_dir:
                image_filename = block.image_path.split("/")[-1]
                image_file_path = assets_dir / image_filename

                if image_file_path.exists():
                    para = doc.add_paragraph()
                    run = para.add_run()
                    run.add_picture(str(image_file_path), width=Inches(DEFAULT_IMAGE_WIDTH_INCHES))
                else:
                    para = doc.add_paragraph(f"[Missing chart: {block.image_path}]")
            else:
                para = doc.add_paragraph("[Chart: no preview available]")
        else:
            # Paragraph and all other block types
            content = block.content
            if COLUMN_BREAK_MARKER in content:
                para = _add_paragraph_with_column_breaks(doc, content)
            elif has_criticmarkup(content):
                para = doc.add_paragraph()
                if default_tc_author is not None:
                    add_text_with_track_changes(para, content, default_author=default_tc_author, default_date=default_tc_date or "")
                else:
                    add_text_with_track_changes(para, content, block.track_changes)
            elif _has_footnote_refs(content) and footnote_defs:
                para = doc.add_paragraph()
                _add_text_with_footnotes(para, content, doc, footnote_defs, footnote_meta)
            elif has_hyperlinks(content):
                para = doc.add_paragraph()
                add_text_with_hyperlinks(para, content)
            else:
                para = doc.add_paragraph(content)

        # Apply styling - use remapped ID if available
        if para is not None:
            if style_id_remap and block.id in style_id_remap:
                lookup_id = style_id_remap[block.id]
            else:
                lookup_id = block.id
            block_style = styles.get("block_styles", {}).get(lookup_id, {})
            _apply_block_formatting(para, block_style)

    # Apply section column properties
    if sections:
        _apply_sections_to_doc(doc, sections)

    return doc


def _populate_header_footer(header_footer: Any, paragraphs: list[dict], assets_dir: Optional[Path] = None) -> None:
    """Populate a header or footer with paragraph content.

    Plain text only — inline formatting (bold, italic, hyperlinks) is not
    currently preserved in header/footer content.
    """
    if not paragraphs:
        return
    header_footer.is_linked_to_previous = False
    for i, para_dict in enumerate(paragraphs):
        if i == 0:
            para = header_footer.paragraphs[0]
        else:
            para = header_footer.add_paragraph()
        if para_dict.get("type") == "image" and para_dict.get("image_path"):
            image_filename = para_dict["image_path"].split("/")[-1]
            if assets_dir:
                image_file_path = assets_dir / image_filename
                if image_file_path.exists():
                    run = para.add_run()
                    run.add_picture(str(image_file_path), width=Inches(DEFAULT_IMAGE_WIDTH_INCHES))
                    continue
            warnings.warn(f"header/footer image not found: {image_filename}", RuntimeWarning, stacklevel=2)
            para.text = f"[Image: {para_dict.get('content', '')}]"
        else:
            para.text = para_dict.get("content", "")


def apply_sections_to_document(doc: DocumentType, sections_data: list[dict], assets_dir: Optional[Path] = None) -> None:
    """Apply section metadata (headers, footers, page setup) to a document."""

    if not sections_data:
        return
    if any(s.get("page_setup", {}).get("odd_and_even_pages") for s in sections_data):
        doc.settings.odd_and_even_pages_header_footer = True
    for section_idx, section_meta in enumerate(sections_data):
        if section_idx < len(doc.sections):
            section = doc.sections[section_idx]
        else:
            # Defensive: metadata has more sections than document (e.g. mismatched structure.json)
            section = doc.add_section()
        page_setup = section_meta.get("page_setup", {})
        if page_setup:
            orientation = page_setup.get("orientation", "portrait")
            if orientation == "landscape":
                section.orientation = WD_ORIENT.LANDSCAPE
            else:
                section.orientation = WD_ORIENT.PORTRAIT
            for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin",
                         "page_width", "page_height", "header_distance", "footer_distance"):
                value = page_setup.get(attr)
                if value is not None:
                    setattr(section, attr, value)
            section.different_first_page_header_footer = page_setup.get("different_first_page", False)
        variants = [
            ("header_default", section.header),
            ("header_first", section.first_page_header),
            ("header_even", section.even_page_header),
            ("footer_default", section.footer),
            ("footer_first", section.first_page_footer),
            ("footer_even", section.even_page_footer),
        ]
        for key, target in variants:
            content = section_meta.get(key, [])
            if content:
                _populate_header_footer(target, content, assets_dir)


def build_docx_from_sidedoc(sidedoc_path: str, output_path: str) -> None:
    """Build a Word document from a sidedoc archive or directory.

    Args:
        sidedoc_path: Path to .sidedoc directory or .sdoc/.sidedoc ZIP
        output_path: Path for output .docx file
    """
    with SidedocStore.open(sidedoc_path) as store:
        if store.is_zip:
            click.echo("Tip: Use `sidedoc unpack` to convert to directory format for editing.", err=True)

        content_md = store.read_text("content.md")
        styles_data = store.read_json("styles.json")

        assets_dir = store.assets_dir if store.list_assets() else None

        blocks = parse_markdown_to_blocks(content_md)

        footnote_defs = _parse_footnote_definitions(content_md)
        footnote_meta: dict[int, dict] = {}
        hf_sections_data: list[dict] = []

        # Read structure.json for track changes, footnotes, and section properties
        sections: list[SectionProperties] | None = None
        if store.has_file("structure.json"):
            structure_data = store.read_json("structure.json")

            # Enrich blocks with track changes data
            structure_blocks = structure_data.get("blocks", [])
            hf_sections_data = structure_data.get("hf_sections", [])

            # Build footnote metadata from structure.json
            footnotes_data = structure_data.get("footnotes", {})
            for fn_id_str, fn_info in footnotes_data.items():
                footnote_meta[int(fn_id_str)] = fn_info

            # Also extract from block footnote_references
            for struct_block in structure_blocks:
                refs = struct_block.get("footnote_references")
                if refs:
                    for ref in refs:
                        mid = ref.get("note_id")
                        if mid and int(mid) not in footnote_meta:
                            footnote_meta[int(mid)] = {"note_type": ref.get("note_type", "footnote")}
            for block, struct_block in zip(blocks, structure_blocks):
                if "track_changes" in struct_block and struct_block["track_changes"]:
                    block.track_changes = [
                        TrackChange(
                            type=tc["type"],
                            start=tc["start"],
                            end=tc["end"],
                            author=tc["author"],
                            date=tc["date"],
                            revision_id=tc.get("revision_id", ""),
                            deleted_text=tc.get("deleted_text"),
                        )
                        for tc in struct_block["track_changes"]
                    ]
                # Transfer text box metadata if present
                if "text_box_metadata" in struct_block and struct_block["text_box_metadata"]:
                    block.text_box_metadata = struct_block["text_box_metadata"]

            # Read section properties (column layouts)
            sections = deserialize_sections(structure_data)

        doc = create_docx_from_blocks(blocks, styles_data, assets_dir, footnote_defs=footnote_defs, footnote_meta=footnote_meta, content_md=content_md, sections=sections)

        if hf_sections_data:
            apply_sections_to_document(doc, hf_sections_data, assets_dir)
        doc.save(output_path)
