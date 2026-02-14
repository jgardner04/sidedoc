"""Sync module for matching blocks and updating documents."""

import hashlib
import json
from pathlib import Path
from typing import Optional, Any
from docx import Document
from sidedoc.models import Block
from sidedoc.utils import get_iso_timestamp, compute_similarity
from sidedoc.constants import (
    SIMILARITY_THRESHOLD,
)
from sidedoc.reconstruct import create_table_from_gfm, apply_inline_formatting, _apply_block_formatting

# Default author for new track changes created during sync
DEFAULT_SYNC_AUTHOR = "Sidedoc AI"


def match_blocks(
    old_blocks: list[Block], new_blocks: list[Block]
) -> dict[str, Block]:
    """Match old blocks to new blocks using hashes, positions, and similarity.

    Matching strategy:
    1. First pass: Match by content hash (unchanged blocks)
    2. Second pass: Match by type, position, and similarity threshold (edited blocks)
       - Blocks at the same position with the same type are only matched if
         their content similarity meets the SIMILARITY_THRESHOLD
       - This prevents incorrectly treating "delete + add" as "edit"

    Args:
        old_blocks: List of blocks from previous version (structure.json)
        new_blocks: List of blocks from edited content.md

    Returns:
        Dictionary mapping old block IDs to their corresponding new blocks.
        Unmatched old blocks indicate deletions.
        Unmatched new blocks indicate additions.
    """
    matches: dict[str, Block] = {}
    used_new_blocks: set[int] = set()

    # First pass: Match by content hash (unchanged blocks)
    # Why this works: Identical content hashes mean the block wasn't edited at all,
    # so we can confidently match it regardless of position changes. This is the
    # fastest and most reliable matching strategy.
    for old_block in old_blocks:
        for i, new_block in enumerate(new_blocks):
            if i in used_new_blocks:
                continue
            if old_block.content_hash == new_block.content_hash:
                matches[old_block.id] = new_block
                used_new_blocks.add(i)
                break  # Move to next old_block once we find a match

    # Second pass: Match by type, position, and similarity (edited blocks)
    # Why we need this: After the first pass, we have blocks that changed but might
    # be edits (modify) rather than delete+add operations. We only match blocks at
    # the SAME position to avoid incorrectly matching unrelated blocks.
    # This handles blocks that were edited but are at the exact same position
    # We only match if:
    # 1. Positions are identical (same index in both lists)
    # 2. Types match (heading stays heading, paragraph stays paragraph)
    # 3. Content similarity meets threshold (distinguishes edit from delete+add)
    remaining_old = [b for b in old_blocks if b.id not in matches]
    remaining_new_indices = [
        i for i in range(len(new_blocks)) if i not in used_new_blocks
    ]

    for old_block in remaining_old:
        old_idx = old_blocks.index(old_block)

        # Only match if there's a new block at the exact same position with same type
        # Why position matters: If a block moved, it's safer to treat it as delete+add
        # rather than risk matching to the wrong block
        if old_idx in remaining_new_indices:
            new_block = new_blocks[old_idx]
            if old_block.type == new_block.type:
                # Check content similarity to distinguish edits from delete+add
                # Why similarity threshold: If someone deleted "Hello" and added "Goodbye"
                # at the same position, similarity is low and we treat it as separate
                # operations. If they changed "Hello world" to "Hello there", similarity
                # is high and we preserve the formatting as an edit.
                similarity = compute_similarity(old_block.content, new_block.content)
                if similarity >= SIMILARITY_THRESHOLD:
                    matches[old_block.id] = new_block
                    used_new_blocks.add(old_idx)
                    remaining_new_indices.remove(old_idx)

    return matches


def remap_styles(styles_data: dict[str, Any], matches: dict[str, Block]) -> dict[str, Any]:
    """Remap block IDs in styles_data based on match mapping.

    When blocks are added/removed during sync, block IDs shift. This function
    updates styles_data so block_styles keys reflect the new block IDs.

    Args:
        styles_data: Original styles dict with block_styles keyed by old block IDs
        matches: Mapping from old block IDs to new Block objects

    Returns:
        Updated styles_data with remapped block_styles keys
    """
    old_block_styles = styles_data.get("block_styles", {})
    new_block_styles: dict[str, Any] = {}

    for old_id, new_block in matches.items():
        if old_id in old_block_styles:
            new_block_styles[new_block.id] = old_block_styles[old_id]

    return {
        **styles_data,
        "block_styles": new_block_styles,
    }


def _create_reverse_mapping(matches: dict[str, Block]) -> dict[str, str]:
    """Create reverse mapping from new block IDs to old block IDs.

    Args:
        matches: Dictionary mapping old block IDs to new blocks

    Returns:
        Dictionary mapping new block IDs to old block IDs
    """
    new_to_old: dict[str, str] = {}
    for old_id, new_block in matches.items():
        new_to_old[new_block.id] = old_id
    return new_to_old


def _get_block_style(
    block: Block,
    new_to_old: dict[str, str],
    styles: dict[str, Any]
) -> dict[str, Any]:
    """Get the style for a block if it was matched to an old block.

    Args:
        block: The new block to get style for
        new_to_old: Mapping from new block IDs to old block IDs
        styles: Style information dictionary with block_styles

    Returns:
        Style dictionary for the block, or empty dict if no style
    """
    old_block_id = new_to_old.get(block.id)
    if old_block_id:
        block_styles: dict[str, Any] = styles.get("block_styles", {})
        result: dict[str, Any] = block_styles.get(old_block_id, {})
        return result
    return {}


def _create_paragraph_from_block(
    doc: Any, block: Block, styles: dict[str, Any], old_block_id: str | None = None
) -> Optional[Any]:
    """Create a paragraph or table from a block based on its type.

    Args:
        doc: python-docx Document object
        block: Block to create paragraph/table from
        styles: Style information dictionary with block_styles
        old_block_id: Resolved old block ID for style lookup during sync

    Returns:
        python-docx Paragraph object, or None for tables
    """
    if block.type == "heading" and block.level:
        text = block.content.lstrip("#").strip()
        style_name = f"Heading {block.level}"
        return doc.add_paragraph(text, style=style_name)
    elif block.type == "table":
        # Use old block ID for style lookup if available (during sync),
        # otherwise fall back to new block ID
        style_id = old_block_id if old_block_id else block.id
        create_table_from_gfm(doc, block.content, styles, style_id, block.table_metadata)
        return None
    elif block.type == "paragraph":
        if "**" in block.content or "*" in block.content:
            para = doc.add_paragraph()
            apply_inline_formatting(para, block.content)
            return para
        else:
            return doc.add_paragraph(block.content)
    else:
        return doc.add_paragraph(block.content)


def generate_updated_docx(
    new_blocks: list[Block],
    matches: dict[str, Block],
    styles: dict[str, Any],
    output_path: str,
) -> None:
    """Generate an updated docx file from new blocks.

    This function creates a new docx with content from new_blocks.
    - Matched blocks preserve their formatting from styles
    - New blocks receive default formatting based on type
    - Deleted blocks are omitted (not in new_blocks)
    - Inline formatting from markdown is applied
    - Table blocks create actual Table objects

    Args:
        new_blocks: List of new Block objects from edited content.md
        matches: Dictionary mapping old block IDs to new blocks
        styles: Style information dictionary with block_styles
        output_path: Path where docx should be saved
    """
    doc = Document()
    new_to_old = _create_reverse_mapping(matches)

    for block in new_blocks:
        old_block_id = new_to_old.get(block.id)
        block_style = _get_block_style(block, new_to_old, styles)
        para = _create_paragraph_from_block(doc, block, styles, old_block_id)
        if para is not None:
            _apply_block_formatting(para, block_style)

    doc.save(output_path)


def update_sidedoc_metadata(
    sidedoc_path: str,
    new_blocks: list[Block],
    new_content: str,
    matches: Optional[dict[str, Block]] = None,
) -> None:
    """Update metadata files in a sidedoc directory after sync.

    Regenerates structure.json, remaps styles.json, and updates manifest.json.
    Only works with .sidedoc/ directories (ZIP archives must be unpacked first).

    Args:
        sidedoc_path: Path to .sidedoc directory
        new_blocks: Updated list of Block objects from edited content
        new_content: New markdown content
        matches: Optional mapping from old block IDs to new blocks for style remapping

    Raises:
        ValueError: If sidedoc_path is not a directory
    """
    path = Path(sidedoc_path)
    if not path.is_dir():
        raise ValueError(
            f"Cannot update metadata for '{sidedoc_path}': not a directory. "
            "Run `sidedoc unpack` to convert to directory format first."
        )
    _update_directory_metadata(sidedoc_path, new_blocks, new_content, matches)


def _build_structure_data(new_blocks: list[Block]) -> dict:
    """Build structure.json data from blocks."""
    from sidedoc.package import block_to_structure_dict

    return {
        "blocks": [block_to_structure_dict(block) for block in new_blocks]
    }


def _update_manifest(old_manifest: dict, new_content: str) -> dict:
    """Build updated manifest data."""
    content_hash = hashlib.sha256(new_content.encode()).hexdigest()
    return {
        "sidedoc_version": old_manifest["sidedoc_version"],
        "created_at": old_manifest["created_at"],
        "modified_at": get_iso_timestamp(),
        "source_file": old_manifest["source_file"],
        "source_hash": old_manifest["source_hash"],
        "content_hash": content_hash,
        "generator": old_manifest["generator"],
    }


def _update_directory_metadata(
    sidedoc_path: str,
    new_blocks: list[Block],
    new_content: str,
    matches: Optional[dict[str, Block]] = None,
) -> None:
    """Update metadata in a sidedoc directory."""
    dir_path = Path(sidedoc_path)

    # Read existing styles
    styles_data = json.loads((dir_path / "styles.json").read_text(encoding="utf-8"))

    # Remap styles if matches provided
    if matches:
        styles_data = remap_styles(styles_data, matches)

    # Read old manifest if it exists
    manifest_path = dir_path / "manifest.json"
    if manifest_path.exists():
        old_manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    else:
        old_manifest = {
            "sidedoc_version": "1.0.0",
            "created_at": get_iso_timestamp(),
            "source_file": "unknown",
            "source_hash": "",
            "generator": "sidedoc-cli",
        }

    structure_data = _build_structure_data(new_blocks)
    manifest_data = _update_manifest(old_manifest, new_content)

    # Write to .tmp files first, then rename for atomicity
    tmp_files = []
    try:
        for name, data in [
            ("content.md", new_content),
            ("structure.json", json.dumps(structure_data, indent=2)),
            ("styles.json", json.dumps(styles_data, indent=2)),
            ("manifest.json", json.dumps(manifest_data, indent=2)),
        ]:
            tmp_path = dir_path / f"{name}.tmp"
            tmp_files.append((tmp_path, dir_path / name))
            tmp_path.write_text(data, encoding="utf-8")

        # Rename all atomically
        for tmp_path, final_path in tmp_files:
            tmp_path.replace(final_path)
    except Exception:
        # Clean up any tmp files on failure
        for tmp_path, _ in tmp_files:
            tmp_path.unlink(missing_ok=True)
        raise


def sync_sidedoc_to_docx(
    sidedoc_path: str,
    output_path: str,
    author: Optional[str] = None,
) -> None:
    """Sync a sidedoc archive to a Word document with CriticMarkup support.

    Reads the content.md from the sidedoc, parses CriticMarkup syntax, and
    generates a docx with proper track changes (w:ins and w:del elements).

    Args:
        sidedoc_path: Path to .sidedoc file
        output_path: Path for output .docx file
        author: Author name for new track changes (default: 'Sidedoc AI')
    """
    from sidedoc.reconstruct import (
        parse_markdown_to_blocks,
        has_criticmarkup,
        add_text_with_track_changes,
    )

    if author is None:
        author = DEFAULT_SYNC_AUTHOR

    # Get current timestamp for new track changes
    sync_date = get_iso_timestamp()

    # Read sidedoc contents (supports both directory and ZIP formats)
    from sidedoc.store import SidedocStore

    with SidedocStore.open(sidedoc_path) as store:
        content_md = store.read_text("content.md")
        styles_data = store.read_json("styles.json")
        structure_data = store.read_json("structure.json")

    # Parse markdown to blocks
    blocks = parse_markdown_to_blocks(content_md)

    # Create document
    doc = Document()

    for block in blocks:
        para = None
        if block.type == "heading" and block.level:
            style_name = f"Heading {block.level}"
            text = block.content.lstrip("#").strip()

            if has_criticmarkup(text):
                para = doc.add_paragraph(style=style_name)
                add_text_with_track_changes(para, text, default_author=author, default_date=sync_date)
            else:
                para = doc.add_paragraph(text, style=style_name)
        elif block.type == "table":
            create_table_from_gfm(doc, block.content, styles_data, block.id, block.table_metadata)
        else:
            content = block.content

            if has_criticmarkup(content):
                para = doc.add_paragraph()
                add_text_with_track_changes(para, content, default_author=author, default_date=sync_date)
            else:
                para = doc.add_paragraph(content)

        # Apply styling if available (only for paragraph-based blocks)
        if para is not None:
            block_style = styles_data.get("block_styles", {}).get(block.id, {})
            _apply_block_formatting(para, block_style)

    doc.save(output_path)
