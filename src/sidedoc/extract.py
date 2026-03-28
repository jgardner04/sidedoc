"""Extract content from Word documents to sidedoc format."""

import hashlib
import io
from pathlib import Path
from typing import Any, Literal, Optional
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from lxml import etree  # type: ignore[import-untyped]
from PIL import Image
from sidedoc.models import Block, ColumnDefinition, SectionProperties, Style, TrackChange
from sidedoc.constants import (
    MAX_IMAGE_SIZE,
    MAX_TABLE_ROWS,
    MAX_TABLE_COLS,
    EMUS_PER_INCH,
    ALIGNMENT_NUMERIC_TO_STRING,
    GFM_ALIGNMENT_TO_SEPARATOR,
    DEFAULT_ALIGNMENT,
    WORDPROCESSINGML_NS,
)


# XML namespaces used in Office Open XML documents
RELATIONSHIPS_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
DRAWINGML_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
WP_DRAWING_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
CHART_NS = "http://schemas.openxmlformats.org/drawingml/2006/chart"


def wrap_formatting(text: str, bold: bool, italic: bool) -> str:
    """Wrap text with markdown bold/italic markers.

    Args:
        text: Plain text to wrap
        bold: Whether to apply bold
        italic: Whether to apply italic

    Returns:
        Text wrapped with appropriate markdown markers
    """
    if bold and italic:
        return f"***{text}***"
    elif bold:
        return f"**{text}**"
    elif italic:
        return f"*{text}*"
    return text


def format_hyperlink_md(text: str, url: str, bold: bool, italic: bool) -> str:
    """Build a markdown hyperlink with optional bold/italic formatting.

    Args:
        text: Display text for the link
        url: URL for the link
        bold: Whether the link text should be bold
        italic: Whether the link text should be italic

    Returns:
        Markdown hyperlink string like [**text**](url)
    """
    escaped_text = escape_markdown_link_text(text)
    encoded_url = encode_url_for_markdown(url)
    formatted_text = wrap_formatting(escaped_text, bold, italic)
    return f"[{formatted_text}]({encoded_url})"


def generate_block_id(index: int) -> str:
    """Generate a unique block ID.

    Args:
        index: Block index in document

    Returns:
        Unique block ID
    """
    return f"block-{index}"


def compute_content_hash(content: str) -> str:
    """Compute SHA256 hash of content.

    Args:
        content: Block content text

    Returns:
        Hex digest of content hash
    """
    return hashlib.sha256(content.encode("utf-8")).hexdigest()


def validate_image(image_bytes: bytes, expected_extension: str) -> tuple[bool, str]:
    """Validate image size, format, and integrity.

    Args:
        image_bytes: Raw image data
        expected_extension: Expected file extension (e.g., 'png', 'jpg')

    Returns:
        Tuple of (is_valid, error_message). If valid, error_message is empty.
    """
    # Check size first - fastest validation, prevents processing huge files
    # Why check size: Protects against ZIP bombs and malicious documents with
    # extremely large embedded images that could cause memory issues
    if len(image_bytes) > MAX_IMAGE_SIZE:
        size_mb = len(image_bytes) / (1024 * 1024)
        max_mb = MAX_IMAGE_SIZE / (1024 * 1024)
        return False, f"Image exceeds maximum size ({size_mb:.1f}MB > {max_mb:.0f}MB)"

    # Pass through EMF/WMF formats without PIL validation.
    # PIL cannot open these metafile formats, but they are valid image types
    # commonly used as cached chart fallback images in Word documents.
    if expected_extension.lower() in ("emf", "wmf"):
        return True, ""

    # Validate image data and format using PIL
    try:
        # First pass: verify image integrity
        # Why verify: Detects corrupted or malicious image data before we try to
        # process it further. This prevents crashes from malformed images.
        # Note: We must load the image twice because PIL's verify() method makes
        # the image object unusable for further operations. This is a known PIL
        # limitation. For large images near the 10MB limit, this doubles memory
        # usage temporarily, but it's necessary to ensure both integrity and
        # format validation.
        img = Image.open(io.BytesIO(image_bytes))
        img.verify()  # Verify that it's a valid image

        # Second pass: reopen to check format (verify() makes the image unusable)
        # Why check format: Detects extension spoofing (e.g., malware.exe renamed
        # to image.png). We verify the actual image format matches the extension.
        img = Image.open(io.BytesIO(image_bytes))

        # Map extensions to PIL formats
        # Why mapping: PIL uses format names like "JPEG" while file extensions are
        # "jpg" or "jpeg". We need to normalize for comparison.
        extension_to_format = {
            'png': 'PNG',
            'jpg': 'JPEG',
            'jpeg': 'JPEG',
            'gif': 'GIF',
            'bmp': 'BMP',
            'tiff': 'TIFF',
            'tif': 'TIFF',
        }

        expected_format = extension_to_format.get(expected_extension.lower())
        if expected_format and img.format != expected_format:
            return False, f"Image format mismatch (extension: {expected_extension}, actual: {img.format})"

        return True, ""

    except (Image.UnidentifiedImageError, OSError, ValueError) as e:
        # UnidentifiedImageError: PIL cannot identify the image format
        # OSError: File-related errors (truncated file, etc.)
        # ValueError: Invalid image data
        # Why catch these: Documents may contain corrupted images or non-image data
        # embedded as images. We want to skip these gracefully with an error message
        # rather than crash the entire extraction.
        return False, f"Invalid or corrupted image data: {str(e)}"


def _extract_blip_image(
    blip: Any, doc_part: Any, filename_prefix: str, counter: int
) -> Optional[tuple[str, str, bytes, str]]:
    """Extract image data from a blip element via relationship lookup.

    Args:
        blip: lxml element for a:blip
        doc_part: Document part for accessing relationships
        filename_prefix: Prefix for the filename (e.g., "image" or "chart")
        counter: Counter for generating unique filenames

    Returns:
        Tuple of (filename, extension, image_bytes, error_message) or None.
    """
    r_embed = blip.get(f'{{{RELATIONSHIPS_NS}}}embed')
    if not r_embed or r_embed not in doc_part.rels:
        return None

    image_part = doc_part.rels[r_embed].target_part
    extension = image_part.partname.split('.')[-1]
    image_bytes = image_part.blob

    is_valid, error_message = validate_image(image_bytes, extension)
    image_filename = f"{filename_prefix}{counter}.{extension}"

    return (image_filename, extension, image_bytes, error_message)


def extract_image_from_paragraph(paragraph: Any, doc_part: Any, image_counter: int) -> Optional[tuple[str, str, bytes, str]]:
    """Check if paragraph contains an image and extract it.

    Args:
        paragraph: python-docx paragraph object
        doc_part: Document part for accessing relationships
        image_counter: Counter for generating unique image names

    Returns:
        Tuple of (image_filename, image_extension, image_bytes, error_message) if image found.
        If image is valid, error_message is empty. If invalid, error_message contains the reason.
        Returns None if no image found.
    """
    # Check for drawing elements (images) in paragraph runs
    # Why check runs: Word documents store images inside run elements within paragraphs.
    # A paragraph may have multiple runs, and any of them could contain an image.
    for run in paragraph.runs:
        drawing_elems = run._element.findall(f'.//{{{WORDPROCESSINGML_NS}}}drawing')
        for drawing in drawing_elems:
            blips = drawing.findall(f'.//{{{DRAWINGML_NS}}}blip')
            for blip in blips:
                result = _extract_blip_image(blip, doc_part, "image", image_counter)
                if result is not None:
                    return result

    return None


def extract_chart_from_paragraph(
    paragraph: Any, doc_part: Any, image_counter: int
) -> Optional[tuple[str, str, bytes, str, str]]:
    """Check if paragraph contains a chart and extract its cached fallback image.

    Charts in OOXML may be embedded in two forms:
    1. mc:AlternateContent — mc:Choice has c:chart, mc:Fallback has cached image blip
    2. Flat w:drawing — c:chart in graphicData with no fallback image

    Args:
        paragraph: python-docx paragraph object
        doc_part: Document part for accessing relationships
        image_counter: Counter for generating unique image names

    Returns:
        Tuple of (image_filename, extension, image_bytes, error_message, chart_rel_id)
        if chart found with cached image. Returns tuple with empty bytes
        ("", "", b"", "", chart_rel_id) if chart found but no cached image.
        Returns None if no chart found.
    """
    for run in paragraph.runs:
        # Case 1: mc:AlternateContent with chart in mc:Choice
        alt_contents = run._element.findall(f'{{{MC_NS}}}AlternateContent')
        for alt in alt_contents:
            choice = alt.find(f'{{{MC_NS}}}Choice')
            if choice is None:
                continue

            charts = choice.findall(f'.//{{{CHART_NS}}}chart')
            if not charts:
                continue

            chart_rel_id = charts[0].get(f'{{{RELATIONSHIPS_NS}}}id', "")

            # Extract cached fallback image from mc:Fallback
            fallback = alt.find(f'{{{MC_NS}}}Fallback')
            if fallback is not None:
                blips = fallback.findall(f'.//{{{DRAWINGML_NS}}}blip')
                for blip in blips:
                    result = _extract_blip_image(blip, doc_part, "chart", image_counter)
                    if result is not None:
                        return (*result, chart_rel_id)

            # Chart found but no cached image
            return ("", "", b"", "", chart_rel_id)

        # Case 2: Flat w:drawing with c:chart in graphicData (no mc:AlternateContent)
        drawing_elems = run._element.findall(f'{{{WORDPROCESSINGML_NS}}}drawing')
        for drawing in drawing_elems:
            charts = drawing.findall(f'.//{{{CHART_NS}}}chart')
            if charts:
                chart_rel_id = charts[0].get(f'{{{RELATIONSHIPS_NS}}}id', "")
                return ("", "", b"", "", chart_rel_id)

    return None


def extract_textbox_from_paragraph(paragraph: Any) -> Optional[list[dict[str, Any]]]:
    """Check if paragraph contains text boxes or shapes with text and extract them.

    Looks for wps:txbxContent elements within drawing elements. These can be
    either text boxes (txBox="1") or shapes with text content.

    Args:
        paragraph: python-docx paragraph object

    Returns:
        List of text box info dicts with keys: texts, anchor_type, width, height,
        position_h, position_v, fill_color, border_color, drawing_xml.
        Returns None if no text boxes found.
    """
    results = []

    for run in paragraph.runs:
        drawing_elems = run._element.findall(f'{{{WORDPROCESSINGML_NS}}}drawing')
        for drawing in drawing_elems:
            txbx_contents = drawing.findall(f'.//{{{WPS_NS}}}txbxContent')
            if not txbx_contents:
                continue

            # Skip images (drawings with blip elements)
            blips = drawing.findall(f'.//{{{DRAWINGML_NS}}}blip')
            if blips:
                continue

            anchor = drawing.find(f'{{{WP_DRAWING_NS}}}anchor')
            inline = drawing.find(f'{{{WP_DRAWING_NS}}}inline')
            anchor_type = "anchor" if anchor is not None else "inline"
            wrapper = anchor if anchor is not None else inline

            if wrapper is None:
                continue

            extent = wrapper.find(f'{{{WP_DRAWING_NS}}}extent')
            width = int(extent.get('cx', '0')) if extent is not None else 0
            height = int(extent.get('cy', '0')) if extent is not None else 0

            position_h = None
            position_v = None
            if anchor_type == "anchor":
                pos_h_elem = wrapper.find(f'{{{WP_DRAWING_NS}}}positionH')
                pos_v_elem = wrapper.find(f'{{{WP_DRAWING_NS}}}positionV')
                if pos_h_elem is not None:
                    offset_h = pos_h_elem.find(f'{{{WP_DRAWING_NS}}}posOffset')
                    if offset_h is not None and offset_h.text:
                        position_h = int(offset_h.text)
                if pos_v_elem is not None:
                    offset_v = pos_v_elem.find(f'{{{WP_DRAWING_NS}}}posOffset')
                    if offset_v is not None and offset_v.text:
                        position_v = int(offset_v.text)

            fill_color = None
            border_color = None
            wsp = drawing.find(f'.//{{{WPS_NS}}}wsp')
            if wsp is not None:
                spPr = wsp.find(f'{{{WPS_NS}}}spPr')
                if spPr is not None:
                    solid_fill = spPr.find(f'{{{DRAWINGML_NS}}}solidFill')
                    if solid_fill is not None:
                        srgb = solid_fill.find(f'{{{DRAWINGML_NS}}}srgbClr')
                        if srgb is not None:
                            fill_color = srgb.get('val')

                    ln = spPr.find(f'{{{DRAWINGML_NS}}}ln')
                    if ln is not None:
                        ln_fill = ln.find(f'{{{DRAWINGML_NS}}}solidFill')
                        if ln_fill is not None:
                            ln_srgb = ln_fill.find(f'{{{DRAWINGML_NS}}}srgbClr')
                            if ln_srgb is not None:
                                border_color = ln_srgb.get('val')

            drawing_xml = etree.tostring(drawing, encoding='unicode')

            for txbx in txbx_contents:
                texts = []
                # txbxContent contains WordProcessingML elements (w:p/w:r/w:t),
                # not DrawingML elements (a:p/a:r/a:t)
                for w_p in txbx.findall(f'{{{WORDPROCESSINGML_NS}}}p'):
                    parts = []
                    for w_r in w_p.findall(f'{{{WORDPROCESSINGML_NS}}}r'):
                        for w_t in w_r.findall(f'{{{WORDPROCESSINGML_NS}}}t'):
                            if w_t.text:
                                parts.append(w_t.text)
                    if parts:
                        texts.append("".join(parts))

                if texts:
                    results.append({
                        "texts": texts,
                        "anchor_type": anchor_type,
                        "width": width,
                        "height": height,
                        "position_h": position_h,
                        "position_v": position_v,
                        "fill_color": fill_color,
                        "border_color": border_color,
                        "drawing_xml": drawing_xml,
                    })

    return results if results else None


def encode_url_for_markdown(url: str) -> str:
    """Encode a URL for safe inclusion in markdown link syntax.

    Percent-encodes parentheses and spaces that would break markdown [text](url) format.

    Args:
        url: The URL to encode

    Returns:
        URL safe for markdown link syntax
    """
    # Encode characters that break markdown link syntax
    # Spaces need encoding
    result = url.replace(" ", "%20")
    # Parentheses MUST be encoded - they break markdown link syntax
    # [text](url_(with)_parens) is ambiguous: first ) closes the link
    result = result.replace("(", "%28")
    result = result.replace(")", "%29")
    return result


def escape_markdown_link_text(text: str) -> str:
    """Escape special markdown characters in link text.

    Args:
        text: The link text to escape

    Returns:
        Escaped text safe for markdown link syntax
    """
    # Escape brackets which would break link syntax
    result = text.replace("[", "\\[").replace("]", "\\]")
    return result


def get_hyperlink_url(hyperlink_elem: Any, doc_part: Any) -> Optional[str]:
    """Extract the URL from a hyperlink XML element.

    Args:
        hyperlink_elem: The w:hyperlink XML element
        doc_part: Document part for accessing relationships

    Returns:
        The URL string, or None if not found
    """
    # Get the relationship ID from the hyperlink element
    r_id = hyperlink_elem.get(qn('r:id'))
    if r_id and r_id in doc_part.rels:
        rel = doc_part.rels[r_id]
        # External hyperlinks have target_ref (the URL)
        if hasattr(rel, '_target') and rel._target:
            return str(rel._target)
        # Some versions use target_ref
        if hasattr(rel, 'target_ref') and rel.target_ref:
            return str(rel.target_ref)
    return None


def is_formatting_enabled(rPr: Any, format_tag: str) -> bool:
    """Check if a formatting element is enabled in run properties.

    In Office Open XML, formatting can be:
    - Present without val attribute: enabled (e.g., <w:b/>)
    - Present with val="0" or val="false": disabled (e.g., <w:b w:val="0"/>)
    - Present with val="1" or val="true": enabled (e.g., <w:b w:val="1"/>)
    - Absent: inherit from style (we treat as disabled)

    Args:
        rPr: The w:rPr run properties element
        format_tag: The tag name without namespace (e.g., 'b', 'i', 'u')

    Returns:
        True if the formatting is enabled, False otherwise
    """
    if rPr is None:
        return False

    elem = rPr.find(f'{{{WORDPROCESSINGML_NS}}}{format_tag}')
    if elem is None:
        return False

    # Check the val attribute
    val = elem.get(qn('w:val'))
    if val is None:
        # Element present without val attribute means enabled
        return True
    # Check for explicit false values
    if val.lower() in ('0', 'false', 'none'):
        return False
    return True


def extract_hyperlink_text_and_formatting(hyperlink_elem: Any) -> tuple[str, bool, bool]:
    """Extract text and formatting from a hyperlink XML element.

    Args:
        hyperlink_elem: The w:hyperlink XML element

    Returns:
        Tuple of (text, is_bold, is_italic)
    """
    text_parts = []
    is_bold = False
    is_italic = False

    # Find all w:r (run) elements within the hyperlink
    for run_elem in hyperlink_elem.findall(f'.//{{{WORDPROCESSINGML_NS}}}r'):
        # Check for bold/italic in run properties
        rPr = run_elem.find(f'{{{WORDPROCESSINGML_NS}}}rPr')
        if is_formatting_enabled(rPr, 'b'):
            is_bold = True
        if is_formatting_enabled(rPr, 'i'):
            is_italic = True

        # Get text from w:t elements
        for text_elem in run_elem.findall(f'{{{WORDPROCESSINGML_NS}}}t'):
            if text_elem.text:
                text_parts.append(text_elem.text)

    return "".join(text_parts), is_bold, is_italic


def detect_track_changes(docx_path: str) -> bool:
    """Detect whether a Word document contains track changes.

    Scans the document for any w:ins (insertion) or w:del (deletion) elements
    to determine if the document has revision tracking.

    Args:
        docx_path: Path to .docx file

    Returns:
        True if the document contains any track changes, False otherwise
    """
    doc = Document(docx_path)

    for paragraph in doc.paragraphs:
        para_elem = paragraph._element

        # Check for any insertion elements
        ins_elements = para_elem.findall(f'.//{{{WORDPROCESSINGML_NS}}}ins')
        if ins_elements:
            return True

        # Check for any deletion elements
        del_elements = para_elem.findall(f'.//{{{WORDPROCESSINGML_NS}}}del')
        if del_elements:
            return True

    return False


def extract_track_change_metadata(element: Any) -> tuple[str, str, str]:
    """Extract author, date, and revision_id from a track change element.

    Args:
        element: The w:ins or w:del XML element

    Returns:
        Tuple of (author, date, revision_id)
    """
    author = element.get(f'{{{WORDPROCESSINGML_NS}}}author') or ""
    date = element.get(f'{{{WORDPROCESSINGML_NS}}}date') or ""
    revision_id = element.get(f'{{{WORDPROCESSINGML_NS}}}id') or ""
    return author, date, revision_id


def _extract_element_text(element: Any, tag: str) -> str:
    """Extract concatenated text from all matching child elements.

    Args:
        element: Parent XML element to search within
        tag: Local tag name (e.g. 't', 'delText')

    Returns:
        Concatenated text content
    """
    return "".join(t.text for t in element.findall(f'.//{{{WORDPROCESSINGML_NS}}}{tag}') if t.text)


def extract_insertion_text(ins_element: Any) -> str:
    """Extract text content from a w:ins element."""
    return _extract_element_text(ins_element, "t")


def extract_deletion_text(del_element: Any) -> str:
    """Extract text content from a w:del element."""
    return _extract_element_text(del_element, "delText")


def extract_paragraph_content(
    para_elem: Any,
    doc_part: Any = None,
    mode: Literal["normal", "accept_all", "track_changes"] = "normal",
    footnote_counter: int = 0,
) -> tuple[str, list[dict[str, Any]] | None, list[TrackChange] | None, list[dict[str, Any]], int]:
    """Extract content from a paragraph XML element.

    Unified function that handles all three extraction modes:
    - "normal": Ignores track change elements, extracts plain content
    - "accept_all": Accepts insertions as plain text, skips deletions
    - "track_changes": Converts insertions/deletions to CriticMarkup syntax

    Args:
        para_elem: The paragraph's XML element
        doc_part: Document part for accessing hyperlink relationships
        mode: Extraction mode - "normal", "accept_all", or "track_changes"

    Returns:
        Tuple of (markdown_content, inline_formatting, track_changes, footnote_refs, footnote_counter)
        - markdown_content: The extracted markdown string
        - inline_formatting: List of inline formatting dicts, or None
        - track_changes: List of TrackChange objects (only populated when mode == "track_changes"), or None
        - footnote_refs: List of footnote reference dicts found in this paragraph
        - footnote_counter: Updated footnote counter after processing this paragraph
    """
    markdown_parts: list[str] = []
    inline_formatting: list[dict[str, Any]] = []
    valid_modes = {"normal", "accept_all", "track_changes"}
    if mode not in valid_modes:
        raise ValueError(f"Unknown mode: {mode!r}. Must be one of {valid_modes}")
    track_changes: list[TrackChange] = []
    footnote_refs: list[dict[str, Any]] = []
    plain_text_position = 0

    for child in para_elem:
        tag = child.tag

        if tag == f'{{{WORDPROCESSINGML_NS}}}ins':
            if mode == "normal":
                pass  # Ignore insertions in normal mode
            elif mode == "accept_all":
                inserted_text = extract_insertion_text(child)
                if inserted_text:
                    markdown_parts.append(inserted_text)
                    plain_text_position += len(inserted_text)
            elif mode == "track_changes":
                inserted_text = extract_insertion_text(child)
                if inserted_text:
                    author, date, revision_id = extract_track_change_metadata(child)
                    markdown_parts.append(f"{{++{inserted_text}++}}")
                    track_changes.append(TrackChange(
                        type="insertion",
                        start=plain_text_position,
                        end=plain_text_position + len(inserted_text),
                        author=author,
                        date=date,
                        revision_id=revision_id,
                    ))
                    plain_text_position += len(inserted_text)

        elif tag == f'{{{WORDPROCESSINGML_NS}}}del':
            if mode == "normal" or mode == "accept_all":
                pass  # Skip deletions in normal and accept_all modes
            elif mode == "track_changes":
                deleted_text = extract_deletion_text(child)
                if deleted_text:
                    author, date, revision_id = extract_track_change_metadata(child)
                    markdown_parts.append(f"{{--{deleted_text}--}}")
                    track_changes.append(TrackChange(
                        type="deletion",
                        start=plain_text_position,
                        end=plain_text_position + len(deleted_text),
                        author=author,
                        date=date,
                        revision_id=revision_id,
                        deleted_text=deleted_text,
                    ))
                    plain_text_position += len(deleted_text)

        elif tag == f'{{{WORDPROCESSINGML_NS}}}hyperlink':
            if doc_part is None:
                continue

            url = get_hyperlink_url(child, doc_part)
            if url is None:
                text, is_bold, is_italic = extract_hyperlink_text_and_formatting(child)
                if text:
                    markdown_parts.append(text)
                    plain_text_position += len(text)
                continue

            text, is_bold, is_italic = extract_hyperlink_text_and_formatting(child)
            if not text:
                continue

            link_md = format_hyperlink_md(text, url, is_bold, is_italic)
            markdown_parts.append(link_md)

            inline_formatting.append({
                "type": "hyperlink",
                "start": plain_text_position,
                "end": plain_text_position + len(text),
                "url": url
            })

            plain_text_position += len(text)

        elif tag == f'{{{WORDPROCESSINGML_NS}}}r':
            # Check for footnote/endnote references in this run
            fn_ref = child.find(f'{{{WORDPROCESSINGML_NS}}}footnoteReference')
            en_ref = child.find(f'{{{WORDPROCESSINGML_NS}}}endnoteReference')
            if fn_ref is not None:
                note_id = fn_ref.get(qn('w:id'))
                if note_id and int(note_id) > 0:
                    footnote_counter += 1
                    marker = f"[^{footnote_counter}]"
                    markdown_parts.append(marker)
                    footnote_refs.append({
                        "note_id": footnote_counter,
                        "note_type": "footnote",
                        "marker": marker,
                        "original_id": note_id,
                    })
                    plain_text_position += len(marker)
                continue
            if en_ref is not None:
                note_id = en_ref.get(qn('w:id'))
                if note_id and int(note_id) > 0:
                    footnote_counter += 1
                    marker = f"[^{footnote_counter}]"
                    markdown_parts.append(marker)
                    footnote_refs.append({
                        "note_id": footnote_counter,
                        "note_type": "endnote",
                        "marker": marker,
                        "original_id": note_id,
                    })
                    plain_text_position += len(marker)
                continue

            text_parts = []
            for run_child in child:
                run_child_tag = run_child.tag
                if run_child_tag == f'{{{WORDPROCESSINGML_NS}}}t':
                    if run_child.text:
                        text_parts.append(run_child.text)
                elif run_child_tag == f'{{{WORDPROCESSINGML_NS}}}br':
                    br_type = run_child.get(qn('w:type'))
                    if br_type == 'column':
                        text_parts.append('\n<!-- column-break -->\n')
                    else:
                        text_parts.append('\n')

            text = "".join(text_parts)
            if not text:
                continue

            rPr = child.find(f'{{{WORDPROCESSINGML_NS}}}rPr')
            is_bold = is_formatting_enabled(rPr, 'b')
            is_italic = is_formatting_enabled(rPr, 'i')
            is_underline = is_formatting_enabled(rPr, 'u')

            markdown_text = wrap_formatting(text, is_bold, is_italic)

            if is_underline:
                inline_formatting.append({
                    "type": "underline",
                    "start": plain_text_position,
                    "end": plain_text_position + len(text),
                    "underline": True
                })

            markdown_parts.append(markdown_text)
            plain_text_position += len(text)

    markdown_content = "".join(markdown_parts)
    return (
        markdown_content,
        inline_formatting if inline_formatting else None,
        track_changes if track_changes else None,
        footnote_refs,
        footnote_counter,
    )


# Backwards-compatible wrappers used by existing tests
def extract_paragraph_accept_all(
    para_elem: Any, doc_part: Any = None
) -> tuple[str, list[dict[str, Any]] | None, list[TrackChange] | None]:
    """Extract content from a paragraph, accepting all track changes."""
    content, fmt, tc, _fn, _counter = extract_paragraph_content(para_elem, doc_part, mode="accept_all")
    return content, fmt, tc


def get_cell_alignment(cell: Any) -> str:
    """Extract horizontal alignment from a cell.

    Args:
        cell: python-docx cell object

    Returns:
        Alignment string: 'left', 'center', or 'right'
    """
    # Check if cell has paragraphs with alignment
    if cell.paragraphs:
        para = cell.paragraphs[0]
        if para.alignment is not None:
            # Note: GFM tables only support left/center/right, so we map
            # justify to left. This is different from ALIGNMENT_NUMERIC_TO_STRING.
            cell_alignment_map = {
                0: "left",      # WD_ALIGN_PARAGRAPH.LEFT
                1: "center",    # WD_ALIGN_PARAGRAPH.CENTER
                2: "right",     # WD_ALIGN_PARAGRAPH.RIGHT
                3: "left",      # WD_ALIGN_PARAGRAPH.JUSTIFY → left for GFM
            }
            return cell_alignment_map.get(para.alignment, DEFAULT_ALIGNMENT)

    return DEFAULT_ALIGNMENT


def get_column_alignments(table: Any) -> list[str]:
    """Extract column alignments from first row cells.

    Args:
        table: python-docx Table object

    Returns:
        List of alignment strings for each column
    """
    if not table.rows:
        return []

    return [get_cell_alignment(cell) for cell in table.rows[0].cells]


def alignment_to_gfm_separator(alignment: str) -> str:
    """Convert alignment to GFM separator indicator.

    Args:
        alignment: 'left', 'center', or 'right'

    Returns:
        GFM alignment indicator like ':---', ':---:', or '---:'
    """
    return GFM_ALIGNMENT_TO_SEPARATOR.get(alignment, GFM_ALIGNMENT_TO_SEPARATOR[DEFAULT_ALIGNMENT])


def escape_cell_content_for_gfm(text: str) -> str:
    """Escape special characters in cell content for GFM table.

    Args:
        text: Cell text content

    Returns:
        Escaped text safe for GFM table cells
    """
    # Escape pipe characters that would break table syntax
    result = text.replace("|", "\\|")
    # Replace newlines with <br> or space (newlines break table rows)
    result = result.replace("\n", " ")
    return result


def _extract_cell_text_with_formatting(cell: Any, doc_part: Any = None) -> str:
    """Extract cell text with inline formatting (bold, italic) and hyperlinks as markdown.

    Args:
        cell: python-docx table cell object
        doc_part: Document part for resolving hyperlink relationships

    Returns:
        Cell text with markdown formatting markers
    """
    if not cell.paragraphs:
        return ""

    parts = []
    for paragraph in cell.paragraphs:
        para_elem = paragraph._element
        para_parts: list[str] = []

        for child in para_elem:
            tag = child.tag
            if tag == f'{{{WORDPROCESSINGML_NS}}}hyperlink':
                if doc_part is None:
                    text, is_bold, is_italic = extract_hyperlink_text_and_formatting(child)
                    if text:
                        para_parts.append(wrap_formatting(text, is_bold, is_italic))
                    continue

                url = get_hyperlink_url(child, doc_part)
                if url is None:
                    text, is_bold, is_italic = extract_hyperlink_text_and_formatting(child)
                    if text:
                        para_parts.append(wrap_formatting(text, is_bold, is_italic))
                    continue

                text, is_bold, is_italic = extract_hyperlink_text_and_formatting(child)
                if text:
                    para_parts.append(format_hyperlink_md(text, url, is_bold, is_italic))

            elif tag == f'{{{WORDPROCESSINGML_NS}}}r':
                text_parts = []
                for run_child in child:
                    run_child_tag = run_child.tag
                    if run_child_tag == f'{{{WORDPROCESSINGML_NS}}}t':
                        if run_child.text:
                            text_parts.append(run_child.text)

                text = "".join(text_parts)
                if not text:
                    continue

                rPr = child.find(f'{{{WORDPROCESSINGML_NS}}}rPr')
                is_bold = is_formatting_enabled(rPr, 'b')
                is_italic = is_formatting_enabled(rPr, 'i')

                para_parts.append(wrap_formatting(text, is_bold, is_italic))

        if para_parts:
            parts.append("".join(para_parts))

    return " ".join(parts).strip()


def table_to_gfm(table: Any, doc_part: Any = None) -> str:
    """Convert a python-docx table to GFM pipe table syntax.

    Args:
        table: python-docx Table object
        doc_part: Document part for resolving hyperlink relationships

    Returns:
        GFM pipe table markdown string
    """
    rows = []

    # Get column alignments
    alignments = get_column_alignments(table)

    for row_idx, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            # Extract text with inline formatting, then escape for GFM
            cell_text = _extract_cell_text_with_formatting(cell, doc_part)
            cell_text = escape_cell_content_for_gfm(cell_text)
            cells.append(cell_text)

        # Create pipe-separated row
        row_line = "| " + " | ".join(cells) + " |"
        rows.append(row_line)

        # Add separator row after header (first row)
        if row_idx == 0:
            # Create separator with alignment indicators
            separator_cells = []
            for i in range(len(cells)):
                alignment = alignments[i] if i < len(alignments) else "left"
                separator_cells.append(alignment_to_gfm_separator(alignment))
            separator_line = "| " + " | ".join(separator_cells) + " |"
            rows.append(separator_line)

    return "\n".join(rows)


def extract_merged_cells(table: Any) -> list[dict[str, int]]:
    """Detect merged cells in a table and return merge regions.

    Detects horizontal merges via gridSpan and vertical merges via vMerge.
    Uses raw XML tc elements to avoid python-docx's merged cell aliasing.

    Args:
        table: python-docx Table object

    Returns:
        List of merge region dicts with start_row, start_col, row_span, col_span
    """
    num_rows = len(table.rows)
    if num_rows == 0:
        return []

    merged_cells: list[dict[str, int]] = []

    # Build a grid of actual tc elements from each tr
    # python-docx's table.cell() returns the primary cell for merged regions,
    # so we iterate raw XML to get the actual tc elements per row.
    for row_idx in range(num_rows):
        tr = table.rows[row_idx]._tr
        tcs = tr.findall(qn('w:tc'))
        for col_idx, tc in enumerate(tcs):
            tcPr = tc.find(qn('w:tcPr'))

            # Check horizontal span (gridSpan)
            col_span = 1
            if tcPr is not None:
                gridSpan = tcPr.find(qn('w:gridSpan'))
                if gridSpan is not None:
                    col_span = int(gridSpan.get(qn('w:val'), '1'))

            # Check if this is a vertical merge start (vMerge=restart)
            row_span = 1
            if tcPr is not None:
                vMerge = tcPr.find(qn('w:vMerge'))
                if vMerge is not None:
                    val = vMerge.get(qn('w:val'))
                    if val == 'restart':
                        # Count continuation cells below by checking raw tc elements
                        for check_row in range(row_idx + 1, num_rows):
                            check_tr = table.rows[check_row]._tr
                            check_tcs = check_tr.findall(qn('w:tc'))
                            if col_idx < len(check_tcs):
                                check_tcPr = check_tcs[col_idx].find(qn('w:tcPr'))
                                if check_tcPr is not None:
                                    check_vMerge = check_tcPr.find(qn('w:vMerge'))
                                    if check_vMerge is not None and check_vMerge.get(qn('w:val')) is None:
                                        row_span += 1
                                        continue
                            break
                    elif val is None:
                        # This is a continuation cell — skip it
                        continue

            if col_span > 1 or row_span > 1:
                merged_cells.append({
                    "start_row": row_idx,
                    "start_col": col_idx,
                    "row_span": row_span,
                    "col_span": col_span,
                })

    return merged_cells


def extract_table_metadata(table: Any, table_index: int) -> dict[str, Any]:
    """Extract metadata from a table for storage in structure.json.

    Args:
        table: python-docx Table object
        table_index: Index of this table in the document

    Returns:
        Dictionary with table metadata including rows, cols, cells, column_alignments, and docx_table_index
    """
    num_rows = len(table.rows)
    num_cols = len(table.columns) if num_rows > 0 else 0

    # Validate table dimensions to prevent memory exhaustion
    if num_rows > MAX_TABLE_ROWS:
        raise ValueError(f"Table has too many rows ({num_rows}), maximum is {MAX_TABLE_ROWS}")
    if num_cols > MAX_TABLE_COLS:
        raise ValueError(f"Table has too many columns ({num_cols}), maximum is {MAX_TABLE_COLS}")

    # Build cells metadata - 2D array of cell info
    cells: list[list[dict[str, Any]]] = []

    for row_idx, row in enumerate(table.rows):
        row_cells: list[dict[str, Any]] = []
        for col_idx, cell in enumerate(row.cells):
            cell_text = cell.text.strip()
            cell_hash = compute_content_hash(cell_text)
            row_cells.append({
                "row": row_idx,
                "col": col_idx,
                "content_hash": cell_hash
            })
        cells.append(row_cells)

    # Extract column alignments
    column_alignments = get_column_alignments(table)

    # Extract header row count
    header_rows = 0
    for row in table.rows:
        trPr = row._tr.find(qn('w:trPr'))
        if trPr is not None:
            tblHeader = trPr.find(qn('w:tblHeader'))
            if tblHeader is not None:
                val = tblHeader.get(qn('w:val'))
                if val is None or val.lower() not in ('0', 'false'):
                    header_rows += 1
                    continue
        break  # Header rows must be contiguous from top
    # GFM always has at least 1 header row
    header_rows = max(header_rows, 1)

    # Extract merged cell regions
    merged = extract_merged_cells(table)

    result: dict[str, Any] = {
        "rows": num_rows,
        "cols": num_cols,
        "cells": cells,
        "column_alignments": column_alignments,
        "docx_table_index": table_index,
        "header_rows": header_rows,
    }

    if merged:
        result["merged_cells"] = merged

    return result


def _process_paragraph(
    paragraph: Any,
    doc_part: Any,
    block_index: int,
    para_index: int,
    content_position: int,
    image_counter: int,
    list_number_counter: int,
    previous_list_type: Optional[str],
    image_data: dict[str, bytes],
    extract_track_changes: bool = False,
    track_changes_explicit: Optional[bool] = None,
    footnote_counter: int = 0,
) -> tuple[Block, int, int, Optional[str], dict[str, bytes], int]:
    """Process a single paragraph element.

    Args:
        paragraph: python-docx paragraph object
        doc_part: Document part for accessing relationships
        block_index: Index for generating block ID
        para_index: Paragraph index in document
        content_position: Current position in content stream
        image_counter: Counter for image numbering
        list_number_counter: Counter for numbered lists
        previous_list_type: Type of previous list item
        image_data: Dictionary to store image data
        extract_track_changes: Whether to extract track changes as CriticMarkup
        track_changes_explicit: The explicit track_changes parameter (None, True, or False)

    Returns:
        Tuple of (block, new_image_counter, new_list_counter, new_list_type, image_data, footnote_counter)
    """
    style_name = paragraph.style.name if paragraph.style else "Normal"
    image_path = None
    block_track_changes = None
    footnote_refs: list[dict[str, Any]] = []
    chart_metadata_val: Optional[dict[str, Any]] = None

    # Ordering matters: chart detection must run BEFORE image extraction.
    # Chart drawings contain both <c:chart> and a cached <a:blip> fallback.
    # If extract_image_from_paragraph() runs first, it finds the fallback blip
    # and silently consumes the chart as a regular image.
    # Order: textbox check (in extract_blocks) → chart check → image check.
    chart_info = extract_chart_from_paragraph(paragraph, doc_part, image_counter)
    if chart_info:
        image_filename, image_extension, image_bytes, error_message, chart_rel_id = chart_info
        if chart_rel_id:
            chart_metadata_val = {"chart_rel_id": chart_rel_id}

        if not image_bytes:
            # Chart found but no cached image available
            markdown_content = "[Chart: no preview available]"
            block_type = "paragraph"
            level_value = None
            inline_formatting = None
        elif error_message:
            markdown_content = f"[Chart {image_counter} skipped: {error_message}]"
            block_type = "paragraph"
            level_value = None
            inline_formatting = None
        else:
            image_path = f"assets/{image_filename}"
            markdown_content = f"![Chart]({image_path})"
            block_type = "chart"
            level_value = None
            inline_formatting = None
            image_data[image_filename] = image_bytes

        image_counter += 1
        list_number_counter = 0
        previous_list_type = None
    elif (image_info := extract_image_from_paragraph(paragraph, doc_part, image_counter)):
        # This is an image paragraph
        image_filename, image_extension, image_bytes, error_message = image_info

        if error_message:
            markdown_content = f"[Image {image_counter} skipped: {error_message}]"
            block_type = "paragraph"
            level_value = None
            inline_formatting = None
        else:
            image_path = f"assets/{image_filename}"
            markdown_content = f"![Image {image_counter}]({image_path})"
            block_type = "image"
            level_value = None
            inline_formatting = None
            image_data[image_filename] = image_bytes

        image_counter += 1
        list_number_counter = 0
        previous_list_type = None
    else:
        # Choose extraction mode based on track changes settings
        mode: Literal["normal", "accept_all", "track_changes"]
        if extract_track_changes:
            mode = "track_changes"
        elif track_changes_explicit is False:
            mode = "accept_all"
        else:
            mode = "normal"

        text_content, inline_formatting, block_track_changes, footnote_refs, footnote_counter = extract_paragraph_content(
            paragraph._element, doc_part, mode=mode, footnote_counter=footnote_counter,
        )

        if style_name.startswith("Heading"):
            try:
                level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                level = 1
            markdown_content = "#" * level + " " + text_content
            block_type = "heading"
            level_value = level
            list_number_counter = 0
            previous_list_type = None
        elif style_name == "List Bullet":
            markdown_content = "- " + text_content
            block_type = "list"
            level_value = None
            if previous_list_type != "bullet":
                list_number_counter = 0
            previous_list_type = "bullet"
        elif style_name == "List Number":
            if previous_list_type != "number":
                list_number_counter = 0
            list_number_counter += 1
            markdown_content = f"{list_number_counter}. " + text_content
            block_type = "list"
            level_value = None
            previous_list_type = "number"
        else:
            markdown_content = text_content
            block_type = "paragraph"
            level_value = None
            list_number_counter = 0
            previous_list_type = None

    content_start = content_position
    content_end = content_position + len(markdown_content)

    block = Block(
        id=generate_block_id(block_index),
        type=block_type,
        content=markdown_content,
        docx_paragraph_index=para_index,
        content_start=content_start,
        content_end=content_end,
        content_hash=compute_content_hash(markdown_content),
        level=level_value,
        inline_formatting=inline_formatting,
        image_path=image_path,
        track_changes=block_track_changes,
        footnote_references=footnote_refs if footnote_refs else None,
        chart_metadata=chart_metadata_val,
    )

    return block, image_counter, list_number_counter, previous_list_type, image_data, footnote_counter


def extract_blocks(
    docx_path: str,
    track_changes: Optional[bool] = None,
    *,
    _doc: Any = None,
) -> tuple[list[Block], dict[str, bytes]]:
    """Extract blocks from a Word document.

    Converts paragraphs, headings, and tables to Block objects with markdown content.
    Processes the document body in order to correctly interleave paragraphs and tables.

    Args:
        docx_path: Path to .docx file
        track_changes: Track changes mode. If None, auto-detect based on document.
                      If True, extract track changes as CriticMarkup.
                      If False, accept all changes and extract plain text.
        _doc: Pre-opened Document object (internal use by extract_document).

    Returns:
        Tuple of (blocks, image_data) where image_data maps filenames to image bytes
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = _doc if _doc is not None else Document(docx_path)
    blocks: list[Block] = []
    image_data: dict[str, bytes] = {}
    content_position = 0
    list_number_counter = 0
    previous_list_type: Optional[str] = None
    image_counter = 1
    block_index = 0
    para_index = 0
    table_index = 0
    footnote_counter = 0

    # Determine track changes mode
    # If not specified (None), auto-detect based on document content
    extract_track_changes = track_changes
    if extract_track_changes is None:
        extract_track_changes = detect_track_changes(docx_path)

    # Iterate over document body elements in order
    # This correctly interleaves paragraphs and tables
    body = doc.element.body
    for child in body:
        tag = child.tag.split('}')[-1]  # Get tag without namespace

        if tag == 'p':
            # Create a Paragraph object from the XML element
            paragraph = Paragraph(child, doc)

            # Check for text boxes/shapes with text before other content
            textbox_infos = extract_textbox_from_paragraph(paragraph)
            if textbox_infos:
                for tb_info in textbox_infos:
                    inner_lines = "\n".join(tb_info["texts"])
                    markdown_content = f"<!-- textbox -->\n{inner_lines}\n<!-- /textbox -->"

                    content_start = content_position
                    content_end = content_position + len(markdown_content)

                    tb_metadata: dict[str, Any] = {
                        "anchor_type": tb_info["anchor_type"],
                        "width": tb_info["width"],
                        "height": tb_info["height"],
                        "drawing_xml": tb_info["drawing_xml"],
                    }
                    if tb_info["position_h"] is not None:
                        tb_metadata["position_h"] = tb_info["position_h"]
                    if tb_info["position_v"] is not None:
                        tb_metadata["position_v"] = tb_info["position_v"]
                    if tb_info["fill_color"] is not None:
                        tb_metadata["fill_color"] = tb_info["fill_color"]
                    if tb_info["border_color"] is not None:
                        tb_metadata["border_color"] = tb_info["border_color"]

                    block = Block(
                        id=generate_block_id(block_index),
                        type="textbox",
                        content=markdown_content,
                        docx_paragraph_index=para_index,
                        content_start=content_start,
                        content_end=content_end,
                        content_hash=compute_content_hash(markdown_content),
                        text_box_metadata=tb_metadata,
                    )

                    blocks.append(block)
                    content_position = content_end + 1
                    block_index += 1

                para_index += 1
                list_number_counter = 0
                previous_list_type = None
            else:
                block, image_counter, list_number_counter, previous_list_type, image_data, footnote_counter = _process_paragraph(
                    paragraph=paragraph,
                    doc_part=doc.part,
                    block_index=block_index,
                    para_index=para_index,
                    content_position=content_position,
                    image_counter=image_counter,
                    list_number_counter=list_number_counter,
                    previous_list_type=previous_list_type,
                    image_data=image_data,
                    extract_track_changes=extract_track_changes,
                    track_changes_explicit=track_changes,
                    footnote_counter=footnote_counter,
                )

                blocks.append(block)
                content_position = block.content_end + 1
                block_index += 1
                para_index += 1

                # Reset list counters for non-list content
                if block.type != "list":
                    list_number_counter = 0
                    previous_list_type = None

        elif tag == 'tbl':
            # Create a Table object from the XML element
            table = Table(child, doc)

            # Convert table to GFM markdown
            markdown_content = table_to_gfm(table, doc.part)

            # Extract table metadata for structure.json
            table_metadata = extract_table_metadata(table, table_index)

            content_start = content_position
            content_end = content_position + len(markdown_content)

            block = Block(
                id=generate_block_id(block_index),
                type="table",
                content=markdown_content,
                docx_paragraph_index=-1,  # Tables don't have paragraph index
                content_start=content_start,
                content_end=content_end,
                content_hash=compute_content_hash(markdown_content),
                level=None,
                inline_formatting=None,
                image_path=None,
                table_metadata=table_metadata
            )

            blocks.append(block)
            content_position = content_end + 1
            block_index += 1
            table_index += 1

            # Reset list counters
            list_number_counter = 0
            previous_list_type = None

    # Append footnote/endnote definitions at the end
    if footnote_counter > 0:
        definitions = _extract_footnote_definitions(doc, blocks)
        if definitions:
            # Add as a special block
            def_content = "\n".join(definitions)
            block = Block(
                id=generate_block_id(block_index),
                type="paragraph",
                content=def_content,
                docx_paragraph_index=-1,
                content_start=content_position,
                content_end=content_position + len(def_content),
                content_hash=compute_content_hash(def_content),
            )
            blocks.append(block)

    return blocks, image_data


def _extract_footnote_definitions(doc: Any, blocks: list[Block]) -> list[str]:
    """Extract footnote/endnote definitions from the document.

    Builds a map from original note IDs to sequential markers, then
    reads the footnotes/endnotes XML parts to extract text.

    Returns:
        List of definition strings like '[^1]: Footnote text.'
    """

    # Build map: (note_type, original_id) -> marker number
    ref_map: dict[tuple[str, str], int] = {}
    for block in blocks:
        if block.footnote_references:
            for ref in block.footnote_references:
                key = (ref["note_type"], ref["original_id"])
                ref_map[key] = ref["note_id"]

    if not ref_map:
        return []

    definitions: list[tuple[int, str]] = []  # (marker_num, definition_string)
    ns = {"w": WORDPROCESSINGML_NS}

    # Access footnotes/endnotes parts via relationship iteration
    FN_RT = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
    EN_RT = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes"

    for rel in doc.part.rels.values():
        if rel.reltype == FN_RT:
            _extract_notes_from_part(rel.target_part, "footnote", ref_map, definitions, ns)
        elif rel.reltype == EN_RT:
            _extract_notes_from_part(rel.target_part, "endnote", ref_map, definitions, ns)

    # Sort by marker number and return
    definitions.sort(key=lambda x: x[0])
    return [d[1] for d in definitions]


def _extract_notes_from_part(
    part: Any,
    note_type: str,
    ref_map: dict[tuple[str, str], int],
    definitions: list[tuple[int, str]],
    ns: dict[str, str],
) -> None:
    """Extract note text from a footnotes or endnotes XML part."""

    # python-docx loads footnotes as generic Part (with .blob) not XmlPart (with ._element)
    if hasattr(part, '_element'):
        root = part._element
    elif hasattr(part, 'blob'):
        root = etree.fromstring(part.blob)
    else:
        return

    tag_name = f"w:{note_type}"
    for note_elem in root.findall(tag_name, ns):
        note_id = note_elem.get(f"{{{ns['w']}}}id")
        if note_id is None:
            continue
        # Skip separator and continuationSeparator (IDs -1 and 0)
        note_type_attr = note_elem.get(f"{{{ns['w']}}}type")
        if note_type_attr in ("separator", "continuationSeparator"):
            continue

        key = (note_type, note_id)
        marker_num = ref_map.get(key)
        if marker_num is None:
            continue

        note_text = _extract_note_text(note_elem, ns)
        if note_text:
            definitions.append((marker_num, f"[^{marker_num}]: {note_text}"))


def _extract_note_text(note_elem: Any, ns: dict[str, str]) -> str:
    """Extract text content from a footnote/endnote element, preserving bold/italic.

    Known limitation: the leading-space strip (OOXML convention after the ref mark)
    only applies to the first run of the first paragraph. Multi-paragraph footnotes
    may retain a leading space on subsequent paragraphs' first runs.
    """
    text_parts: list[str] = []
    for para in note_elem.findall("w:p", ns):
        for run in para.findall("w:r", ns):
            # Skip the footnoteRef/endnoteRef marker run
            if run.find("w:footnoteRef", ns) is not None:
                continue
            if run.find("w:endnoteRef", ns) is not None:
                continue

            run_text_parts = []
            for t in run.findall("w:t", ns):
                if t.text:
                    run_text_parts.append(t.text)
            run_text = "".join(run_text_parts)
            if not run_text:
                continue

            # Strip leading space (OOXML convention: space after ref mark)
            if not text_parts and run_text.startswith(" "):
                run_text = run_text[1:]
                if not run_text:
                    continue

            # Check formatting
            rPr = run.find("w:rPr", ns)
            is_bold = False
            is_italic = False
            if rPr is not None:
                b = rPr.find("w:b", ns)
                if b is not None and b.get(f"{{{ns['w']}}}val", "true") != "false":
                    is_bold = True
                i = rPr.find("w:i", ns)
                if i is not None and i.get(f"{{{ns['w']}}}val", "true") != "false":
                    is_italic = True

            formatted = wrap_formatting(run_text, is_bold, is_italic)
            text_parts.append(formatted)

    return "".join(text_parts)


def extract_document(
    docx_path: str,
    track_changes: Optional[bool] = None,
) -> tuple[list[Block], dict[str, bytes], list[SectionProperties]]:
    """Extract blocks and section properties from a Word document in one pass.

    Opens the document once and shares the Document/body element between block
    extraction and section extraction, ensuring block indices are consistent.

    Args:
        docx_path: Path to .docx file
        track_changes: Track changes mode (see extract_blocks).

    Returns:
        Tuple of (blocks, image_data, sections)
    """
    doc = Document(docx_path)
    blocks, image_data = extract_blocks(docx_path, track_changes=track_changes, _doc=doc)
    sections = extract_sections(body=doc.element.body)
    return blocks, image_data, sections


def _safe_int(value: str | None, default: int | None = None) -> int | None:
    """Convert a string to int, returning default on None or ValueError."""
    if value is None:
        return default
    try:
        return int(value)
    except ValueError:
        return default


def _parse_cols_element(cols_elem: Any) -> dict[str, Any]:
    """Parse a w:cols element into a dict of column properties.

    Args:
        cols_elem: The w:cols XML element

    Returns:
        Dict with column_count, column_spacing, equal_width, and columns
    """
    col_count = _safe_int(cols_elem.get(qn('w:num'), '1'), 1) or 1
    equal_width_val = cols_elem.get(qn('w:equalWidth'))

    # equalWidth defaults to True when not specified
    equal_width = equal_width_val != '0'

    column_spacing = _safe_int(cols_elem.get(qn('w:space')))

    columns = None
    if not equal_width:
        col_elems = cols_elem.findall(qn('w:col'))
        if col_elems:
            columns = []
            for col_el in col_elems:
                width = _safe_int(col_el.get(qn('w:w'), '0'), 0) or 0
                space_val = _safe_int(col_el.get(qn('w:space')))
                columns.append(ColumnDefinition(
                    width=width,
                    space=space_val,
                ))

    return {
        'column_count': col_count,
        'column_spacing': column_spacing,
        'equal_width': equal_width,
        'columns': columns,
    }


def _parse_sect_pr(sect_pr: Any) -> SectionProperties:
    """Parse a w:sectPr element into SectionProperties.

    Args:
        sect_pr: The w:sectPr XML element

    Returns:
        SectionProperties with column data populated
    """
    props = SectionProperties()
    cols_elem = sect_pr.find(qn('w:cols'))
    if cols_elem is not None:
        parsed = _parse_cols_element(cols_elem)
        props.column_count = parsed['column_count']
        props.column_spacing = parsed['column_spacing']
        props.equal_width = parsed['equal_width']
        props.columns = parsed['columns']
    return props


def extract_sections(docx_path: str | None = None, *, body: Any = None) -> list[SectionProperties]:
    """Extract section properties (column layouts) from a Word document.

    OOXML sections are defined by:
    - Mid-document: w:sectPr inside the last w:pPr of a section's final paragraph
    - Final section: w:sectPr as direct child of w:body

    Args:
        docx_path: Path to .docx file (used if body is not provided)
        body: Pre-opened document body element. When provided, avoids
              re-opening the document and ensures block indices stay
              consistent with extract_blocks.

    Returns:
        List of SectionProperties in document order
    """
    if body is None:
        if docx_path is None:
            raise ValueError("Either docx_path or body must be provided")
        doc = Document(docx_path)
        body = doc.element.body
    sections: list[SectionProperties] = []
    block_index = 0

    for child in body:
        tag = child.tag.split('}')[-1]

        if tag == 'p':
            # Check for mid-document section break (sectPr inside pPr)
            pPr = child.find(qn('w:pPr'))
            if pPr is not None:
                sect_pr = pPr.find(qn('w:sectPr'))
                if sect_pr is not None:
                    props = _parse_sect_pr(sect_pr)
                    props.end_block_index = block_index
                    prev_end = sections[-1].end_block_index if sections else None
                    props.start_block_index = (prev_end + 1) if prev_end is not None else 0
                    sections.append(props)
            block_index += 1

        elif tag == 'tbl':
            block_index += 1

        elif tag == 'sectPr':
            # Final section (body-level sectPr)
            props = _parse_sect_pr(child)
            props.end_block_index = block_index - 1 if block_index > 0 else 0
            prev_end = sections[-1].end_block_index if sections else None
            props.start_block_index = (prev_end + 1) if prev_end is not None else 0
            sections.append(props)

    # If no sections found, create a default single-column section
    if not sections:
        props = SectionProperties()
        props.start_block_index = 0
        props.end_block_index = max(block_index - 1, 0)
        sections.append(props)

    return sections


def extract_cell_shading(cell: Any) -> str | None:
    """Extract background color from a cell's tcPr/shd element.

    Args:
        cell: python-docx table cell object

    Returns:
        Hex color string like 'D9E2F3' or None if no shading
    """
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return None
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        return None
    fill = shd.get(qn('w:fill'))
    if fill and str(fill).upper() != 'AUTO':
        return str(fill).upper()
    return None


def extract_cell_pattern_fill(cell: Any) -> str | None:
    """Extract pattern fill type from a cell's tcPr/shd element.

    Args:
        cell: python-docx table cell object

    Returns:
        Pattern name like 'diagStripe' or None if no pattern (clear/nil/solid)
    """
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return None
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        return None
    val = shd.get(qn('w:val'))
    if val and str(val) not in ('clear', 'nil', 'solid'):
        return str(val)
    return None


def extract_cell_borders(cell: Any) -> dict[str, dict[str, Any]] | None:
    """Extract border styles from all four sides of a cell.

    Args:
        cell: python-docx table cell object

    Returns:
        Dictionary with border data per side, or None if no borders
    """
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return None
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        return None
    borders: dict[str, dict[str, Any]] = {}
    for side in ['top', 'bottom', 'left', 'right']:
        border_elem = tcBorders.find(qn(f'w:{side}'))
        if border_elem is not None:
            val = border_elem.get(qn('w:val'))
            if val and val != 'nil':
                borders[side] = {
                    'style': val,
                    'width': int(border_elem.get(qn('w:sz'), '0')),
                    'color': border_elem.get(qn('w:color'), 'auto'),
                }
    return borders if borders else None


def extract_table_formatting(table: Any) -> dict[str, Any]:
    """Extract formatting information from a table.

    Args:
        table: python-docx Table object

    Returns:
        Dictionary with table formatting: column_widths, table_alignment,
        table_style, cell_styles
    """
    from docx.shared import Inches

    # Extract column widths
    column_widths: list[float] = []
    for col in table.columns:
        # Width is in EMUs (English Metric Units) or can be None
        width = col.width
        if width:
            # Convert to inches
            width_inches = width / EMUS_PER_INCH
            column_widths.append(round(width_inches, 2))
        else:
            # Default width if not specified
            column_widths.append(1.0)

    # Extract table alignment
    # Default is left
    table_alignment = "left"
    # python-docx doesn't have direct table alignment property
    # It's stored in the table's XML properties
    tblPr = table._tbl.tblPr
    if tblPr is not None:
        jc = tblPr.find(f'{{{WORDPROCESSINGML_NS}}}jc')
        if jc is not None:
            val = jc.get(qn('w:val'))
            if val:
                alignment_map = {
                    'left': 'left',
                    'center': 'center',
                    'right': 'right',
                    'start': 'left',
                    'end': 'right'
                }
                table_alignment = alignment_map.get(val, 'left')

    # Extract table style name if present
    table_style_name = None
    if table.style:
        table_style_name = table.style.name

    # Extract cell-level styles (shading and borders)
    cell_styles: dict[str, dict[str, Any]] = {}
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell_data: dict[str, Any] = {}

            shading = extract_cell_shading(cell)
            if shading:
                cell_data["background_color"] = shading

            pattern = extract_cell_pattern_fill(cell)
            if pattern:
                cell_data["pattern_fill"] = pattern

            borders = extract_cell_borders(cell)
            if borders:
                cell_data["borders"] = borders

            if cell_data:
                cell_styles[f"{row_idx},{col_idx}"] = cell_data

    result: dict[str, Any] = {
        "column_widths": column_widths,
        "table_alignment": table_alignment,
        "table_style": table_style_name,
    }

    if cell_styles:
        result["cell_styles"] = cell_styles

    return result


def extract_styles(docx_path: str, blocks: list[Block]) -> list[Style]:
    """Extract style information from Word document.

    Args:
        docx_path: Path to .docx file
        blocks: List of Block objects

    Returns:
        List of Style objects
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = Document(docx_path)
    styles: list[Style] = []

    # Iterate over document body in order to match blocks correctly
    body = doc.element.body
    block_index = 0

    for child in body:
        if block_index >= len(blocks):
            break

        tag = child.tag.split('}')[-1]
        block = blocks[block_index]

        if tag == 'p':
            paragraph = Paragraph(child, doc)

            # Check if this paragraph produced text box blocks
            if block.type == "textbox":
                while block_index < len(blocks) and blocks[block_index].type == "textbox":
                    style = Style(
                        block_id=blocks[block_index].id,
                        docx_style="TextBox",
                        font_name="Calibri",
                        font_size=11,
                        alignment="left",
                    )
                    styles.append(style)
                    block_index += 1
                continue

            # Get font properties
            font_name = "Calibri"
            font_size = 11
            alignment = "left"

            if paragraph.style and paragraph.style.font.name:
                font_name = paragraph.style.font.name

            if paragraph.style and paragraph.style.font.size:
                font_size = int(paragraph.style.font.size.pt)

            if paragraph.alignment is not None:
                alignment = ALIGNMENT_NUMERIC_TO_STRING.get(
                    paragraph.alignment, DEFAULT_ALIGNMENT
                )

            style = Style(
                block_id=block.id,
                docx_style=paragraph.style.name if paragraph.style else "Normal",
                font_name=font_name,
                font_size=font_size,
                alignment=alignment,
            )
            styles.append(style)
            block_index += 1

        elif tag == 'tbl':
            table = Table(child, doc)

            # Extract table-specific formatting
            table_formatting = extract_table_formatting(table)

            # Create style with table_formatting
            style = Style(
                block_id=block.id,
                docx_style="Table",
                font_name="Calibri",  # Default for tables
                font_size=11,
                alignment="left",
                table_formatting=table_formatting
            )
            styles.append(style)
            block_index += 1

    return styles


def _extract_header_footer_paragraphs(
    header_footer: Any, image_data: dict[str, bytes], image_counter: int,
) -> tuple[list[dict], dict[str, bytes], int]:
    """Extract paragraphs from a header or footer element.

    Skips extraction if the header/footer is linked to the previous section.
    Images are extracted using the header/footer's own part reference.

    Args:
        header_footer: A python-docx header or footer object.
        image_data: Mutable dict mapping image filenames to their bytes.
        image_counter: Current image counter for generating unique filenames.

    Returns:
        Tuple of (paragraphs, image_data, image_counter) where paragraphs is a
        list of dicts with type/content/image_path keys.
    """
    paragraphs: list[dict] = []
    if header_footer.is_linked_to_previous:
        return paragraphs, image_data, image_counter
    for para in header_footer.paragraphs:
        text = para.text.strip()
        image_info = extract_image_from_paragraph(
            para, header_footer.part, image_counter
        )
        if image_info:
            image_filename, image_extension, image_bytes, error_message = image_info
            if error_message:
                paragraphs.append({
                    "type": "paragraph",
                    "content": f"[Image {image_counter} skipped: {error_message}]",
                })
                image_counter += 1
                continue
            # Prefix prevents filename collision with body images (which use
            # image1.png, image2.png, ...) regardless of counter values.
            hf_filename = f"hf_{image_filename}"
            image_data[hf_filename] = image_bytes
            paragraphs.append({
                "type": "image",
                "content": "",
                "image_path": f"assets/{hf_filename}",
            })
            image_counter += 1
            continue
        if not text:
            continue
        paragraphs.append({"type": "paragraph", "content": text})
    return paragraphs, image_data, image_counter


def extract_section_metadata(docx_path: str) -> tuple[list[dict], dict[str, bytes]]:
    """Extract section metadata (headers, footers, page setup) from a Word document.

    Returns:
        Tuple of (sections_data, image_data) where image_data maps filenames to bytes.
    """
    doc = Document(docx_path)
    sections_data = []
    image_data: dict[str, bytes] = {}
    image_counter = 1
    odd_and_even = doc.settings.odd_and_even_pages_header_footer
    for section in doc.sections:
        section_dict: dict[str, Any] = {}
        variants = [
            ("header_default", section.header),
            ("footer_default", section.footer),
        ]
        if section.different_first_page_header_footer:
            variants.extend([
                ("header_first", section.first_page_header),
                ("footer_first", section.first_page_footer),
            ])
        if odd_and_even:
            even_header = section.even_page_header
            if not even_header.is_linked_to_previous:
                variants.append(("header_even", even_header))
            even_footer = section.even_page_footer
            if not even_footer.is_linked_to_previous:
                variants.append(("footer_even", even_footer))
        for key, hf in variants:
            paras, image_data, image_counter = _extract_header_footer_paragraphs(
                hf, image_data, image_counter
            )
            section_dict[key] = paras
        section_dict.setdefault("header_default", [])
        section_dict.setdefault("footer_default", [])
        if section.different_first_page_header_footer:
            section_dict.setdefault("header_first", [])
            section_dict.setdefault("footer_first", [])
        if odd_and_even:
            section_dict.setdefault("header_even", [])
            section_dict.setdefault("footer_even", [])
        orientation_val = section.orientation
        orientation_str = "landscape" if orientation_val == WD_ORIENT.LANDSCAPE else "portrait"
        section_dict["page_setup"] = {
            "orientation": orientation_str,
            "top_margin": section.top_margin,
            "bottom_margin": section.bottom_margin,
            "left_margin": section.left_margin,
            "right_margin": section.right_margin,
            "page_width": section.page_width,
            "page_height": section.page_height,
            "header_distance": section.header_distance,
            "footer_distance": section.footer_distance,
            "different_first_page": section.different_first_page_header_footer,
            "odd_and_even_pages": odd_and_even,
        }
        sections_data.append(section_dict)
    return sections_data, image_data


def blocks_to_markdown(blocks: list[Block]) -> str:
    """Convert blocks to markdown content.

    Args:
        blocks: List of Block objects

    Returns:
        Markdown content string
    """
    lines = [block.content for block in blocks]
    return "\n".join(lines)
