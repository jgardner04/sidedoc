"""Script to create test fixtures with footnotes and endnotes.

Footnotes/endnotes in Word use w:footnoteReference/w:endnoteReference in body
paragraphs, with content stored in separate word/footnotes.xml and word/endnotes.xml
parts. python-docx has no native API for creating these, so this script
manipulates the OPC layer and XML directly.
"""

from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.part import XmlPart
from docx.opc.packuri import PackURI
from docx.oxml.parser import parse_xml

# Get the directory where this script is located
FIXTURES_DIR = Path(__file__).parent

# XML namespace for preserving whitespace
XML_SPACE_NS = "{http://www.w3.org/XML/1998/namespace}space"

# Relationship types
FOOTNOTES_RT = (
    "http://schemas.openxmlformats.org/officeDocument"
    "/2006/relationships/footnotes"
)
ENDNOTES_RT = (
    "http://schemas.openxmlformats.org/officeDocument"
    "/2006/relationships/endnotes"
)

# Content types
FOOTNOTES_CT = (
    "application/vnd.openxmlformats-officedocument"
    ".wordprocessingml.footnotes+xml"
)
ENDNOTES_CT = (
    "application/vnd.openxmlformats-officedocument"
    ".wordprocessingml.endnotes+xml"
)

# Default XML templates for footnotes/endnotes parts
DEFAULT_FOOTNOTES_XML = b"""\
<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<w:footnotes xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
             xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:footnote w:type="separator" w:id="-1">
    <w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>
         <w:r><w:separator/></w:r></w:p>
  </w:footnote>
  <w:footnote w:type="continuationSeparator" w:id="0">
    <w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>
         <w:r><w:continuationSeparator/></w:r></w:p>
  </w:footnote>
</w:footnotes>"""

DEFAULT_ENDNOTES_XML = b"""\
<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<w:endnotes xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
            xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:endnote w:type="separator" w:id="-1">
    <w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>
         <w:r><w:separator/></w:r></w:p>
  </w:endnote>
  <w:endnote w:type="continuationSeparator" w:id="0">
    <w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>
         <w:r><w:continuationSeparator/></w:r></w:p>
  </w:endnote>
</w:endnotes>"""


def _get_or_create_footnotes_part(doc):
    """Get existing or create new footnotes part."""
    doc_part = doc.part
    try:
        return doc_part.part_related_by(FOOTNOTES_RT)
    except KeyError:
        footnotes_part = XmlPart(
            PackURI("/word/footnotes.xml"),
            FOOTNOTES_CT,
            parse_xml(DEFAULT_FOOTNOTES_XML),
            doc_part.package,
        )
        doc_part.relate_to(footnotes_part, FOOTNOTES_RT)
        return footnotes_part


def _get_or_create_endnotes_part(doc):
    """Get existing or create new endnotes part."""
    doc_part = doc.part
    try:
        return doc_part.part_related_by(ENDNOTES_RT)
    except KeyError:
        endnotes_part = XmlPart(
            PackURI("/word/endnotes.xml"),
            ENDNOTES_CT,
            parse_xml(DEFAULT_ENDNOTES_XML),
            doc_part.package,
        )
        doc_part.relate_to(endnotes_part, ENDNOTES_RT)
        return endnotes_part


def _add_note_to_part(part, note_text, tag_name, style_name, ref_tag_name,
                      bold=False, italic=False):
    """Add a footnote or endnote entry to its XML part.

    Returns the assigned ID.
    """
    root = part._element
    ns_tag = qn(f"w:{tag_name}")

    existing_ids = [
        int(el.get(qn("w:id")))
        for el in root.findall(ns_tag)
        if el.get(qn("w:id")) is not None
    ]
    next_id = max((i for i in existing_ids if i >= 1), default=0) + 1

    note = OxmlElement(f"w:{tag_name}")
    note.set(qn("w:id"), str(next_id))

    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), style_name)
    pPr.append(pStyle)
    p.append(pPr)

    # Run with footnote/endnote reference mark
    r_ref = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    # Use FootnoteReference style for both (Word convention)
    ref_style = "FootnoteReference" if tag_name == "footnote" else "EndnoteReference"
    rStyle.set(qn("w:val"), ref_style)
    rPr.append(rStyle)
    r_ref.append(rPr)
    r_ref.append(OxmlElement(f"w:{ref_tag_name}"))
    p.append(r_ref)

    # Run with actual text
    r_text = OxmlElement("w:r")
    if bold or italic:
        rPr2 = OxmlElement("w:rPr")
        if bold:
            rPr2.append(OxmlElement("w:b"))
        if italic:
            rPr2.append(OxmlElement("w:i"))
        r_text.append(rPr2)
    t = OxmlElement("w:t")
    t.text = " " + note_text
    t.set(XML_SPACE_NS, "preserve")
    r_text.append(t)
    p.append(r_text)

    note.append(p)
    root.append(note)
    return next_id


def add_footnote(paragraph, footnote_text, doc, bold=False, italic=False):
    """Add a footnote reference to a paragraph and its content to footnotes.xml."""
    part = _get_or_create_footnotes_part(doc)
    fn_id = _add_note_to_part(
        part, footnote_text, "footnote", "FootnoteText", "footnoteRef",
        bold=bold, italic=italic,
    )

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "FootnoteReference")
    rPr.append(rStyle)
    r.append(rPr)
    ref = OxmlElement("w:footnoteReference")
    ref.set(qn("w:id"), str(fn_id))
    r.append(ref)
    paragraph._p.append(r)


def add_endnote(paragraph, endnote_text, doc, bold=False, italic=False):
    """Add an endnote reference to a paragraph and its content to endnotes.xml."""
    part = _get_or_create_endnotes_part(doc)
    en_id = _add_note_to_part(
        part, endnote_text, "endnote", "EndnoteText", "endnoteRef",
        bold=bold, italic=italic,
    )

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "EndnoteReference")
    rPr.append(rStyle)
    r.append(rPr)
    ref = OxmlElement("w:endnoteReference")
    ref.set(qn("w:id"), str(en_id))
    r.append(ref)
    paragraph._p.append(r)


def create_footnotes_simple():
    """Create footnotes_simple.docx with basic footnotes.

    Contents:
    - Paragraph 1: "This is a claim that needs citation[^1]. Another point[^2]."
    - Footnote 1: "Smith, J. (2024). Document Processing. Journal of AI, 15(3), 42-58."
    - Footnote 2: "This is an explanatory footnote with more detail."
    """
    doc = Document()

    p = doc.add_paragraph("This is a claim that needs citation")
    add_footnote(p, 'Smith, J. (2024). "Document Processing." Journal of AI, 15(3), 42-58.', doc)
    p.add_run(". Another point")
    add_footnote(p, "This is an explanatory footnote with more detail.", doc)
    p.add_run(".")

    doc.save(str(FIXTURES_DIR / "footnotes_simple.docx"))
    print("Created footnotes_simple.docx")


def create_footnotes_endnotes_mixed():
    """Create footnotes_endnotes_mixed.docx with both footnotes and endnotes.

    Contents:
    - Paragraph 1: "Body text with a footnote[^1] and an endnote[^2]."
    - Footnote 1: "This is a footnote."
    - Endnote 1: "This is an endnote."
    - Paragraph 2: "More text with another endnote[^3]."
    - Endnote 2: "Second endnote content."
    """
    doc = Document()

    p1 = doc.add_paragraph("Body text with a footnote")
    add_footnote(p1, "This is a footnote.", doc)
    p1.add_run(" and an endnote")
    add_endnote(p1, "This is an endnote.", doc)
    p1.add_run(".")

    p2 = doc.add_paragraph("More text with another endnote")
    add_endnote(p2, "Second endnote content.", doc)
    p2.add_run(".")

    doc.save(str(FIXTURES_DIR / "footnotes_endnotes_mixed.docx"))
    print("Created footnotes_endnotes_mixed.docx")


def create_footnotes_formatted():
    """Create footnotes_formatted.docx with inline formatting in footnotes.

    Contents:
    - Paragraph 1: "Text with formatted footnote[^1]."
    - Footnote 1: "**This footnote has bold text.**" (bold)
    - Paragraph 2: "Another paragraph with footnote[^2]."
    - Footnote 2: "*This footnote has italic text.*" (italic)
    """
    doc = Document()

    p1 = doc.add_paragraph("Text with formatted footnote")
    add_footnote(p1, "This footnote has bold text.", doc, bold=True)
    p1.add_run(".")

    p2 = doc.add_paragraph("Another paragraph with footnote")
    add_footnote(p2, "This footnote has italic text.", doc, italic=True)
    p2.add_run(".")

    doc.save(str(FIXTURES_DIR / "footnotes_formatted.docx"))
    print("Created footnotes_formatted.docx")


def create_footnotes_consecutive():
    """Create footnotes_consecutive.docx with consecutive footnotes (no text between).

    Contents:
    - Paragraph 1: "Multiple citations here[^1][^2][^3]."
    - Footnote 1: "First source."
    - Footnote 2: "Second source."
    - Footnote 3: "Third source."
    """
    doc = Document()

    p = doc.add_paragraph("Multiple citations here")
    add_footnote(p, "First source.", doc)
    add_footnote(p, "Second source.", doc)
    add_footnote(p, "Third source.", doc)
    p.add_run(".")

    doc.save(str(FIXTURES_DIR / "footnotes_consecutive.docx"))
    print("Created footnotes_consecutive.docx")


def create_footnotes_in_heading():
    """Create footnotes_in_heading.docx with a footnote in a heading.

    Contents:
    - Heading 1: "Chapter Title[^1]"
    - Footnote 1: "Chapter footnote content."
    - Paragraph: "Body text follows."
    """
    doc = Document()

    h = doc.add_heading("Chapter Title", level=1)
    add_footnote(h, "Chapter footnote content.", doc)

    doc.add_paragraph("Body text follows.")

    doc.save(str(FIXTURES_DIR / "footnotes_in_heading.docx"))
    print("Created footnotes_in_heading.docx")


def create_footnotes_multiblock():
    """Create footnotes_multiblock.docx with footnotes across multiple paragraphs.

    Contents:
    - Heading 1: "Introduction"
    - Paragraph 1: "First paragraph with a footnote[^1]."
    - Paragraph 2: "Second paragraph, no footnotes."
    - Heading 2: "Methods"
    - Paragraph 3: "Third paragraph with footnote[^2] and another[^3]."
    - Footnote 1: "Introduction source."
    - Footnote 2: "Methods reference one."
    - Footnote 3: "Methods reference two."
    """
    doc = Document()

    doc.add_heading("Introduction", level=1)

    p1 = doc.add_paragraph("First paragraph with a footnote")
    add_footnote(p1, "Introduction source.", doc)
    p1.add_run(".")

    doc.add_paragraph("Second paragraph, no footnotes.")

    doc.add_heading("Methods", level=2)

    p3 = doc.add_paragraph("Third paragraph with footnote")
    add_footnote(p3, "Methods reference one.", doc)
    p3.add_run(" and another")
    add_footnote(p3, "Methods reference two.", doc)
    p3.add_run(".")

    doc.save(str(FIXTURES_DIR / "footnotes_multiblock.docx"))
    print("Created footnotes_multiblock.docx")


if __name__ == "__main__":
    print("Creating footnote/endnote test fixtures...")
    create_footnotes_simple()
    create_footnotes_endnotes_mixed()
    create_footnotes_formatted()
    create_footnotes_consecutive()
    create_footnotes_in_heading()
    create_footnotes_multiblock()
    print("\nAll footnote/endnote fixtures created successfully!")
