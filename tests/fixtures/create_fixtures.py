"""Script to create test fixture docx files."""

from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Get the directory where this script is located
FIXTURES_DIR = Path(__file__).parent


def _apply_cell_shading(cell: object, fill_color: str) -> None:
    """Apply background shading to a cell."""
    tc = cell._tc  # type: ignore[attr-defined]
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)


def _apply_cell_borders(cell: object, color: str, width: str = '4') -> None:
    """Apply borders to all four sides of a cell."""
    tc = cell._tc  # type: ignore[attr-defined]
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'bottom', 'left', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), width)
        border.set(qn('w:color'), color)
        border.set(qn('w:space'), '0')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def create_simple_docx() -> None:
    """Create simple.docx with headings and paragraphs only."""
    doc = Document()

    doc.add_heading("Main Title", level=1)
    doc.add_paragraph("This is a simple paragraph with plain text.")

    doc.add_heading("Section 1", level=2)
    doc.add_paragraph("First paragraph in section 1.")
    doc.add_paragraph("Second paragraph in section 1.")

    doc.add_heading("Section 2", level=2)
    doc.add_paragraph("A paragraph in section 2.")

    doc.save(str(FIXTURES_DIR / "simple.docx"))
    print("✓ Created simple.docx")


def create_lists_docx() -> None:
    """Create lists.docx with bulleted and numbered lists."""
    doc = Document()

    doc.add_heading("Lists Example", level=1)

    doc.add_heading("Bulleted List", level=2)
    doc.add_paragraph("First bullet item", style="List Bullet")
    doc.add_paragraph("Second bullet item", style="List Bullet")
    doc.add_paragraph("Third bullet item", style="List Bullet")

    doc.add_heading("Numbered List", level=2)
    doc.add_paragraph("First numbered item", style="List Number")
    doc.add_paragraph("Second numbered item", style="List Number")
    doc.add_paragraph("Third numbered item", style="List Number")

    doc.add_heading("Mixed Lists", level=2)
    doc.add_paragraph("Bullet one", style="List Bullet")
    doc.add_paragraph("Bullet two", style="List Bullet")
    doc.add_paragraph("Then a numbered item", style="List Number")
    doc.add_paragraph("Another numbered item", style="List Number")

    doc.save(str(FIXTURES_DIR / "lists.docx"))
    print("✓ Created lists.docx")


def create_formatted_docx() -> None:
    """Create formatted.docx with bold, italic, mixed inline formatting."""
    doc = Document()

    doc.add_heading("Formatted Text Example", level=1)

    # Bold text
    p = doc.add_paragraph()
    p.add_run("This text is ").bold = False
    p.add_run("bold").bold = True
    p.add_run(".")

    # Italic text
    p = doc.add_paragraph()
    p.add_run("This text is ").italic = False
    p.add_run("italic").italic = True
    p.add_run(".")

    # Bold and italic
    p = doc.add_paragraph()
    p.add_run("This text is ").bold = False
    run = p.add_run("bold and italic")
    run.bold = True
    run.italic = True
    p.add_run(".")

    # Underline
    p = doc.add_paragraph()
    p.add_run("This text is ").underline = False
    p.add_run("underlined").underline = True
    p.add_run(".")

    # Mixed formatting in a sentence
    p = doc.add_paragraph()
    p.add_run("A sentence with ")
    p.add_run("bold").bold = True
    p.add_run(", ")
    p.add_run("italic").italic = True
    p.add_run(", and ")
    run = p.add_run("both")
    run.bold = True
    run.italic = True
    p.add_run(" mixed together.")

    doc.save(str(FIXTURES_DIR / "formatted.docx"))
    print("✓ Created formatted.docx")


def create_test_images() -> None:
    """Create simple PNG image files for testing."""
    from PIL import Image as PILImage

    # Create a simple red square
    img = PILImage.new('RGB', (100, 100), color='red')
    img.save(str(FIXTURES_DIR / 'test_image1.png'))

    # Create a simple blue square
    img2 = PILImage.new('RGB', (100, 100), color='blue')
    img2.save(str(FIXTURES_DIR / 'test_image2.png'))


def create_images_docx() -> None:
    """Create images.docx with multiple images."""
    # First create the test images
    create_test_images()

    doc = Document()

    doc.add_heading("Document with Images", level=1)

    doc.add_paragraph("First image:")
    doc.add_picture(str(FIXTURES_DIR / 'test_image1.png'), width=Inches(1.5))

    doc.add_paragraph("Second image:")
    doc.add_picture(str(FIXTURES_DIR / 'test_image2.png'), width=Inches(2.0))

    doc.add_paragraph("Text after images.")

    doc.save(str(FIXTURES_DIR / "images.docx"))
    print("✓ Created images.docx")


def create_tables_simple_docx() -> None:
    """Create tables_simple.docx with a basic 3x3 table.

    Table structure:
    | Name  | Role     | Start Date |
    |-------|----------|------------|
    | Alice | Engineer | 2024-01-15 |
    | Bob   | Designer | 2024-02-01 |
    """
    doc = Document()

    doc.add_heading("Simple Table Example", level=1)
    doc.add_paragraph("This document contains a basic table.")

    # Create a 3x3 table (3 rows, 3 columns)
    table = doc.add_table(rows=3, cols=3)

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Name"
    header_cells[1].text = "Role"
    header_cells[2].text = "Start Date"

    # Data row 1
    row1_cells = table.rows[1].cells
    row1_cells[0].text = "Alice"
    row1_cells[1].text = "Engineer"
    row1_cells[2].text = "2024-01-15"

    # Data row 2
    row2_cells = table.rows[2].cells
    row2_cells[0].text = "Bob"
    row2_cells[1].text = "Designer"
    row2_cells[2].text = "2024-02-01"

    doc.add_paragraph("Text after the table.")

    doc.save(str(FIXTURES_DIR / "tables_simple.docx"))
    print("✓ Created tables_simple.docx")


def create_tables_formatted_docx() -> None:
    """Create tables_formatted.docx with styled table (colors, borders, shading).

    Table structure (4 rows x 3 columns):
    | Department | Q1 Revenue | Q2 Revenue |   <- Blue header row with accent borders
    |------------|------------|------------|
    | Sales      | $1,200,000 | $1,350,000 |   <- White background
    | Marketing  | $800,000   | $920,000   |   <- Light gray alternating shading
    | Engineering| $950,000   | $1,100,000 |   <- White background
    """
    doc = Document()
    doc.add_heading("Formatted Table Example", level=1)
    doc.add_paragraph("This document contains a styled table with colors and borders.")

    table = doc.add_table(rows=4, cols=3)

    # Header row
    headers = ["Department", "Q1 Revenue", "Q2 Revenue"]
    data = [
        ["Sales", "$1,200,000", "$1,350,000"],
        ["Marketing", "$800,000", "$920,000"],
        ["Engineering", "$950,000", "$1,100,000"],
    ]

    # Populate cells
    for col_idx, header in enumerate(headers):
        table.cell(0, col_idx).text = header
    for row_idx, row_data in enumerate(data):
        for col_idx, value in enumerate(row_data):
            table.cell(row_idx + 1, col_idx).text = value

    # Apply header row formatting: blue background + accent borders
    for col_idx in range(3):
        cell = table.cell(0, col_idx)
        _apply_cell_shading(cell, 'D9E2F3')
        _apply_cell_borders(cell, '4472C4', '8')
        # Bold header text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Apply data row formatting
    for row_idx in range(1, 4):
        for col_idx in range(3):
            cell = table.cell(row_idx, col_idx)
            # Alternating shading: gray on even data rows (row 2 = index 2)
            if row_idx == 2:
                _apply_cell_shading(cell, 'F2F2F2')
            # Thin gray borders on all data cells
            _apply_cell_borders(cell, 'D9D9D9', '4')

    doc.add_paragraph("Text after the formatted table.")
    doc.save(str(FIXTURES_DIR / "tables_formatted.docx"))
    print("✓ Created tables_formatted.docx")


def create_tables_merged_docx() -> None:
    """Create tables_merged.docx with horizontally and vertically merged cells.

    Table structure (4 rows x 3 columns):
    | Report Title (merged across 3 cols)      |
    |------------|------------|-----------------|
    | Category   | Q1         | Q2              |
    | Revenue (merged vertically with row 3)  | $1M | $1.2M |
    | (continuation of vertical merge)        | $0.8M | $0.9M |
    """
    doc = Document()
    doc.add_heading("Merged Cells Example", level=1)
    doc.add_paragraph("This document contains a table with merged cells.")

    table = doc.add_table(rows=4, cols=3)

    # Populate all cells first
    table.cell(0, 0).text = "Report Title"
    table.cell(0, 1).text = ""
    table.cell(0, 2).text = ""
    table.cell(1, 0).text = "Category"
    table.cell(1, 1).text = "Q1"
    table.cell(1, 2).text = "Q2"
    table.cell(2, 0).text = "Revenue"
    table.cell(2, 1).text = "$1M"
    table.cell(2, 2).text = "$1.2M"
    table.cell(3, 0).text = ""
    table.cell(3, 1).text = "$0.8M"
    table.cell(3, 2).text = "$0.9M"

    # Apply horizontal merge: row 0, cols 0-2
    table.cell(0, 0).merge(table.cell(0, 2))

    # Apply vertical merge: rows 2-3, col 0
    table.cell(2, 0).merge(table.cell(3, 0))

    doc.add_paragraph("Text after the merged table.")
    doc.save(str(FIXTURES_DIR / "tables_merged.docx"))
    print("✓ Created tables_merged.docx")


def create_complex_docx() -> None:
    """Create complex.docx with all supported elements combined."""
    doc = Document()

    doc.add_heading("Complex Document", level=1)

    doc.add_paragraph("This document contains all supported elements.")

    doc.add_heading("Section with Formatting", level=2)
    p = doc.add_paragraph()
    p.add_run("This paragraph has ")
    p.add_run("bold").bold = True
    p.add_run(", ")
    p.add_run("italic").italic = True
    p.add_run(", and ")
    run = p.add_run("both")
    run.bold = True
    run.italic = True
    p.add_run(" formatting.")

    doc.add_heading("Lists Section", level=2)
    doc.add_paragraph("Bullet list:")
    doc.add_paragraph("First item", style="List Bullet")
    doc.add_paragraph("Second item", style="List Bullet")

    doc.add_paragraph("Numbered list:")
    doc.add_paragraph("First", style="List Number")
    doc.add_paragraph("Second", style="List Number")

    doc.add_heading("Image Section", level=2)
    doc.add_paragraph("An embedded image:")
    doc.add_picture(str(FIXTURES_DIR / 'test_image1.png'), width=Inches(1.0))

    doc.add_heading("Final Section", level=2)
    doc.add_paragraph("A concluding paragraph with regular text.")

    doc.save(str(FIXTURES_DIR / "complex.docx"))
    print("✓ Created complex.docx")


def create_tables_complex_docx() -> None:
    """Create tables_complex.docx with all table features combined.

    Table structure (5 rows x 4 columns):
    | Project Status (merged cols 0-1) | Priority | Owner       |  <- Blue header, accent borders, bold, center
    |----------------------------------|----------|-------------|
    | **Phase 1**                      | High     | Alice       |  <- White
    | Completed tasks                  | Medium   | **Bob**     |  <- Alternating gray (F2F2F2)
    | Summary (merged rows 3-4, col 0) | Low      | *Charlie*   |  <- White
    | (continuation)                   |          | Dave        |  <- Alternating gray (F2F2F2)

    Mixed column alignments: left, center, right, left
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading("Complex Table Example", level=1)
    doc.add_paragraph("This document contains a table with all supported features.")

    table = doc.add_table(rows=5, cols=4)

    # Populate cells
    data = [
        ["Project Status", "Phase", "Priority", "Owner"],
        ["Phase 1", "Design", "High", "Alice"],
        ["Completed tasks", "Build", "Medium", "Bob"],
        ["Summary", "Test", "Low", "Charlie"],
        ["", "Deploy", "", "Dave"],
    ]

    for row_idx, row_data in enumerate(data):
        for col_idx, value in enumerate(row_data):
            table.cell(row_idx, col_idx).text = value

    # Apply header row formatting: blue background + accent borders + bold + center
    for col_idx in range(4):
        cell = table.cell(0, col_idx)
        _apply_cell_shading(cell, 'D9E2F3')
        _apply_cell_borders(cell, '4472C4', '8')
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True

    # Apply column alignments: left, center, right, left
    alignment_map = [
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.RIGHT,
        WD_ALIGN_PARAGRAPH.LEFT,
    ]
    for row_idx in range(1, 5):
        for col_idx in range(4):
            cell = table.cell(row_idx, col_idx)
            for paragraph in cell.paragraphs:
                paragraph.alignment = alignment_map[col_idx]

    # Apply alternating row shading (rows 2 and 4, i.e. data rows 2 and 4)
    for col_idx in range(4):
        _apply_cell_shading(table.cell(2, col_idx), 'F2F2F2')
        _apply_cell_shading(table.cell(4, col_idx), 'F2F2F2')

    # Apply inline formatting: bold in cell (1,0), bold in cell (3,3) italic
    cell_1_0 = table.cell(1, 0)
    cell_1_0.text = ""
    p = cell_1_0.paragraphs[0]
    p.alignment = alignment_map[0]
    run = p.add_run("Phase 1")
    run.bold = True

    cell_3_3 = table.cell(3, 3)
    cell_3_3.text = ""
    p = cell_3_3.paragraphs[0]
    p.alignment = alignment_map[3]
    run = p.add_run("Charlie")
    run.italic = True

    # Apply horizontal merge: row 0, cols 0-1
    table.cell(0, 0).merge(table.cell(0, 1))

    # Apply vertical merge: rows 3-4, col 0
    table.cell(3, 0).merge(table.cell(4, 0))

    doc.add_paragraph("Text after the complex table.")
    doc.save(str(FIXTURES_DIR / "tables_complex.docx"))
    print("✓ Created tables_complex.docx")


def _make_textbox_drawing_xml(
    content_paragraphs: list[str],
    *,
    anchor: bool = True,
    width_emu: int = 3657600,
    height_emu: int = 1371600,
    pos_x: int = 914400,
    pos_y: int = 914400,
    border_color: str | None = None,
    fill_color: str | None = None,
    bold_runs: set[int] | None = None,
) -> str:
    """Build DrawingML XML string for a text box."""
    bold_runs = bold_runs or set()

    paras_xml = ""
    for idx, text in enumerate(content_paragraphs):
        rpr = '<a:rPr lang="en-US" b="1"/>' if idx in bold_runs else '<a:rPr lang="en-US"/>'
        paras_xml += f'<a:p><a:r>{rpr}<a:t>{text}</a:t></a:r></a:p>'

    fill_xml = '<a:noFill/>' if fill_color is None else (
        f'<a:solidFill><a:srgbClr val="{fill_color}"/></a:solidFill>'
    )
    line_xml = (
        f'<a:ln w="12700"><a:solidFill><a:srgbClr val="{border_color}"/></a:solidFill></a:ln>'
        if border_color else '<a:ln><a:noFill/></a:ln>'
    )

    wsp_xml = (
        '<wps:wsp>'
        '<wps:cNvSpPr txBox="1"/>'
        f'<wps:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{width_emu}" cy="{height_emu}"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        f'{fill_xml}{line_xml}'
        '</wps:spPr>'
        f'<wps:txbxContent>{paras_xml}</wps:txbxContent>'
        '</wps:wsp>'
    )

    graphic_xml = (
        '<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        f'{wsp_xml}'
        '</a:graphicData>'
        '</a:graphic>'
    )

    ns_attrs = (
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" '
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
    )

    if anchor:
        return (
            f'<wp:anchor distT="0" distB="0" distL="114300" distR="114300" '
            f'simplePos="0" relativeHeight="0" behindDoc="0" locked="0" '
            f'layoutInCell="1" allowOverlap="1" {ns_attrs}>'
            f'<wp:simplePos x="0" y="0"/>'
            f'<wp:positionH relativeFrom="column"><wp:posOffset>{pos_x}</wp:posOffset></wp:positionH>'
            f'<wp:positionV relativeFrom="paragraph"><wp:posOffset>{pos_y}</wp:posOffset></wp:positionV>'
            f'<wp:extent cx="{width_emu}" cy="{height_emu}"/>'
            f'<wp:wrapNone/>'
            f'{graphic_xml}'
            f'</wp:anchor>'
        )
    else:
        return (
            f'<wp:inline distT="0" distB="0" distL="0" distR="0" {ns_attrs}>'
            f'<wp:extent cx="{width_emu}" cy="{height_emu}"/>'
            f'{graphic_xml}'
            f'</wp:inline>'
        )


def create_textboxes_docx() -> None:
    """Create textboxes.docx with various text box configurations."""
    from lxml import etree

    doc = Document()

    doc.add_heading("Document with Text Boxes", level=1)
    doc.add_paragraph("Text before the first text box.")

    # Text box 1: simple anchored text box
    p1 = doc.add_paragraph()
    run1 = p1.add_run()
    drawing_xml = _make_textbox_drawing_xml(
        ["This is a simple text box."],
        anchor=True, pos_x=914400, pos_y=914400,
        width_emu=3657600, height_emu=914400,
    )
    drawing_elem = etree.fromstring(
        f'<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'{drawing_xml}</w:drawing>'
    )
    run1._element.append(drawing_elem)

    doc.add_paragraph("Text between text boxes.")

    # Text box 2: inline text box with formatting
    p2 = doc.add_paragraph()
    run2 = p2.add_run()
    drawing_xml2 = _make_textbox_drawing_xml(
        ["Bold heading inside text box", "Regular paragraph in text box."],
        anchor=False, width_emu=4572000, height_emu=1371600,
        border_color="000000", fill_color="FFFF00", bold_runs={0},
    )
    drawing_elem2 = etree.fromstring(
        f'<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'{drawing_xml2}</w:drawing>'
    )
    run2._element.append(drawing_elem2)

    # Text box 3: shape with text (not txBox flag, roundRect shape)
    p3 = doc.add_paragraph()
    run3 = p3.add_run()
    shape_xml = (
        '<wp:anchor distT="0" distB="0" distL="114300" distR="114300" '
        'simplePos="0" relativeHeight="0" behindDoc="0" locked="0" '
        'layoutInCell="1" allowOverlap="1" '
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" '
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        '<wp:simplePos x="0" y="0"/>'
        '<wp:positionH relativeFrom="column"><wp:posOffset>914400</wp:posOffset></wp:positionH>'
        '<wp:positionV relativeFrom="paragraph"><wp:posOffset>2286000</wp:posOffset></wp:positionV>'
        '<wp:extent cx="2743200" cy="914400"/>'
        '<wp:wrapNone/>'
        '<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wps:wsp>'
        '<wps:cNvSpPr/>'
        '<wps:spPr>'
        '<a:xfrm><a:off x="0" y="0"/><a:ext cx="2743200" cy="914400"/></a:xfrm>'
        '<a:prstGeom prst="roundRect"><a:avLst/></a:prstGeom>'
        '<a:solidFill><a:srgbClr val="4472C4"/></a:solidFill>'
        '</wps:spPr>'
        '<wps:txbxContent>'
        '<a:p><a:r><a:rPr lang="en-US"/><a:t>Text inside a shape.</a:t></a:r></a:p>'
        '</wps:txbxContent>'
        '</wps:wsp>'
        '</a:graphicData>'
        '</a:graphic>'
        '</wp:anchor>'
    )
    drawing_elem3 = etree.fromstring(
        f'<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'{shape_xml}</w:drawing>'
    )
    run3._element.append(drawing_elem3)

    doc.add_paragraph("Text after all text boxes.")

    doc.save(str(FIXTURES_DIR / "textboxes.docx"))
    print("✓ Created textboxes.docx")


if __name__ == "__main__":
    print("Creating test fixtures...")
    create_simple_docx()
    create_lists_docx()
    create_formatted_docx()
    create_images_docx()
    create_complex_docx()
    create_tables_simple_docx()
    create_tables_formatted_docx()
    create_tables_merged_docx()
    create_tables_complex_docx()
    create_textboxes_docx()
    print("\nAll fixtures created successfully!")
