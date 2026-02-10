"""Script to create test fixtures with track changes.

Track changes in Word use w:ins and w:del XML elements that python-docx
doesn't directly support for creation. This script manipulates the XML
directly to create proper track change fixtures.
"""

from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

# Get the directory where this script is located
FIXTURES_DIR = Path(__file__).parent

# XML namespace for preserving whitespace
XML_SPACE_NS = "{http://www.w3.org/XML/1998/namespace}space"


def clear_paragraph_runs(p_element: etree._Element) -> None:
    """Remove all default w:r (run) elements from a paragraph element.

    When adding track changes to a paragraph, we need to clear any default
    runs that python-docx adds so we can control the exact structure.

    Args:
        p_element: The paragraph's underlying XML element
    """
    for child in list(p_element):
        if child.tag == qn("w:r"):
            p_element.remove(child)


def create_ins_element(text: str, author: str, date: str, revision_id: str) -> etree._Element:
    """Create a w:ins element with proper structure.

    Args:
        text: The inserted text
        author: Author name
        date: ISO 8601 date string
        revision_id: Unique revision ID

    Returns:
        lxml Element representing the insertion
    """
    # Create w:ins element
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), revision_id)
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), date)

    # Create w:r (run) inside the insertion
    run = OxmlElement("w:r")
    ins.append(run)

    # Create w:t (text) inside the run
    t = OxmlElement("w:t")
    t.text = text
    # Preserve whitespace
    t.set(XML_SPACE_NS, "preserve")
    run.append(t)

    return ins


def create_del_element(text: str, author: str, date: str, revision_id: str) -> etree._Element:
    """Create a w:del element with proper structure.

    Args:
        text: The deleted text
        author: Author name
        date: ISO 8601 date string
        revision_id: Unique revision ID

    Returns:
        lxml Element representing the deletion
    """
    # Create w:del element
    del_elem = OxmlElement("w:del")
    del_elem.set(qn("w:id"), revision_id)
    del_elem.set(qn("w:author"), author)
    del_elem.set(qn("w:date"), date)

    # Create w:r (run) inside the deletion
    run = OxmlElement("w:r")
    del_elem.append(run)

    # Create w:delText (deleted text) inside the run
    del_text = OxmlElement("w:delText")
    del_text.text = text
    # Preserve whitespace
    del_text.set(XML_SPACE_NS, "preserve")
    run.append(del_text)

    return del_elem


def create_run_element(text: str) -> etree._Element:
    """Create a w:r (run) element with text.

    Args:
        text: The text content

    Returns:
        lxml Element representing the run
    """
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(XML_SPACE_NS, "preserve")
    run.append(t)
    return run


def create_track_changes_simple() -> None:
    """Create track_changes_simple.docx with basic insertions and deletions.

    Contents:
    - Paragraph 1: "This is some text." (no changes)
    - Paragraph 2: "Hello inserted world" - "inserted" is a tracked insertion
    - Paragraph 3: "Hello world" with "deleted " tracked as deletion
    """
    doc = Document()

    # Use a fixed date for reproducibility
    test_date = "2026-01-15T10:30:00Z"
    author = "Test Author"

    # Paragraph 1: No changes
    doc.add_paragraph("This is some text.")

    # Paragraph 2: With insertion
    p2 = doc.add_paragraph()
    p2_element = p2._element
    clear_paragraph_runs(p2_element)

    # Add "Hello "
    p2_element.append(create_run_element("Hello "))

    # Add insertion "{++inserted++}"
    p2_element.append(create_ins_element("inserted", author, test_date, "1"))

    # Add " world"
    p2_element.append(create_run_element(" world"))

    # Paragraph 3: With deletion
    p3 = doc.add_paragraph()
    p3_element = p3._element
    clear_paragraph_runs(p3_element)

    # Add "Hello "
    p3_element.append(create_run_element("Hello "))

    # Add deletion "{--deleted--}"
    p3_element.append(create_del_element("deleted ", author, test_date, "2"))

    # Add "world"
    p3_element.append(create_run_element("world"))

    doc.save(str(FIXTURES_DIR / "track_changes_simple.docx"))
    print("✓ Created track_changes_simple.docx")


def create_track_changes_paragraph() -> None:
    """Create track_changes_paragraph.docx with multiple changes in one paragraph.

    Contents:
    - Paragraph 1: "This has multiple removed changes added here"
      - "has" is inserted
      - "removed" is deleted
      - "added" is inserted
    """
    doc = Document()

    test_date = "2026-01-15T11:00:00Z"
    author = "Test Author"

    # Paragraph with multiple track changes
    p = doc.add_paragraph()
    p_element = p._element
    clear_paragraph_runs(p_element)

    # Build: "This {++has++} multiple {--removed--} changes {++added++} here"
    p_element.append(create_run_element("This "))
    p_element.append(create_ins_element("has", author, test_date, "1"))
    p_element.append(create_run_element(" multiple "))
    p_element.append(create_del_element("removed", author, test_date, "2"))
    p_element.append(create_run_element(" changes "))
    p_element.append(create_ins_element("added", author, test_date, "3"))
    p_element.append(create_run_element(" here"))

    doc.save(str(FIXTURES_DIR / "track_changes_paragraph.docx"))
    print("✓ Created track_changes_paragraph.docx")


def create_track_changes_multiauthor() -> None:
    """Create track_changes_multiauthor.docx with changes from multiple authors.

    Contents:
    - Paragraph 1: "Alice added this" - "added this" inserted by Alice
    - Paragraph 2: "Bob world" - "removed this " deleted by Bob
    - Paragraph 3: "Charlie added and removed" - mixed authors
    """
    doc = Document()

    date_alice = "2026-01-15T09:00:00Z"
    date_bob = "2026-01-15T10:00:00Z"
    date_charlie = "2026-01-15T11:00:00Z"
    date_dave = "2026-01-15T12:00:00Z"

    # Paragraph 1: Alice's insertion
    p1 = doc.add_paragraph()
    p1_element = p1._element
    clear_paragraph_runs(p1_element)

    p1_element.append(create_run_element("Alice "))
    p1_element.append(create_ins_element("added this", "Alice", date_alice, "1"))

    # Paragraph 2: Bob's deletion
    p2 = doc.add_paragraph()
    p2_element = p2._element
    clear_paragraph_runs(p2_element)

    p2_element.append(create_run_element("Bob "))
    p2_element.append(create_del_element("removed this ", "Bob", date_bob, "2"))
    p2_element.append(create_run_element("world"))

    # Paragraph 3: Mixed - Charlie insertion and Dave deletion
    p3 = doc.add_paragraph()
    p3_element = p3._element
    clear_paragraph_runs(p3_element)

    p3_element.append(create_ins_element("Charlie added", "Charlie", date_charlie, "3"))
    p3_element.append(create_run_element(" and "))
    p3_element.append(create_del_element("Dave removed", "Dave", date_dave, "4"))

    doc.save(str(FIXTURES_DIR / "track_changes_multiauthor.docx"))
    print("✓ Created track_changes_multiauthor.docx")


def create_track_changes_headings() -> None:
    """Create track_changes_headings.docx with track changes in headings.

    Contents:
    - Heading 1: "Title with {++inserted++} text"
    - Heading 2: "Subtitle {--removed--} here"
    """
    doc = Document()

    test_date = "2026-01-15T10:30:00Z"
    author = "Test Author"

    # Heading 1 with insertion
    h1 = doc.add_heading("", level=1)
    h1_element = h1._element
    clear_paragraph_runs(h1_element)

    h1_element.append(create_run_element("Title with "))
    h1_element.append(create_ins_element("inserted", author, test_date, "1"))
    h1_element.append(create_run_element(" text"))

    # Heading 2 with deletion
    h2 = doc.add_heading("", level=2)
    h2_element = h2._element
    clear_paragraph_runs(h2_element)

    h2_element.append(create_run_element("Subtitle "))
    h2_element.append(create_del_element("removed", author, test_date, "2"))
    h2_element.append(create_run_element(" here"))

    doc.save(str(FIXTURES_DIR / "track_changes_headings.docx"))
    print("✓ Created track_changes_headings.docx")


def create_track_changes_lists() -> None:
    """Create track_changes_lists.docx with track changes in list items.

    Contents:
    - Bullet 1: "First {++added++} item"
    - Bullet 2: "Second {--removed--} item"
    - Number 1: "Numbered {++inserted++} item"
    """
    doc = Document()

    test_date = "2026-01-15T10:30:00Z"
    author = "Test Author"

    # Bulleted list with insertion
    bullet1 = doc.add_paragraph("", style="List Bullet")
    b1_element = bullet1._element
    clear_paragraph_runs(b1_element)

    b1_element.append(create_run_element("First "))
    b1_element.append(create_ins_element("added", author, test_date, "1"))
    b1_element.append(create_run_element(" item"))

    # Bulleted list with deletion
    bullet2 = doc.add_paragraph("", style="List Bullet")
    b2_element = bullet2._element
    clear_paragraph_runs(b2_element)

    b2_element.append(create_run_element("Second "))
    b2_element.append(create_del_element("removed", author, test_date, "2"))
    b2_element.append(create_run_element(" item"))

    # Numbered list with insertion
    number1 = doc.add_paragraph("", style="List Number")
    n1_element = number1._element
    clear_paragraph_runs(n1_element)

    n1_element.append(create_run_element("Numbered "))
    n1_element.append(create_ins_element("inserted", author, test_date, "3"))
    n1_element.append(create_run_element(" item"))

    doc.save(str(FIXTURES_DIR / "track_changes_lists.docx"))
    print("✓ Created track_changes_lists.docx")


if __name__ == "__main__":
    print("Creating track changes test fixtures...")
    create_track_changes_simple()
    create_track_changes_paragraph()
    create_track_changes_multiauthor()
    create_track_changes_headings()
    create_track_changes_lists()
    print("\nAll track changes fixtures created successfully!")
