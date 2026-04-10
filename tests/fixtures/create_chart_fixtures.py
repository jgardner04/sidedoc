"""Script to create test fixture docx files with charts and SmartArt.

Charts and SmartArt require manual OOXML construction since python-docx
has no API for them. We build the XML parts directly and inject them
into the docx ZIP archive.
"""

import io
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET
from docx import Document

FIXTURES_DIR = Path(__file__).parent

# Namespace URIs
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
A = "http://schemas.openxmlformats.org/drawingml/2006/main"
C = "http://schemas.openxmlformats.org/drawingml/2006/chart"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
DGM = "http://schemas.openxmlformats.org/drawingml/2006/diagram"

# Relationship types
CHART_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart"
IMAGE_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
DGM_DATA_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/diagramData"
DGM_LAYOUT_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/diagramLayout"
DGM_STYLE_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/diagramStyle"
DGM_COLORS_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/diagramColors"

# Content types
CHART_CT = "application/vnd.openxmlformats-officedocument.drawingml.chart+xml"
DGM_DATA_CT = "application/vnd.openxmlformats-officedocument.drawingml.diagramData+xml"
DGM_LAYOUT_CT = "application/vnd.openxmlformats-officedocument.drawingml.diagramLayout+xml"
DGM_STYLE_CT = "application/vnd.openxmlformats-officedocument.drawingml.diagramStyle+xml"
DGM_COLORS_CT = "application/vnd.openxmlformats-officedocument.drawingml.diagramColors+xml"

# Register namespaces for clean serialization
for prefix, uri in [
    ("a", A), ("c", C), ("r", R), ("wp", WP), ("w", W), ("dgm", DGM),
]:
    ET.register_namespace(prefix, uri)


def _el(tag: str, text: str | None = None, **attribs: str) -> ET.Element:
    """Shorthand to create an Element."""
    elem = ET.Element(tag, attribs)
    if text is not None:
        elem.text = text
    return elem


def _sub(parent: ET.Element, tag: str, text: str | None = None, **attribs: str) -> ET.Element:
    """Shorthand to create a SubElement."""
    elem = ET.SubElement(parent, tag, attribs)
    if text is not None:
        elem.text = text
    return elem


def _to_xml(elem: ET.Element) -> bytes:
    """Serialize element to XML bytes."""
    return ET.tostring(elem, xml_declaration=True, encoding="UTF-8")


def _add_chart_series(parent: ET.Element, idx: str, name: str,
                      categories: list[str], values: list[str]) -> None:
    """Add a data series to a chart element."""
    ser = _sub(parent, f"{{{C}}}ser")
    _sub(ser, f"{{{C}}}idx", val=idx)
    _sub(ser, f"{{{C}}}order", val=idx)

    # Series name
    tx = _sub(ser, f"{{{C}}}tx")
    str_ref = _sub(tx, f"{{{C}}}strRef")
    str_cache = _sub(str_ref, f"{{{C}}}strCache")
    _sub(str_cache, f"{{{C}}}ptCount", val="1")
    pt = _sub(str_cache, f"{{{C}}}pt", idx="0")
    _sub(pt, f"{{{C}}}v", text=name)

    # Categories
    cat = _sub(ser, f"{{{C}}}cat")
    str_ref_cat = _sub(cat, f"{{{C}}}strRef")
    str_cache_cat = _sub(str_ref_cat, f"{{{C}}}strCache")
    _sub(str_cache_cat, f"{{{C}}}ptCount", val=str(len(categories)))
    for i, label in enumerate(categories):
        pt = _sub(str_cache_cat, f"{{{C}}}pt", idx=str(i))
        _sub(pt, f"{{{C}}}v", text=label)

    # Values
    val_elem = _sub(ser, f"{{{C}}}val")
    num_ref = _sub(val_elem, f"{{{C}}}numRef")
    num_cache = _sub(num_ref, f"{{{C}}}numCache")
    _sub(num_cache, f"{{{C}}}ptCount", val=str(len(values)))
    for i, value in enumerate(values):
        pt = _sub(num_cache, f"{{{C}}}pt", idx=str(i))
        _sub(pt, f"{{{C}}}v", text=value)


def _create_bar_chart_xml(title: str = "Q4 Revenue") -> bytes:
    """Create a bar chart XML part with one series."""
    chart_space = _el(f"{{{C}}}chartSpace")
    chart = _sub(chart_space, f"{{{C}}}chart")

    # Title
    title_el = _sub(chart, f"{{{C}}}title")
    tx = _sub(title_el, f"{{{C}}}tx")
    rich = _sub(tx, f"{{{C}}}rich")
    _sub(rich, f"{{{A}}}bodyPr")
    p = _sub(rich, f"{{{A}}}p")
    r = _sub(p, f"{{{A}}}r")
    _sub(r, f"{{{A}}}t", text=title)

    # Plot area with bar chart
    plot_area = _sub(chart, f"{{{C}}}plotArea")
    bar_chart = _sub(plot_area, f"{{{C}}}barChart")
    _sub(bar_chart, f"{{{C}}}barDir", val="col")

    _add_chart_series(
        bar_chart, "0", "Revenue",
        categories=["Oct", "Nov", "Dec"],
        values=["1200000", "1350000", "1500000"],
    )

    return _to_xml(chart_space)


def _create_pie_chart_xml(title: str = "Market Share") -> bytes:
    """Create a pie chart XML part with one series."""
    chart_space = _el(f"{{{C}}}chartSpace")
    chart = _sub(chart_space, f"{{{C}}}chart")

    title_el = _sub(chart, f"{{{C}}}title")
    tx = _sub(title_el, f"{{{C}}}tx")
    rich = _sub(tx, f"{{{C}}}rich")
    _sub(rich, f"{{{A}}}bodyPr")
    p = _sub(rich, f"{{{A}}}p")
    r = _sub(p, f"{{{A}}}r")
    _sub(r, f"{{{A}}}t", text=title)

    plot_area = _sub(chart, f"{{{C}}}plotArea")
    pie_chart = _sub(plot_area, f"{{{C}}}pieChart")

    _add_chart_series(
        pie_chart, "0", "Share",
        categories=["Product A", "Product B", "Product C"],
        values=["45", "30", "25"],
    )

    return _to_xml(chart_space)


def _create_smartart_data_xml() -> bytes:
    """Create SmartArt diagram data XML (org chart with 3 nodes)."""
    data_model = _el(f"{{{DGM}}}dataModel")
    pt_lst = _sub(data_model, f"{{{DGM}}}ptLst")

    # Root node
    pt_root = _sub(pt_lst, f"{{{DGM}}}pt", modelId="0", type="doc")
    _sub(pt_root, f"{{{DGM}}}t", text="")

    # Child nodes
    for model_id, text in [("1", "CEO"), ("2", "VP Engineering"), ("3", "VP Sales")]:
        pt = _sub(pt_lst, f"{{{DGM}}}pt", modelId=model_id, type="node")
        _sub(pt, f"{{{DGM}}}t", text=text)

    # Connections
    cxn_lst = _sub(data_model, f"{{{DGM}}}cxnLst")
    for child_id in ["1", "2", "3"]:
        _sub(cxn_lst, f"{{{DGM}}}cxn",
             modelId=f"c{child_id}", type="parOf",
             srcId="0", destId=child_id, sibTransId="0")

    return _to_xml(data_model)


def _create_smartart_layout_xml() -> bytes:
    layout = _el(f"{{{DGM}}}layoutDef")
    _sub(layout, f"{{{DGM}}}title", val="Organization Chart")
    return _to_xml(layout)


def _create_smartart_style_xml() -> bytes:
    return _to_xml(_el(f"{{{DGM}}}styleDef"))


def _create_smartart_colors_xml() -> bytes:
    return _to_xml(_el(f"{{{DGM}}}colorsDef"))


def _create_fallback_image() -> bytes:
    """Create a small PNG for chart/SmartArt fallback rendering."""
    from PIL import Image as PILImage
    img = PILImage.new("RGB", (200, 150), color="#4472C4")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _make_chart_drawing_xml(chart_rel_id: str, fallback_img_rel_id: str | None = None) -> ET.Element:
    """Build the w:drawing element that references a chart."""
    r_elem = _el(f"{{{W}}}r")
    drawing = _sub(r_elem, f"{{{W}}}drawing")
    inline = _sub(drawing, f"{{{WP}}}inline")
    _sub(inline, f"{{{WP}}}extent", cx="5486400", cy="3200400")

    graphic = _sub(inline, f"{{{A}}}graphic")
    gd = _sub(graphic, f"{{{A}}}graphicData",
              uri="http://schemas.openxmlformats.org/drawingml/2006/chart")
    _sub(gd, f"{{{C}}}chart", **{f"{{{R}}}id": chart_rel_id})

    if fallback_img_rel_id:
        blip_fill = _sub(inline, f"{{{A}}}blipFill")
        _sub(blip_fill, f"{{{A}}}blip", **{f"{{{R}}}embed": fallback_img_rel_id})

    return r_elem


def _make_smartart_drawing_xml(
    dm_id: str, lo_id: str, qs_id: str, cs_id: str,
    fallback_img_rel_id: str | None = None,
) -> ET.Element:
    """Build the w:drawing element that references SmartArt."""
    r_elem = _el(f"{{{W}}}r")
    drawing = _sub(r_elem, f"{{{W}}}drawing")
    inline = _sub(drawing, f"{{{WP}}}inline")
    _sub(inline, f"{{{WP}}}extent", cx="5486400", cy="3200400")

    graphic = _sub(inline, f"{{{A}}}graphic")
    gd = _sub(graphic, f"{{{A}}}graphicData",
              uri="http://schemas.openxmlformats.org/drawingml/2006/diagram")
    _sub(gd, f"{{{DGM}}}relIds", **{
        f"{{{R}}}dm": dm_id,
        f"{{{R}}}lo": lo_id,
        f"{{{R}}}qs": qs_id,
        f"{{{R}}}cs": cs_id,
    })

    if fallback_img_rel_id:
        blip_fill = _sub(inline, f"{{{A}}}blipFill")
        _sub(blip_fill, f"{{{A}}}blip", **{f"{{{R}}}embed": fallback_img_rel_id})

    return r_elem


def _inject_into_docx(
    docx_bytes: bytes,
    drawing_elem: ET.Element,
    paragraph_index: int,
    extra_parts: dict[str, bytes],
    extra_rels: list[tuple[str, str, str]],
    extra_content_types: list[tuple[str, str]],
) -> bytes:
    """Inject drawing XML and parts into a docx ZIP.

    Args:
        docx_bytes: Base docx bytes
        drawing_elem: The w:r element containing w:drawing to inject
        paragraph_index: Which paragraph to replace
        extra_parts: Map of part path -> XML bytes to add
        extra_rels: List of (Id, Type, Target) for document.xml.rels
        extra_content_types: List of (PartName, ContentType) overrides
    """
    src = zipfile.ZipFile(io.BytesIO(docx_bytes), "r")
    out_buf = io.BytesIO()
    dst = zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED)

    # Also register the relationships namespace used in .rels files
    ET.register_namespace("", "http://schemas.openxmlformats.org/package/2006/relationships")

    for item in src.infolist():
        data = src.read(item.filename)

        if item.filename == "word/document.xml":
            doc_tree = ET.fromstring(data)
            paragraphs = doc_tree.findall(f".//{{{W}}}p")
            if paragraph_index < len(paragraphs):
                para = paragraphs[paragraph_index]
                for child in list(para):
                    para.remove(child)
                para.append(drawing_elem)
            data = _to_xml(doc_tree)

        elif item.filename == "word/_rels/document.xml.rels":
            rels_tree = ET.fromstring(data)
            for rel_id, rel_type, target in extra_rels:
                rel = ET.SubElement(rels_tree, "Relationship")
                rel.set("Id", rel_id)
                rel.set("Type", rel_type)
                rel.set("Target", target)
            data = _to_xml(rels_tree)

        elif item.filename == "[Content_Types].xml":
            # Register the content types namespace
            ET.register_namespace("", "http://schemas.openxmlformats.org/package/2006/content-types")
            ct_tree = ET.fromstring(data)
            for part_name, content_type in extra_content_types:
                override = ET.SubElement(ct_tree, "Override")
                override.set("PartName", part_name)
                override.set("ContentType", content_type)
            data = _to_xml(ct_tree)

        dst.writestr(item, data)

    for part_path, part_bytes in extra_parts.items():
        dst.writestr(part_path, part_bytes)

    src.close()
    dst.close()
    return out_buf.getvalue()


def _make_base_docx(heading: str, placeholder: str, after: str) -> bytes:
    """Create a simple 3-paragraph docx."""
    doc = Document()
    doc.add_heading(heading, level=1)
    doc.add_paragraph(placeholder)
    doc.add_paragraph(after)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def create_chart_bar_docx() -> None:
    """Create chart_bar.docx with a bar chart and fallback image."""
    base = _make_base_docx("Bar Chart Example", "Chart placeholder.", "Text after the chart.")
    chart_xml = _create_bar_chart_xml("Q4 Revenue")
    fallback = _create_fallback_image()

    drawing = _make_chart_drawing_xml("rIdChart1", "rIdChartImg1")
    result = _inject_into_docx(
        base, drawing, paragraph_index=1,
        extra_parts={
            "word/charts/chart1.xml": chart_xml,
            "word/media/chart_fallback1.png": fallback,
        },
        extra_rels=[
            ("rIdChart1", CHART_REL_TYPE, "charts/chart1.xml"),
            ("rIdChartImg1", IMAGE_REL_TYPE, "media/chart_fallback1.png"),
        ],
        extra_content_types=[
            ("/word/charts/chart1.xml", CHART_CT),
        ],
    )
    (FIXTURES_DIR / "chart_bar.docx").write_bytes(result)
    print("Created chart_bar.docx")


def create_chart_pie_docx() -> None:
    """Create chart_pie.docx with a pie chart but NO fallback image."""
    base = _make_base_docx("Pie Chart Example", "Chart placeholder.", "Text after the chart.")
    chart_xml = _create_pie_chart_xml("Market Share")

    drawing = _make_chart_drawing_xml("rIdChart1", fallback_img_rel_id=None)
    result = _inject_into_docx(
        base, drawing, paragraph_index=1,
        extra_parts={"word/charts/chart1.xml": chart_xml},
        extra_rels=[("rIdChart1", CHART_REL_TYPE, "charts/chart1.xml")],
        extra_content_types=[("/word/charts/chart1.xml", CHART_CT)],
    )
    (FIXTURES_DIR / "chart_pie.docx").write_bytes(result)
    print("Created chart_pie.docx (no cached image)")


def create_smartart_orgchart_docx() -> None:
    """Create smartart_orgchart.docx with a SmartArt org chart and fallback image."""
    base = _make_base_docx("Org Chart Example", "SmartArt placeholder.", "Text after the SmartArt.")
    fallback = _create_fallback_image()

    drawing = _make_smartart_drawing_xml(
        "rIdDgmData1", "rIdDgmLayout1", "rIdDgmStyle1", "rIdDgmColors1",
        fallback_img_rel_id="rIdDgmImg1",
    )
    result = _inject_into_docx(
        base, drawing, paragraph_index=1,
        extra_parts={
            "word/diagrams/data1.xml": _create_smartart_data_xml(),
            "word/diagrams/layout1.xml": _create_smartart_layout_xml(),
            "word/diagrams/style1.xml": _create_smartart_style_xml(),
            "word/diagrams/colors1.xml": _create_smartart_colors_xml(),
            "word/media/smartart_fallback1.png": fallback,
        },
        extra_rels=[
            ("rIdDgmData1", DGM_DATA_REL, "diagrams/data1.xml"),
            ("rIdDgmLayout1", DGM_LAYOUT_REL, "diagrams/layout1.xml"),
            ("rIdDgmStyle1", DGM_STYLE_REL, "diagrams/style1.xml"),
            ("rIdDgmColors1", DGM_COLORS_REL, "diagrams/colors1.xml"),
            ("rIdDgmImg1", IMAGE_REL_TYPE, "media/smartart_fallback1.png"),
        ],
        extra_content_types=[
            ("/word/diagrams/data1.xml", DGM_DATA_CT),
            ("/word/diagrams/layout1.xml", DGM_LAYOUT_CT),
            ("/word/diagrams/style1.xml", DGM_STYLE_CT),
            ("/word/diagrams/colors1.xml", DGM_COLORS_CT),
        ],
    )
    (FIXTURES_DIR / "smartart_orgchart.docx").write_bytes(result)
    print("Created smartart_orgchart.docx")


if __name__ == "__main__":
    print("Creating chart and SmartArt test fixtures...")
    create_chart_bar_docx()
    create_chart_pie_docx()
    create_smartart_orgchart_docx()
    print("\nAll chart/SmartArt fixtures created!")
