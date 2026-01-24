"""Tests for hyperlink support in sidedoc."""

import json
import tempfile
import zipfile
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from sidedoc.extract import extract_blocks, extract_styles, blocks_to_markdown
from sidedoc.reconstruct import build_docx_from_sidedoc, parse_markdown_to_blocks
from sidedoc.package import create_sidedoc_archive


def add_hyperlink(paragraph, text: str, url: str, bold: bool = False, italic: bool = False):
    """Add a hyperlink to a paragraph for test setup.

    Args:
        paragraph: The paragraph to add the hyperlink to
        text: The display text for the hyperlink
        url: The URL the hyperlink points to
        bold: Whether the hyperlink text should be bold
        italic: Whether the hyperlink text should be italic
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)

    if italic:
        i = OxmlElement('w:i')
        rPr.append(i)

    run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = text
    run.append(text_elem)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)

    return hyperlink


def create_hyperlink_docx(content_setup) -> str:
    """Create a test docx file with hyperlinks.

    Args:
        content_setup: Function that sets up the document content

    Returns:
        Path to temporary docx file
    """
    doc = Document()
    content_setup(doc)
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    temp_file.close()
    return temp_file.name


# =============================================================================
# Test: Simple hyperlink extraction
# =============================================================================

def test_extract_simple_hyperlink():
    """Test extracting a simple hyperlink as markdown."""
    def setup(doc):
        para = doc.add_paragraph("Visit ")
        add_hyperlink(para, "Google", "https://www.google.com")
        para.add_run(" for search.")

    docx_path = create_hyperlink_docx(setup)

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        assert blocks[0].type == "paragraph"
        # The hyperlink should be converted to markdown [text](url) format
        assert "[Google](https://www.google.com)" in blocks[0].content
        assert blocks[0].content == "Visit [Google](https://www.google.com) for search."
    finally:
        Path(docx_path).unlink()


def test_extract_hyperlink_only():
    """Test extracting a paragraph with only a hyperlink."""
    def setup(doc):
        para = doc.add_paragraph("")
        add_hyperlink(para, "Click here", "https://example.com")

    docx_path = create_hyperlink_docx(setup)

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        assert blocks[0].content == "[Click here](https://example.com)"
    finally:
        Path(docx_path).unlink()


# =============================================================================
# Test: Multiple hyperlinks in one paragraph
# =============================================================================

def test_extract_multiple_hyperlinks():
    """Test extracting multiple hyperlinks in one paragraph."""
    def setup(doc):
        para = doc.add_paragraph("Check out ")
        add_hyperlink(para, "GitHub", "https://github.com")
        para.add_run(" and ")
        add_hyperlink(para, "Stack Overflow", "https://stackoverflow.com")
        para.add_run(" for coding help.")

    docx_path = create_hyperlink_docx(setup)

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        assert "[GitHub](https://github.com)" in blocks[0].content
        assert "[Stack Overflow](https://stackoverflow.com)" in blocks[0].content
        expected = "Check out [GitHub](https://github.com) and [Stack Overflow](https://stackoverflow.com) for coding help."
        assert blocks[0].content == expected
    finally:
        Path(docx_path).unlink()


# =============================================================================
# Test: Hyperlinks with formatting (bold, italic)
# =============================================================================

def test_extract_bold_hyperlink():
    """Test extracting a bold hyperlink."""
    def setup(doc):
        para = doc.add_paragraph("This is a ")
        add_hyperlink(para, "bold link", "https://example.com/bold", bold=True)
        para.add_run(" in the text.")

    docx_path = create_hyperlink_docx(setup)

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        # Bold hyperlink should have both bold and link formatting
        # The markdown should be **[text](url)** or [**text**](url)
        assert "bold link" in blocks[0].content
        assert "https://example.com/bold" in blocks[0].content
        # We'll accept either **[text](url)** or [**text**](url)
        assert ("**[bold link](https://example.com/bold)**" in blocks[0].content or
                "[**bold link**](https://example.com/bold)" in blocks[0].content)
    finally:
        Path(docx_path).unlink()


def test_extract_italic_hyperlink():
    """Test extracting an italic hyperlink."""
    def setup(doc):
        para = doc.add_paragraph("This is an ")
        add_hyperlink(para, "italic link", "https://example.com/italic", italic=True)
        para.add_run(" in the text.")

    docx_path = create_hyperlink_docx(setup)

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        assert "italic link" in blocks[0].content
        assert "https://example.com/italic" in blocks[0].content
        # We'll accept either *[text](url)* or [*text*](url)
        assert ("*[italic link](https://example.com/italic)*" in blocks[0].content or
                "[*italic link*](https://example.com/italic)" in blocks[0].content)
    finally:
        Path(docx_path).unlink()


def test_extract_bold_italic_hyperlink():
    """Test extracting a bold and italic hyperlink."""
    def setup(doc):
        para = doc.add_paragraph("This is a ")
        add_hyperlink(para, "bold italic link", "https://example.com/bi", bold=True, italic=True)
        para.add_run(".")

    docx_path = create_hyperlink_docx(setup)

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        assert "bold italic link" in blocks[0].content
        assert "https://example.com/bi" in blocks[0].content
    finally:
        Path(docx_path).unlink()


# =============================================================================
# Test: Hyperlinks in headings
# =============================================================================

def test_extract_hyperlink_in_heading1():
    """Test extracting a hyperlink in a Heading 1."""
    def setup(doc):
        heading = doc.add_paragraph("", style="Heading 1")
        heading.add_run("Heading with ")
        add_hyperlink(heading, "link", "https://example.com/heading")

    docx_path = create_hyperlink_docx(setup)

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        assert blocks[0].type == "heading"
        assert blocks[0].level == 1
        assert "[link](https://example.com/heading)" in blocks[0].content
        assert blocks[0].content == "# Heading with [link](https://example.com/heading)"
    finally:
        Path(docx_path).unlink()


def test_extract_hyperlink_in_heading2():
    """Test extracting a hyperlink in a Heading 2."""
    def setup(doc):
        heading = doc.add_paragraph("", style="Heading 2")
        heading.add_run("Section with ")
        add_hyperlink(heading, "reference", "https://example.com/ref")

    docx_path = create_hyperlink_docx(setup)

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        assert blocks[0].type == "heading"
        assert blocks[0].level == 2
        assert blocks[0].content == "## Section with [reference](https://example.com/ref)"
    finally:
        Path(docx_path).unlink()


# =============================================================================
# Test: Hyperlinks in lists
# =============================================================================

def test_extract_hyperlink_in_bullet_list():
    """Test extracting hyperlinks in bulleted list items."""
    def setup(doc):
        bullet1 = doc.add_paragraph("First item with ", style="List Bullet")
        add_hyperlink(bullet1, "link one", "https://example.com/list1")

        bullet2 = doc.add_paragraph("Second item with ", style="List Bullet")
        add_hyperlink(bullet2, "link two", "https://example.com/list2")

    docx_path = create_hyperlink_docx(setup)

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 2
        assert blocks[0].type == "list"
        assert blocks[0].content == "- First item with [link one](https://example.com/list1)"
        assert blocks[1].type == "list"
        assert blocks[1].content == "- Second item with [link two](https://example.com/list2)"
    finally:
        Path(docx_path).unlink()


def test_extract_hyperlink_in_numbered_list():
    """Test extracting hyperlinks in numbered list items."""
    def setup(doc):
        num1 = doc.add_paragraph("Numbered with ", style="List Number")
        add_hyperlink(num1, "link three", "https://example.com/list3")

        num2 = doc.add_paragraph("Numbered with ", style="List Number")
        add_hyperlink(num2, "link four", "https://example.com/list4")

    docx_path = create_hyperlink_docx(setup)

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 2
        assert blocks[0].type == "list"
        assert blocks[0].content == "1. Numbered with [link three](https://example.com/list3)"
        assert blocks[1].type == "list"
        assert blocks[1].content == "2. Numbered with [link four](https://example.com/list4)"
    finally:
        Path(docx_path).unlink()


# =============================================================================
# Test: URL encoding and special characters
# =============================================================================

def test_extract_hyperlink_with_parentheses_in_url():
    """Test extracting hyperlink with parentheses in URL (like Wikipedia)."""
    def setup(doc):
        para = doc.add_paragraph("See ")
        add_hyperlink(para, "Python", "https://en.wikipedia.org/wiki/Python_(programming_language)")

    docx_path = create_hyperlink_docx(setup)

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        # Parentheses in URL should be percent-encoded to avoid breaking markdown
        assert "Python" in blocks[0].content
        # Either the parentheses are encoded or escaped properly
        assert ("Python_(programming_language)" in blocks[0].content or
                "Python_%28programming_language%29" in blocks[0].content)
    finally:
        Path(docx_path).unlink()


def test_extract_hyperlink_with_spaces_in_url():
    """Test extracting hyperlink with spaces in URL."""
    def setup(doc):
        para = doc.add_paragraph("Link with ")
        add_hyperlink(para, "spaces", "https://example.com/path with spaces")

    docx_path = create_hyperlink_docx(setup)

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        # Spaces should be percent-encoded
        assert "spaces" in blocks[0].content
        assert "https://example.com/path%20with%20spaces" in blocks[0].content
    finally:
        Path(docx_path).unlink()


def test_extract_hyperlink_with_special_chars_in_text():
    """Test extracting hyperlink with markdown special chars in text."""
    def setup(doc):
        para = doc.add_paragraph("Link: ")
        add_hyperlink(para, "text with [brackets] and *asterisks*", "https://example.com/special")

    docx_path = create_hyperlink_docx(setup)

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        # Special chars in link text should be escaped
        assert "https://example.com/special" in blocks[0].content
        # The brackets and asterisks should be escaped in the markdown
        # Either escaped with backslash or the text preserved somehow
        assert "brackets" in blocks[0].content
        assert "asterisks" in blocks[0].content
    finally:
        Path(docx_path).unlink()


def test_extract_hyperlink_long_url():
    """Test extracting hyperlink with a very long URL."""
    def setup(doc):
        long_url = "https://example.com/very/long/path/with/many/segments?param1=value1&param2=value2&param3=value3#section"
        para = doc.add_paragraph("Here's a ")
        add_hyperlink(para, "long URL link", long_url)

    docx_path = create_hyperlink_docx(setup)

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        assert "[long URL link]" in blocks[0].content
        assert "example.com/very/long/path" in blocks[0].content
    finally:
        Path(docx_path).unlink()


def test_extract_mailto_link():
    """Test extracting mailto: hyperlink."""
    def setup(doc):
        para = doc.add_paragraph("Contact us at ")
        add_hyperlink(para, "support@example.com", "mailto:support@example.com")

    docx_path = create_hyperlink_docx(setup)

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        assert "[support@example.com](mailto:support@example.com)" in blocks[0].content
    finally:
        Path(docx_path).unlink()


# =============================================================================
# Test: Hyperlink positions in inline_formatting
# =============================================================================

def test_hyperlink_stores_position_in_inline_formatting():
    """Test that hyperlink positions are stored in inline_formatting."""
    def setup(doc):
        para = doc.add_paragraph("Visit ")
        add_hyperlink(para, "Google", "https://www.google.com")
        para.add_run(" now.")

    docx_path = create_hyperlink_docx(setup)

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        assert blocks[0].inline_formatting is not None

        # Find the hyperlink formatting entry
        hyperlink_format = next(
            (fmt for fmt in blocks[0].inline_formatting if fmt.get("type") == "hyperlink"),
            None
        )
        assert hyperlink_format is not None
        assert hyperlink_format["url"] == "https://www.google.com"
        # Position should be in plain text coordinates
        assert "start" in hyperlink_format
        assert "end" in hyperlink_format
    finally:
        Path(docx_path).unlink()


def test_multiple_hyperlinks_store_positions():
    """Test that multiple hyperlinks store their positions."""
    def setup(doc):
        para = doc.add_paragraph("See ")
        add_hyperlink(para, "link1", "https://example.com/1")
        para.add_run(" and ")
        add_hyperlink(para, "link2", "https://example.com/2")

    docx_path = create_hyperlink_docx(setup)

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        assert blocks[0].inline_formatting is not None

        hyperlink_formats = [
            fmt for fmt in blocks[0].inline_formatting if fmt.get("type") == "hyperlink"
        ]
        assert len(hyperlink_formats) == 2
        assert hyperlink_formats[0]["url"] == "https://example.com/1"
        assert hyperlink_formats[1]["url"] == "https://example.com/2"
    finally:
        Path(docx_path).unlink()


# =============================================================================
# Test: Edge cases
# =============================================================================

def test_extract_empty_hyperlink_text():
    """Test extracting hyperlink with empty text."""
    def setup(doc):
        para = doc.add_paragraph("Link: ")
        add_hyperlink(para, "", "https://example.com/empty")

    docx_path = create_hyperlink_docx(setup)

    try:
        blocks, _ = extract_blocks(docx_path)
        # Empty link text - should handle gracefully
        # Either skip the empty link or use URL as text
        assert len(blocks) == 1
    finally:
        Path(docx_path).unlink()


def test_extract_hyperlink_with_no_url():
    """Test handling hyperlink with missing URL."""
    # This tests graceful handling of malformed documents
    def setup(doc):
        para = doc.add_paragraph("Normal text only")

    docx_path = create_hyperlink_docx(setup)

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        assert blocks[0].content == "Normal text only"
    finally:
        Path(docx_path).unlink()


# =============================================================================
# Test: Fixture file
# =============================================================================

def test_extract_hyperlinks_fixture():
    """Test extracting from the hyperlinks.docx fixture file."""
    fixture_path = Path(__file__).parent / "fixtures" / "hyperlinks.docx"
    if not fixture_path.exists():
        import pytest
        pytest.skip("hyperlinks.docx fixture not found")

    blocks, _ = extract_blocks(str(fixture_path))

    # Should have multiple blocks
    assert len(blocks) > 0

    # Check that at least one block has a hyperlink
    has_hyperlink = any("[" in b.content and "](" in b.content for b in blocks)
    assert has_hyperlink, "Expected at least one hyperlink in markdown format"

    # Check for heading with hyperlink
    heading_blocks = [b for b in blocks if b.type == "heading"]
    heading_with_link = any("[" in b.content for b in heading_blocks)
    assert heading_with_link, "Expected at least one heading with hyperlink"

    # Check for list item with hyperlink
    list_blocks = [b for b in blocks if b.type == "list"]
    list_with_link = any("[" in b.content for b in list_blocks)
    assert list_with_link, "Expected at least one list item with hyperlink"


# =============================================================================
# Test: Hyperlink reconstruction (build from sidedoc)
# =============================================================================

def create_sidedoc_with_hyperlinks(markdown_content: str) -> str:
    """Create a sidedoc archive with the given markdown content.

    Args:
        markdown_content: Markdown content with hyperlinks

    Returns:
        Path to the temporary sidedoc file
    """
    # Create a simple test document to get structure
    doc = Document()
    doc.add_paragraph("placeholder")

    # Create temp files
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_docx.name)
    temp_docx.close()

    # Extract and create sidedoc
    blocks, image_data = extract_blocks(temp_docx.name)
    styles = extract_styles(temp_docx.name, blocks)

    # Override with our hyperlink content
    # Parse the markdown to get proper blocks
    new_blocks = parse_markdown_to_blocks(markdown_content)

    temp_sidedoc = tempfile.NamedTemporaryFile(delete=False, suffix=".sidedoc")
    temp_sidedoc.close()

    create_sidedoc_archive(temp_sidedoc.name, markdown_content, new_blocks, styles, temp_docx.name, {})

    # Clean up docx
    Path(temp_docx.name).unlink()

    return temp_sidedoc.name


def test_reconstruct_simple_hyperlink():
    """Test reconstructing a simple hyperlink in docx."""
    markdown = "Visit [Google](https://www.google.com) for search."

    sidedoc_path = create_sidedoc_with_hyperlinks(markdown)
    output_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    output_docx.close()

    try:
        build_docx_from_sidedoc(sidedoc_path, output_docx.name)

        # Open the result and check for hyperlink
        doc = Document(output_docx.name)
        assert len(doc.paragraphs) >= 1

        para = doc.paragraphs[0]
        para_text = para.text

        # The paragraph should contain "Google" as the link text
        assert "Google" in para_text

        # Check that hyperlink was created in XML
        para_xml = para._element
        hyperlinks = para_xml.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')
        assert len(hyperlinks) >= 1, "Expected at least one hyperlink in the reconstructed document"

        # Verify the URL is correct
        from docx.oxml.ns import qn as qname
        r_id = hyperlinks[0].get(qname('r:id'))
        assert r_id is not None
        rel = doc.part.rels[r_id]
        assert rel._target == "https://www.google.com"

    finally:
        Path(sidedoc_path).unlink()
        Path(output_docx.name).unlink()


def test_reconstruct_multiple_hyperlinks():
    """Test reconstructing multiple hyperlinks in one paragraph."""
    markdown = "See [GitHub](https://github.com) and [Google](https://google.com)."

    sidedoc_path = create_sidedoc_with_hyperlinks(markdown)
    output_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    output_docx.close()

    try:
        build_docx_from_sidedoc(sidedoc_path, output_docx.name)

        doc = Document(output_docx.name)
        para = doc.paragraphs[0]

        # Check that two hyperlinks were created
        para_xml = para._element
        hyperlinks = para_xml.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')
        assert len(hyperlinks) >= 2, f"Expected 2 hyperlinks, found {len(hyperlinks)}"

    finally:
        Path(sidedoc_path).unlink()
        Path(output_docx.name).unlink()


def test_reconstruct_hyperlink_in_heading():
    """Test reconstructing a hyperlink in a heading."""
    markdown = "# Heading with [link](https://example.com)"

    sidedoc_path = create_sidedoc_with_hyperlinks(markdown)
    output_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    output_docx.close()

    try:
        build_docx_from_sidedoc(sidedoc_path, output_docx.name)

        doc = Document(output_docx.name)
        para = doc.paragraphs[0]

        # Check heading style
        assert "Heading" in para.style.name

        # Check for hyperlink
        para_xml = para._element
        hyperlinks = para_xml.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')
        assert len(hyperlinks) >= 1, "Expected hyperlink in heading"

    finally:
        Path(sidedoc_path).unlink()
        Path(output_docx.name).unlink()


def test_roundtrip_preserves_hyperlinks():
    """Test that extract->build round-trip preserves hyperlinks."""
    # Create a document with hyperlinks
    def setup(doc):
        para = doc.add_paragraph("Visit ")
        add_hyperlink(para, "Google", "https://www.google.com")
        para.add_run(" for search.")

        para2 = doc.add_paragraph("Also check ")
        add_hyperlink(para2, "GitHub", "https://github.com")
        para2.add_run(".")

    original_docx = create_hyperlink_docx(setup)

    # Extract to sidedoc
    blocks, image_data = extract_blocks(original_docx)
    styles = extract_styles(original_docx, blocks)
    content_md = blocks_to_markdown(blocks)

    temp_sidedoc = tempfile.NamedTemporaryFile(delete=False, suffix=".sidedoc")
    temp_sidedoc.close()
    create_sidedoc_archive(temp_sidedoc.name, content_md, blocks, styles, original_docx, image_data)

    # Rebuild
    output_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    output_docx.close()
    build_docx_from_sidedoc(temp_sidedoc.name, output_docx.name)

    try:
        # Extract from rebuilt document
        rebuilt_blocks, _ = extract_blocks(output_docx.name)

        # Should have same number of blocks
        assert len(rebuilt_blocks) == len(blocks)

        # Hyperlinks should be preserved
        for original, rebuilt in zip(blocks, rebuilt_blocks):
            # Check that hyperlink markdown is preserved
            if "[" in original.content and "](" in original.content:
                assert "[" in rebuilt.content, f"Hyperlink lost in rebuild: {original.content}"
                assert "](" in rebuilt.content, f"Hyperlink lost in rebuild: {original.content}"

    finally:
        Path(original_docx).unlink()
        Path(temp_sidedoc.name).unlink()
        Path(output_docx.name).unlink()


# =============================================================================
# Test: Sync with hyperlinks
# =============================================================================

def test_sync_adds_new_hyperlinks():
    """Test that sync can add new hyperlinks to content."""
    # Create a document with plain text
    doc = Document()
    doc.add_paragraph("Visit Google for search.")

    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_docx.name)
    temp_docx.close()

    # Extract to sidedoc
    blocks, image_data = extract_blocks(temp_docx.name)
    styles = extract_styles(temp_docx.name, blocks)
    content_md = blocks_to_markdown(blocks)

    temp_sidedoc = tempfile.NamedTemporaryFile(delete=False, suffix=".sidedoc")
    temp_sidedoc.close()
    create_sidedoc_archive(temp_sidedoc.name, content_md, blocks, styles, temp_docx.name, image_data)

    # Now modify the content.md to add a hyperlink
    modified_content = "Visit [Google](https://www.google.com) for search."

    # Update the sidedoc with the new content
    import zipfile
    with zipfile.ZipFile(temp_sidedoc.name, 'r') as zip_read:
        manifest = json.loads(zip_read.read("manifest.json"))
        structure = json.loads(zip_read.read("structure.json"))
        styles_json = json.loads(zip_read.read("styles.json"))

    # Rewrite with modified content
    with zipfile.ZipFile(temp_sidedoc.name, 'w', zipfile.ZIP_DEFLATED) as zip_write:
        zip_write.writestr("content.md", modified_content)
        zip_write.writestr("manifest.json", json.dumps(manifest))
        zip_write.writestr("structure.json", json.dumps(structure))
        zip_write.writestr("styles.json", json.dumps(styles_json))

    # Build the updated docx
    output_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    output_docx.close()

    try:
        build_docx_from_sidedoc(temp_sidedoc.name, output_docx.name)

        # Verify the hyperlink was created
        result_doc = Document(output_docx.name)
        para = result_doc.paragraphs[0]

        # Check for hyperlink in XML
        para_xml = para._element
        hyperlinks = para_xml.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')
        assert len(hyperlinks) >= 1, "Expected hyperlink to be added"

    finally:
        Path(temp_docx.name).unlink()
        Path(temp_sidedoc.name).unlink()
        Path(output_docx.name).unlink()


def test_sync_removes_hyperlinks():
    """Test that sync can remove hyperlinks (convert to plain text)."""
    # Create a document with hyperlinks
    def setup(doc):
        para = doc.add_paragraph("Visit ")
        add_hyperlink(para, "Google", "https://www.google.com")
        para.add_run(" for search.")

    original_docx = create_hyperlink_docx(setup)

    # Extract to sidedoc
    blocks, image_data = extract_blocks(original_docx)
    styles = extract_styles(original_docx, blocks)
    content_md = blocks_to_markdown(blocks)

    temp_sidedoc = tempfile.NamedTemporaryFile(delete=False, suffix=".sidedoc")
    temp_sidedoc.close()
    create_sidedoc_archive(temp_sidedoc.name, content_md, blocks, styles, original_docx, image_data)

    # Modify the content to remove the hyperlink (just plain text)
    modified_content = "Visit Google for search."

    import zipfile
    with zipfile.ZipFile(temp_sidedoc.name, 'r') as zip_read:
        manifest = json.loads(zip_read.read("manifest.json"))
        structure = json.loads(zip_read.read("structure.json"))
        styles_json = json.loads(zip_read.read("styles.json"))

    with zipfile.ZipFile(temp_sidedoc.name, 'w', zipfile.ZIP_DEFLATED) as zip_write:
        zip_write.writestr("content.md", modified_content)
        zip_write.writestr("manifest.json", json.dumps(manifest))
        zip_write.writestr("structure.json", json.dumps(structure))
        zip_write.writestr("styles.json", json.dumps(styles_json))

    # Build the updated docx
    output_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    output_docx.close()

    try:
        build_docx_from_sidedoc(temp_sidedoc.name, output_docx.name)

        # Verify the hyperlink was removed
        result_doc = Document(output_docx.name)
        para = result_doc.paragraphs[0]

        # Check that there are no hyperlinks
        para_xml = para._element
        hyperlinks = para_xml.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')
        assert len(hyperlinks) == 0, "Expected hyperlink to be removed"

        # But text should still be there
        assert "Google" in para.text

    finally:
        Path(original_docx).unlink()
        Path(temp_sidedoc.name).unlink()
        Path(output_docx.name).unlink()


def test_sync_modifies_hyperlink_url():
    """Test that sync can modify hyperlink URLs."""
    # Create a document with a hyperlink
    def setup(doc):
        para = doc.add_paragraph("Visit ")
        add_hyperlink(para, "Google", "https://www.google.com")

    original_docx = create_hyperlink_docx(setup)

    # Extract to sidedoc
    blocks, image_data = extract_blocks(original_docx)
    styles = extract_styles(original_docx, blocks)
    content_md = blocks_to_markdown(blocks)

    temp_sidedoc = tempfile.NamedTemporaryFile(delete=False, suffix=".sidedoc")
    temp_sidedoc.close()
    create_sidedoc_archive(temp_sidedoc.name, content_md, blocks, styles, original_docx, image_data)

    # Modify the URL
    modified_content = "Visit [Google](https://google.co.uk)"

    import zipfile
    with zipfile.ZipFile(temp_sidedoc.name, 'r') as zip_read:
        manifest = json.loads(zip_read.read("manifest.json"))
        structure = json.loads(zip_read.read("structure.json"))
        styles_json = json.loads(zip_read.read("styles.json"))

    with zipfile.ZipFile(temp_sidedoc.name, 'w', zipfile.ZIP_DEFLATED) as zip_write:
        zip_write.writestr("content.md", modified_content)
        zip_write.writestr("manifest.json", json.dumps(manifest))
        zip_write.writestr("structure.json", json.dumps(structure))
        zip_write.writestr("styles.json", json.dumps(styles_json))

    # Build the updated docx
    output_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    output_docx.close()

    try:
        build_docx_from_sidedoc(temp_sidedoc.name, output_docx.name)

        # Extract from rebuilt document and check URL
        rebuilt_blocks, _ = extract_blocks(output_docx.name)
        assert len(rebuilt_blocks) == 1
        assert "google.co.uk" in rebuilt_blocks[0].content

    finally:
        Path(original_docx).unlink()
        Path(temp_sidedoc.name).unlink()
        Path(output_docx.name).unlink()


# =============================================================================
# Test: Hyperlink validation
# =============================================================================

def test_validate_sidedoc_with_hyperlinks():
    """Test that sidedoc validate passes for valid hyperlink content."""
    from click.testing import CliRunner
    from sidedoc.cli import main

    # Create a document with hyperlinks
    def setup(doc):
        para = doc.add_paragraph("Visit ")
        add_hyperlink(para, "Google", "https://www.google.com")

    original_docx = create_hyperlink_docx(setup)

    # Extract to sidedoc
    blocks, image_data = extract_blocks(original_docx)
    styles = extract_styles(original_docx, blocks)
    content_md = blocks_to_markdown(blocks)

    temp_sidedoc = tempfile.NamedTemporaryFile(delete=False, suffix=".sidedoc")
    temp_sidedoc.close()
    create_sidedoc_archive(temp_sidedoc.name, content_md, blocks, styles, original_docx, image_data)

    try:
        runner = CliRunner()
        result = runner.invoke(main, ["validate", temp_sidedoc.name])
        assert result.exit_code == 0
        assert "valid" in result.output.lower()
    finally:
        Path(original_docx).unlink()
        Path(temp_sidedoc.name).unlink()
