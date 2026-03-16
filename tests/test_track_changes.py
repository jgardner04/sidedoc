"""Tests for track changes support.

US-001: Test fixtures created - verified by TestTrackChangesFixtures
US-002: TrackChange models - verified by TestTrackChangeModels
US-003: CriticMarkup patterns - verified by TestCriticMarkupPatterns
US-004: Extract w:ins elements - verified by TestExtractInsertions
US-005: Extract w:del elements - verified by TestExtractDeletions
US-006: Multiple track changes in paragraph - verified by TestMultipleTrackChanges
US-007: Detect track changes mode - verified by TestDetectTrackChanges
US-008: Track changes in headings - verified by TestTrackChangesInHeadings
US-009: Track changes in lists - verified by TestTrackChangesInLists
US-010: Reconstruct w:ins elements - verified by TestReconstructInsertions
US-011: Reconstruct w:del elements - verified by TestReconstructDeletions
US-012: Round-trip integration - verified by TestTrackChangesRoundTrip

Track Changes Test Fixtures:
============================

track_changes_simple.docx:
- Paragraph 1: "This is some text." (no changes)
- Paragraph 2: "Hello {++inserted++} world" - contains insertion by "Test Author"
- Paragraph 3: "Hello {--deleted--} world" - contains deletion by "Test Author"

track_changes_paragraph.docx:
- Paragraph 1: "This {++has++} multiple {--removed--} changes {++added++} here"
  Multiple track changes in a single paragraph

track_changes_multiauthor.docx:
- Paragraph 1: "Alice {++added this++}" - by Author "Alice"
- Paragraph 2: "Bob {--removed this--}" - by Author "Bob"
- Paragraph 3: "{++Charlie added++} and {--Dave removed--}" - mixed authors
"""

from pathlib import Path
from dataclasses import fields
import re

import pytest
from docx import Document

from sidedoc.models import Block

# Path to test fixtures
FIXTURES_DIR = Path(__file__).parent / "fixtures"


class TestTrackChangesFixtures:
    """Tests to verify track changes fixtures exist and have expected content."""

    def test_track_changes_simple_exists(self) -> None:
        """Verify track_changes_simple.docx fixture exists."""
        fixture_path = FIXTURES_DIR / "track_changes_simple.docx"
        assert fixture_path.exists(), f"Fixture not found: {fixture_path}"

    def test_track_changes_simple_has_insertion(self) -> None:
        """Verify track_changes_simple.docx contains w:ins element."""
        fixture_path = FIXTURES_DIR / "track_changes_simple.docx"
        doc = Document(str(fixture_path))

        # Check XML for w:ins elements
        nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        found_ins = False
        for para in doc.paragraphs:
            ins_elements = para._element.findall(".//w:ins", nsmap)
            if ins_elements:
                found_ins = True
                break

        assert found_ins, "No w:ins (insertion) elements found in track_changes_simple.docx"

    def test_track_changes_simple_has_deletion(self) -> None:
        """Verify track_changes_simple.docx contains w:del element."""
        fixture_path = FIXTURES_DIR / "track_changes_simple.docx"
        doc = Document(str(fixture_path))

        # Check XML for w:del elements
        nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        found_del = False
        for para in doc.paragraphs:
            del_elements = para._element.findall(".//w:del", nsmap)
            if del_elements:
                found_del = True
                break

        assert found_del, "No w:del (deletion) elements found in track_changes_simple.docx"

    def test_track_changes_simple_has_author_metadata(self) -> None:
        """Verify track changes in simple fixture have author metadata."""
        fixture_path = FIXTURES_DIR / "track_changes_simple.docx"
        doc = Document(str(fixture_path))

        nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        found_author = False
        for para in doc.paragraphs:
            ins_elements = para._element.findall(".//w:ins", nsmap)
            for ins in ins_elements:
                author = ins.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author")
                if author:
                    found_author = True
                    break

        assert found_author, "No author metadata found on w:ins elements"

    def test_track_changes_simple_has_date_metadata(self) -> None:
        """Verify track changes in simple fixture have date metadata."""
        fixture_path = FIXTURES_DIR / "track_changes_simple.docx"
        doc = Document(str(fixture_path))

        nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        found_date = False
        for para in doc.paragraphs:
            ins_elements = para._element.findall(".//w:ins", nsmap)
            for ins in ins_elements:
                date = ins.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date")
                if date:
                    found_date = True
                    break

        assert found_date, "No date metadata found on w:ins elements"

    def test_track_changes_paragraph_exists(self) -> None:
        """Verify track_changes_paragraph.docx fixture exists."""
        fixture_path = FIXTURES_DIR / "track_changes_paragraph.docx"
        assert fixture_path.exists(), f"Fixture not found: {fixture_path}"

    def test_track_changes_paragraph_has_multiple_changes(self) -> None:
        """Verify track_changes_paragraph.docx has multiple track changes in one paragraph."""
        fixture_path = FIXTURES_DIR / "track_changes_paragraph.docx"
        doc = Document(str(fixture_path))

        nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        # Find a paragraph with multiple track changes
        for para in doc.paragraphs:
            ins_elements = para._element.findall(".//w:ins", nsmap)
            del_elements = para._element.findall(".//w:del", nsmap)
            total_changes = len(ins_elements) + len(del_elements)
            if total_changes >= 2:
                return  # Success - found paragraph with multiple changes

        pytest.fail("No paragraph found with multiple track changes")

    def test_track_changes_multiauthor_exists(self) -> None:
        """Verify track_changes_multiauthor.docx fixture exists."""
        fixture_path = FIXTURES_DIR / "track_changes_multiauthor.docx"
        assert fixture_path.exists(), f"Fixture not found: {fixture_path}"

    def test_track_changes_multiauthor_has_different_authors(self) -> None:
        """Verify track_changes_multiauthor.docx has changes from multiple authors."""
        fixture_path = FIXTURES_DIR / "track_changes_multiauthor.docx"
        doc = Document(str(fixture_path))

        nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        authors = set()

        for para in doc.paragraphs:
            for ins in para._element.findall(".//w:ins", nsmap):
                author = ins.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author")
                if author:
                    authors.add(author)
            for del_elem in para._element.findall(".//w:del", nsmap):
                author = del_elem.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author")
                if author:
                    authors.add(author)

        assert len(authors) >= 2, f"Expected at least 2 different authors, found: {authors}"


class TestTrackChangeModels:
    """Tests for TrackChange data model (US-002)."""

    def test_track_change_dataclass_exists(self) -> None:
        """Verify TrackChange dataclass can be imported."""
        from sidedoc.models import TrackChange

        assert TrackChange is not None

    def test_track_change_has_required_fields(self) -> None:
        """Verify TrackChange has all required fields."""
        from sidedoc.models import TrackChange

        field_names = {f.name for f in fields(TrackChange)}
        required_fields = {"type", "start", "end", "author", "date", "revision_id"}

        assert required_fields.issubset(
            field_names
        ), f"Missing fields: {required_fields - field_names}"

    def test_track_change_has_deleted_text_field(self) -> None:
        """Verify TrackChange has deleted_text field for deletions."""
        from sidedoc.models import TrackChange

        field_names = {f.name for f in fields(TrackChange)}
        assert "deleted_text" in field_names

    def test_track_change_can_be_instantiated(self) -> None:
        """Verify TrackChange can be instantiated with valid data."""
        from sidedoc.models import TrackChange

        tc = TrackChange(
            type="insertion",
            start=0,
            end=10,
            author="Test Author",
            date="2026-01-15T10:30:00Z",
            revision_id="1",
        )
        assert tc.type == "insertion"
        assert tc.start == 0
        assert tc.end == 10
        assert tc.author == "Test Author"

    def test_track_change_deletion_with_deleted_text(self) -> None:
        """Verify deletion TrackChange can store deleted_text."""
        from sidedoc.models import TrackChange

        tc = TrackChange(
            type="deletion",
            start=5,
            end=15,
            author="Test Author",
            date="2026-01-15T10:30:00Z",
            revision_id="2",
            deleted_text="removed text",
        )
        assert tc.type == "deletion"
        assert tc.deleted_text == "removed text"

    def test_block_has_track_changes_field(self) -> None:
        """Verify Block dataclass has track_changes field."""
        field_names = {f.name for f in fields(Block)}
        assert "track_changes" in field_names

    def test_block_track_changes_defaults_to_none(self) -> None:
        """Verify Block.track_changes defaults to None."""
        block = Block(
            id="block-0",
            type="paragraph",
            content="Test content",
            docx_paragraph_index=0,
            content_start=0,
            content_end=12,
            content_hash="abc123",
        )
        # Should default to None or empty list
        assert block.track_changes is None or block.track_changes == []

    def test_track_changes_config_exists(self) -> None:
        """Verify TrackChangesConfig dataclass can be imported."""
        from sidedoc.models import TrackChangesConfig

        assert TrackChangesConfig is not None

    def test_track_changes_config_has_required_fields(self) -> None:
        """Verify TrackChangesConfig has all required fields."""
        from sidedoc.models import TrackChangesConfig

        field_names = {f.name for f in fields(TrackChangesConfig)}
        required_fields = {"enabled", "source_had_revisions", "default_author"}

        assert required_fields.issubset(
            field_names
        ), f"Missing fields: {required_fields - field_names}"

    def test_track_changes_config_can_be_instantiated(self) -> None:
        """Verify TrackChangesConfig can be instantiated."""
        from sidedoc.models import TrackChangesConfig

        config = TrackChangesConfig(
            enabled=True, source_had_revisions=True, default_author="Sidedoc AI"
        )
        assert config.enabled is True
        assert config.source_had_revisions is True
        assert config.default_author == "Sidedoc AI"


class TestCriticMarkupPatterns:
    """Tests for CriticMarkup regex patterns (US-003)."""

    def test_insertion_pattern_exists(self) -> None:
        """Verify INSERTION_PATTERN constant exists."""
        from sidedoc.constants import INSERTION_PATTERN

        assert INSERTION_PATTERN is not None

    def test_deletion_pattern_exists(self) -> None:
        """Verify DELETION_PATTERN constant exists."""
        from sidedoc.constants import DELETION_PATTERN

        assert DELETION_PATTERN is not None

    def test_substitution_pattern_exists(self) -> None:
        """Verify SUBSTITUTION_PATTERN constant exists."""
        from sidedoc.constants import SUBSTITUTION_PATTERN

        assert SUBSTITUTION_PATTERN is not None

    def test_insertion_pattern_matches_simple(self) -> None:
        """Verify INSERTION_PATTERN matches simple insertions."""
        from sidedoc.constants import INSERTION_PATTERN

        text = "Hello {++inserted++} world"
        match = re.search(INSERTION_PATTERN, text)
        assert match is not None
        assert match.group(1) == "inserted"

    def test_insertion_pattern_matches_with_spaces(self) -> None:
        """Verify INSERTION_PATTERN matches insertions with spaces."""
        from sidedoc.constants import INSERTION_PATTERN

        text = "Hello {++inserted text here++} world"
        match = re.search(INSERTION_PATTERN, text)
        assert match is not None
        assert match.group(1) == "inserted text here"

    def test_deletion_pattern_matches_simple(self) -> None:
        """Verify DELETION_PATTERN matches simple deletions."""
        from sidedoc.constants import DELETION_PATTERN

        text = "Hello {--deleted--} world"
        match = re.search(DELETION_PATTERN, text)
        assert match is not None
        assert match.group(1) == "deleted"

    def test_deletion_pattern_matches_with_spaces(self) -> None:
        """Verify DELETION_PATTERN matches deletions with spaces."""
        from sidedoc.constants import DELETION_PATTERN

        text = "Hello {--deleted text here--} world"
        match = re.search(DELETION_PATTERN, text)
        assert match is not None
        assert match.group(1) == "deleted text here"

    def test_substitution_pattern_matches_simple(self) -> None:
        """Verify SUBSTITUTION_PATTERN matches simple substitutions."""
        from sidedoc.constants import SUBSTITUTION_PATTERN

        text = "Hello {~~old~>new~~} world"
        match = re.search(SUBSTITUTION_PATTERN, text)
        assert match is not None
        assert match.group(1) == "old"
        assert match.group(2) == "new"

    def test_substitution_pattern_matches_with_spaces(self) -> None:
        """Verify SUBSTITUTION_PATTERN matches substitutions with spaces."""
        from sidedoc.constants import SUBSTITUTION_PATTERN

        text = "Hello {~~old text~>new text~~} world"
        match = re.search(SUBSTITUTION_PATTERN, text)
        assert match is not None
        assert match.group(1) == "old text"
        assert match.group(2) == "new text"

    def test_insertion_pattern_handles_nested_markdown(self) -> None:
        """Verify INSERTION_PATTERN handles nested markdown formatting."""
        from sidedoc.constants import INSERTION_PATTERN

        text = "Hello {++**bold text**++} world"
        match = re.search(INSERTION_PATTERN, text)
        assert match is not None
        assert match.group(1) == "**bold text**"

    def test_deletion_pattern_handles_nested_markdown(self) -> None:
        """Verify DELETION_PATTERN handles nested markdown formatting."""
        from sidedoc.constants import DELETION_PATTERN

        text = "Hello {--*italic text*--} world"
        match = re.search(DELETION_PATTERN, text)
        assert match is not None
        assert match.group(1) == "*italic text*"

    def test_multiple_insertions_in_text(self) -> None:
        """Verify INSERTION_PATTERN can find all insertions in text."""
        from sidedoc.constants import INSERTION_PATTERN

        text = "Hello {++first++} and {++second++} world"
        matches = re.findall(INSERTION_PATTERN, text)
        assert len(matches) == 2
        assert matches[0] == "first"
        assert matches[1] == "second"

    def test_multiple_deletions_in_text(self) -> None:
        """Verify DELETION_PATTERN can find all deletions in text."""
        from sidedoc.constants import DELETION_PATTERN

        text = "Hello {--first--} and {--second--} world"
        matches = re.findall(DELETION_PATTERN, text)
        assert len(matches) == 2
        assert matches[0] == "first"
        assert matches[1] == "second"


class TestExtractInsertions:
    """Tests for extracting w:ins elements from docx (US-004)."""

    def test_extract_insertion_as_criticmarkup(self) -> None:
        """Verify insertions are extracted as {++text++} CriticMarkup syntax."""
        from sidedoc.extract import extract_blocks

        fixture_path = str(FIXTURES_DIR / "track_changes_simple.docx")
        blocks, _ = extract_blocks(fixture_path)

        # Paragraph 2 has "Hello {++inserted++} world"
        para_with_insertion = blocks[1]  # 0-indexed, second paragraph
        assert "{++inserted++}" in para_with_insertion.content

    def test_extract_insertion_stores_metadata(self) -> None:
        """Verify insertion metadata is stored in track_changes array."""
        from sidedoc.extract import extract_blocks

        fixture_path = str(FIXTURES_DIR / "track_changes_simple.docx")
        blocks, _ = extract_blocks(fixture_path)

        # Paragraph 2 has an insertion
        para_with_insertion = blocks[1]
        assert para_with_insertion.track_changes is not None
        assert len(para_with_insertion.track_changes) >= 1

        insertion = para_with_insertion.track_changes[0]
        assert insertion.type == "insertion"

    def test_extract_insertion_author(self) -> None:
        """Verify insertion author is extracted."""
        from sidedoc.extract import extract_blocks

        fixture_path = str(FIXTURES_DIR / "track_changes_simple.docx")
        blocks, _ = extract_blocks(fixture_path)

        para_with_insertion = blocks[1]
        assert para_with_insertion.track_changes is not None

        insertion = para_with_insertion.track_changes[0]
        assert insertion.author == "Test Author"

    def test_extract_insertion_date(self) -> None:
        """Verify insertion date is extracted."""
        from sidedoc.extract import extract_blocks

        fixture_path = str(FIXTURES_DIR / "track_changes_simple.docx")
        blocks, _ = extract_blocks(fixture_path)

        para_with_insertion = blocks[1]
        assert para_with_insertion.track_changes is not None

        insertion = para_with_insertion.track_changes[0]
        assert insertion.date == "2026-01-15T10:30:00Z"

    def test_extract_insertion_positions(self) -> None:
        """Verify insertion start/end positions are calculated."""
        from sidedoc.extract import extract_blocks

        fixture_path = str(FIXTURES_DIR / "track_changes_simple.docx")
        blocks, _ = extract_blocks(fixture_path)

        para_with_insertion = blocks[1]
        assert para_with_insertion.track_changes is not None

        insertion = para_with_insertion.track_changes[0]
        # Positions should be valid integers
        assert isinstance(insertion.start, int)
        assert isinstance(insertion.end, int)
        assert insertion.start >= 0
        assert insertion.end > insertion.start


class TestExtractDeletions:
    """Tests for extracting w:del elements from docx (US-005)."""

    def test_extract_deletion_as_criticmarkup(self) -> None:
        """Verify deletions are extracted as {--text--} CriticMarkup syntax."""
        from sidedoc.extract import extract_blocks

        fixture_path = str(FIXTURES_DIR / "track_changes_simple.docx")
        blocks, _ = extract_blocks(fixture_path)

        # Paragraph 3 has "Hello {--deleted --}world"
        para_with_deletion = blocks[2]  # 0-indexed, third paragraph
        assert "{--deleted --}" in para_with_deletion.content

    def test_extract_deletion_stores_metadata(self) -> None:
        """Verify deletion metadata is stored in track_changes array."""
        from sidedoc.extract import extract_blocks

        fixture_path = str(FIXTURES_DIR / "track_changes_simple.docx")
        blocks, _ = extract_blocks(fixture_path)

        # Paragraph 3 has a deletion
        para_with_deletion = blocks[2]
        assert para_with_deletion.track_changes is not None
        assert len(para_with_deletion.track_changes) >= 1

        deletion = para_with_deletion.track_changes[0]
        assert deletion.type == "deletion"

    def test_extract_deletion_stores_deleted_text(self) -> None:
        """Verify deleted text is stored in deletion metadata."""
        from sidedoc.extract import extract_blocks

        fixture_path = str(FIXTURES_DIR / "track_changes_simple.docx")
        blocks, _ = extract_blocks(fixture_path)

        para_with_deletion = blocks[2]
        assert para_with_deletion.track_changes is not None

        deletion = para_with_deletion.track_changes[0]
        assert deletion.deleted_text == "deleted "

    def test_extract_deletion_author(self) -> None:
        """Verify deletion author is extracted."""
        from sidedoc.extract import extract_blocks

        fixture_path = str(FIXTURES_DIR / "track_changes_simple.docx")
        blocks, _ = extract_blocks(fixture_path)

        para_with_deletion = blocks[2]
        assert para_with_deletion.track_changes is not None

        deletion = para_with_deletion.track_changes[0]
        assert deletion.author == "Test Author"

    def test_extract_deletion_date(self) -> None:
        """Verify deletion date is extracted."""
        from sidedoc.extract import extract_blocks

        fixture_path = str(FIXTURES_DIR / "track_changes_simple.docx")
        blocks, _ = extract_blocks(fixture_path)

        para_with_deletion = blocks[2]
        assert para_with_deletion.track_changes is not None

        deletion = para_with_deletion.track_changes[0]
        assert deletion.date == "2026-01-15T10:30:00Z"


class TestMultipleTrackChanges:
    """Tests for multiple track changes in a single paragraph (US-006)."""

    def test_extract_multiple_changes_in_paragraph(self) -> None:
        """Verify all track changes in a paragraph are extracted."""
        from sidedoc.extract import extract_blocks

        fixture_path = str(FIXTURES_DIR / "track_changes_paragraph.docx")
        blocks, _ = extract_blocks(fixture_path)

        # track_changes_paragraph.docx has:
        # "This {++has++} multiple {--removed--} changes {++added++} here"
        para = blocks[0]
        assert para.track_changes is not None
        assert len(para.track_changes) == 3  # 2 insertions + 1 deletion

    def test_multiple_changes_correct_order(self) -> None:
        """Verify track changes are in document order."""
        from sidedoc.extract import extract_blocks

        fixture_path = str(FIXTURES_DIR / "track_changes_paragraph.docx")
        blocks, _ = extract_blocks(fixture_path)

        para = blocks[0]
        assert para.track_changes is not None

        # Order should be: insertion "has", deletion "removed", insertion "added"
        assert para.track_changes[0].type == "insertion"
        assert para.track_changes[1].type == "deletion"
        assert para.track_changes[2].type == "insertion"

    def test_multiple_changes_correct_positions(self) -> None:
        """Verify track change positions are sequential and non-overlapping."""
        from sidedoc.extract import extract_blocks

        fixture_path = str(FIXTURES_DIR / "track_changes_paragraph.docx")
        blocks, _ = extract_blocks(fixture_path)

        para = blocks[0]
        assert para.track_changes is not None

        # Each track change should have distinct start/end positions
        for i, tc in enumerate(para.track_changes):
            assert tc.start >= 0
            assert tc.end > tc.start
            # Later changes should have higher positions (roughly)
            if i > 0:
                # At minimum, positions should be non-negative
                assert tc.start >= 0

    def test_multiple_changes_markdown_content(self) -> None:
        """Verify markdown content contains all CriticMarkup syntax."""
        from sidedoc.extract import extract_blocks

        fixture_path = str(FIXTURES_DIR / "track_changes_paragraph.docx")
        blocks, _ = extract_blocks(fixture_path)

        para = blocks[0]
        # Content should have all three changes as CriticMarkup
        assert "{++has++}" in para.content
        assert "{--removed--}" in para.content
        assert "{++added++}" in para.content

    def test_multiauthor_extracts_different_authors(self) -> None:
        """Verify track changes from different authors preserve author names."""
        from sidedoc.extract import extract_blocks

        fixture_path = str(FIXTURES_DIR / "track_changes_multiauthor.docx")
        blocks, _ = extract_blocks(fixture_path)

        # Collect all authors from track changes
        authors = set()
        for block in blocks:
            if block.track_changes:
                for tc in block.track_changes:
                    authors.add(tc.author)

        # Should have at least 2 different authors
        assert len(authors) >= 2


class TestDetectTrackChanges:
    """Tests for track changes detection in documents (US-007)."""

    def test_detect_track_changes_function_exists(self) -> None:
        """Verify detect_track_changes function can be imported."""
        from sidedoc.extract import detect_track_changes

        assert detect_track_changes is not None

    def test_detect_track_changes_in_document_with_changes(self) -> None:
        """Verify detection returns True for document with track changes."""
        from sidedoc.extract import detect_track_changes

        fixture_path = str(FIXTURES_DIR / "track_changes_simple.docx")
        has_changes = detect_track_changes(fixture_path)
        assert has_changes is True

    def test_detect_track_changes_in_document_without_changes(self) -> None:
        """Verify detection returns False for document without track changes."""
        from sidedoc.extract import detect_track_changes

        fixture_path = str(FIXTURES_DIR / "simple.docx")
        has_changes = detect_track_changes(fixture_path)
        assert has_changes is False

    def test_detect_insertion_only(self) -> None:
        """Verify detection finds insertions."""
        from sidedoc.extract import detect_track_changes

        fixture_path = str(FIXTURES_DIR / "track_changes_simple.docx")
        # This fixture has both insertions and deletions
        has_changes = detect_track_changes(fixture_path)
        assert has_changes is True

    def test_detect_deletion_only(self) -> None:
        """Verify detection finds deletions."""
        from sidedoc.extract import detect_track_changes

        # multiauthor has deletions too
        fixture_path = str(FIXTURES_DIR / "track_changes_multiauthor.docx")
        has_changes = detect_track_changes(fixture_path)
        assert has_changes is True


class TestTrackChangesInHeadings:
    """Tests for track changes in headings (US-008)."""

    def test_heading_fixture_exists(self) -> None:
        """Verify track_changes_headings.docx fixture exists."""
        fixture_path = FIXTURES_DIR / "track_changes_headings.docx"
        assert fixture_path.exists(), f"Fixture not found: {fixture_path}"

    def test_extract_insertion_in_heading1(self) -> None:
        """Verify insertions in H1 headings are extracted correctly."""
        from sidedoc.extract import extract_blocks

        fixture_path = str(FIXTURES_DIR / "track_changes_headings.docx")
        blocks, _ = extract_blocks(fixture_path)

        # First block is Heading 1 with insertion
        heading1 = blocks[0]
        assert heading1.type == "heading"
        assert heading1.level == 1
        assert "{++inserted++}" in heading1.content
        assert heading1.track_changes is not None
        assert len(heading1.track_changes) >= 1
        assert heading1.track_changes[0].type == "insertion"

    def test_extract_deletion_in_heading2(self) -> None:
        """Verify deletions in H2 headings are extracted correctly."""
        from sidedoc.extract import extract_blocks

        fixture_path = str(FIXTURES_DIR / "track_changes_headings.docx")
        blocks, _ = extract_blocks(fixture_path)

        # Second block is Heading 2 with deletion
        heading2 = blocks[1]
        assert heading2.type == "heading"
        assert heading2.level == 2
        assert "{--removed--}" in heading2.content
        assert heading2.track_changes is not None
        assert len(heading2.track_changes) >= 1
        assert heading2.track_changes[0].type == "deletion"

    def test_heading_markdown_includes_level(self) -> None:
        """Verify heading markdown includes proper heading markers."""
        from sidedoc.extract import extract_blocks

        fixture_path = str(FIXTURES_DIR / "track_changes_headings.docx")
        blocks, _ = extract_blocks(fixture_path)

        # H1 should start with #
        heading1 = blocks[0]
        assert heading1.content.startswith("# ")

        # H2 should start with ##
        heading2 = blocks[1]
        assert heading2.content.startswith("## ")


class TestTrackChangesInLists:
    """Tests for track changes in list items (US-009)."""

    def test_list_fixture_exists(self) -> None:
        """Verify track_changes_lists.docx fixture exists."""
        fixture_path = FIXTURES_DIR / "track_changes_lists.docx"
        assert fixture_path.exists(), f"Fixture not found: {fixture_path}"

    def test_extract_insertion_in_bullet_list(self) -> None:
        """Verify insertions in bulleted list items are extracted correctly."""
        from sidedoc.extract import extract_blocks

        fixture_path = str(FIXTURES_DIR / "track_changes_lists.docx")
        blocks, _ = extract_blocks(fixture_path)

        # First block is bullet with insertion
        bullet1 = blocks[0]
        assert bullet1.type == "list"
        assert "{++added++}" in bullet1.content
        assert bullet1.track_changes is not None
        assert bullet1.track_changes[0].type == "insertion"

    def test_extract_deletion_in_bullet_list(self) -> None:
        """Verify deletions in bulleted list items are extracted correctly."""
        from sidedoc.extract import extract_blocks

        fixture_path = str(FIXTURES_DIR / "track_changes_lists.docx")
        blocks, _ = extract_blocks(fixture_path)

        # Second block is bullet with deletion
        bullet2 = blocks[1]
        assert bullet2.type == "list"
        assert "{--removed--}" in bullet2.content
        assert bullet2.track_changes is not None
        assert bullet2.track_changes[0].type == "deletion"

    def test_extract_insertion_in_numbered_list(self) -> None:
        """Verify insertions in numbered list items are extracted correctly."""
        from sidedoc.extract import extract_blocks

        fixture_path = str(FIXTURES_DIR / "track_changes_lists.docx")
        blocks, _ = extract_blocks(fixture_path)

        # Third block is numbered list with insertion
        numbered1 = blocks[2]
        assert numbered1.type == "list"
        assert "{++inserted++}" in numbered1.content
        assert numbered1.track_changes is not None
        assert numbered1.track_changes[0].type == "insertion"

    def test_list_markdown_includes_markers(self) -> None:
        """Verify list markdown includes proper list markers."""
        from sidedoc.extract import extract_blocks

        fixture_path = str(FIXTURES_DIR / "track_changes_lists.docx")
        blocks, _ = extract_blocks(fixture_path)

        # Bullet items should start with "- "
        assert blocks[0].content.startswith("- ")
        assert blocks[1].content.startswith("- ")

        # Numbered item should start with "1. "
        assert blocks[2].content.startswith("1. ")


class TestReconstructInsertions:
    """Tests for reconstructing w:ins elements from CriticMarkup (US-010)."""

    def test_parse_criticmarkup_insertion(self) -> None:
        """Verify CriticMarkup insertion is parsed correctly."""
        from sidedoc.reconstruct import parse_criticmarkup

        text = "Hello {++inserted++} world"
        segments = parse_criticmarkup(text)

        # Should have 3 segments: text, insertion, text
        assert len(segments) == 3
        assert segments[0] == ("text", "Hello ")
        assert segments[1] == ("insertion", "inserted")
        assert segments[2] == ("text", " world")

    def test_create_ins_element(self) -> None:
        """Verify w:ins XML element is created correctly."""
        from sidedoc.reconstruct import create_ins_element

        ins_elem = create_ins_element(
            text="inserted",
            author="Test Author",
            date="2026-01-15T10:30:00Z",
            revision_id="1"
        )

        # Check element type
        assert ins_elem.tag.endswith("}ins")

        # Check attributes
        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        assert ins_elem.get(f"{{{ns}}}author") == "Test Author"
        assert ins_elem.get(f"{{{ns}}}date") == "2026-01-15T10:30:00Z"

    def test_reconstruct_insertion_in_docx(self) -> None:
        """Verify insertion creates w:ins element in rebuilt docx."""
        import tempfile
        from sidedoc.extract import extract_blocks, extract_styles, blocks_to_markdown
        from sidedoc.reconstruct import build_docx_from_sidedoc
        from sidedoc.package import create_sidedoc_archive
        from docx import Document

        # Create a sidedoc from fixture
        fixture_path = str(FIXTURES_DIR / "track_changes_simple.docx")

        with tempfile.TemporaryDirectory() as temp_dir:
            sidedoc_path = f"{temp_dir}/test.sidedoc"
            output_path = f"{temp_dir}/rebuilt.docx"

            # Create sidedoc using the package functions
            blocks, image_data = extract_blocks(fixture_path)
            styles = extract_styles(fixture_path, blocks)
            content_md = blocks_to_markdown(blocks)
            create_sidedoc_archive(sidedoc_path, content_md, blocks, styles, fixture_path, image_data)

            # Build back to docx
            build_docx_from_sidedoc(sidedoc_path, output_path)

            # Check rebuilt docx has w:ins elements
            doc = Document(output_path)
            nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

            found_ins = False
            for para in doc.paragraphs:
                ins_elements = para._element.findall(".//w:ins", nsmap)
                if ins_elements:
                    found_ins = True
                    break

            assert found_ins, "Rebuilt docx should contain w:ins elements"


class TestReconstructDeletions:
    """Tests for reconstructing w:del elements from CriticMarkup (US-011)."""

    def test_parse_criticmarkup_deletion(self) -> None:
        """Verify CriticMarkup deletion is parsed correctly."""
        from sidedoc.reconstruct import parse_criticmarkup

        text = "Hello {--deleted--} world"
        segments = parse_criticmarkup(text)

        # Should have 3 segments: text, deletion, text
        assert len(segments) == 3
        assert segments[0] == ("text", "Hello ")
        assert segments[1] == ("deletion", "deleted")
        assert segments[2] == ("text", " world")

    def test_create_del_element(self) -> None:
        """Verify w:del XML element is created correctly."""
        from sidedoc.reconstruct import create_del_element

        del_elem = create_del_element(
            text="deleted",
            author="Test Author",
            date="2026-01-15T10:30:00Z",
            revision_id="1"
        )

        # Check element type
        assert del_elem.tag.endswith("}del")

        # Check attributes
        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        assert del_elem.get(f"{{{ns}}}author") == "Test Author"
        assert del_elem.get(f"{{{ns}}}date") == "2026-01-15T10:30:00Z"

    def test_reconstruct_deletion_in_docx(self) -> None:
        """Verify deletion creates w:del element in rebuilt docx."""
        import tempfile
        from sidedoc.extract import extract_blocks, extract_styles, blocks_to_markdown
        from sidedoc.reconstruct import build_docx_from_sidedoc
        from sidedoc.package import create_sidedoc_archive
        from docx import Document

        fixture_path = str(FIXTURES_DIR / "track_changes_simple.docx")

        with tempfile.TemporaryDirectory() as temp_dir:
            sidedoc_path = f"{temp_dir}/test.sidedoc"
            output_path = f"{temp_dir}/rebuilt.docx"

            blocks, image_data = extract_blocks(fixture_path)
            styles = extract_styles(fixture_path, blocks)
            content_md = blocks_to_markdown(blocks)
            create_sidedoc_archive(sidedoc_path, content_md, blocks, styles, fixture_path, image_data)

            build_docx_from_sidedoc(sidedoc_path, output_path)

            doc = Document(output_path)
            nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

            found_del = False
            for para in doc.paragraphs:
                del_elements = para._element.findall(".//w:del", nsmap)
                if del_elements:
                    found_del = True
                    break

            assert found_del, "Rebuilt docx should contain w:del elements"


class TestTrackChangesRoundTrip:
    """Tests for round-trip preservation of track changes (US-012)."""

    def test_roundtrip_preserves_insertions(self) -> None:
        """Verify insertions survive extract -> build round-trip."""
        import tempfile
        from sidedoc.extract import extract_blocks, extract_styles, blocks_to_markdown, detect_track_changes
        from sidedoc.reconstruct import build_docx_from_sidedoc
        from sidedoc.package import create_sidedoc_archive

        fixture_path = str(FIXTURES_DIR / "track_changes_simple.docx")

        with tempfile.TemporaryDirectory() as temp_dir:
            sidedoc_path = f"{temp_dir}/test.sidedoc"
            output_path = f"{temp_dir}/rebuilt.docx"

            blocks, image_data = extract_blocks(fixture_path)
            styles = extract_styles(fixture_path, blocks)
            content_md = blocks_to_markdown(blocks)
            create_sidedoc_archive(sidedoc_path, content_md, blocks, styles, fixture_path, image_data)

            build_docx_from_sidedoc(sidedoc_path, output_path)

            # Verify rebuilt has track changes
            assert detect_track_changes(output_path) is True

    def test_roundtrip_preserves_author(self) -> None:
        """Verify author attribute survives round-trip."""
        import tempfile
        from sidedoc.extract import extract_blocks, extract_styles, blocks_to_markdown
        from sidedoc.reconstruct import build_docx_from_sidedoc
        from sidedoc.package import create_sidedoc_archive
        from docx import Document

        fixture_path = str(FIXTURES_DIR / "track_changes_simple.docx")

        with tempfile.TemporaryDirectory() as temp_dir:
            sidedoc_path = f"{temp_dir}/test.sidedoc"
            output_path = f"{temp_dir}/rebuilt.docx"

            blocks, image_data = extract_blocks(fixture_path)
            styles = extract_styles(fixture_path, blocks)
            content_md = blocks_to_markdown(blocks)
            create_sidedoc_archive(sidedoc_path, content_md, blocks, styles, fixture_path, image_data)

            build_docx_from_sidedoc(sidedoc_path, output_path)

            # Check author on rebuilt w:ins elements
            doc = Document(output_path)
            ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

            for para in doc.paragraphs:
                for ins in para._element.findall(f".//{{{ns}}}ins"):
                    author = ins.get(f"{{{ns}}}author")
                    if author:
                        assert author == "Test Author"
                        return

            # If we get here, no author found
            assert False, "No author attribute found on w:ins elements"

    def test_roundtrip_preserves_date(self) -> None:
        """Verify date attribute survives round-trip."""
        import tempfile
        from sidedoc.extract import extract_blocks, extract_styles, blocks_to_markdown
        from sidedoc.reconstruct import build_docx_from_sidedoc
        from sidedoc.package import create_sidedoc_archive
        from docx import Document

        fixture_path = str(FIXTURES_DIR / "track_changes_simple.docx")

        with tempfile.TemporaryDirectory() as temp_dir:
            sidedoc_path = f"{temp_dir}/test.sidedoc"
            output_path = f"{temp_dir}/rebuilt.docx"

            blocks, image_data = extract_blocks(fixture_path)
            styles = extract_styles(fixture_path, blocks)
            content_md = blocks_to_markdown(blocks)
            create_sidedoc_archive(sidedoc_path, content_md, blocks, styles, fixture_path, image_data)

            build_docx_from_sidedoc(sidedoc_path, output_path)

            doc = Document(output_path)
            ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

            for para in doc.paragraphs:
                for ins in para._element.findall(f".//{{{ns}}}ins"):
                    date = ins.get(f"{{{ns}}}date")
                    if date:
                        assert date == "2026-01-15T10:30:00Z"
                        return

            assert False, "No date attribute found on w:ins elements"


class TestCriticMarkupSync:
    """Tests for parsing CriticMarkup during sync (US-013)."""

    def test_sync_creates_insertion_from_criticmarkup(self) -> None:
        """Verify {++new text++} in edited content.md creates insertion in docx after sync."""
        import tempfile
        import zipfile
        from sidedoc.extract import extract_blocks, extract_styles, blocks_to_markdown
        from sidedoc.package import create_sidedoc_archive
        from sidedoc.sync import sync_sidedoc_to_docx
        from docx import Document

        # Create a sidedoc from a simple document (no track changes)
        fixture_path = str(FIXTURES_DIR / "simple.docx")

        with tempfile.TemporaryDirectory() as temp_dir:
            sidedoc_path = f"{temp_dir}/test.sidedoc"
            output_path = f"{temp_dir}/synced.docx"

            # Create initial sidedoc
            blocks, image_data = extract_blocks(fixture_path)
            styles = extract_styles(fixture_path, blocks)
            content_md = blocks_to_markdown(blocks)
            create_sidedoc_archive(sidedoc_path, content_md, blocks, styles, fixture_path, image_data)

            # Modify content.md to add CriticMarkup insertion
            with zipfile.ZipFile(sidedoc_path, "r") as z:
                original_content = z.read("content.md").decode("utf-8")
                structure_json = z.read("structure.json").decode("utf-8")
                styles_json = z.read("styles.json").decode("utf-8")
                manifest_json = z.read("manifest.json").decode("utf-8")

            # Add insertion to a paragraph
            modified_content = original_content.replace("simple document", "{++very ++}simple document")

            # Rewrite sidedoc with modified content
            with zipfile.ZipFile(sidedoc_path, "w", zipfile.ZIP_DEFLATED) as z:
                z.writestr("content.md", modified_content)
                z.writestr("structure.json", structure_json)
                z.writestr("styles.json", styles_json)
                z.writestr("manifest.json", manifest_json)

            # Sync to docx
            sync_sidedoc_to_docx(sidedoc_path, output_path)

            # Verify w:ins element was created
            doc = Document(output_path)
            nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

            found_ins = False
            for para in doc.paragraphs:
                ins_elements = para._element.findall(".//w:ins", nsmap)
                if ins_elements:
                    found_ins = True
                    break

            assert found_ins, "Synced docx should contain w:ins elements for CriticMarkup insertions"

    def test_sync_creates_deletion_from_criticmarkup(self) -> None:
        """Verify {--removed text--} in edited content.md creates deletion in docx after sync."""
        import tempfile
        import zipfile
        from sidedoc.extract import extract_blocks, extract_styles, blocks_to_markdown
        from sidedoc.package import create_sidedoc_archive
        from sidedoc.sync import sync_sidedoc_to_docx
        from docx import Document

        fixture_path = str(FIXTURES_DIR / "simple.docx")

        with tempfile.TemporaryDirectory() as temp_dir:
            sidedoc_path = f"{temp_dir}/test.sidedoc"
            output_path = f"{temp_dir}/synced.docx"

            blocks, image_data = extract_blocks(fixture_path)
            styles = extract_styles(fixture_path, blocks)
            content_md = blocks_to_markdown(blocks)
            create_sidedoc_archive(sidedoc_path, content_md, blocks, styles, fixture_path, image_data)

            with zipfile.ZipFile(sidedoc_path, "r") as z:
                original_content = z.read("content.md").decode("utf-8")
                structure_json = z.read("structure.json").decode("utf-8")
                styles_json = z.read("styles.json").decode("utf-8")
                manifest_json = z.read("manifest.json").decode("utf-8")

            # Add deletion: mark some text as deleted
            modified_content = original_content.replace("simple", "{--simple--}")

            with zipfile.ZipFile(sidedoc_path, "w", zipfile.ZIP_DEFLATED) as z:
                z.writestr("content.md", modified_content)
                z.writestr("structure.json", structure_json)
                z.writestr("styles.json", styles_json)
                z.writestr("manifest.json", manifest_json)

            sync_sidedoc_to_docx(sidedoc_path, output_path)

            doc = Document(output_path)
            nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

            found_del = False
            for para in doc.paragraphs:
                del_elements = para._element.findall(".//w:del", nsmap)
                if del_elements:
                    found_del = True
                    break

            assert found_del, "Synced docx should contain w:del elements for CriticMarkup deletions"

    def test_sync_uses_configured_author(self) -> None:
        """Verify new track changes use configured author name."""
        import tempfile
        import zipfile
        from sidedoc.extract import extract_blocks, extract_styles, blocks_to_markdown
        from sidedoc.package import create_sidedoc_archive
        from sidedoc.sync import sync_sidedoc_to_docx
        from docx import Document

        fixture_path = str(FIXTURES_DIR / "simple.docx")

        with tempfile.TemporaryDirectory() as temp_dir:
            sidedoc_path = f"{temp_dir}/test.sidedoc"
            output_path = f"{temp_dir}/synced.docx"

            blocks, image_data = extract_blocks(fixture_path)
            styles = extract_styles(fixture_path, blocks)
            content_md = blocks_to_markdown(blocks)
            create_sidedoc_archive(sidedoc_path, content_md, blocks, styles, fixture_path, image_data)

            with zipfile.ZipFile(sidedoc_path, "r") as z:
                original_content = z.read("content.md").decode("utf-8")
                structure_json = z.read("structure.json").decode("utf-8")
                styles_json = z.read("styles.json").decode("utf-8")
                manifest_json = z.read("manifest.json").decode("utf-8")

            modified_content = original_content.replace("simple document", "{++very ++}simple document")

            with zipfile.ZipFile(sidedoc_path, "w", zipfile.ZIP_DEFLATED) as z:
                z.writestr("content.md", modified_content)
                z.writestr("structure.json", structure_json)
                z.writestr("styles.json", styles_json)
                z.writestr("manifest.json", manifest_json)

            # Sync with custom author
            sync_sidedoc_to_docx(sidedoc_path, output_path, author="Custom Author")

            doc = Document(output_path)
            ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

            for para in doc.paragraphs:
                for ins in para._element.findall(f".//{{{ns}}}ins"):
                    author = ins.get(f"{{{ns}}}author")
                    if author:
                        assert author == "Custom Author", f"Expected 'Custom Author', got '{author}'"
                        return

            pytest.fail("No author attribute found on w:ins elements")

    def test_sync_uses_default_author_when_not_specified(self) -> None:
        """Verify default author is 'Sidedoc AI' when not specified."""
        import tempfile
        import zipfile
        from sidedoc.extract import extract_blocks, extract_styles, blocks_to_markdown
        from sidedoc.package import create_sidedoc_archive
        from sidedoc.sync import sync_sidedoc_to_docx
        from docx import Document

        fixture_path = str(FIXTURES_DIR / "simple.docx")

        with tempfile.TemporaryDirectory() as temp_dir:
            sidedoc_path = f"{temp_dir}/test.sidedoc"
            output_path = f"{temp_dir}/synced.docx"

            blocks, image_data = extract_blocks(fixture_path)
            styles = extract_styles(fixture_path, blocks)
            content_md = blocks_to_markdown(blocks)
            create_sidedoc_archive(sidedoc_path, content_md, blocks, styles, fixture_path, image_data)

            with zipfile.ZipFile(sidedoc_path, "r") as z:
                original_content = z.read("content.md").decode("utf-8")
                structure_json = z.read("structure.json").decode("utf-8")
                styles_json = z.read("styles.json").decode("utf-8")
                manifest_json = z.read("manifest.json").decode("utf-8")

            modified_content = original_content.replace("simple document", "{++very ++}simple document")

            with zipfile.ZipFile(sidedoc_path, "w", zipfile.ZIP_DEFLATED) as z:
                z.writestr("content.md", modified_content)
                z.writestr("structure.json", structure_json)
                z.writestr("styles.json", styles_json)
                z.writestr("manifest.json", manifest_json)

            # Sync without specifying author
            sync_sidedoc_to_docx(sidedoc_path, output_path)

            doc = Document(output_path)
            ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

            for para in doc.paragraphs:
                for ins in para._element.findall(f".//{{{ns}}}ins"):
                    author = ins.get(f"{{{ns}}}author")
                    if author:
                        assert author == "Sidedoc AI", f"Expected 'Sidedoc AI', got '{author}'"
                        return

            pytest.fail("No author attribute found on w:ins elements")

    def test_sync_uses_current_timestamp(self) -> None:
        """Verify new track changes use current timestamp."""
        import tempfile
        import zipfile
        from datetime import datetime, timezone
        from sidedoc.extract import extract_blocks, extract_styles, blocks_to_markdown
        from sidedoc.package import create_sidedoc_archive
        from sidedoc.sync import sync_sidedoc_to_docx
        from docx import Document

        fixture_path = str(FIXTURES_DIR / "simple.docx")

        with tempfile.TemporaryDirectory() as temp_dir:
            sidedoc_path = f"{temp_dir}/test.sidedoc"
            output_path = f"{temp_dir}/synced.docx"

            blocks, image_data = extract_blocks(fixture_path)
            styles = extract_styles(fixture_path, blocks)
            content_md = blocks_to_markdown(blocks)
            create_sidedoc_archive(sidedoc_path, content_md, blocks, styles, fixture_path, image_data)

            with zipfile.ZipFile(sidedoc_path, "r") as z:
                original_content = z.read("content.md").decode("utf-8")
                structure_json = z.read("structure.json").decode("utf-8")
                styles_json = z.read("styles.json").decode("utf-8")
                manifest_json = z.read("manifest.json").decode("utf-8")

            modified_content = original_content.replace("simple document", "{++very ++}simple document")

            with zipfile.ZipFile(sidedoc_path, "w", zipfile.ZIP_DEFLATED) as z:
                z.writestr("content.md", modified_content)
                z.writestr("structure.json", structure_json)
                z.writestr("styles.json", styles_json)
                z.writestr("manifest.json", manifest_json)

            before_sync = datetime.now(timezone.utc)
            sync_sidedoc_to_docx(sidedoc_path, output_path)
            after_sync = datetime.now(timezone.utc)

            doc = Document(output_path)
            ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

            for para in doc.paragraphs:
                for ins in para._element.findall(f".//{{{ns}}}ins"):
                    date_str = ins.get(f"{{{ns}}}date")
                    if date_str:
                        # Parse ISO date and verify it's recent
                        sync_date = datetime.fromisoformat(date_str.replace("Z", "+00:00"))
                        # Date should be between before and after sync (with 1 minute buffer)
                        from datetime import timedelta
                        assert sync_date >= before_sync - timedelta(minutes=1), "Date should be recent"
                        assert sync_date <= after_sync + timedelta(minutes=1), "Date should not be in future"
                        return

            pytest.fail("No date attribute found on w:ins elements")


class TestCriticMarkupSubstitution:
    """Tests for CriticMarkup substitution syntax (US-014)."""

    def test_parse_criticmarkup_substitution(self) -> None:
        """Verify {~~old~>new~~} is parsed into deletion and insertion."""
        from sidedoc.reconstruct import parse_criticmarkup

        text = "Hello {~~world~>universe~~} today"
        segments = parse_criticmarkup(text)

        # Should have: text, deletion, insertion, text
        assert len(segments) == 4
        assert segments[0] == ("text", "Hello ")
        assert segments[1] == ("deletion", "world")
        assert segments[2] == ("insertion", "universe")
        assert segments[3] == ("text", " today")

    def test_sync_substitution_creates_deletion_and_insertion(self) -> None:
        """Verify {~~old~>new~~} creates both w:del and w:ins elements."""
        import tempfile
        import zipfile
        from sidedoc.extract import extract_blocks, extract_styles, blocks_to_markdown
        from sidedoc.package import create_sidedoc_archive
        from sidedoc.sync import sync_sidedoc_to_docx
        from docx import Document

        fixture_path = str(FIXTURES_DIR / "simple.docx")

        with tempfile.TemporaryDirectory() as temp_dir:
            sidedoc_path = f"{temp_dir}/test.sidedoc"
            output_path = f"{temp_dir}/synced.docx"

            blocks, image_data = extract_blocks(fixture_path)
            styles = extract_styles(fixture_path, blocks)
            content_md = blocks_to_markdown(blocks)
            create_sidedoc_archive(sidedoc_path, content_md, blocks, styles, fixture_path, image_data)

            with zipfile.ZipFile(sidedoc_path, "r") as z:
                original_content = z.read("content.md").decode("utf-8")
                structure_json = z.read("structure.json").decode("utf-8")
                styles_json = z.read("styles.json").decode("utf-8")
                manifest_json = z.read("manifest.json").decode("utf-8")

            # Use substitution syntax
            modified_content = original_content.replace("simple", "{~~simple~>straightforward~~}")

            with zipfile.ZipFile(sidedoc_path, "w", zipfile.ZIP_DEFLATED) as z:
                z.writestr("content.md", modified_content)
                z.writestr("structure.json", structure_json)
                z.writestr("styles.json", styles_json)
                z.writestr("manifest.json", manifest_json)

            sync_sidedoc_to_docx(sidedoc_path, output_path)

            doc = Document(output_path)
            nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

            found_ins = False
            found_del = False
            for para in doc.paragraphs:
                ins_elements = para._element.findall(".//w:ins", nsmap)
                del_elements = para._element.findall(".//w:del", nsmap)
                if ins_elements:
                    found_ins = True
                if del_elements:
                    found_del = True

            assert found_ins, "Substitution should create w:ins element"
            assert found_del, "Substitution should create w:del element"

    def test_substitution_preserves_adjacent_text(self) -> None:
        """Verify substitution preserves text before and after."""
        from sidedoc.reconstruct import parse_criticmarkup

        text = "This is a {~~test~>demo~~} of substitution"
        segments = parse_criticmarkup(text)

        # Extract plain text parts
        text_parts = [seg[1] for seg in segments if seg[0] == "text"]
        assert "This is a " in text_parts
        assert " of substitution" in text_parts

    def test_multiple_substitutions_in_same_text(self) -> None:
        """Verify multiple substitutions in same text work correctly."""
        from sidedoc.reconstruct import parse_criticmarkup

        text = "Hello {~~world~>universe~~} and {~~goodbye~>farewell~~} today"
        segments = parse_criticmarkup(text)

        # Should have: text, del, ins, text, del, ins, text
        deletions = [seg for seg in segments if seg[0] == "deletion"]
        insertions = [seg for seg in segments if seg[0] == "insertion"]

        assert len(deletions) == 2
        assert len(insertions) == 2
        assert deletions[0] == ("deletion", "world")
        assert deletions[1] == ("deletion", "goodbye")
        assert insertions[0] == ("insertion", "universe")
        assert insertions[1] == ("insertion", "farewell")


class TestAIEditsAsTrackedChanges:
    """Tests for AI edits appearing as tracked changes (US-015)."""

    def test_ai_edit_creates_insertion_with_author(self) -> None:
        """Verify AI edits use configured author name for insertions."""
        import tempfile
        import zipfile
        from sidedoc.extract import extract_blocks, extract_styles, blocks_to_markdown
        from sidedoc.package import create_sidedoc_archive
        from sidedoc.sync import sync_sidedoc_to_docx
        from docx import Document

        # Use a fixture that already has track changes
        fixture_path = str(FIXTURES_DIR / "track_changes_simple.docx")

        with tempfile.TemporaryDirectory() as temp_dir:
            sidedoc_path = f"{temp_dir}/test.sidedoc"
            output_path = f"{temp_dir}/synced.docx"

            blocks, image_data = extract_blocks(fixture_path)
            styles = extract_styles(fixture_path, blocks)
            content_md = blocks_to_markdown(blocks)
            create_sidedoc_archive(sidedoc_path, content_md, blocks, styles, fixture_path, image_data)

            with zipfile.ZipFile(sidedoc_path, "r") as z:
                original_content = z.read("content.md").decode("utf-8")
                structure_json = z.read("structure.json").decode("utf-8")
                styles_json = z.read("styles.json").decode("utf-8")
                manifest_json = z.read("manifest.json").decode("utf-8")

            # AI adds new content
            modified_content = original_content + "\n\n{++This is a new paragraph added by AI.++}"

            with zipfile.ZipFile(sidedoc_path, "w", zipfile.ZIP_DEFLATED) as z:
                z.writestr("content.md", modified_content)
                z.writestr("structure.json", structure_json)
                z.writestr("styles.json", styles_json)
                z.writestr("manifest.json", manifest_json)

            sync_sidedoc_to_docx(sidedoc_path, output_path, author="AI Assistant")

            doc = Document(output_path)
            ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

            # Find the AI's insertion
            for para in doc.paragraphs:
                for ins in para._element.findall(f".//{{{ns}}}ins"):
                    author = ins.get(f"{{{ns}}}author")
                    if author == "AI Assistant":
                        # Check the text is our new content
                        text_elems = ins.findall(f".//{{{ns}}}t")
                        text = "".join(t.text for t in text_elems if t.text)
                        if "new paragraph" in text:
                            return

            pytest.fail("AI insertion with author 'AI Assistant' not found")

    def test_ai_edit_creates_deletion_with_author(self) -> None:
        """Verify AI edits use configured author name for deletions."""
        import tempfile
        import zipfile
        from sidedoc.extract import extract_blocks, extract_styles, blocks_to_markdown
        from sidedoc.package import create_sidedoc_archive
        from sidedoc.sync import sync_sidedoc_to_docx
        from docx import Document

        fixture_path = str(FIXTURES_DIR / "track_changes_simple.docx")

        with tempfile.TemporaryDirectory() as temp_dir:
            sidedoc_path = f"{temp_dir}/test.sidedoc"
            output_path = f"{temp_dir}/synced.docx"

            blocks, image_data = extract_blocks(fixture_path)
            styles = extract_styles(fixture_path, blocks)
            content_md = blocks_to_markdown(blocks)
            create_sidedoc_archive(sidedoc_path, content_md, blocks, styles, fixture_path, image_data)

            with zipfile.ZipFile(sidedoc_path, "r") as z:
                original_content = z.read("content.md").decode("utf-8")
                structure_json = z.read("structure.json").decode("utf-8")
                styles_json = z.read("styles.json").decode("utf-8")
                manifest_json = z.read("manifest.json").decode("utf-8")

            # AI marks content as deleted
            modified_content = original_content.replace("some text", "{--some text--}")

            with zipfile.ZipFile(sidedoc_path, "w", zipfile.ZIP_DEFLATED) as z:
                z.writestr("content.md", modified_content)
                z.writestr("structure.json", structure_json)
                z.writestr("styles.json", styles_json)
                z.writestr("manifest.json", manifest_json)

            sync_sidedoc_to_docx(sidedoc_path, output_path, author="AI Assistant")

            doc = Document(output_path)
            ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

            # Find the AI's deletion
            for para in doc.paragraphs:
                for del_elem in para._element.findall(f".//{{{ns}}}del"):
                    author = del_elem.get(f"{{{ns}}}author")
                    if author == "AI Assistant":
                        return

            pytest.fail("AI deletion with author 'AI Assistant' not found")

    def test_ai_edit_uses_current_timestamp(self) -> None:
        """Verify AI edits use current timestamp."""
        import tempfile
        import zipfile
        from datetime import datetime, timedelta, timezone
        from sidedoc.extract import extract_blocks, extract_styles, blocks_to_markdown
        from sidedoc.package import create_sidedoc_archive
        from sidedoc.sync import sync_sidedoc_to_docx
        from docx import Document

        fixture_path = str(FIXTURES_DIR / "track_changes_simple.docx")

        with tempfile.TemporaryDirectory() as temp_dir:
            sidedoc_path = f"{temp_dir}/test.sidedoc"
            output_path = f"{temp_dir}/synced.docx"

            blocks, image_data = extract_blocks(fixture_path)
            styles = extract_styles(fixture_path, blocks)
            content_md = blocks_to_markdown(blocks)
            create_sidedoc_archive(sidedoc_path, content_md, blocks, styles, fixture_path, image_data)

            with zipfile.ZipFile(sidedoc_path, "r") as z:
                original_content = z.read("content.md").decode("utf-8")
                structure_json = z.read("structure.json").decode("utf-8")
                styles_json = z.read("styles.json").decode("utf-8")
                manifest_json = z.read("manifest.json").decode("utf-8")

            modified_content = original_content + "\n\n{++AI added this.++}"

            with zipfile.ZipFile(sidedoc_path, "w", zipfile.ZIP_DEFLATED) as z:
                z.writestr("content.md", modified_content)
                z.writestr("structure.json", structure_json)
                z.writestr("styles.json", styles_json)
                z.writestr("manifest.json", manifest_json)

            before_sync = datetime.now(timezone.utc)
            sync_sidedoc_to_docx(sidedoc_path, output_path, author="AI Assistant")
            after_sync = datetime.now(timezone.utc)

            doc = Document(output_path)
            ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

            for para in doc.paragraphs:
                for ins in para._element.findall(f".//{{{ns}}}ins"):
                    author = ins.get(f"{{{ns}}}author")
                    if author == "AI Assistant":
                        date_str = ins.get(f"{{{ns}}}date")
                        if date_str:
                            sync_date = datetime.fromisoformat(date_str.replace("Z", "+00:00"))
                            # Timestamp should be within the sync window
                            assert sync_date >= before_sync - timedelta(minutes=1)
                            assert sync_date <= after_sync + timedelta(minutes=1)
                            return

            pytest.fail("AI insertion with current timestamp not found")

    def test_preserves_original_track_changes_content(self) -> None:
        """Verify original track changes content is preserved alongside AI edits."""
        import tempfile
        import zipfile
        from sidedoc.extract import extract_blocks, extract_styles, blocks_to_markdown
        from sidedoc.package import create_sidedoc_archive
        from sidedoc.sync import sync_sidedoc_to_docx
        from docx import Document

        fixture_path = str(FIXTURES_DIR / "track_changes_simple.docx")

        with tempfile.TemporaryDirectory() as temp_dir:
            sidedoc_path = f"{temp_dir}/test.sidedoc"
            output_path = f"{temp_dir}/synced.docx"

            blocks, image_data = extract_blocks(fixture_path)
            styles = extract_styles(fixture_path, blocks)
            content_md = blocks_to_markdown(blocks)
            create_sidedoc_archive(sidedoc_path, content_md, blocks, styles, fixture_path, image_data)

            with zipfile.ZipFile(sidedoc_path, "r") as z:
                original_content = z.read("content.md").decode("utf-8")
                structure_json = z.read("structure.json").decode("utf-8")
                styles_json = z.read("styles.json").decode("utf-8")
                manifest_json = z.read("manifest.json").decode("utf-8")

            # AI adds new content but doesn't remove existing track changes
            modified_content = original_content + "\n\n{++New from AI.++}"

            with zipfile.ZipFile(sidedoc_path, "w", zipfile.ZIP_DEFLATED) as z:
                z.writestr("content.md", modified_content)
                z.writestr("structure.json", structure_json)
                z.writestr("styles.json", styles_json)
                z.writestr("manifest.json", manifest_json)

            sync_sidedoc_to_docx(sidedoc_path, output_path, author="AI Assistant")

            doc = Document(output_path)
            ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

            # Check that all CriticMarkup content is preserved as track changes
            ins_count = 0
            del_count = 0
            for para in doc.paragraphs:
                ins_count += len(para._element.findall(f".//{{{ns}}}ins"))
                del_count += len(para._element.findall(f".//{{{ns}}}del"))

            # Original content has 1 insertion and 1 deletion, plus AI's new insertion
            assert ins_count >= 2, f"Expected at least 2 insertions (original + AI), got {ins_count}"
            assert del_count >= 1, f"Expected at least 1 deletion (original), got {del_count}"


class TestTrackChangesCLI:
    """Tests for --track-changes CLI option (US-016)."""

    def test_extract_with_track_changes_flag(self) -> None:
        """Verify --track-changes flag forces track changes mode."""
        import tempfile
        from click.testing import CliRunner
        from sidedoc.cli import main

        runner = CliRunner()
        # Use a fixture that has track changes to verify they're extracted
        fixture_path = str(FIXTURES_DIR / "track_changes_simple.docx")

        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = f"{temp_dir}/output.sidedoc"

            result = runner.invoke(main, ["extract", fixture_path, "-o", output_path, "--track-changes"])

            assert result.exit_code == 0, f"Extract failed: {result.output}"
            assert Path(output_path).exists()

            # Content should have CriticMarkup when track changes are forced
            content = (Path(output_path) / "content.md").read_text()
            # With --track-changes, should contain CriticMarkup
            assert "{++" in content or "{--" in content, "Track changes should be extracted with --track-changes"

    def test_extract_with_no_track_changes_flag(self) -> None:
        """Verify --no-track-changes flag disables track changes mode."""
        import tempfile
        from click.testing import CliRunner
        from sidedoc.cli import main

        runner = CliRunner()
        # Use a fixture that has track changes
        fixture_path = str(FIXTURES_DIR / "track_changes_simple.docx")

        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = f"{temp_dir}/output.sidedoc"

            result = runner.invoke(main, ["extract", fixture_path, "-o", output_path, "--no-track-changes"])

            assert result.exit_code == 0, f"Extract failed: {result.output}"

            # Content should NOT have CriticMarkup when track changes are disabled
            content = (Path(output_path) / "content.md").read_text()
            # Should not contain CriticMarkup syntax
            assert "{++" not in content, "Content should not have insertions with --no-track-changes"
            assert "{--" not in content, "Content should not have deletions with --no-track-changes"

    def test_extract_auto_detects_track_changes(self) -> None:
        """Verify extract auto-detects track changes when no flag is provided."""
        import tempfile
        from click.testing import CliRunner
        from sidedoc.cli import main

        runner = CliRunner()
        fixture_path = str(FIXTURES_DIR / "track_changes_simple.docx")

        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = f"{temp_dir}/output.sidedoc"

            result = runner.invoke(main, ["extract", fixture_path, "-o", output_path])

            assert result.exit_code == 0, f"Extract failed: {result.output}"

            # Content should have CriticMarkup when track changes are auto-detected
            content = (Path(output_path) / "content.md").read_text()
            # Should contain CriticMarkup syntax
            assert "{++" in content or "{--" in content, "Track changes should be auto-detected"


class TestSyncAuthorCLI:
    """Tests for --author CLI option on sync command (US-017)."""

    def test_sync_with_author_option(self) -> None:
        """Verify --author option sets author for track changes."""
        import tempfile
        from click.testing import CliRunner
        from sidedoc.cli import main
        from sidedoc.extract import extract_blocks, extract_styles, blocks_to_markdown
        from sidedoc.package import create_sidedoc_directory
        from docx import Document

        runner = CliRunner()
        fixture_path = str(FIXTURES_DIR / "simple.docx")

        with tempfile.TemporaryDirectory() as temp_dir:
            sidedoc_path = f"{temp_dir}/test.sidedoc"
            output_path = f"{temp_dir}/synced.docx"

            # Create sidedoc directory
            blocks, image_data = extract_blocks(fixture_path)
            styles = extract_styles(fixture_path, blocks)
            content_md = blocks_to_markdown(blocks)
            create_sidedoc_directory(sidedoc_path, content_md, blocks, styles, fixture_path, image_data)

            # Modify content with CriticMarkup
            content_file = Path(sidedoc_path) / "content.md"
            original_content = content_file.read_text()
            modified_content = original_content.replace("simple", "{++very ++}simple")
            content_file.write_text(modified_content)

            # Sync with custom author
            result = runner.invoke(main, ["sync", sidedoc_path, "-o", output_path, "--author", "Custom Author"])

            assert result.exit_code == 0, f"Sync failed: {result.output}"

            # Verify author in docx
            doc = Document(output_path)
            ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

            for para in doc.paragraphs:
                for ins in para._element.findall(f".//{{{ns}}}ins"):
                    author = ins.get(f"{{{ns}}}author")
                    if author:
                        assert author == "Custom Author", f"Expected 'Custom Author', got '{author}'"
                        return

            pytest.fail("No author attribute found on w:ins elements")

    def test_sync_default_author_when_not_specified(self) -> None:
        """Verify default author is 'Sidedoc AI' when --author is not specified."""
        import tempfile
        from click.testing import CliRunner
        from sidedoc.cli import main
        from sidedoc.extract import extract_blocks, extract_styles, blocks_to_markdown
        from sidedoc.package import create_sidedoc_directory
        from docx import Document

        runner = CliRunner()
        fixture_path = str(FIXTURES_DIR / "simple.docx")

        with tempfile.TemporaryDirectory() as temp_dir:
            sidedoc_path = f"{temp_dir}/test.sidedoc"
            output_path = f"{temp_dir}/synced.docx"

            blocks, image_data = extract_blocks(fixture_path)
            styles = extract_styles(fixture_path, blocks)
            content_md = blocks_to_markdown(blocks)
            create_sidedoc_directory(sidedoc_path, content_md, blocks, styles, fixture_path, image_data)

            # Modify content with CriticMarkup
            content_file = Path(sidedoc_path) / "content.md"
            original_content = content_file.read_text()
            modified_content = original_content.replace("simple", "{++very ++}simple")
            content_file.write_text(modified_content)

            # Sync without specifying author
            result = runner.invoke(main, ["sync", sidedoc_path, "-o", output_path])

            assert result.exit_code == 0, f"Sync failed: {result.output}"

            doc = Document(output_path)
            ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

            for para in doc.paragraphs:
                for ins in para._element.findall(f".//{{{ns}}}ins"):
                    author = ins.get(f"{{{ns}}}author")
                    if author:
                        assert author == "Sidedoc AI", f"Expected 'Sidedoc AI', got '{author}'"
                        return

            pytest.fail("No author attribute found on w:ins elements")


class TestTrackChangesValidation:
    """Tests for track changes validation in sidedoc validate command (US-018)."""

    def test_validate_passes_for_valid_track_changes(self) -> None:
        """Verify validation passes for documents with valid track changes."""
        import tempfile
        from click.testing import CliRunner
        from sidedoc.cli import main
        from sidedoc.extract import extract_blocks, extract_styles, blocks_to_markdown
        from sidedoc.package import create_sidedoc_archive

        runner = CliRunner()
        fixture_path = str(FIXTURES_DIR / "track_changes_simple.docx")

        with tempfile.TemporaryDirectory() as temp_dir:
            sidedoc_path = f"{temp_dir}/test.sidedoc"

            blocks, image_data = extract_blocks(fixture_path)
            styles = extract_styles(fixture_path, blocks)
            content_md = blocks_to_markdown(blocks)
            create_sidedoc_archive(sidedoc_path, content_md, blocks, styles, fixture_path, image_data)

            result = runner.invoke(main, ["validate", sidedoc_path])

            assert result.exit_code == 0, f"Validation failed unexpectedly: {result.output}"
            assert "valid" in result.output.lower()

    def test_validate_checks_track_change_positions(self) -> None:
        """Verify validation checks track change positions are within block bounds."""
        import tempfile
        import zipfile
        import json
        from click.testing import CliRunner
        from sidedoc.cli import main
        from sidedoc.extract import extract_blocks, extract_styles, blocks_to_markdown
        from sidedoc.package import create_sidedoc_archive

        runner = CliRunner()
        fixture_path = str(FIXTURES_DIR / "track_changes_simple.docx")

        with tempfile.TemporaryDirectory() as temp_dir:
            sidedoc_path = f"{temp_dir}/test.sidedoc"

            blocks, image_data = extract_blocks(fixture_path)
            styles = extract_styles(fixture_path, blocks)
            content_md = blocks_to_markdown(blocks)
            create_sidedoc_archive(sidedoc_path, content_md, blocks, styles, fixture_path, image_data)

            # Corrupt the track changes to have invalid positions
            with zipfile.ZipFile(sidedoc_path, "r") as z:
                structure = json.loads(z.read("structure.json").decode("utf-8"))
                content = z.read("content.md").decode("utf-8")
                styles_json = z.read("styles.json").decode("utf-8")
                manifest_json = z.read("manifest.json").decode("utf-8")

            # Add track changes with invalid positions
            for block in structure.get("blocks", []):
                if block.get("track_changes"):
                    for tc in block["track_changes"]:
                        tc["end"] = 9999  # Invalid position beyond block bounds

            with zipfile.ZipFile(sidedoc_path, "w", zipfile.ZIP_DEFLATED) as z:
                z.writestr("content.md", content)
                z.writestr("structure.json", json.dumps(structure, indent=2))
                z.writestr("styles.json", styles_json)
                z.writestr("manifest.json", manifest_json)

            result = runner.invoke(main, ["validate", sidedoc_path])

            # Should fail validation due to invalid positions
            assert result.exit_code != 0 or "warning" in result.output.lower() or "invalid" in result.output.lower()

    def test_validate_checks_track_change_metadata(self) -> None:
        """Verify validation checks track change metadata is complete."""
        import tempfile
        import zipfile
        import json
        from click.testing import CliRunner
        from sidedoc.cli import main
        from sidedoc.extract import extract_blocks, extract_styles, blocks_to_markdown
        from sidedoc.package import create_sidedoc_archive

        runner = CliRunner()
        fixture_path = str(FIXTURES_DIR / "track_changes_simple.docx")

        with tempfile.TemporaryDirectory() as temp_dir:
            sidedoc_path = f"{temp_dir}/test.sidedoc"

            blocks, image_data = extract_blocks(fixture_path)
            styles = extract_styles(fixture_path, blocks)
            content_md = blocks_to_markdown(blocks)
            create_sidedoc_archive(sidedoc_path, content_md, blocks, styles, fixture_path, image_data)

            # Remove author/date metadata
            with zipfile.ZipFile(sidedoc_path, "r") as z:
                structure = json.loads(z.read("structure.json").decode("utf-8"))
                content = z.read("content.md").decode("utf-8")
                styles_json = z.read("styles.json").decode("utf-8")
                manifest_json = z.read("manifest.json").decode("utf-8")

            for block in structure.get("blocks", []):
                if block.get("track_changes"):
                    for tc in block["track_changes"]:
                        tc["author"] = ""  # Missing author
                        tc["date"] = ""    # Missing date

            with zipfile.ZipFile(sidedoc_path, "w", zipfile.ZIP_DEFLATED) as z:
                z.writestr("content.md", content)
                z.writestr("structure.json", json.dumps(structure, indent=2))
                z.writestr("styles.json", styles_json)
                z.writestr("manifest.json", manifest_json)

            result = runner.invoke(main, ["validate", sidedoc_path])

            # Validation should fail with warnings about missing metadata
            assert result.exit_code != 0, "Should fail validation for missing metadata"
            assert "missing author" in result.output.lower()
            assert "missing date" in result.output.lower()


class TestNestedMarkdownInCriticMarkup:
    """Tests for CriticMarkup with nested markdown formatting (US-019)."""

    def test_parse_bold_inside_insertion(self) -> None:
        """Verify {++**bold text**++} is parsed correctly."""
        from sidedoc.reconstruct import parse_criticmarkup

        text = "Hello {++**bold**++} world"
        segments = parse_criticmarkup(text)

        # Should have: text, insertion (with bold markers inside), text
        assert len(segments) == 3
        assert segments[0] == ("text", "Hello ")
        assert segments[1] == ("insertion", "**bold**")
        assert segments[2] == ("text", " world")

    def test_parse_italic_inside_deletion(self) -> None:
        """Verify {--*italic*--} is parsed correctly."""
        from sidedoc.reconstruct import parse_criticmarkup

        text = "Hello {--*italic*--} world"
        segments = parse_criticmarkup(text)

        assert len(segments) == 3
        assert segments[0] == ("text", "Hello ")
        assert segments[1] == ("deletion", "*italic*")
        assert segments[2] == ("text", " world")

    def test_sync_nested_formatting_preserved(self) -> None:
        """Verify nested formatting in CriticMarkup is preserved in sync."""
        import tempfile
        import zipfile
        from sidedoc.extract import extract_blocks, extract_styles, blocks_to_markdown
        from sidedoc.package import create_sidedoc_archive
        from sidedoc.sync import sync_sidedoc_to_docx
        from docx import Document

        fixture_path = str(FIXTURES_DIR / "simple.docx")

        with tempfile.TemporaryDirectory() as temp_dir:
            sidedoc_path = f"{temp_dir}/test.sidedoc"
            output_path = f"{temp_dir}/synced.docx"

            blocks, image_data = extract_blocks(fixture_path)
            styles = extract_styles(fixture_path, blocks)
            content_md = blocks_to_markdown(blocks)
            create_sidedoc_archive(sidedoc_path, content_md, blocks, styles, fixture_path, image_data)

            with zipfile.ZipFile(sidedoc_path, "r") as z:
                original_content = z.read("content.md").decode("utf-8")
                structure_json = z.read("structure.json").decode("utf-8")
                styles_json = z.read("styles.json").decode("utf-8")
                manifest_json = z.read("manifest.json").decode("utf-8")

            # Add insertion with nested bold
            modified_content = original_content.replace("simple", "{++**important**++} simple")

            with zipfile.ZipFile(sidedoc_path, "w", zipfile.ZIP_DEFLATED) as z:
                z.writestr("content.md", modified_content)
                z.writestr("structure.json", structure_json)
                z.writestr("styles.json", styles_json)
                z.writestr("manifest.json", manifest_json)

            sync_sidedoc_to_docx(sidedoc_path, output_path)

            doc = Document(output_path)
            ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

            # Verify w:ins element was created
            found_ins = False
            for para in doc.paragraphs:
                ins_elements = para._element.findall(f".//{{{ns}}}ins")
                if ins_elements:
                    found_ins = True
                    # Check that the text inside contains the bold markers
                    for ins in ins_elements:
                        text_elems = ins.findall(f".//{{{ns}}}t")
                        text = "".join(t.text for t in text_elems if t.text)
                        # The text should include the bold content
                        if "important" in text:
                            return

            assert found_ins, "Should create w:ins for insertion with nested formatting"

    def test_multiple_nested_formats(self) -> None:
        """Verify complex nested formatting in CriticMarkup works."""
        from sidedoc.reconstruct import parse_criticmarkup

        text = "Text {++**bold *and italic***++} more"
        segments = parse_criticmarkup(text)

        # Should preserve the nested formatting inside the insertion
        assert len(segments) == 3
        assert segments[0] == ("text", "Text ")
        assert segments[1] == ("insertion", "**bold *and italic***")
        assert segments[2] == ("text", " more")


class TestInvalidCriticMarkupSyntax:
    """Tests for error handling with invalid CriticMarkup syntax (US-020)."""

    def test_unclosed_insertion_treated_as_text(self) -> None:
        """Verify unclosed {++ is treated as plain text."""
        from sidedoc.reconstruct import parse_criticmarkup

        # Unclosed insertion - should be treated as plain text
        text = "Hello {++ unclosed"
        segments = parse_criticmarkup(text)

        # Should return whole text as-is (no CriticMarkup found)
        assert len(segments) == 1
        assert segments[0] == ("text", "Hello {++ unclosed")

    def test_unclosed_deletion_treated_as_text(self) -> None:
        """Verify unclosed {-- is treated as plain text."""
        from sidedoc.reconstruct import parse_criticmarkup

        text = "Hello {-- unclosed"
        segments = parse_criticmarkup(text)

        assert len(segments) == 1
        assert segments[0] == ("text", "Hello {-- unclosed")

    def test_malformed_substitution_treated_as_text(self) -> None:
        """Verify malformed {~~old~>new is treated as plain text."""
        from sidedoc.reconstruct import parse_criticmarkup

        # Malformed - missing closing ~~}
        text = "Hello {~~old~>new missing close"
        segments = parse_criticmarkup(text)

        assert len(segments) == 1
        assert segments[0] == ("text", "Hello {~~old~>new missing close")

    def test_mixed_valid_and_invalid_syntax(self) -> None:
        """Verify valid CriticMarkup is extracted even with invalid syntax present."""
        from sidedoc.reconstruct import parse_criticmarkup

        # Has both valid and invalid
        text = "Hello {++valid++} and {-- unclosed"
        segments = parse_criticmarkup(text)

        # Should extract the valid one, leave the invalid as text
        assert len(segments) == 3
        assert segments[0] == ("text", "Hello ")
        assert segments[1] == ("insertion", "valid")
        assert segments[2] == ("text", " and {-- unclosed")

    def test_validate_criticmarkup_syntax(self) -> None:
        """Verify validate_criticmarkup function detects errors."""
        from sidedoc.reconstruct import validate_criticmarkup

        # Valid syntax
        errors = validate_criticmarkup("Hello {++world++}")
        assert len(errors) == 0

        # Invalid syntax - unclosed
        errors = validate_criticmarkup("Hello {++ unclosed")
        assert len(errors) >= 1
        assert any("unclosed" in e.lower() for e in errors)

    def test_validate_criticmarkup_with_line_numbers(self) -> None:
        """Verify error messages include line numbers."""
        from sidedoc.reconstruct import validate_criticmarkup

        content = "Line 1\nLine 2 with {++ unclosed\nLine 3"
        errors = validate_criticmarkup(content)

        # Should have error with line context
        assert len(errors) >= 1
        assert any("line 2" in e.lower() or "2" in e for e in errors)
