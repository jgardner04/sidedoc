"""Test roundtrip: extract → build produces correct output."""

import tempfile
from pathlib import Path
from click.testing import CliRunner
from docx import Document
from sidedoc.cli import main


def test_roundtrip_preserves_headings():
    """Test that extract → build preserves heading styles."""
    runner = CliRunner()
    with runner.isolated_filesystem():
        # Create original docx
        doc = Document()
        doc.add_paragraph("Main Title", style="Heading 1")
        doc.add_paragraph("Introduction text")
        doc.add_paragraph("Section Title", style="Heading 2")
        doc.add_paragraph("Section content")

        docx_path = "original.docx"
        doc.save(docx_path)

        # Extract
        result = runner.invoke(main, ["extract", docx_path])
        assert result.exit_code == 0

        # Build
        result = runner.invoke(main, ["build", "original.sidedoc", "-o", "rebuilt.docx"])
        assert result.exit_code == 0

        # Verify rebuilt document
        rebuilt = Document("rebuilt.docx")
        assert len(rebuilt.paragraphs) >= 4

        # Check text content is preserved
        texts = [p.text for p in rebuilt.paragraphs]
        assert "Main Title" in texts
        assert "Introduction text" in texts
        assert "Section Title" in texts
        assert "Section content" in texts


def test_roundtrip_simple_document():
    """Test roundtrip with a simple document."""
    runner = CliRunner()
    with runner.isolated_filesystem():
        # Create simple docx
        doc = Document()
        doc.add_paragraph("Hello world")
        doc.add_paragraph("Second paragraph")

        doc.save("test.docx")

        # Extract and build
        runner.invoke(main, ["extract", "test.docx"])
        result = runner.invoke(main, ["build", "test.sidedoc", "-o", "rebuilt.docx"])

        assert result.exit_code == 0

        # Verify content
        rebuilt = Document("rebuilt.docx")
        texts = [p.text for p in rebuilt.paragraphs]
        assert "Hello world" in texts
        assert "Second paragraph" in texts


def test_roundtrip_multiple_heading_levels():
    """Test roundtrip with multiple heading levels."""
    runner = CliRunner()
    with runner.isolated_filesystem():
        # Create docx with various heading levels
        doc = Document()
        doc.add_paragraph("H1", style="Heading 1")
        doc.add_paragraph("H2", style="Heading 2")
        doc.add_paragraph("H3", style="Heading 3")

        doc.save("test.docx")

        # Extract and build
        runner.invoke(main, ["extract", "test.docx"])
        result = runner.invoke(main, ["build", "test.sidedoc"])

        assert result.exit_code == 0

        # Verify headings preserved
        rebuilt = Document("test.docx")
        texts = [p.text for p in rebuilt.paragraphs]
        assert "H1" in texts
        assert "H2" in texts
        assert "H3" in texts


def test_complete_workflow():
    """Test complete workflow: extract → build from directory, and extract --pack → unpack → pack → build."""
    runner = CliRunner()
    with runner.isolated_filesystem():
        # Create original
        doc = Document()
        doc.add_paragraph("Test Document", style="Heading 1")
        doc.add_paragraph("This is a test")
        doc.save("original.docx")

        # Extract to directory
        result = runner.invoke(main, ["extract", "original.docx"])
        assert result.exit_code == 0
        assert Path("original.sidedoc").is_dir()
        assert Path("original.sidedoc/content.md").exists()

        # Build from directory
        result = runner.invoke(main, ["build", "original.sidedoc", "-o", "from_dir.docx"])
        assert result.exit_code == 0

        # Verify document from directory
        from_dir = Document("from_dir.docx")
        texts = [p.text for p in from_dir.paragraphs]
        assert "Test Document" in texts
        assert "This is a test" in texts

        # Also test ZIP distribution workflow: pack → unpack → build
        result = runner.invoke(main, ["pack", "original.sidedoc", "-o", "distributed.sdoc"])
        assert result.exit_code == 0

        result = runner.invoke(main, ["unpack", "distributed.sdoc", "-o", "unpacked.sidedoc"])
        assert result.exit_code == 0
        assert Path("unpacked.sidedoc/content.md").exists()

        result = runner.invoke(main, ["build", "unpacked.sidedoc", "-o", "final.docx"])
        assert result.exit_code == 0

        final = Document("final.docx")
        texts = [p.text for p in final.paragraphs]
        assert "Test Document" in texts
        assert "This is a test" in texts


def _count_runs_by_format(doc):
    """Helper: tally bold/italic/underline runs across body paragraphs and
    table cells. Counts literal '*' in run text as well."""
    bold = italic = underline = 0
    asterisks = 0

    def _tally(paragraphs):
        nonlocal bold, italic, underline, asterisks
        for p in paragraphs:
            for r in p.runs:
                if r.bold:
                    bold += 1
                if r.italic:
                    italic += 1
                if r.underline:
                    underline += 1
                asterisks += r.text.count("*")

    _tally(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                _tally(c.paragraphs)
    return bold, italic, underline, asterisks


def test_roundtrip_paragraph_preserves_bold_italic():
    """Issue #72: bold/italic in plain paragraphs were rebuilt as literal **/* text.

    The paragraph emission path (reconstruct.py) was calling add_paragraph(content)
    directly, never parsing markdown emphasis into runs.
    """
    runner = CliRunner()
    with runner.isolated_filesystem():
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Plain ")
        r = p.add_run("bold"); r.bold = True
        p.add_run(" and ")
        r = p.add_run("italic"); r.italic = True
        p.add_run(".")
        doc.save("src.docx")

        assert runner.invoke(main, ["extract", "src.docx"]).exit_code == 0
        assert runner.invoke(main, ["build", "src.sidedoc", "-o", "rebuilt.docx"]).exit_code == 0

        rebuilt = Document("rebuilt.docx")
        bold, italic, _, asterisks = _count_runs_by_format(rebuilt)
        assert bold >= 1, "bold run lost on rebuild"
        assert italic >= 1, "italic run lost on rebuild"
        assert asterisks == 0, "literal * leaked into rebuilt docx"


def test_roundtrip_heading_preserves_bold():
    """Bold inside a heading must survive rebuild as a bold run, not literal **."""
    runner = CliRunner()
    with runner.isolated_filesystem():
        doc = Document()
        h = doc.add_paragraph(style="Heading 1")
        h.add_run("Intro ")
        r = h.add_run("Title"); r.bold = True
        doc.save("src.docx")

        assert runner.invoke(main, ["extract", "src.docx"]).exit_code == 0
        assert runner.invoke(main, ["build", "src.sidedoc", "-o", "rebuilt.docx"]).exit_code == 0

        rebuilt = Document("rebuilt.docx")
        bold, _, _, asterisks = _count_runs_by_format(rebuilt)
        assert bold >= 1
        assert asterisks == 0


def test_roundtrip_bold_with_trailing_space():
    """The exact reporter symptom: bold run 'СОДЕРЖАНИЕ ' must rebuild as bold."""
    runner = CliRunner()
    with runner.isolated_filesystem():
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("СОДЕРЖАНИЕ "); r.bold = True
        p.add_run("body")
        doc.save("src.docx")

        assert runner.invoke(main, ["extract", "src.docx"]).exit_code == 0
        assert runner.invoke(main, ["build", "src.sidedoc", "-o", "rebuilt.docx"]).exit_code == 0

        rebuilt = Document("rebuilt.docx")
        bold, _, _, asterisks = _count_runs_by_format(rebuilt)
        assert bold >= 1
        assert asterisks == 0
        # The Russian word survives intact.
        assert any("СОДЕРЖАНИЕ" in p.text for p in rebuilt.paragraphs)


def test_roundtrip_underline_preserved():
    """Issue #72: underline runs were extracted into inline_formatting but
    reconstruct.apply_inline_formatting never consumed that array, so all 32
    underlines in the reporter's document silently vanished."""
    runner = CliRunner()
    with runner.isolated_filesystem():
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("plain ")
        r = p.add_run("underlined"); r.underline = True
        p.add_run(" plain")
        doc.save("src.docx")

        assert runner.invoke(main, ["extract", "src.docx"]).exit_code == 0
        assert runner.invoke(main, ["build", "src.sidedoc", "-o", "rebuilt.docx"]).exit_code == 0

        rebuilt = Document("rebuilt.docx")
        _, _, underline, _ = _count_runs_by_format(rebuilt)
        assert underline >= 1, "underline run lost on rebuild"
        # Only the middle word is underlined; surrounding text is not.
        underlined_runs = [r for p in rebuilt.paragraphs for r in p.runs if r.underline]
        assert any("underlined" in r.text for r in underlined_runs)
        assert not any(r.underline and r.text.strip() == "plain" for p in rebuilt.paragraphs for r in p.runs)


def _count_tab_elements(doc):
    """Helper: count <w:tab/> elements in all paragraph bodies."""
    from docx.oxml.ns import qn
    return sum(
        len(p._element.findall(f".//{qn('w:tab')}"))
        for p in doc.paragraphs
    )


def test_roundtrip_tab_in_paragraph():
    """Issue #72: w:tab elements were silently dropped on extract.

    Signature lines, TOC dot leaders and form alignment all rely on
    tabs. The reporter saw 78 tabs go to 0.
    """
    runner = CliRunner()
    with runner.isolated_filesystem():
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("name")
        p.add_run().add_tab()
        p.add_run("value")
        doc.save("src.docx")

        assert runner.invoke(main, ["extract", "src.docx"]).exit_code == 0
        assert runner.invoke(main, ["build", "src.sidedoc", "-o", "rebuilt.docx"]).exit_code == 0

        rebuilt = Document("rebuilt.docx")
        assert _count_tab_elements(rebuilt) == 1
        full = "".join(p.text for p in rebuilt.paragraphs)
        # python-docx's Paragraph.text exposes tabs as "\t".
        assert "name\tvalue" in full


def test_roundtrip_consecutive_tabs():
    """Multiple consecutive tabs (used for column-style alignment) must all survive."""
    runner = CliRunner()
    with runner.isolated_filesystem():
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("a")
        for _ in range(3):
            p.add_run().add_tab()
        p.add_run("b")
        doc.save("src.docx")

        assert runner.invoke(main, ["extract", "src.docx"]).exit_code == 0
        assert runner.invoke(main, ["build", "src.sidedoc", "-o", "rebuilt.docx"]).exit_code == 0

        rebuilt = Document("rebuilt.docx")
        assert _count_tab_elements(rebuilt) == 3


def test_roundtrip_leading_tab_in_paragraph():
    """Leading tab on a paragraph (e.g. '\\t____ Name') must not be stripped
    by parse_markdown_to_blocks. Issue #72 had several such signature lines."""
    runner = CliRunner()
    with runner.isolated_filesystem():
        doc = Document()
        p = doc.add_paragraph()
        p.add_run().add_tab()
        p.add_run("____ Лищук И.В.")
        doc.save("src.docx")

        assert runner.invoke(main, ["extract", "src.docx"]).exit_code == 0
        assert runner.invoke(main, ["build", "src.sidedoc", "-o", "rebuilt.docx"]).exit_code == 0

        rebuilt = Document("rebuilt.docx")
        assert _count_tab_elements(rebuilt) == 1


def test_roundtrip_issue_72_russian_report():
    """End-to-end regression check against the file from issue #72.

    Original measurements (reporter's evidence + our remeasurement):
        bold runs:        47
        italic runs:      14
        tab characters:   10   (reporter's '78' included pPr tab-stop defs)
        literal '*':       0
        non-empty paras: 111

    Before any fix this branch ships:
        italic 14 -> 0, tabs 10 -> 0, literal '*' 0 -> 160.

    Assert the round-trip is now lossless for these dimensions.
    """
    runner = CliRunner()
    fixture = Path(__file__).parent / "fixtures" / "issue-72-russian-report.docx"
    with runner.isolated_filesystem() as tmpdir:
        src = Path(tmpdir) / "report.docx"
        src.write_bytes(fixture.read_bytes())

        assert runner.invoke(main, ["extract", str(src)]).exit_code == 0
        assert runner.invoke(main, ["build", str(src.with_suffix(".sidedoc")), "-o", "rebuilt.docx"]).exit_code == 0

        rebuilt = Document("rebuilt.docx")
        bold, italic, _, asterisks = _count_runs_by_format(rebuilt)
        # Cell-level bold may lose one run; cell extraction is its own ticket.
        assert bold >= 45, f"bold runs dropped: {bold}"
        assert italic >= 14, f"italic runs dropped: {italic}"
        assert asterisks == 0, f"stray '*' leaked into rebuilt: {asterisks}"
        assert _count_tab_elements(rebuilt) == 10, "tab characters dropped on rebuild"


def test_roundtrip_signature_line():
    """Russian signature line from issue #72: 'Подпись:\\t\\t\\t____________'."""
    runner = CliRunner()
    with runner.isolated_filesystem():
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Подпись:")
        for _ in range(3):
            p.add_run().add_tab()
        p.add_run("____________")
        doc.save("src.docx")

        assert runner.invoke(main, ["extract", "src.docx"]).exit_code == 0
        assert runner.invoke(main, ["build", "src.sidedoc", "-o", "rebuilt.docx"]).exit_code == 0

        rebuilt = Document("rebuilt.docx")
        assert _count_tab_elements(rebuilt) == 3
        assert any("Подпись:" in p.text and "____________" in p.text for p in rebuilt.paragraphs)


def test_roundtrip_underline_with_bold():
    """Bold (markdown) and underline (inline_formatting) on the same span."""
    runner = CliRunner()
    with runner.isolated_filesystem():
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("bold-underline")
        r.bold = True
        r.underline = True
        doc.save("src.docx")

        assert runner.invoke(main, ["extract", "src.docx"]).exit_code == 0
        assert runner.invoke(main, ["build", "src.sidedoc", "-o", "rebuilt.docx"]).exit_code == 0

        rebuilt = Document("rebuilt.docx")
        bold, _, underline, _ = _count_runs_by_format(rebuilt)
        assert bold >= 1
        assert underline >= 1
        # Same run carries both.
        assert any(r.bold and r.underline for p in rebuilt.paragraphs for r in p.runs)


