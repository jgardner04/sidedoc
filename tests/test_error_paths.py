"""Test error handling and edge cases for sidedoc operations."""

import json
import tempfile
import zipfile
from pathlib import Path
from io import BytesIO
import pytest
from docx import Document
from sidedoc.extract import extract_blocks
from sidedoc import cli
from click.testing import CliRunner


# ============================================================================
# Corrupt ZIP File Tests
# ============================================================================

def test_extract_handles_corrupt_zip_file():
    """Test that extract command handles corrupt ZIP files gracefully."""
    # Create a file that's not a valid ZIP
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".sidedoc")
    temp_file.write(b"This is not a valid ZIP file content")
    temp_file.close()

    runner = CliRunner()
    result = runner.invoke(cli.unpack, [temp_file.name, "-o", tempfile.mkdtemp()])

    try:
        # Should fail with appropriate error message
        assert result.exit_code != 0
        assert "invalid" in result.output.lower() or "zip" in result.output.lower()
    finally:
        Path(temp_file.name).unlink()


def test_validate_handles_corrupt_zip_file():
    """Test that validate command handles corrupt ZIP files gracefully."""
    # Create a file that's not a valid ZIP
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".sidedoc")
    temp_file.write(b"Not a ZIP file at all!")
    temp_file.close()

    runner = CliRunner()
    result = runner.invoke(cli.validate, [temp_file.name])

    try:
        # Should fail with appropriate error message
        assert result.exit_code != 0
        assert "invalid" in result.output.lower() or "zip" in result.output.lower()
    finally:
        Path(temp_file.name).unlink()


def test_build_handles_corrupt_zip_file():
    """Test that build command handles corrupt ZIP files gracefully."""
    # Create a file that's not a valid ZIP
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".sidedoc")
    temp_file.write(b"Corrupted data")
    temp_file.close()

    runner = CliRunner()
    result = runner.invoke(cli.build, [temp_file.name])

    try:
        # Should fail with appropriate error message
        assert result.exit_code != 0
        assert "invalid" in result.output.lower() or "zip" in result.output.lower()
    finally:
        Path(temp_file.name).unlink()


def test_sync_handles_corrupt_zip_file():
    """Test that sync command handles corrupt ZIP files gracefully."""
    # Create a file that's not a valid ZIP
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".sidedoc")
    temp_file.write(b"Invalid ZIP data")
    temp_file.close()

    runner = CliRunner()
    result = runner.invoke(cli.sync, [temp_file.name])

    try:
        # Should fail with appropriate error message
        assert result.exit_code != 0
        assert "invalid" in result.output.lower() or "zip" in result.output.lower()
    finally:
        Path(temp_file.name).unlink()


# ============================================================================
# Invalid JSON Metadata Tests
# ============================================================================

def create_sidedoc_with_invalid_json(invalid_file: str) -> str:
    """Create a sidedoc file with invalid JSON in specified metadata file.

    Args:
        invalid_file: Name of the file to make invalid ('manifest.json', 'structure.json', or 'styles.json')

    Returns:
        Path to created sidedoc file
    """
    temp_sidedoc = tempfile.NamedTemporaryFile(delete=False, suffix=".sidedoc")
    temp_sidedoc.close()

    with zipfile.ZipFile(temp_sidedoc.name, 'w') as zf:
        # Add valid content.md
        zf.writestr("content.md", "# Test Document\n\nSome content.")

        # Add manifest.json (valid or invalid based on parameter)
        if invalid_file == "manifest.json":
            zf.writestr("manifest.json", "{invalid json syntax")
        else:
            manifest = {
                "version": "1.0",
                "created_at": "2024-01-01T00:00:00",
                "source_file": "test.docx"
            }
            zf.writestr("manifest.json", json.dumps(manifest))

        # Add structure.json (valid or invalid based on parameter)
        if invalid_file == "structure.json":
            zf.writestr("structure.json", "not valid json at all")
        else:
            structure = {"blocks": [{"id": "block-0", "type": "heading", "content_hash": "abc123"}]}
            zf.writestr("structure.json", json.dumps(structure))

        # Add styles.json (valid or invalid based on parameter)
        if invalid_file == "styles.json":
            zf.writestr("styles.json", "[this is, broken json")
        else:
            styles = {"block-0": {"font_size": 14, "bold": False}}
            zf.writestr("styles.json", json.dumps(styles))

    return temp_sidedoc.name


def test_validate_handles_invalid_manifest_json():
    """Test that validate command handles invalid manifest.json gracefully."""
    sidedoc_path = create_sidedoc_with_invalid_json("manifest.json")

    runner = CliRunner()
    result = runner.invoke(cli.validate, [sidedoc_path])

    try:
        # Should fail with appropriate error message
        assert result.exit_code != 0
        # Error message should indicate JSON parsing issue
        assert any(word in result.output.lower() for word in ["json", "decode", "expecting", "value"])
    finally:
        Path(sidedoc_path).unlink()


def test_validate_handles_invalid_structure_json():
    """Test that validate command handles invalid structure.json gracefully."""
    sidedoc_path = create_sidedoc_with_invalid_json("structure.json")

    runner = CliRunner()
    result = runner.invoke(cli.validate, [sidedoc_path])

    try:
        # Should fail with appropriate error message
        assert result.exit_code != 0
        # Error message should indicate JSON parsing issue
        assert any(word in result.output.lower() for word in ["json", "decode", "expecting", "value"])
    finally:
        Path(sidedoc_path).unlink()


def test_validate_handles_invalid_styles_json():
    """Test that validate command handles invalid styles.json gracefully."""
    sidedoc_path = create_sidedoc_with_invalid_json("styles.json")

    runner = CliRunner()
    result = runner.invoke(cli.validate, [sidedoc_path])

    try:
        # Should fail with appropriate error message
        assert result.exit_code != 0
        # Error message should indicate JSON parsing issue
        assert any(word in result.output.lower() for word in ["json", "decode", "expecting", "value"])
    finally:
        Path(sidedoc_path).unlink()


def test_build_handles_invalid_structure_json():
    """Test that build command handles invalid structure.json gracefully."""
    sidedoc_path = create_sidedoc_with_invalid_json("structure.json")

    runner = CliRunner()
    result = runner.invoke(cli.build, [sidedoc_path])

    try:
        # Should fail with appropriate error message
        assert result.exit_code != 0
        # Error message should indicate JSON parsing issue
        assert any(word in result.output.lower() for word in ["json", "decode", "expecting", "value"])
    finally:
        Path(sidedoc_path).unlink()


def test_sync_handles_invalid_structure_json():
    """Test that sync command handles invalid structure.json gracefully."""
    sidedoc_path = create_sidedoc_with_invalid_json("structure.json")

    runner = CliRunner()
    result = runner.invoke(cli.sync, [sidedoc_path])

    try:
        # Should fail with appropriate error message
        assert result.exit_code != 0
        # Error message should indicate JSON parsing issue
        assert any(word in result.output.lower() for word in ["json", "decode", "expecting", "value"])
    finally:
        Path(sidedoc_path).unlink()


# ============================================================================
# Missing Required Files Tests
# ============================================================================

def create_incomplete_sidedoc(missing_file: str) -> str:
    """Create a sidedoc file missing a required file.

    Args:
        missing_file: Name of the file to omit ('content.md', 'manifest.json', 'structure.json', or 'styles.json')

    Returns:
        Path to created sidedoc file
    """
    temp_sidedoc = tempfile.NamedTemporaryFile(delete=False, suffix=".sidedoc")
    temp_sidedoc.close()

    with zipfile.ZipFile(temp_sidedoc.name, 'w') as zf:
        if missing_file != "content.md":
            zf.writestr("content.md", "# Test Document\n\nSome content.")

        if missing_file != "manifest.json":
            manifest = {
                "version": "1.0",
                "created_at": "2024-01-01T00:00:00",
                "source_file": "test.docx"
            }
            zf.writestr("manifest.json", json.dumps(manifest))

        if missing_file != "structure.json":
            structure = {"blocks": [{"id": "block-0", "type": "heading", "content_hash": "abc123"}]}
            zf.writestr("structure.json", json.dumps(structure))

        if missing_file != "styles.json":
            styles = {"block-0": {"font_size": 14, "bold": False}}
            zf.writestr("styles.json", json.dumps(styles))

    return temp_sidedoc.name


def test_validate_handles_missing_content_md():
    """Test that validate command handles missing content.md gracefully."""
    sidedoc_path = create_incomplete_sidedoc("content.md")

    runner = CliRunner()
    result = runner.invoke(cli.validate, [sidedoc_path])

    try:
        # Should fail with appropriate error message
        assert result.exit_code != 0
        assert "content.md" in result.output.lower() or "missing" in result.output.lower()
    finally:
        Path(sidedoc_path).unlink()


def test_validate_handles_missing_manifest():
    """Test that validate command handles missing manifest.json gracefully."""
    sidedoc_path = create_incomplete_sidedoc("manifest.json")

    runner = CliRunner()
    result = runner.invoke(cli.validate, [sidedoc_path])

    try:
        # Should fail with appropriate error message
        assert result.exit_code != 0
        assert "manifest" in result.output.lower() or "missing" in result.output.lower()
    finally:
        Path(sidedoc_path).unlink()


def test_validate_handles_missing_structure():
    """Test that validate command handles missing structure.json gracefully."""
    sidedoc_path = create_incomplete_sidedoc("structure.json")

    runner = CliRunner()
    result = runner.invoke(cli.validate, [sidedoc_path])

    try:
        # Should fail with appropriate error message
        assert result.exit_code != 0
        assert "structure" in result.output.lower() or "missing" in result.output.lower()
    finally:
        Path(sidedoc_path).unlink()


def test_validate_handles_missing_styles():
    """Test that validate command handles missing styles.json gracefully."""
    sidedoc_path = create_incomplete_sidedoc("styles.json")

    runner = CliRunner()
    result = runner.invoke(cli.validate, [sidedoc_path])

    try:
        # Should fail with appropriate error message
        assert result.exit_code != 0
        assert "styles" in result.output.lower() or "missing" in result.output.lower()
    finally:
        Path(sidedoc_path).unlink()


# ============================================================================
# Large Document Performance Tests
# ============================================================================

def create_large_docx(num_paragraphs: int = 1000) -> str:
    """Create a large docx file for performance testing.

    Args:
        num_paragraphs: Number of paragraphs to create

    Returns:
        Path to temporary docx file
    """
    doc = Document()

    # Add a variety of content types
    doc.add_heading("Large Document Test", level=1)

    for i in range(num_paragraphs):
        if i % 100 == 0:
            # Add headings every 100 paragraphs
            doc.add_heading(f"Section {i // 100}", level=2)
        elif i % 50 == 0:
            # Add bullet list every 50 paragraphs
            doc.add_paragraph(f"List item {i}", style="List Bullet")
        else:
            # Regular paragraph with varied content
            doc.add_paragraph(f"This is paragraph {i} with some content that varies in length. " * (i % 3 + 1))

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    temp_file.close()

    return temp_file.name


def test_extract_large_document_performance():
    """Test that extract can handle large documents (1000+ paragraphs) efficiently."""
    docx_path = create_large_docx(1000)

    try:
        # This should complete without hanging or running out of memory
        import time
        start_time = time.time()

        blocks, _ = extract_blocks(docx_path)

        elapsed_time = time.time() - start_time

        # Should extract all blocks
        assert len(blocks) > 1000

        # Should complete in reasonable time (less than 30 seconds for 1000 paragraphs)
        assert elapsed_time < 30, f"Extract took too long: {elapsed_time:.2f} seconds"

        # Verify blocks have required fields
        for block in blocks[:10]:  # Check first 10 blocks
            assert block.id is not None
            assert block.content is not None
            assert block.content_hash is not None
    finally:
        Path(docx_path).unlink()


def test_roundtrip_large_document():
    """Test complete roundtrip (extract -> build) for large document."""
    docx_path = create_large_docx(500)  # Use smaller size for full roundtrip

    try:
        # Extract to sidedoc
        temp_sidedoc = tempfile.NamedTemporaryFile(delete=False, suffix=".sidedoc")
        temp_sidedoc.close()

        runner = CliRunner()
        result = runner.invoke(cli.extract, [docx_path, "-o", temp_sidedoc.name])
        assert result.exit_code == 0

        # Build back to docx
        temp_output = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        temp_output.close()

        result = runner.invoke(cli.build, [temp_sidedoc.name, "-o", temp_output.name])
        assert result.exit_code == 0

        # Verify output exists and is valid
        assert Path(temp_output.name).exists()

        # Should be able to open rebuilt document
        rebuilt_doc = Document(temp_output.name)
        assert len(rebuilt_doc.paragraphs) > 500

        # Cleanup
        Path(temp_sidedoc.name).unlink()
        Path(temp_output.name).unlink()
    finally:
        Path(docx_path).unlink()


# ============================================================================
# Multiple Image Format Tests
# ============================================================================

def create_minimal_jpg() -> bytes:
    """Create a minimal valid JPEG image.

    Returns:
        JPEG file bytes
    """
    # Minimal JPEG: SOI marker + minimal frame + EOI marker
    # This is a 1x1 black JPEG
    return bytes([
        0xFF, 0xD8,  # SOI (Start of Image)
        0xFF, 0xE0,  # APP0
        0x00, 0x10,  # APP0 length
        0x4A, 0x46, 0x49, 0x46, 0x00,  # JFIF
        0x01, 0x01,  # Version
        0x00,  # No units
        0x00, 0x01, 0x00, 0x01,  # X/Y density
        0x00, 0x00,  # No thumbnail
        0xFF, 0xDB,  # DQT (quantization table)
        0x00, 0x43,  # Length
        0x00,  # Table 0
    ] + [1] * 64 +  # Quantization values
    [
        0xFF, 0xC0,  # SOF0 (Start of Frame)
        0x00, 0x0B,  # Length
        0x08,  # 8-bit precision
        0x00, 0x01, 0x00, 0x01,  # 1x1 dimensions
        0x01,  # 1 component
        0x01, 0x11, 0x00,  # Component info
        0xFF, 0xC4,  # DHT (Huffman table)
        0x00, 0x1F,  # Length
        0x00,  # Table 0
    ] + [0] * 28 +  # Huffman values
    [
        0xFF, 0xDA,  # SOS (Start of Scan)
        0x00, 0x08,  # Length
        0x01,  # 1 component
        0x01, 0x00,  # Component info
        0x00, 0x3F, 0x00,  # Spectral selection
        0xFF, 0xD9  # EOI (End of Image)
    ])


def create_minimal_gif() -> bytes:
    """Create a minimal valid GIF image.

    Returns:
        GIF file bytes
    """
    # Minimal GIF87a: 1x1 black image
    return bytes([
        # Header
        0x47, 0x49, 0x46, 0x38, 0x37, 0x61,  # "GIF87a"
        # Logical Screen Descriptor
        0x01, 0x00,  # Width = 1
        0x01, 0x00,  # Height = 1
        0x80,  # Global color table flag
        0x00,  # Background color index
        0x00,  # Pixel aspect ratio
        # Global Color Table (2 colors)
        0x00, 0x00, 0x00,  # Black
        0xFF, 0xFF, 0xFF,  # White
        # Image Descriptor
        0x2C,  # Image separator
        0x00, 0x00, 0x00, 0x00,  # Left, Top
        0x01, 0x00, 0x01, 0x00,  # Width, Height
        0x00,  # No local color table
        # Image Data
        0x02,  # LZW minimum code size
        0x02,  # Data sub-block length
        0x4C, 0x01,  # Compressed data
        0x00,  # Block terminator
        # Trailer
        0x3B  # GIF trailer
    ])


def test_extract_document_with_jpg_image():
    """Test extracting a document with a JPEG image."""
    # Create temp JPG file
    temp_jpg = tempfile.NamedTemporaryFile(delete=False, suffix='.jpg')
    temp_jpg.write(create_minimal_jpg())
    temp_jpg.close()

    # Create docx with JPG image
    doc = Document()
    doc.add_paragraph('Before JPG image')
    doc.add_picture(temp_jpg.name, width=1)
    doc.add_paragraph('After JPG image')

    temp_doc = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_doc.name)
    temp_doc.close()

    try:
        blocks, _ = extract_blocks(temp_doc.name)

        # Should have 3 blocks: paragraph, image, paragraph
        assert len(blocks) == 3
        assert blocks[1].type == "image"
        assert ".jp" in blocks[1].content.lower()  # Should reference JPG/JPEG
        assert blocks[1].image_path is not None
    finally:
        Path(temp_doc.name).unlink()
        Path(temp_jpg.name).unlink()


def test_extract_document_with_gif_image():
    """Test extracting a document with a GIF image."""
    # Create temp GIF file
    temp_gif = tempfile.NamedTemporaryFile(delete=False, suffix='.gif')
    temp_gif.write(create_minimal_gif())
    temp_gif.close()

    # Create docx with GIF image
    doc = Document()
    doc.add_paragraph('Before GIF image')
    doc.add_picture(temp_gif.name, width=1)
    doc.add_paragraph('After GIF image')

    temp_doc = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_doc.name)
    temp_doc.close()

    try:
        blocks, _ = extract_blocks(temp_doc.name)

        # Should have 3 blocks: paragraph, image, paragraph
        assert len(blocks) == 3
        assert blocks[1].type == "image"
        assert ".gif" in blocks[1].content.lower()
        assert blocks[1].image_path is not None
    finally:
        Path(temp_doc.name).unlink()
        Path(temp_gif.name).unlink()


def test_extract_document_with_mixed_image_formats():
    """Test extracting a document with multiple different image formats."""
    # Create temp image files
    temp_jpg = tempfile.NamedTemporaryFile(delete=False, suffix='.jpg')
    temp_jpg.write(create_minimal_jpg())
    temp_jpg.close()

    temp_gif = tempfile.NamedTemporaryFile(delete=False, suffix='.gif')
    temp_gif.write(create_minimal_gif())
    temp_gif.close()

    # Create docx with multiple images
    doc = Document()
    doc.add_paragraph('JPG image:')
    doc.add_picture(temp_jpg.name, width=1)
    doc.add_paragraph('GIF image:')
    doc.add_picture(temp_gif.name, width=1)

    temp_doc = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_doc.name)
    temp_doc.close()

    try:
        blocks, _ = extract_blocks(temp_doc.name)

        # Should have 4 blocks
        assert len(blocks) == 4

        # Find image blocks
        image_blocks = [b for b in blocks if b.type == "image"]
        assert len(image_blocks) == 2

        # Images should have different formats
        formats = [b.image_path.split('.')[-1] if b.image_path else '' for b in image_blocks]
        assert len(set(formats)) > 1  # At least 2 different formats
    finally:
        Path(temp_doc.name).unlink()
        Path(temp_jpg.name).unlink()
        Path(temp_gif.name).unlink()


def test_extract_document_with_many_images():
    """Test extracting a document with 10+ images."""
    # Create a single temp image we can reuse
    temp_jpg = tempfile.NamedTemporaryFile(delete=False, suffix='.jpg')
    temp_jpg.write(create_minimal_jpg())
    temp_jpg.close()

    # Create docx with many images
    doc = Document()
    for i in range(12):
        doc.add_paragraph(f'Image {i + 1}:')
        doc.add_picture(temp_jpg.name, width=1)

    temp_doc = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_doc.name)
    temp_doc.close()

    try:
        blocks, _ = extract_blocks(temp_doc.name)

        # Should have 24 blocks (12 paragraphs + 12 images)
        assert len(blocks) == 24

        # Count image blocks
        image_blocks = [b for b in blocks if b.type == "image"]
        assert len(image_blocks) == 12

        # Each image should have unique identifier
        image_paths = [b.image_path for b in image_blocks if b.image_path]
        assert len(image_paths) == len(set(image_paths))  # All unique
    finally:
        Path(temp_doc.name).unlink()
        Path(temp_jpg.name).unlink()


# ============================================================================
# Disk Full / Write Error Tests
# ============================================================================

def test_extract_handles_write_errors():
    """Test that extract command handles write errors gracefully."""
    # Create a valid docx
    doc = Document()
    doc.add_paragraph("Test content")
    temp_doc = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_doc.name)
    temp_doc.close()

    try:
        # Try to write to an invalid/inaccessible location
        runner = CliRunner()
        result = runner.invoke(cli.extract, [temp_doc.name, "-o", "/nonexistent/path/output.sidedoc"])

        # Should fail with appropriate error message
        assert result.exit_code != 0
        assert "error" in result.output.lower() or "permission" in result.output.lower() or "not found" in result.output.lower()
    finally:
        Path(temp_doc.name).unlink()


def test_build_handles_write_errors():
    """Test that build command handles write errors gracefully."""
    # Create a minimal valid sidedoc
    temp_sidedoc = tempfile.NamedTemporaryFile(delete=False, suffix=".sidedoc")
    temp_sidedoc.close()

    with zipfile.ZipFile(temp_sidedoc.name, 'w') as zf:
        zf.writestr("content.md", "# Test\n\nContent.")
        manifest = {"version": "1.0", "created_at": "2024-01-01T00:00:00", "source_file": "test.docx"}
        zf.writestr("manifest.json", json.dumps(manifest))
        structure = {"blocks": [{"id": "block-0", "type": "heading", "content_hash": "abc123"}]}
        zf.writestr("structure.json", json.dumps(structure))
        styles = {"block-0": {"font_size": 14}}
        zf.writestr("styles.json", json.dumps(styles))

    try:
        # Try to write to an invalid/inaccessible location
        runner = CliRunner()
        result = runner.invoke(cli.build, [temp_sidedoc.name, "-o", "/nonexistent/path/output.docx"])

        # Should fail with appropriate error message
        assert result.exit_code != 0
        assert "error" in result.output.lower() or "permission" in result.output.lower() or "not found" in result.output.lower()
    finally:
        Path(temp_sidedoc.name).unlink()
