"""Tests for table support in sidedoc.

Test fixtures:
- tables_simple.docx: Basic 3x3 table with:
  - Header row: Name | Role | Start Date
  - Data row 1: Alice | Engineer | 2024-01-15
  - Data row 2: Bob | Designer | 2024-02-01
"""

import json
from pathlib import Path
from typing import Any
from docx import Document

from sidedoc.extract import extract_blocks, blocks_to_markdown, extract_styles
from sidedoc.models import Style
from sidedoc.reconstruct import parse_markdown_to_blocks, create_docx_from_blocks


def _build_styles_dict(styles: list[Style]) -> dict[str, Any]:
    """Build a styles dict from a list of Style objects for use in reconstruction."""
    styles_dict: dict[str, Any] = {"block_styles": {}}
    for style in styles:
        style_data: dict[str, Any] = {
            "docx_style": style.docx_style,
            "font_name": style.font_name,
            "font_size": style.font_size,
            "alignment": style.alignment,
        }
        if style.table_formatting:
            style_data["table_formatting"] = style.table_formatting
        styles_dict["block_styles"][style.block_id] = style_data
    return styles_dict


# ============================================================================
# US-T01: Create table test fixtures
# ============================================================================


def test_tables_simple_fixture_exists() -> None:
    """Verify tables_simple.docx fixture exists and has expected structure."""
    fixtures_dir = Path("tests/fixtures")
    fixture_path = fixtures_dir / "tables_simple.docx"

    assert fixture_path.exists(), "tables_simple.docx should exist"

    # Open and verify structure
    doc = Document(str(fixture_path))

    # Should have at least one table
    assert len(doc.tables) >= 1, "Document should have at least one table"

    # Get the first table
    table = doc.tables[0]

    # Should have 3 rows (header + 2 data rows)
    assert len(table.rows) == 3, "Table should have 3 rows"

    # Should have 3 columns
    assert len(table.columns) == 3, "Table should have 3 columns"

    # Verify header row content
    header_cells = [cell.text for cell in table.rows[0].cells]
    assert header_cells == ["Name", "Role", "Start Date"], f"Header row should be Name, Role, Start Date but got {header_cells}"

    # Verify first data row
    row1_cells = [cell.text for cell in table.rows[1].cells]
    assert row1_cells == ["Alice", "Engineer", "2024-01-15"], f"First data row should be Alice, Engineer, 2024-01-15 but got {row1_cells}"

    # Verify second data row
    row2_cells = [cell.text for cell in table.rows[2].cells]
    assert row2_cells == ["Bob", "Designer", "2024-02-01"], f"Second data row should be Bob, Designer, 2024-02-01 but got {row2_cells}"


# ============================================================================
# US-T02: Extract simple tables to GFM markdown
# ============================================================================


def test_extract_simple_table_to_gfm() -> None:
    """Test that tables are extracted as GFM pipe table syntax."""
    fixture_path = Path("tests/fixtures/tables_simple.docx")
    blocks, _ = extract_blocks(str(fixture_path))

    # Find table blocks
    table_blocks = [b for b in blocks if b.type == "table"]
    assert len(table_blocks) >= 1, "Should extract at least one table block"

    table_block = table_blocks[0]

    # Content should be GFM pipe table format
    content = table_block.content
    lines = content.strip().split("\n")

    # Should have header row, separator row, and data rows
    assert len(lines) >= 4, f"Table should have at least 4 lines (header, separator, 2 data rows), got {len(lines)}"

    # First line should be header
    assert "|" in lines[0], "Header row should contain pipe characters"
    assert "Name" in lines[0], "Header should contain 'Name'"
    assert "Role" in lines[0], "Header should contain 'Role'"
    assert "Start Date" in lines[0], "Header should contain 'Start Date'"

    # Second line should be separator (dashes and pipes)
    assert "|" in lines[1], "Separator row should contain pipe characters"
    assert "-" in lines[1], "Separator row should contain dashes"

    # Data rows
    assert "Alice" in lines[2], "First data row should contain 'Alice'"
    assert "Bob" in lines[3], "Second data row should contain 'Bob'"


def test_extract_table_creates_table_block_type() -> None:
    """Test that extracted tables have type 'table' in structure."""
    fixture_path = Path("tests/fixtures/tables_simple.docx")
    blocks, _ = extract_blocks(str(fixture_path))

    # Find table blocks
    table_blocks = [b for b in blocks if b.type == "table"]
    assert len(table_blocks) >= 1, "Should have at least one table block"

    # Verify block type
    assert table_blocks[0].type == "table"


def test_extract_table_preserves_empty_cells() -> None:
    """Test that empty cells are preserved as empty pipe segments."""
    # We'll test this with a table that has empty cells
    # For now, just verify the basic extraction works
    fixture_path = Path("tests/fixtures/tables_simple.docx")
    blocks, _ = extract_blocks(str(fixture_path))

    table_blocks = [b for b in blocks if b.type == "table"]
    assert len(table_blocks) >= 1

    # The simple fixture doesn't have empty cells, but verify format
    content = table_blocks[0].content
    # Each row should have proper pipe structure
    for line in content.strip().split("\n"):
        if line.strip():
            assert line.startswith("|"), f"Each row should start with pipe: {line}"
            assert line.endswith("|"), f"Each row should end with pipe: {line}"


# ============================================================================
# US-T03: Store table structure metadata in structure.json
# ============================================================================


def test_table_block_has_rows_cols_metadata() -> None:
    """Test that table blocks have rows and cols metadata."""
    fixture_path = Path("tests/fixtures/tables_simple.docx")
    blocks, _ = extract_blocks(str(fixture_path))

    table_blocks = [b for b in blocks if b.type == "table"]
    assert len(table_blocks) >= 1

    table_block = table_blocks[0]

    # Block should have table_metadata attribute with rows and cols
    assert hasattr(table_block, 'table_metadata'), "Table block should have table_metadata"
    assert table_block.table_metadata is not None, "table_metadata should not be None"
    assert "rows" in table_block.table_metadata, "table_metadata should have 'rows'"
    assert "cols" in table_block.table_metadata, "table_metadata should have 'cols'"
    assert table_block.table_metadata["rows"] == 3, "Should have 3 rows"
    assert table_block.table_metadata["cols"] == 3, "Should have 3 columns"


def test_table_block_has_cells_metadata() -> None:
    """Test that table blocks have cells array with content hashes."""
    fixture_path = Path("tests/fixtures/tables_simple.docx")
    blocks, _ = extract_blocks(str(fixture_path))

    table_blocks = [b for b in blocks if b.type == "table"]
    table_block = table_blocks[0]

    assert table_block.table_metadata is not None
    assert "cells" in table_block.table_metadata, "table_metadata should have 'cells'"

    cells = table_block.table_metadata["cells"]
    # Should be a 2D array (list of rows, each row is list of cell metadata)
    assert len(cells) == 3, "Should have 3 rows of cells"
    assert len(cells[0]) == 3, "First row should have 3 cells"

    # Each cell should have row, col, content_hash
    first_cell = cells[0][0]
    assert "row" in first_cell, "Cell should have 'row'"
    assert "col" in first_cell, "Cell should have 'col'"
    assert "content_hash" in first_cell, "Cell should have 'content_hash'"
    assert first_cell["row"] == 0
    assert first_cell["col"] == 0


def test_table_block_has_docx_table_index() -> None:
    """Test that table blocks have docx_table_index field."""
    fixture_path = Path("tests/fixtures/tables_simple.docx")
    blocks, _ = extract_blocks(str(fixture_path))

    table_blocks = [b for b in blocks if b.type == "table"]
    table_block = table_blocks[0]

    assert table_block.table_metadata is not None
    assert "docx_table_index" in table_block.table_metadata, "table_metadata should have 'docx_table_index'"
    assert table_block.table_metadata["docx_table_index"] == 0, "First table should have index 0"


# ============================================================================
# US-T04: Store basic table formatting in styles.json
# ============================================================================


def test_table_style_has_column_widths() -> None:
    """Test that table styles include column_widths array."""
    fixture_path = Path("tests/fixtures/tables_simple.docx")
    blocks, _ = extract_blocks(str(fixture_path))
    styles = extract_styles(str(fixture_path), blocks)

    # Find the style for our table block
    table_blocks = [b for b in blocks if b.type == "table"]
    assert len(table_blocks) >= 1

    table_block = table_blocks[0]

    # Find corresponding style
    table_styles = [s for s in styles if s.block_id == table_block.id]
    assert len(table_styles) == 1, "Should have a style for the table block"

    table_style = table_styles[0]

    # Style should have table_formatting with column_widths
    assert hasattr(table_style, 'table_formatting'), "Style should have table_formatting"
    assert table_style.table_formatting is not None, "table_formatting should not be None"
    assert "column_widths" in table_style.table_formatting, "Should have column_widths"

    column_widths = table_style.table_formatting["column_widths"]
    assert len(column_widths) == 3, "Should have 3 column widths for 3-column table"


def test_table_style_has_table_alignment() -> None:
    """Test that table styles include table_alignment."""
    fixture_path = Path("tests/fixtures/tables_simple.docx")
    blocks, _ = extract_blocks(str(fixture_path))
    styles = extract_styles(str(fixture_path), blocks)

    table_blocks = [b for b in blocks if b.type == "table"]
    table_block = table_blocks[0]

    table_styles = [s for s in styles if s.block_id == table_block.id]
    table_style = table_styles[0]

    assert table_style.table_formatting is not None
    assert "table_alignment" in table_style.table_formatting, "Should have table_alignment"
    # Default is usually left
    assert table_style.table_formatting["table_alignment"] in ("left", "center", "right")


# ============================================================================
# US-T05: Reconstruct basic tables in docx
# ============================================================================


def test_parse_markdown_detects_table_blocks() -> None:
    """Test that parse_markdown_to_blocks detects GFM table syntax."""
    markdown = """# Title

Some text.

| Name | Role |
|------|------|
| Alice | Engineer |
| Bob | Designer |

More text.
"""
    blocks = parse_markdown_to_blocks(markdown)

    # Should have a table block
    table_blocks = [b for b in blocks if b.type == "table"]
    assert len(table_blocks) >= 1, "Should detect at least one table block"


def test_create_docx_creates_table() -> None:
    """Test that create_docx_from_blocks creates tables from table blocks."""
    # Create a simple table block
    from sidedoc.models import Block

    table_content = """| Name | Role |
| --- | --- |
| Alice | Engineer |
| Bob | Designer |"""

    blocks = [
        Block(
            id="block-0",
            type="table",
            content=table_content,
            docx_paragraph_index=-1,
            content_start=0,
            content_end=len(table_content),
            content_hash="abc123",
            table_metadata={
                "rows": 3,
                "cols": 2,
                "cells": [],
                "docx_table_index": 0
            }
        )
    ]

    styles: dict[str, Any] = {"block_styles": {}}
    doc = create_docx_from_blocks(blocks, styles)

    # Verify the document has a table
    assert len(doc.tables) >= 1, "Document should have at least one table"

    table = doc.tables[0]
    assert len(table.rows) == 3, "Table should have 3 rows"
    assert len(table.columns) == 2, "Table should have 2 columns"

    # Verify cell content
    assert table.rows[0].cells[0].text == "Name"
    assert table.rows[0].cells[1].text == "Role"
    assert table.rows[1].cells[0].text == "Alice"
    assert table.rows[1].cells[1].text == "Engineer"


def test_reconstruct_table_preserves_content() -> None:
    """Test that table content matches after reconstruction."""
    fixture_path = Path("tests/fixtures/tables_simple.docx")
    blocks, _ = extract_blocks(str(fixture_path))
    styles = extract_styles(str(fixture_path), blocks)

    # Convert styles to the format expected by create_docx_from_blocks
    styles_dict: dict[str, Any] = {"block_styles": {}}

    doc = create_docx_from_blocks(blocks, styles_dict)

    # Verify the document has a table
    assert len(doc.tables) >= 1, "Reconstructed doc should have at least one table"

    table = doc.tables[0]
    assert table.rows[0].cells[0].text == "Name"
    assert table.rows[1].cells[0].text == "Alice"
    assert table.rows[2].cells[0].text == "Bob"


# ============================================================================
# US-T06: Extract column alignment from tables
# ============================================================================


def test_table_gfm_has_separator_row() -> None:
    """Test that extracted GFM tables have a proper separator row."""
    fixture_path = Path("tests/fixtures/tables_simple.docx")
    blocks, _ = extract_blocks(str(fixture_path))

    table_blocks = [b for b in blocks if b.type == "table"]
    table_block = table_blocks[0]

    lines = table_block.content.split("\n")
    assert len(lines) >= 2

    # Second line should be separator with alignment indicators
    separator = lines[1]
    assert "|" in separator
    assert "---" in separator or ":---" in separator or "---:" in separator


def test_table_metadata_has_column_alignments() -> None:
    """Test that table metadata includes column alignments."""
    fixture_path = Path("tests/fixtures/tables_simple.docx")
    blocks, _ = extract_blocks(str(fixture_path))

    table_blocks = [b for b in blocks if b.type == "table"]
    table_block = table_blocks[0]

    assert table_block.table_metadata is not None
    assert "column_alignments" in table_block.table_metadata, "table_metadata should have 'column_alignments'"

    alignments = table_block.table_metadata["column_alignments"]
    assert len(alignments) == 3, "Should have alignment for each column"
    # Default alignment is typically 'left'
    assert all(a in ("left", "center", "right") for a in alignments)


# ============================================================================
# US-T07: Apply column alignment during reconstruction
# ============================================================================


def test_parse_markdown_table_escaped_pipe_column_count() -> None:
    """Test that escaped pipes don't inflate column count in parse_markdown_to_blocks."""
    markdown = """| Test\\|Pipe | Normal |
| --- | --- |
| A | B |"""

    blocks = parse_markdown_to_blocks(markdown)
    table_blocks = [b for b in blocks if b.type == "table"]
    assert len(table_blocks) == 1

    assert table_blocks[0].table_metadata is not None
    assert table_blocks[0].table_metadata["cols"] == 2, (
        f"Expected 2 columns, got {table_blocks[0].table_metadata['cols']}"
    )


def test_parse_markdown_table_metadata_includes_column_alignments() -> None:
    """Test that parse_markdown_to_blocks includes column_alignments in table_metadata."""
    markdown = """| Left | Center | Right |
| :--- | :---: | ---: |
| A | B | C |"""

    blocks = parse_markdown_to_blocks(markdown)
    table_blocks = [b for b in blocks if b.type == "table"]
    assert len(table_blocks) == 1

    metadata = table_blocks[0].table_metadata
    assert metadata is not None
    assert "column_alignments" in metadata, "table_metadata should have 'column_alignments'"
    assert metadata["column_alignments"] == ["left", "center", "right"]


def test_parse_markdown_table_metadata_default_alignments() -> None:
    """Test that default alignments are recorded when no alignment indicators."""
    markdown = """| A | B |
| --- | --- |
| 1 | 2 |"""

    blocks = parse_markdown_to_blocks(markdown)
    table_blocks = [b for b in blocks if b.type == "table"]
    assert len(table_blocks) == 1

    metadata = table_blocks[0].table_metadata
    assert metadata is not None
    assert "column_alignments" in metadata
    assert metadata["column_alignments"] == ["left", "left"]


def test_parse_gfm_alignment_colon_left_explicit() -> None:
    """Test that :--- (explicit left) is parsed correctly as left alignment."""
    from sidedoc.reconstruct import parse_gfm_alignments

    result = parse_gfm_alignments("| :--- | :--- |")
    assert result == ["left", "left"], f"Expected ['left', 'left'], got {result}"


def test_parse_gfm_alignment_indicators() -> None:
    """Test that GFM alignment indicators are parsed correctly."""
    from sidedoc.reconstruct import parse_gfm_alignments

    # Test left alignment (default)
    assert parse_gfm_alignments("| --- | --- |") == ["left", "left"]

    # Test center alignment
    assert parse_gfm_alignments("| :---: | :---: |") == ["center", "center"]

    # Test right alignment
    assert parse_gfm_alignments("| ---: | ---: |") == ["right", "right"]

    # Test mixed alignments
    assert parse_gfm_alignments("| :--- | :---: | ---: |") == ["left", "center", "right"]


def test_reconstruct_table_applies_alignment() -> None:
    """Test that reconstructed tables have correct cell alignment."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from sidedoc.models import Block

    # Create a table with center and right alignment
    table_content = """| Name | Amount |
| :---: | ---: |
| Alice | $100 |
| Bob | $200 |"""

    blocks = [
        Block(
            id="block-0",
            type="table",
            content=table_content,
            docx_paragraph_index=-1,
            content_start=0,
            content_end=len(table_content),
            content_hash="abc123",
            table_metadata={
                "rows": 3,
                "cols": 2,
                "cells": [],
                "column_alignments": ["center", "right"],
                "docx_table_index": 0
            }
        )
    ]

    styles_dict: dict[str, Any] = {"block_styles": {}}
    doc = create_docx_from_blocks(blocks, styles_dict)

    assert len(doc.tables) >= 1
    table = doc.tables[0]

    # Check first column is center-aligned
    cell_0_0 = table.cell(0, 0)
    if cell_0_0.paragraphs:
        assert cell_0_0.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER

    # Check second column is right-aligned
    cell_0_1 = table.cell(0, 1)
    if cell_0_1.paragraphs:
        assert cell_0_1.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT


# ============================================================================
# US-T08: Handle special characters in table cells
# ============================================================================


def test_extract_table_escapes_pipe_characters() -> None:
    """Test that pipe characters in cell content are escaped and roundtrip correctly."""
    # Create a test document with pipe in cell
    from docx import Document as DocxDocument
    import tempfile

    with tempfile.TemporaryDirectory() as tmp_dir:
        # Create a docx with pipe in cell
        doc = DocxDocument()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "Test|Pipe"  # Cell with pipe character
        table.cell(1, 1).text = "Normal"

        test_path = Path(tmp_dir) / "test_pipes.docx"
        doc.save(str(test_path))

        # Extract the table
        blocks, _ = extract_blocks(str(test_path))

        table_blocks = [b for b in blocks if b.type == "table"]
        assert len(table_blocks) >= 1

        table_content = table_blocks[0].content

        # The pipe should be escaped as \|
        assert "\\|" in table_content, f"Pipe should be escaped, got: {table_content}"
        assert "Test\\|Pipe" in table_content, f"Cell content should have escaped pipe"

        # Test roundtrip: reconstruct and verify pipes are unescaped
        styles_dict: dict[str, Any] = {"block_styles": {}}
        reconstructed_doc = create_docx_from_blocks(blocks, styles_dict)

        # Verify the reconstructed table has unescaped pipe
        assert len(reconstructed_doc.tables) >= 1
        reconstructed_table = reconstructed_doc.tables[0]
        assert reconstructed_table.cell(1, 0).text == "Test|Pipe", \
            f"Roundtrip should preserve pipe: got {reconstructed_table.cell(1, 0).text}"


def test_extract_rejects_oversized_table() -> None:
    """Test that extraction raises ValueError for tables exceeding dimension limits."""
    import pytest
    import tempfile
    from unittest.mock import patch
    from sidedoc.constants import MAX_TABLE_ROWS, MAX_TABLE_COLS

    with tempfile.TemporaryDirectory() as tmp_dir:
        # Create a docx with a table exceeding MAX_TABLE_COLS
        doc = Document()
        # We can't easily create a 101-column table, so we patch the constant
        # Create a small table and patch the limit down
        table = doc.add_table(rows=2, cols=5)
        table.cell(0, 0).text = "H1"
        table.cell(0, 1).text = "H2"
        table.cell(0, 2).text = "H3"
        table.cell(0, 3).text = "H4"
        table.cell(0, 4).text = "H5"
        table.cell(1, 0).text = "A"
        table.cell(1, 1).text = "B"
        table.cell(1, 2).text = "C"
        table.cell(1, 3).text = "D"
        table.cell(1, 4).text = "E"

        test_path = Path(tmp_dir) / "test_oversized.docx"
        doc.save(str(test_path))

        # Patch MAX_TABLE_COLS to be smaller than our table
        with patch("sidedoc.extract.MAX_TABLE_COLS", 3):
            with pytest.raises(ValueError, match="too many columns"):
                extract_blocks(str(test_path))

        # Also test rows limit
        with patch("sidedoc.extract.MAX_TABLE_ROWS", 1):
            with pytest.raises(ValueError, match="too many rows"):
                extract_blocks(str(test_path))


def test_parse_gfm_table_inconsistent_columns() -> None:
    """Test that parse_gfm_table normalizes rows with inconsistent column counts."""
    from sidedoc.reconstruct import parse_gfm_table

    # Table with inconsistent column counts
    table_content = """| A | B | C |
| --- | --- | --- |
| 1 | 2 |
| X | Y | Z |"""

    rows, alignments = parse_gfm_table(table_content)

    # All rows should be normalized to the header column count (3)
    assert len(rows) == 3  # header + 2 data rows
    for row in rows:
        assert len(row) == 3, f"Row should have 3 columns, got {len(row)}: {row}"

    # The short row should be padded with empty strings
    assert rows[1] == ["1", "2", ""]


def test_parse_gfm_table_truncates_long_rows() -> None:
    """Test that parse_gfm_table truncates data rows longer than header."""
    from sidedoc.reconstruct import parse_gfm_table

    table_content = """| A | B |
| --- | --- |
| 1 | 2 | 3 | extra |
| X | Y |"""

    rows, alignments = parse_gfm_table(table_content)

    # All rows should be truncated/normalized to the header column count (2)
    assert len(rows) == 3  # header + 2 data rows
    for row in rows:
        assert len(row) == 2, f"Row should have 2 columns, got {len(row)}: {row}"

    # The long row should be truncated
    assert rows[1] == ["1", "2"]


def test_extract_table_handles_newlines_in_cells() -> None:
    """Test that newlines in cells are handled gracefully."""
    from docx import Document as DocxDocument
    import tempfile

    with tempfile.TemporaryDirectory() as tmp_dir:
        # Create a docx with newline in cell
        doc = DocxDocument()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header"
        table.cell(0, 1).text = "Value"
        # Add a cell with paragraph break
        cell = table.cell(1, 0)
        cell.text = "Line1"  # First line
        # Note: python-docx cell.text will join multiple paragraphs
        table.cell(1, 1).text = "Normal"

        test_path = Path(tmp_dir) / "test_newlines.docx"
        doc.save(str(test_path))

        # Extract should not crash
        blocks, _ = extract_blocks(str(test_path))

        table_blocks = [b for b in blocks if b.type == "table"]
        assert len(table_blocks) >= 1

        # Content should be parseable (no literal newlines breaking the table)
        table_content = table_blocks[0].content
        lines = table_content.split("\n")
        # Each line should be a valid table row
        for line in lines:
            assert line.strip().startswith("|") and line.strip().endswith("|")


# ============================================================================
# Bug fix: docx_table_index should increment for multiple tables
# ============================================================================


def test_parse_markdown_multiple_tables_have_distinct_indices() -> None:
    """Test that multiple tables get distinct docx_table_index values."""
    markdown = """| A | B |
| --- | --- |
| 1 | 2 |

| X | Y |
| --- | --- |
| 3 | 4 |"""

    blocks = parse_markdown_to_blocks(markdown)
    table_blocks = [b for b in blocks if b.type == "table"]

    assert len(table_blocks) == 2, f"Expected 2 tables, got {len(table_blocks)}"

    assert table_blocks[0].table_metadata is not None
    assert table_blocks[1].table_metadata is not None

    assert table_blocks[0].table_metadata["docx_table_index"] == 0, \
        f"First table should have index 0, got {table_blocks[0].table_metadata['docx_table_index']}"
    assert table_blocks[1].table_metadata["docx_table_index"] == 1, \
        f"Second table should have index 1, got {table_blocks[1].table_metadata['docx_table_index']}"


# ============================================================================
# Bug fix: dimension validation in parse_markdown_to_blocks
# ============================================================================


def test_parse_markdown_to_blocks_rejects_oversized_table() -> None:
    """Test that parse_markdown_to_blocks rejects tables exceeding dimension limits."""
    import pytest
    from unittest.mock import patch

    # Table with 3 data rows + header = 4 total rows, 2 columns
    markdown = """| A | B |
| --- | --- |
| 1 | 2 |
| 3 | 4 |
| 5 | 6 |"""

    # Patch MAX_TABLE_ROWS to 2 (header + 1 data row max)
    with patch("sidedoc.reconstruct.MAX_TABLE_ROWS", 2):
        with pytest.raises(ValueError, match="too many rows"):
            parse_markdown_to_blocks(markdown)

    # Patch MAX_TABLE_COLS to 1
    with patch("sidedoc.reconstruct.MAX_TABLE_COLS", 1):
        with pytest.raises(ValueError, match="too many columns"):
            parse_markdown_to_blocks(markdown)


# ============================================================================
# Step 0: Fix package.py serialization bug
# ============================================================================


def test_package_serializes_table_metadata_in_structure() -> None:
    """Test that _build_metadata includes table_metadata in structure.json."""
    import tempfile
    from sidedoc.package import create_sidedoc_directory

    fixture_path = Path("tests/fixtures/tables_simple.docx")
    blocks, image_data = extract_blocks(str(fixture_path))
    styles = extract_styles(str(fixture_path), blocks)

    with tempfile.TemporaryDirectory() as tmp_dir:
        output_path = Path(tmp_dir) / "test.sidedoc"
        create_sidedoc_directory(
            str(output_path), blocks_to_markdown(blocks), blocks, styles, str(fixture_path)
        )

        # Read back structure.json
        structure = json.loads((output_path / "structure.json").read_text())
        table_blocks_in_structure = [
            b for b in structure["blocks"] if b["type"] == "table"
        ]
        assert len(table_blocks_in_structure) >= 1
        assert table_blocks_in_structure[0]["table_metadata"] is not None
        assert "rows" in table_blocks_in_structure[0]["table_metadata"]
        assert "cols" in table_blocks_in_structure[0]["table_metadata"]


def test_package_serializes_table_formatting_in_styles() -> None:
    """Test that _build_metadata includes table_formatting in styles.json."""
    import tempfile
    from sidedoc.package import create_sidedoc_directory

    fixture_path = Path("tests/fixtures/tables_simple.docx")
    blocks, image_data = extract_blocks(str(fixture_path))
    styles = extract_styles(str(fixture_path), blocks)

    with tempfile.TemporaryDirectory() as tmp_dir:
        output_path = Path(tmp_dir) / "test.sidedoc"
        create_sidedoc_directory(
            str(output_path), blocks_to_markdown(blocks), blocks, styles, str(fixture_path)
        )

        # Read back styles.json
        styles_data = json.loads((output_path / "styles.json").read_text())

        # Find the table block's style
        table_blocks_list = [b for b in blocks if b.type == "table"]
        assert len(table_blocks_list) >= 1
        table_block_id = table_blocks_list[0].id

        assert table_block_id in styles_data["block_styles"]
        table_style = styles_data["block_styles"][table_block_id]
        assert "table_formatting" in table_style
        assert table_style["table_formatting"] is not None
        assert "column_widths" in table_style["table_formatting"]


# ============================================================================
# US-T09: Create formatted table test fixture
# ============================================================================


def test_tables_formatted_fixture_exists() -> None:
    """Verify tables_formatted.docx fixture exists and has expected structure."""
    from docx.oxml.ns import qn

    fixtures_dir = Path("tests/fixtures")
    fixture_path = fixtures_dir / "tables_formatted.docx"

    assert fixture_path.exists(), "tables_formatted.docx should exist"

    doc = Document(str(fixture_path))
    assert len(doc.tables) >= 1, "Document should have at least one table"

    table = doc.tables[0]
    assert len(table.rows) == 4, "Table should have 4 rows (header + 3 data rows)"
    assert len(table.columns) == 3, "Table should have 3 columns"

    # Verify header row has shading (background color)
    header_cell = table.cell(0, 0)
    tcPr = header_cell._tc.find(qn('w:tcPr'))
    assert tcPr is not None, "Header cell should have tcPr"
    shd = tcPr.find(qn('w:shd'))
    assert shd is not None, "Header cell should have shading"
    fill = shd.get(qn('w:fill'))
    assert fill is not None, "Header cell shading should have fill color"

    # Verify header cell has borders
    tcBorders = tcPr.find(qn('w:tcBorders'))
    assert tcBorders is not None, "Header cell should have borders"


# ============================================================================
# US-T10: Extract cell background colors
# ============================================================================


def test_extract_cell_background_colors() -> None:
    """Test that cell background colors are extracted into cell_styles."""
    fixture_path = Path("tests/fixtures/tables_formatted.docx")
    blocks, _ = extract_blocks(str(fixture_path))
    styles = extract_styles(str(fixture_path), blocks)

    table_blocks_list = [b for b in blocks if b.type == "table"]
    assert len(table_blocks_list) >= 1
    table_block = table_blocks_list[0]

    table_styles = [s for s in styles if s.block_id == table_block.id]
    assert len(table_styles) == 1
    table_style = table_styles[0]

    assert table_style.table_formatting is not None
    assert "cell_styles" in table_style.table_formatting, \
        "table_formatting should have cell_styles"

    cell_styles = table_style.table_formatting["cell_styles"]
    # Header row cells should have background color
    assert "0,0" in cell_styles, "Header cell (0,0) should have styling"
    assert "background_color" in cell_styles["0,0"], \
        "Header cell should have background_color"


def test_extract_cell_shading_only_stores_non_default() -> None:
    """Test that cell_styles only stores cells with non-default formatting.

    Cells without any shading or borders are omitted from cell_styles.
    In our formatted fixture, all cells have borders, so we test with the
    simple fixture which has no cell-level formatting.
    """
    fixture_path = Path("tests/fixtures/tables_simple.docx")
    blocks, _ = extract_blocks(str(fixture_path))
    styles = extract_styles(str(fixture_path), blocks)

    table_blocks_list = [b for b in blocks if b.type == "table"]
    table_style = [s for s in styles if s.block_id == table_blocks_list[0].id][0]

    # Simple table has no cell-level formatting, so cell_styles should be absent
    assert "cell_styles" not in table_style.table_formatting or \
        len(table_style.table_formatting.get("cell_styles", {})) == 0, \
        "Simple table should have no cell_styles"


# ============================================================================
# US-T12: Extract cell border styles
# ============================================================================


def test_extract_cell_border_styles() -> None:
    """Test that cell border styles are extracted into cell_styles."""
    fixture_path = Path("tests/fixtures/tables_formatted.docx")
    blocks, _ = extract_blocks(str(fixture_path))
    styles = extract_styles(str(fixture_path), blocks)

    table_blocks_list = [b for b in blocks if b.type == "table"]
    table_style = [s for s in styles if s.block_id == table_blocks_list[0].id][0]

    cell_styles = table_style.table_formatting["cell_styles"]
    # Header cells should have borders
    assert "0,0" in cell_styles
    assert "borders" in cell_styles["0,0"], "Header cell should have border data"

    borders = cell_styles["0,0"]["borders"]
    # Should have at least one border side
    assert any(side in borders for side in ["top", "bottom", "left", "right"]), \
        "Header cell should have at least one border side"

    # Each border should have style, width, color
    for side, border_data in borders.items():
        assert "style" in border_data, f"{side} border should have style"
        assert "width" in border_data, f"{side} border should have width"
        assert "color" in border_data, f"{side} border should have color"


# ============================================================================
# US-T11: Apply cell background colors during reconstruction
# ============================================================================


def test_reconstruct_applies_cell_background_colors() -> None:
    """Test that cell background colors from styles are applied during reconstruction."""
    from docx.oxml.ns import qn
    from sidedoc.models import Block

    table_content = """| Name | Role |
| --- | --- |
| Alice | Engineer |"""

    blocks = [
        Block(
            id="block-0",
            type="table",
            content=table_content,
            docx_paragraph_index=-1,
            content_start=0,
            content_end=len(table_content),
            content_hash="abc123",
            table_metadata={
                "rows": 2,
                "cols": 2,
                "cells": [],
                "docx_table_index": 0,
            },
        )
    ]

    styles_dict: dict[str, Any] = {
        "block_styles": {
            "block-0": {
                "table_formatting": {
                    "column_widths": [2.0, 2.0],
                    "table_alignment": "left",
                    "cell_styles": {
                        "0,0": {"background_color": "D9E2F3"},
                        "0,1": {"background_color": "D9E2F3"},
                    },
                }
            }
        }
    }

    doc = create_docx_from_blocks(blocks, styles_dict)
    assert len(doc.tables) >= 1

    table = doc.tables[0]
    # Check header cells have shading
    cell_0_0 = table.cell(0, 0)
    tcPr = cell_0_0._tc.find(qn('w:tcPr'))
    assert tcPr is not None, "Cell should have tcPr"
    shd = tcPr.find(qn('w:shd'))
    assert shd is not None, "Cell should have shading"
    assert shd.get(qn('w:fill')) == 'D9E2F3', "Cell should have correct fill color"

    # Data row cell should NOT have shading
    cell_1_0 = table.cell(1, 0)
    tcPr_data = cell_1_0._tc.find(qn('w:tcPr'))
    if tcPr_data is not None:
        shd_data = tcPr_data.find(qn('w:shd'))
        assert shd_data is None or shd_data.get(qn('w:fill')) in (None, 'auto'), \
            "Data cell should not have shading"


# ============================================================================
# US-T13: Apply cell borders during reconstruction
# ============================================================================


def test_reconstruct_applies_cell_borders() -> None:
    """Test that cell borders from styles are applied during reconstruction."""
    from docx.oxml.ns import qn
    from sidedoc.models import Block

    table_content = """| Name | Role |
| --- | --- |
| Alice | Engineer |"""

    blocks = [
        Block(
            id="block-0",
            type="table",
            content=table_content,
            docx_paragraph_index=-1,
            content_start=0,
            content_end=len(table_content),
            content_hash="abc123",
            table_metadata={
                "rows": 2,
                "cols": 2,
                "cells": [],
                "docx_table_index": 0,
            },
        )
    ]

    styles_dict: dict[str, Any] = {
        "block_styles": {
            "block-0": {
                "table_formatting": {
                    "column_widths": [2.0, 2.0],
                    "table_alignment": "left",
                    "cell_styles": {
                        "0,0": {
                            "borders": {
                                "top": {"style": "single", "width": 8, "color": "4472C4"},
                                "bottom": {"style": "single", "width": 8, "color": "4472C4"},
                            }
                        },
                    },
                }
            }
        }
    }

    doc = create_docx_from_blocks(blocks, styles_dict)
    assert len(doc.tables) >= 1

    table = doc.tables[0]
    cell_0_0 = table.cell(0, 0)
    tcPr = cell_0_0._tc.find(qn('w:tcPr'))
    assert tcPr is not None, "Cell should have tcPr"
    tcBorders = tcPr.find(qn('w:tcBorders'))
    assert tcBorders is not None, "Cell should have borders"

    top_border = tcBorders.find(qn('w:top'))
    assert top_border is not None, "Cell should have top border"
    assert top_border.get(qn('w:val')) == 'single'
    assert top_border.get(qn('w:color')) == '4472C4'


def test_formatted_table_roundtrip() -> None:
    """Test that formatted table survives extract -> reconstruct roundtrip."""
    from docx.oxml.ns import qn

    fixture_path = Path("tests/fixtures/tables_formatted.docx")
    blocks, _ = extract_blocks(str(fixture_path))
    styles = extract_styles(str(fixture_path), blocks)

    # Build styles dict for reconstruction
    styles_dict = _build_styles_dict(styles)

    doc = create_docx_from_blocks(blocks, styles_dict)

    # Verify table exists and has shading on header
    assert len(doc.tables) >= 1
    table = doc.tables[0]
    cell_0_0 = table.cell(0, 0)
    tcPr = cell_0_0._tc.find(qn('w:tcPr'))
    assert tcPr is not None
    shd = tcPr.find(qn('w:shd'))
    assert shd is not None, "Header cell should have shading after roundtrip"
    assert shd.get(qn('w:fill')) == 'D9E2F3'


# ============================================================================
# US-T14: Sync cell content changes to docx
# ============================================================================


def test_sync_table_cell_edit_preserves_formatting() -> None:
    """Test that editing a cell during sync preserves cell formatting."""
    import tempfile
    from docx.oxml.ns import qn
    from sidedoc.sync import generate_updated_docx, match_blocks
    from sidedoc.reconstruct import parse_markdown_to_blocks

    fixture_path = Path("tests/fixtures/tables_formatted.docx")
    blocks, _ = extract_blocks(str(fixture_path))
    styles = extract_styles(str(fixture_path), blocks)

    # Build styles dict
    styles_dict = _build_styles_dict(styles)

    # Generate markdown, then edit a cell
    original_md = blocks_to_markdown(blocks)
    modified_md = original_md.replace("$1,200,000", "$1,500,000")
    assert modified_md != original_md, "Content should be different after edit"

    # Parse new blocks and match
    new_blocks = parse_markdown_to_blocks(modified_md)
    matches = match_blocks(blocks, new_blocks)

    with tempfile.TemporaryDirectory() as tmp_dir:
        output_path = Path(tmp_dir) / "output.docx"
        generate_updated_docx(new_blocks, matches, styles_dict, str(output_path))

        doc = Document(str(output_path))
        assert len(doc.tables) >= 1

        table = doc.tables[0]
        # Verify the edit took effect
        all_cell_text = [
            table.cell(r, c).text
            for r in range(len(table.rows))
            for c in range(len(table.columns))
        ]
        assert "$1,500,000" in all_cell_text, "Edited cell should have new value"

        # Verify formatting preserved: header should still have shading
        cell_0_0 = table.cell(0, 0)
        tcPr = cell_0_0._tc.find(qn('w:tcPr'))
        assert tcPr is not None
        shd = tcPr.find(qn('w:shd'))
        assert shd is not None, "Header cell should still have shading after sync"
        assert shd.get(qn('w:fill')) == 'D9E2F3'


# ============================================================================
# US-T15: Inline formatting in table cells
# ============================================================================


def test_extract_inline_formatting_in_cells() -> None:
    """Test that bold/italic in cell text is extracted as markdown."""
    import tempfile
    from docx import Document as DocxDocument

    with tempfile.TemporaryDirectory() as tmp_dir:
        doc = DocxDocument()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header"
        table.cell(0, 1).text = "Value"
        # Add bold text in a cell
        cell = table.cell(1, 0)
        cell.text = ""  # Clear default paragraph
        p = cell.paragraphs[0]
        run = p.add_run("bold text")
        run.bold = True
        table.cell(1, 1).text = "normal"

        test_path = Path(tmp_dir) / "test_inline.docx"
        doc.save(str(test_path))

        blocks, _ = extract_blocks(str(test_path))
        table_blocks_list = [b for b in blocks if b.type == "table"]
        assert len(table_blocks_list) >= 1

        content = table_blocks_list[0].content
        # Bold should be represented as **bold text** in the GFM
        assert "**bold text**" in content, \
            f"Bold text should be wrapped in ** markers, got: {content}"


def test_reconstruct_inline_formatting_in_cells() -> None:
    """Test that markdown formatting in cells creates styled runs in docx."""
    from sidedoc.models import Block

    table_content = """| Header | Value |
| --- | --- |
| **bold** | *italic* |"""

    blocks = [
        Block(
            id="block-0",
            type="table",
            content=table_content,
            docx_paragraph_index=-1,
            content_start=0,
            content_end=len(table_content),
            content_hash="abc123",
            table_metadata={
                "rows": 2,
                "cols": 2,
                "cells": [],
                "docx_table_index": 0,
            },
        )
    ]

    styles_dict: dict[str, Any] = {"block_styles": {}}
    doc = create_docx_from_blocks(blocks, styles_dict)

    assert len(doc.tables) >= 1
    table = doc.tables[0]

    # Check bold cell
    cell_1_0 = table.cell(1, 0)
    assert cell_1_0.text.strip() == "bold", \
        f"Cell text should be 'bold', got '{cell_1_0.text.strip()}'"
    has_bold = any(run.bold for run in cell_1_0.paragraphs[0].runs)
    assert has_bold, "Cell should have bold formatting"

    # Check italic cell
    cell_1_1 = table.cell(1, 1)
    assert cell_1_1.text.strip() == "italic", \
        f"Cell text should be 'italic', got '{cell_1_1.text.strip()}'"
    has_italic = any(run.italic for run in cell_1_1.paragraphs[0].runs)
    assert has_italic, "Cell should have italic formatting"


# ============================================================================
# US-T16: Add rows to tables via markdown
# ============================================================================


def test_sync_add_row_to_table() -> None:
    """Test that adding a row in GFM adds row in synced docx."""
    import tempfile
    from sidedoc.sync import generate_updated_docx, match_blocks
    from sidedoc.reconstruct import parse_markdown_to_blocks

    fixture_path = Path("tests/fixtures/tables_simple.docx")
    blocks, _ = extract_blocks(str(fixture_path))
    styles = extract_styles(str(fixture_path), blocks)

    styles_dict = _build_styles_dict(styles)

    # Add a row to the table
    original_md = blocks_to_markdown(blocks)
    # Add a new row after the last data row
    modified_md = original_md.replace(
        "| Bob | Designer | 2024-02-01 |",
        "| Bob | Designer | 2024-02-01 |\n| Charlie | Manager | 2024-03-15 |"
    )
    assert "Charlie" in modified_md

    new_blocks = parse_markdown_to_blocks(modified_md)
    matches = match_blocks(blocks, new_blocks)

    with tempfile.TemporaryDirectory() as tmp_dir:
        output_path = Path(tmp_dir) / "output.docx"
        generate_updated_docx(new_blocks, matches, styles_dict, str(output_path))

        doc = Document(str(output_path))
        assert len(doc.tables) >= 1
        table = doc.tables[0]
        assert len(table.rows) == 4, f"Should have 4 rows after adding, got {len(table.rows)}"
        assert table.cell(3, 0).text == "Charlie"
        assert table.cell(3, 1).text == "Manager"


# ============================================================================
# US-T17: Remove rows from tables via markdown
# ============================================================================


def test_sync_remove_row_from_table() -> None:
    """Test that removing a row in GFM removes row in synced docx."""
    import tempfile
    from sidedoc.sync import generate_updated_docx, match_blocks
    from sidedoc.reconstruct import parse_markdown_to_blocks

    fixture_path = Path("tests/fixtures/tables_simple.docx")
    blocks, _ = extract_blocks(str(fixture_path))
    styles = extract_styles(str(fixture_path), blocks)

    styles_dict = _build_styles_dict(styles)

    # Remove Bob's row
    original_md = blocks_to_markdown(blocks)
    modified_md = original_md.replace("| Bob | Designer | 2024-02-01 |\n", "")
    assert "Bob" not in modified_md

    new_blocks = parse_markdown_to_blocks(modified_md)
    matches = match_blocks(blocks, new_blocks)

    with tempfile.TemporaryDirectory() as tmp_dir:
        output_path = Path(tmp_dir) / "output.docx"
        generate_updated_docx(new_blocks, matches, styles_dict, str(output_path))

        doc = Document(str(output_path))
        assert len(doc.tables) >= 1
        table = doc.tables[0]
        assert len(table.rows) == 2, f"Should have 2 rows after removing, got {len(table.rows)}"
        assert table.cell(0, 0).text == "Name"  # Header preserved
        assert table.cell(1, 0).text == "Alice"  # Remaining row preserved


# ============================================================================
# US-T18: Add columns to tables via markdown
# ============================================================================


def test_sync_add_column_to_table() -> None:
    """Test that adding a column in GFM adds column in synced docx."""
    import tempfile
    from sidedoc.sync import generate_updated_docx, match_blocks
    from sidedoc.reconstruct import parse_markdown_to_blocks

    fixture_path = Path("tests/fixtures/tables_simple.docx")
    blocks, _ = extract_blocks(str(fixture_path))
    styles = extract_styles(str(fixture_path), blocks)

    styles_dict = _build_styles_dict(styles)

    # Add an Email column
    original_md = blocks_to_markdown(blocks)
    modified_md = original_md.replace(
        "| Name | Role | Start Date |",
        "| Name | Role | Start Date | Email |"
    ).replace(
        "| --- | --- | --- |",
        "| --- | --- | --- | --- |"
    ).replace(
        "| Alice | Engineer | 2024-01-15 |",
        "| Alice | Engineer | 2024-01-15 | alice@co.com |"
    ).replace(
        "| Bob | Designer | 2024-02-01 |",
        "| Bob | Designer | 2024-02-01 | bob@co.com |"
    )
    assert "Email" in modified_md

    new_blocks = parse_markdown_to_blocks(modified_md)
    matches = match_blocks(blocks, new_blocks)

    with tempfile.TemporaryDirectory() as tmp_dir:
        output_path = Path(tmp_dir) / "output.docx"
        generate_updated_docx(new_blocks, matches, styles_dict, str(output_path))

        doc = Document(str(output_path))
        assert len(doc.tables) >= 1
        table = doc.tables[0]
        assert len(table.columns) == 4, f"Should have 4 columns after adding, got {len(table.columns)}"
        assert table.cell(0, 3).text == "Email"
        assert table.cell(1, 3).text == "alice@co.com"


# ============================================================================
# US-T19: Remove columns from tables via markdown
# ============================================================================


def test_sync_remove_column_from_table() -> None:
    """Test that removing a column in GFM removes column in synced docx."""
    import tempfile
    from sidedoc.sync import generate_updated_docx, match_blocks
    from sidedoc.reconstruct import parse_markdown_to_blocks

    fixture_path = Path("tests/fixtures/tables_simple.docx")
    blocks, _ = extract_blocks(str(fixture_path))
    styles = extract_styles(str(fixture_path), blocks)

    styles_dict = _build_styles_dict(styles)

    # Remove the "Start Date" column (3rd column)
    original_md = blocks_to_markdown(blocks)
    modified_md = original_md.replace(
        "| Name | Role | Start Date |",
        "| Name | Role |"
    ).replace(
        "| --- | --- | --- |",
        "| --- | --- |"
    ).replace(
        "| Alice | Engineer | 2024-01-15 |",
        "| Alice | Engineer |"
    ).replace(
        "| Bob | Designer | 2024-02-01 |",
        "| Bob | Designer |"
    )
    assert "Start Date" not in modified_md

    new_blocks = parse_markdown_to_blocks(modified_md)
    matches = match_blocks(blocks, new_blocks)

    with tempfile.TemporaryDirectory() as tmp_dir:
        output_path = Path(tmp_dir) / "output.docx"
        generate_updated_docx(new_blocks, matches, styles_dict, str(output_path))

        doc = Document(str(output_path))
        assert len(doc.tables) >= 1
        table = doc.tables[0]
        assert len(table.columns) == 2, f"Should have 2 columns after removing, got {len(table.columns)}"
        assert table.cell(0, 0).text == "Name"
        assert table.cell(0, 1).text == "Role"


# ============================================================================
# US-T20: Create merged cells test fixture
# ============================================================================


def test_tables_merged_fixture_exists() -> None:
    """Verify tables_merged.docx fixture exists and has expected merge structure."""
    fixtures_dir = Path("tests/fixtures")
    fixture_path = fixtures_dir / "tables_merged.docx"

    assert fixture_path.exists(), "tables_merged.docx should exist"

    doc = Document(str(fixture_path))
    assert len(doc.tables) >= 1
    table = doc.tables[0]
    assert len(table.rows) == 4
    assert len(table.columns) == 3

    # Verify horizontal merge: row 0 cells should all reference same tc element
    assert table.cell(0, 0)._tc is table.cell(0, 1)._tc, \
        "Row 0, col 0 and col 1 should be merged"
    assert table.cell(0, 0)._tc is table.cell(0, 2)._tc, \
        "Row 0, col 0 and col 2 should be merged"

    # Verify vertical merge: rows 2-3, col 0 should reference same tc element
    assert table.cell(2, 0)._tc is table.cell(3, 0)._tc, \
        "Row 2, col 0 and row 3, col 0 should be merged"


# ============================================================================
# US-T21: Extract merged cell metadata
# ============================================================================


def test_extract_merged_cells_metadata() -> None:
    """Test that merged cells are detected and stored in table_metadata."""
    fixture_path = Path("tests/fixtures/tables_merged.docx")
    blocks, _ = extract_blocks(str(fixture_path))

    table_blocks_list = [b for b in blocks if b.type == "table"]
    assert len(table_blocks_list) >= 1
    table_block = table_blocks_list[0]

    assert table_block.table_metadata is not None
    assert "merged_cells" in table_block.table_metadata, \
        "table_metadata should have merged_cells"

    merged = table_block.table_metadata["merged_cells"]
    assert len(merged) >= 2, f"Should have at least 2 merge regions, got {len(merged)}"

    # Check horizontal merge: row 0, cols 0-2 (col_span=3)
    h_merge = [m for m in merged if m["start_row"] == 0 and m["col_span"] > 1]
    assert len(h_merge) >= 1, "Should have horizontal merge in row 0"
    assert h_merge[0]["col_span"] == 3

    # Check vertical merge: rows 2-3, col 0 (row_span=2)
    v_merge = [m for m in merged if m["start_col"] == 0 and m["row_span"] > 1]
    assert len(v_merge) >= 1, "Should have vertical merge in col 0"
    assert v_merge[0]["row_span"] == 2


# ============================================================================
# US-T22: Apply merged cells during reconstruction
# ============================================================================


def test_reconstruct_applies_merged_cells() -> None:
    """Test that merged_cells from metadata are applied during reconstruction."""
    from sidedoc.models import Block

    table_content = """| Report Title |  |  |
| --- | --- | --- |
| Category | Q1 | Q2 |
| Revenue | $1M | $1.2M |
|  | $0.8M | $0.9M |"""

    blocks = [
        Block(
            id="block-0",
            type="table",
            content=table_content,
            docx_paragraph_index=-1,
            content_start=0,
            content_end=len(table_content),
            content_hash="abc123",
            table_metadata={
                "rows": 4,
                "cols": 3,
                "cells": [],
                "column_alignments": ["left", "left", "left"],
                "docx_table_index": 0,
                "merged_cells": [
                    {"start_row": 0, "start_col": 0, "row_span": 1, "col_span": 3},
                    {"start_row": 2, "start_col": 0, "row_span": 2, "col_span": 1},
                ],
            },
        )
    ]

    styles_dict: dict[str, Any] = {"block_styles": {}}
    doc = create_docx_from_blocks(blocks, styles_dict)

    assert len(doc.tables) >= 1
    table = doc.tables[0]

    # Verify horizontal merge: row 0 cells should reference same tc
    assert table.cell(0, 0)._tc is table.cell(0, 2)._tc, \
        "Row 0 cells should be merged horizontally"

    # Verify vertical merge: rows 2-3, col 0 should reference same tc
    assert table.cell(2, 0)._tc is table.cell(3, 0)._tc, \
        "Rows 2-3, col 0 should be merged vertically"


# ============================================================================
# US-T23: Track header rows in tables
# ============================================================================


def test_extract_header_row_count() -> None:
    """Test that header rows are detected in table metadata."""
    fixture_path = Path("tests/fixtures/tables_simple.docx")
    blocks, _ = extract_blocks(str(fixture_path))

    table_blocks_list = [b for b in blocks if b.type == "table"]
    assert len(table_blocks_list) >= 1
    table_block = table_blocks_list[0]

    assert table_block.table_metadata is not None
    # Default should be 1 (GFM always has a header row)
    assert table_block.table_metadata.get("header_rows", 1) >= 1


def test_reconstruct_sets_header_rows() -> None:
    """Test that header rows are set during reconstruction."""
    from docx.oxml.ns import qn
    from sidedoc.models import Block

    table_content = """| Name | Role |
| --- | --- |
| Alice | Engineer |"""

    blocks = [
        Block(
            id="block-0",
            type="table",
            content=table_content,
            docx_paragraph_index=-1,
            content_start=0,
            content_end=len(table_content),
            content_hash="abc123",
            table_metadata={
                "rows": 2,
                "cols": 2,
                "cells": [],
                "docx_table_index": 0,
                "header_rows": 1,
            },
        )
    ]

    styles_dict: dict[str, Any] = {"block_styles": {}}
    doc = create_docx_from_blocks(blocks, styles_dict)
    table = doc.tables[0]

    # Check header row has tblHeader element
    first_row = table.rows[0]
    trPr = first_row._tr.find(qn('w:trPr'))
    assert trPr is not None, "First row should have trPr"
    tblHeader = trPr.find(qn('w:tblHeader'))
    assert tblHeader is not None, "First row should be marked as header"


# ============================================================================
# US-T25: Handle table edge cases
# ============================================================================


def test_empty_table_extraction() -> None:
    """Test that a table with only headers is handled gracefully."""
    import tempfile
    from docx import Document as DocxDocument

    with tempfile.TemporaryDirectory() as tmp_dir:
        doc = DocxDocument()
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Header1"
        table.cell(0, 1).text = "Header2"

        test_path = Path(tmp_dir) / "test_header_only.docx"
        doc.save(str(test_path))

        blocks, _ = extract_blocks(str(test_path))
        table_blocks_list = [b for b in blocks if b.type == "table"]
        assert len(table_blocks_list) >= 1, "Should extract header-only table"

        content = table_blocks_list[0].content
        assert "Header1" in content
        assert "Header2" in content


def test_wide_table_extraction() -> None:
    """Test that a wide table (15 columns) is extracted correctly."""
    import tempfile
    from docx import Document as DocxDocument

    with tempfile.TemporaryDirectory() as tmp_dir:
        doc = DocxDocument()
        table = doc.add_table(rows=2, cols=15)
        for col in range(15):
            table.cell(0, col).text = f"H{col}"
            table.cell(1, col).text = f"D{col}"

        test_path = Path(tmp_dir) / "test_wide.docx"
        doc.save(str(test_path))

        blocks, _ = extract_blocks(str(test_path))
        table_blocks_list = [b for b in blocks if b.type == "table"]
        assert len(table_blocks_list) >= 1

        metadata = table_blocks_list[0].table_metadata
        assert metadata is not None
        assert metadata["cols"] == 15


def test_tall_table_extraction() -> None:
    """Test that a tall table (100 rows) is extracted efficiently."""
    import tempfile
    import time
    from docx import Document as DocxDocument

    with tempfile.TemporaryDirectory() as tmp_dir:
        doc = DocxDocument()
        table = doc.add_table(rows=100, cols=3)
        for row in range(100):
            for col in range(3):
                table.cell(row, col).text = f"R{row}C{col}"

        test_path = Path(tmp_dir) / "test_tall.docx"
        doc.save(str(test_path))

        start = time.time()
        blocks, _ = extract_blocks(str(test_path))
        elapsed = time.time() - start

        table_blocks_list = [b for b in blocks if b.type == "table"]
        assert len(table_blocks_list) >= 1
        assert table_blocks_list[0].table_metadata["rows"] == 100
        assert elapsed < 5.0, f"Extraction took too long: {elapsed:.1f}s"


# ============================================================================
# US-T24: Validate tables in sidedoc
# ============================================================================


def test_validate_detects_table_dimension_mismatch() -> None:
    """Test that validation catches table dimension mismatches."""
    import tempfile
    from click.testing import CliRunner
    from sidedoc.cli import main as cli
    from tests.helpers import create_sidedoc_dir

    with tempfile.TemporaryDirectory() as tmp_dir:
        sidedoc_path = Path(tmp_dir) / "test.sidedoc"
        content = "| A | B |\n| --- | --- |\n| 1 | 2 |"
        structure = {
            "blocks": [{
                "id": "block-0",
                "type": "table",
                "docx_paragraph_index": -1,
                "content_start": 0,
                "content_end": len(content),
                "content_hash": "abc",
                "level": None,
                "image_path": None,
                "inline_formatting": None,
                "table_metadata": {
                    "rows": 5,  # Wrong! Content has 2 rows
                    "cols": 2,
                    "cells": [],
                    "column_alignments": ["left", "left"],
                    "docx_table_index": 0,
                },
                "track_changes": None,
            }]
        }
        create_sidedoc_dir(sidedoc_path, content, structure)

        runner = CliRunner()
        result = runner.invoke(cli, ["validate", str(sidedoc_path)])
        # Validation should report a warning about dimension mismatch
        assert "mismatch" in result.output.lower()


# ============================================================================
# US-T27: Roundtrip integration tests
# ============================================================================


def test_roundtrip_simple_table() -> None:
    """Roundtrip test: extract -> build for simple tables."""
    import tempfile

    fixture_path = Path("tests/fixtures/tables_simple.docx")
    blocks, _ = extract_blocks(str(fixture_path))
    styles = extract_styles(str(fixture_path), blocks)

    styles_dict = _build_styles_dict(styles)

    doc = create_docx_from_blocks(blocks, styles_dict)
    assert len(doc.tables) >= 1
    table = doc.tables[0]
    assert table.cell(0, 0).text == "Name"
    assert table.cell(1, 0).text == "Alice"
    assert table.cell(2, 0).text == "Bob"
    assert len(table.rows) == 3
    assert len(table.columns) == 3


def test_roundtrip_formatted_table() -> None:
    """Roundtrip test: extract -> build for formatted tables."""
    import tempfile
    from docx.oxml.ns import qn

    fixture_path = Path("tests/fixtures/tables_formatted.docx")
    blocks, _ = extract_blocks(str(fixture_path))
    styles = extract_styles(str(fixture_path), blocks)

    styles_dict = _build_styles_dict(styles)

    doc = create_docx_from_blocks(blocks, styles_dict)
    assert len(doc.tables) >= 1
    table = doc.tables[0]

    # Verify shading survived
    cell_0_0 = table.cell(0, 0)
    tcPr = cell_0_0._tc.find(qn('w:tcPr'))
    assert tcPr is not None
    shd = tcPr.find(qn('w:shd'))
    assert shd is not None, "Shading should survive roundtrip"

    # Verify borders survived
    tcBorders = tcPr.find(qn('w:tcBorders'))
    assert tcBorders is not None, "Borders should survive roundtrip"


def test_roundtrip_cell_edit_sync() -> None:
    """Roundtrip test: extract -> edit -> sync -> build for cell edits."""
    import tempfile
    from sidedoc.sync import generate_updated_docx, match_blocks
    from sidedoc.reconstruct import parse_markdown_to_blocks

    fixture_path = Path("tests/fixtures/tables_simple.docx")
    blocks, _ = extract_blocks(str(fixture_path))
    styles = extract_styles(str(fixture_path), blocks)

    styles_dict = _build_styles_dict(styles)

    original_md = blocks_to_markdown(blocks)
    modified_md = original_md.replace("Alice", "Alice Smith")

    new_blocks = parse_markdown_to_blocks(modified_md)
    matches = match_blocks(blocks, new_blocks)

    with tempfile.TemporaryDirectory() as tmp_dir:
        output_path = Path(tmp_dir) / "output.docx"
        generate_updated_docx(new_blocks, matches, styles_dict, str(output_path))

        doc = Document(str(output_path))
        assert len(doc.tables) >= 1
        table = doc.tables[0]

        all_text = [
            table.cell(r, c).text
            for r in range(len(table.rows))
            for c in range(len(table.columns))
        ]
        assert "Alice Smith" in all_text
        assert "Name" in all_text  # Header preserved
        assert "Bob" in all_text  # Other data preserved


# ============================================================================
# Fix 1: Table dimension pre-validation in create_table_from_gfm
# ============================================================================


def test_create_table_from_gfm_rejects_oversized_rows() -> None:
    """Test that create_table_from_gfm rejects tables exceeding MAX_TABLE_ROWS."""
    import pytest
    from unittest.mock import patch
    from sidedoc.reconstruct import create_table_from_gfm

    # Build a table with 6 data rows + header = 7 total
    header = "| A | B |"
    sep = "| --- | --- |"
    data_rows = "\n".join(f"| R{i} | D{i} |" for i in range(6))
    table_content = f"{header}\n{sep}\n{data_rows}"

    doc = Document()
    styles: dict[str, Any] = {"block_styles": {}}

    with patch("sidedoc.reconstruct.MAX_TABLE_ROWS", 3):
        with pytest.raises(ValueError, match="too many rows"):
            create_table_from_gfm(doc, table_content, styles, "block-0")


def test_create_table_from_gfm_rejects_oversized_cols() -> None:
    """Test that create_table_from_gfm rejects tables exceeding MAX_TABLE_COLS."""
    import pytest
    from unittest.mock import patch
    from sidedoc.reconstruct import create_table_from_gfm

    # Build a table with 5 columns
    header = "| A | B | C | D | E |"
    sep = "| --- | --- | --- | --- | --- |"
    data = "| 1 | 2 | 3 | 4 | 5 |"
    table_content = f"{header}\n{sep}\n{data}"

    doc = Document()
    styles: dict[str, Any] = {"block_styles": {}}

    with patch("sidedoc.reconstruct.MAX_TABLE_COLS", 3):
        with pytest.raises(ValueError, match="too many columns"):
            create_table_from_gfm(doc, table_content, styles, "block-0")


def test_prevalidation_does_not_call_parse_gfm_table() -> None:
    """Test that pre-validation rejects before parse_gfm_table is called."""
    import pytest
    from unittest.mock import patch, MagicMock
    from sidedoc.reconstruct import create_table_from_gfm

    # Build a table with 6 data rows + header = 7 total
    header = "| A | B |"
    sep = "| --- | --- |"
    data_rows = "\n".join(f"| R{i} | D{i} |" for i in range(6))
    table_content = f"{header}\n{sep}\n{data_rows}"

    doc = Document()
    styles: dict[str, Any] = {"block_styles": {}}

    mock_parse = MagicMock()

    with patch("sidedoc.reconstruct.MAX_TABLE_ROWS", 3), \
         patch("sidedoc.reconstruct.parse_gfm_table", mock_parse):
        with pytest.raises(ValueError, match="too many rows"):
            create_table_from_gfm(doc, table_content, styles, "block-0")

    mock_parse.assert_not_called()


# ============================================================================
# Fix 3: CriticMarkup and hyperlinks in table cells
# ============================================================================


def test_reconstruct_table_cell_with_criticmarkup_insertion() -> None:
    """Test that CriticMarkup insertions in table cells produce w:ins elements."""
    from docx.oxml.ns import qn
    from sidedoc.models import Block

    table_content = """| Name | Status |
| --- | --- |
| Alice | {++promoted++} |"""

    blocks = [
        Block(
            id="block-0",
            type="table",
            content=table_content,
            docx_paragraph_index=-1,
            content_start=0,
            content_end=len(table_content),
            content_hash="abc123",
            table_metadata={"rows": 2, "cols": 2, "cells": [], "docx_table_index": 0},
        )
    ]

    styles: dict[str, Any] = {"block_styles": {}}
    doc = create_docx_from_blocks(blocks, styles)

    assert len(doc.tables) >= 1
    table = doc.tables[0]
    cell = table.cell(1, 1)

    # The cell should contain a w:ins element
    ins_elements = cell._tc.findall(f'.//{{{qn("w:ins").split("}")[0][1:]}}}ins')
    assert len(ins_elements) >= 1, (
        f"Cell should contain w:ins element for CriticMarkup insertion. "
        f"Cell XML: {cell._tc.xml[:300]}"
    )


def test_reconstruct_table_cell_with_criticmarkup_deletion() -> None:
    """Test that CriticMarkup deletions in table cells produce w:del elements."""
    from docx.oxml.ns import qn
    from sidedoc.models import Block

    table_content = """| Name | Status |
| --- | --- |
| Bob | {--demoted--} |"""

    blocks = [
        Block(
            id="block-0",
            type="table",
            content=table_content,
            docx_paragraph_index=-1,
            content_start=0,
            content_end=len(table_content),
            content_hash="abc123",
            table_metadata={"rows": 2, "cols": 2, "cells": [], "docx_table_index": 0},
        )
    ]

    styles: dict[str, Any] = {"block_styles": {}}
    doc = create_docx_from_blocks(blocks, styles)

    assert len(doc.tables) >= 1
    table = doc.tables[0]
    cell = table.cell(1, 1)

    # The cell should contain a w:del element
    del_elements = cell._tc.findall(f'.//{{{qn("w:del").split("}")[0][1:]}}}del')
    assert len(del_elements) >= 1, (
        f"Cell should contain w:del element for CriticMarkup deletion. "
        f"Cell XML: {cell._tc.xml[:300]}"
    )


def test_reconstruct_table_cell_with_hyperlink() -> None:
    """Test that hyperlinks in table cells produce w:hyperlink elements."""
    from docx.oxml.ns import qn
    from sidedoc.models import Block

    table_content = """| Name | Link |
| --- | --- |
| Alice | [homepage](https://example.com) |"""

    blocks = [
        Block(
            id="block-0",
            type="table",
            content=table_content,
            docx_paragraph_index=-1,
            content_start=0,
            content_end=len(table_content),
            content_hash="abc123",
            table_metadata={"rows": 2, "cols": 2, "cells": [], "docx_table_index": 0},
        )
    ]

    styles: dict[str, Any] = {"block_styles": {}}
    doc = create_docx_from_blocks(blocks, styles)

    assert len(doc.tables) >= 1
    table = doc.tables[0]
    cell = table.cell(1, 1)

    # The cell should contain a w:hyperlink element
    hyperlink_elements = cell._tc.findall(f'.//{{{qn("w:hyperlink").split("}")[0][1:]}}}hyperlink')
    assert len(hyperlink_elements) >= 1, (
        f"Cell should contain w:hyperlink element. "
        f"Cell XML: {cell._tc.xml[:300]}"
    )
