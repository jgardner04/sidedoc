"""Tests for multi-column layout support (JON-95)."""

import tempfile
from pathlib import Path

from click.testing import CliRunner
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from sidedoc.cli import main
from sidedoc.extract import extract_blocks, extract_sections
from sidedoc.models import ColumnDefinition, SectionProperties


# =============================================================================
# Model Tests
# =============================================================================


class TestColumnDefinitionModel:
    """Tests for the ColumnDefinition dataclass."""

    def test_default_column_definition(self):
        col = ColumnDefinition(width=4320, space=720)
        assert col.width == 4320
        assert col.space == 720

    def test_column_definition_no_space(self):
        col = ColumnDefinition(width=4320)
        assert col.width == 4320
        assert col.space is None


class TestSectionPropertiesModel:
    """Tests for the SectionProperties dataclass."""

    def test_default_single_column(self):
        section = SectionProperties()
        assert section.column_count == 1
        assert section.column_spacing is None
        assert section.equal_width is True
        assert section.columns is None

    def test_two_column_equal_width(self):
        section = SectionProperties(
            column_count=2,
            column_spacing=720,
            equal_width=True,
        )
        assert section.column_count == 2
        assert section.column_spacing == 720
        assert section.equal_width is True
        assert section.columns is None

    def test_unequal_columns(self):
        cols = [
            ColumnDefinition(width=4320, space=720),
            ColumnDefinition(width=4320),
        ]
        section = SectionProperties(
            column_count=2,
            column_spacing=None,
            equal_width=False,
            columns=cols,
        )
        assert section.column_count == 2
        assert section.equal_width is False
        assert len(section.columns) == 2
        assert section.columns[0].width == 4320
        assert section.columns[0].space == 720
        assert section.columns[1].space is None

    def test_section_with_block_range(self):
        section = SectionProperties(
            column_count=2,
            column_spacing=720,
            start_block_index=0,
            end_block_index=5,
        )
        assert section.start_block_index == 0
        assert section.end_block_index == 5


# =============================================================================
# Extraction Tests
# =============================================================================


def _add_section_break(doc, col_count=1, col_space=720, equal_width=True, columns=None):
    """Add a section break with column properties to a document.

    Inserts a paragraph with sectPr to create a mid-document section break.
    """
    p = doc.add_paragraph()
    p_elem = p._element
    pPr = p_elem.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')

    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), str(col_count))
    if col_space is not None:
        cols.set(qn('w:space'), str(col_space))
    if equal_width:
        cols.set(qn('w:equalWidth'), '1')
    else:
        cols.set(qn('w:equalWidth'), '0')
        if columns:
            for col_def in columns:
                col_elem = OxmlElement('w:col')
                col_elem.set(qn('w:w'), str(col_def['width']))
                if col_def.get('space') is not None:
                    col_elem.set(qn('w:space'), str(col_def['space']))
                cols.append(col_elem)

    sectPr.append(cols)
    pPr.append(sectPr)
    return p


def _set_body_columns(doc, col_count=1, col_space=720, equal_width=True, columns=None):
    """Set column properties on the document's body-level sectPr."""
    body = doc.element.body
    # Get or create body sectPr
    sectPr = body.find(qn('w:sectPr'))
    if sectPr is None:
        sectPr = OxmlElement('w:sectPr')
        body.append(sectPr)

    # Remove existing cols elements
    for existing_cols in sectPr.findall(qn('w:cols')):
        sectPr.remove(existing_cols)

    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), str(col_count))
    if col_space is not None:
        cols.set(qn('w:space'), str(col_space))
    if equal_width:
        cols.set(qn('w:equalWidth'), '1')
    else:
        cols.set(qn('w:equalWidth'), '0')
        if columns:
            for col_def in columns:
                col_elem = OxmlElement('w:col')
                col_elem.set(qn('w:w'), str(col_def['width']))
                if col_def.get('space') is not None:
                    col_elem.set(qn('w:space'), str(col_def['space']))
                cols.append(col_elem)

    sectPr.append(cols)


def _add_column_break(paragraph):
    """Add a column break to an existing paragraph."""
    run = paragraph.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'column')
    run._element.append(br)


class TestExtractColumns:
    """Tests for extracting column properties from docx."""

    def test_extract_single_column_default(self):
        """Documents with no column settings should have 1-column sections."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            doc.add_paragraph("Simple paragraph")
            doc.save(f.name)

            sections = extract_sections(f.name)
            assert len(sections) == 1
            assert sections[0].column_count == 1

    def test_extract_two_column_equal_width(self):
        """Extract 2-column equal-width layout from body sectPr."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            doc.add_paragraph("Column content")
            _set_body_columns(doc, col_count=2, col_space=720, equal_width=True)
            doc.save(f.name)

            sections = extract_sections(f.name)
            assert len(sections) == 1
            assert sections[0].column_count == 2
            assert sections[0].column_spacing == 720
            assert sections[0].equal_width is True

    def test_extract_three_column(self):
        """Extract 3-column layout."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            doc.add_paragraph("Three column content")
            _set_body_columns(doc, col_count=3, col_space=360, equal_width=True)
            doc.save(f.name)

            sections = extract_sections(f.name)
            assert len(sections) == 1
            assert sections[0].column_count == 3
            assert sections[0].column_spacing == 360

    def test_extract_unequal_columns(self):
        """Extract unequal column widths."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            doc.add_paragraph("Unequal columns")
            _set_body_columns(
                doc, col_count=2, equal_width=False,
                columns=[
                    {'width': 5760, 'space': 720},
                    {'width': 3600},
                ]
            )
            doc.save(f.name)

            sections = extract_sections(f.name)
            assert len(sections) == 1
            assert sections[0].column_count == 2
            assert sections[0].equal_width is False
            assert len(sections[0].columns) == 2
            assert sections[0].columns[0].width == 5760
            assert sections[0].columns[0].space == 720
            assert sections[0].columns[1].width == 3600
            assert sections[0].columns[1].space is None

    def test_extract_multi_section_with_columns(self):
        """Extract column properties from multi-section document."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            doc.add_paragraph("Section 1 content")
            _add_section_break(doc, col_count=2, col_space=720)
            doc.add_paragraph("Section 2 content")
            _set_body_columns(doc, col_count=1)
            doc.save(f.name)

            sections = extract_sections(f.name)
            assert len(sections) == 2
            assert sections[0].column_count == 2
            assert sections[0].column_spacing == 720
            assert sections[1].column_count == 1


class TestExtractColumnBreaks:
    """Tests for column break detection in content."""

    def test_column_break_in_content(self):
        """Column breaks should appear as <!-- column-break --> in content."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            p = doc.add_paragraph("Before break")
            _add_column_break(p)
            doc.add_paragraph("After break")
            _set_body_columns(doc, col_count=2)
            doc.save(f.name)

            blocks, _ = extract_blocks(f.name)
            # The paragraph with a column break should contain the marker
            assert any("<!-- column-break -->" in b.content for b in blocks)


# =============================================================================
# Round-trip Tests
# =============================================================================


class TestColumnRoundtrip:
    """Tests for extract → build round-trip preservation of columns."""

    def test_roundtrip_two_column_equal(self):
        """Round-trip preserves 2-column equal-width layout."""
        runner = CliRunner()
        with runner.isolated_filesystem():
            doc = Document()
            doc.add_paragraph("Two column text")
            doc.add_paragraph("More content")
            _set_body_columns(doc, col_count=2, col_space=720, equal_width=True)
            doc.save("original.docx")

            result = runner.invoke(main, ["extract", "original.docx"])
            assert result.exit_code == 0

            result = runner.invoke(main, ["build", "original.sidedoc", "-o", "rebuilt.docx"])
            assert result.exit_code == 0

            rebuilt = Document("rebuilt.docx")
            body = rebuilt.element.body
            sectPr = body.find(qn('w:sectPr'))
            assert sectPr is not None
            cols = sectPr.find(qn('w:cols'))
            assert cols is not None
            assert cols.get(qn('w:num')) == '2'
            assert cols.get(qn('w:space')) == '720'

    def test_roundtrip_three_column(self):
        """Round-trip preserves 3-column layout."""
        runner = CliRunner()
        with runner.isolated_filesystem():
            doc = Document()
            doc.add_paragraph("Three column text")
            _set_body_columns(doc, col_count=3, col_space=360, equal_width=True)
            doc.save("original.docx")

            result = runner.invoke(main, ["extract", "original.docx"])
            assert result.exit_code == 0

            result = runner.invoke(main, ["build", "original.sidedoc", "-o", "rebuilt.docx"])
            assert result.exit_code == 0

            rebuilt = Document("rebuilt.docx")
            body = rebuilt.element.body
            sectPr = body.find(qn('w:sectPr'))
            cols = sectPr.find(qn('w:cols'))
            assert cols.get(qn('w:num')) == '3'

    def test_roundtrip_unequal_columns(self):
        """Round-trip preserves unequal column widths."""
        runner = CliRunner()
        with runner.isolated_filesystem():
            doc = Document()
            doc.add_paragraph("Unequal column text")
            _set_body_columns(
                doc, col_count=2, equal_width=False,
                columns=[
                    {'width': 5760, 'space': 720},
                    {'width': 3600},
                ]
            )
            doc.save("original.docx")

            result = runner.invoke(main, ["extract", "original.docx"])
            assert result.exit_code == 0

            result = runner.invoke(main, ["build", "original.sidedoc", "-o", "rebuilt.docx"])
            assert result.exit_code == 0

            rebuilt = Document("rebuilt.docx")
            body = rebuilt.element.body
            sectPr = body.find(qn('w:sectPr'))
            cols = sectPr.find(qn('w:cols'))
            assert cols.get(qn('w:equalWidth')) == '0'
            col_elems = cols.findall(qn('w:col'))
            assert len(col_elems) == 2
            assert col_elems[0].get(qn('w:w')) == '5760'
            assert col_elems[0].get(qn('w:space')) == '720'
            assert col_elems[1].get(qn('w:w')) == '3600'

    def test_roundtrip_multi_section(self):
        """Round-trip preserves column properties across multiple sections."""
        runner = CliRunner()
        with runner.isolated_filesystem():
            doc = Document()
            doc.add_paragraph("Section 1 - two columns")
            _add_section_break(doc, col_count=2, col_space=720)
            doc.add_paragraph("Section 2 - single column")
            _set_body_columns(doc, col_count=1)
            doc.save("original.docx")

            result = runner.invoke(main, ["extract", "original.docx"])
            assert result.exit_code == 0

            result = runner.invoke(main, ["build", "original.sidedoc", "-o", "rebuilt.docx"])
            assert result.exit_code == 0

            rebuilt = Document("rebuilt.docx")
            body = rebuilt.element.body
            # Check that section properties exist in rebuilt doc
            all_sectPr = body.findall('.//' + qn('w:sectPr'))
            assert len(all_sectPr) >= 2

    def test_roundtrip_column_break(self):
        """Round-trip preserves column breaks."""
        runner = CliRunner()
        with runner.isolated_filesystem():
            doc = Document()
            p = doc.add_paragraph("Before column break")
            _add_column_break(p)
            doc.add_paragraph("After column break")
            _set_body_columns(doc, col_count=2, col_space=720)
            doc.save("original.docx")

            result = runner.invoke(main, ["extract", "original.docx"])
            assert result.exit_code == 0

            result = runner.invoke(main, ["build", "original.sidedoc", "-o", "rebuilt.docx"])
            assert result.exit_code == 0

            rebuilt = Document("rebuilt.docx")
            # Find column break in rebuilt document
            breaks = rebuilt.element.body.findall('.//' + qn('w:br'))
            column_breaks = [br for br in breaks if br.get(qn('w:type')) == 'column']
            assert len(column_breaks) >= 1


# =============================================================================
# Serialization Tests
# =============================================================================


class TestSectionSerialization:
    """Tests for section metadata in structure.json."""

    def test_sections_in_structure_json(self):
        """Sections with column properties should appear in structure.json."""
        runner = CliRunner()
        with runner.isolated_filesystem():
            import json

            doc = Document()
            doc.add_paragraph("Two column content")
            _set_body_columns(doc, col_count=2, col_space=720, equal_width=True)
            doc.save("test.docx")

            result = runner.invoke(main, ["extract", "test.docx"])
            assert result.exit_code == 0

            structure = json.loads(Path("test.sidedoc/structure.json").read_text())
            assert "sections" in structure
            sections = structure["sections"]
            assert len(sections) == 1
            assert sections[0]["column_count"] == 2
            assert sections[0]["column_spacing"] == 720
            assert sections[0]["equal_width"] is True

    def test_unequal_columns_in_structure_json(self):
        """Unequal column definitions should be serialized."""
        runner = CliRunner()
        with runner.isolated_filesystem():
            import json

            doc = Document()
            doc.add_paragraph("Unequal columns")
            _set_body_columns(
                doc, col_count=2, equal_width=False,
                columns=[
                    {'width': 5760, 'space': 720},
                    {'width': 3600},
                ]
            )
            doc.save("test.docx")

            result = runner.invoke(main, ["extract", "test.docx"])
            assert result.exit_code == 0

            structure = json.loads(Path("test.sidedoc/structure.json").read_text())
            sections = structure["sections"]
            assert sections[0]["equal_width"] is False
            assert len(sections[0]["columns"]) == 2
            assert sections[0]["columns"][0]["width"] == 5760
            assert sections[0]["columns"][0]["space"] == 720
            assert sections[0]["columns"][1]["width"] == 3600
            assert sections[0]["columns"][1].get("space") is None
