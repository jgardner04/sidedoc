"""Tests for docx_style application and paragraph format preservation (JON-91)."""

import json
import warnings
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt, Inches
from click.testing import CliRunner

from sidedoc.cli import main
from sidedoc.extract import extract_blocks, extract_styles


class TestDocxStyleApplication:
    """Test that docx_style is applied to paragraphs during reconstruction."""

    def test_roundtrip_preserves_paragraph_style(self):
        """Rebuilt document should preserve paragraph styles, not fallback to Normal."""
        runner = CliRunner()
        with runner.isolated_filesystem():
            doc = Document()
            doc.add_paragraph("A quote paragraph", style="Quote")
            doc.add_paragraph("Normal paragraph")
            doc.add_paragraph("Title text", style="Title")
            doc.save("original.docx")

            extract_result = runner.invoke(main, ["extract", "original.docx"])
            assert extract_result.exit_code == 0, extract_result.output
            result = runner.invoke(main, ["build", "original.sidedoc", "-o", "rebuilt.docx"])
            assert result.exit_code == 0

            rebuilt = Document("rebuilt.docx")
            style_names = [p.style.name for p in rebuilt.paragraphs]
            assert "Quote" in style_names
            assert "Title" in style_names

    def test_roundtrip_preserves_list_bullet_style(self):
        """List Bullet style should survive round-trip."""
        runner = CliRunner()
        with runner.isolated_filesystem():
            doc = Document()
            doc.add_paragraph("Bullet item", style="List Bullet")
            doc.save("original.docx")

            extract_result = runner.invoke(main, ["extract", "original.docx"])
            assert extract_result.exit_code == 0, extract_result.output
            result = runner.invoke(main, ["build", "original.sidedoc", "-o", "rebuilt.docx"])
            assert result.exit_code == 0

            rebuilt = Document("rebuilt.docx")
            # Find the paragraph with the bullet content
            bullet_paras = [p for p in rebuilt.paragraphs if "Bullet item" in p.text]
            assert len(bullet_paras) > 0
            assert bullet_paras[0].style.name == "List Bullet"

    def test_roundtrip_preserves_heading_styles(self):
        """Heading styles should be applied (not just Normal)."""
        runner = CliRunner()
        with runner.isolated_filesystem():
            doc = Document()
            doc.add_paragraph("Chapter One", style="Heading 1")
            doc.add_paragraph("Section A", style="Heading 2")
            doc.add_paragraph("Body text")
            doc.save("original.docx")

            extract_result = runner.invoke(main, ["extract", "original.docx"])
            assert extract_result.exit_code == 0, extract_result.output
            result = runner.invoke(main, ["build", "original.sidedoc", "-o", "rebuilt.docx"])
            assert result.exit_code == 0

            rebuilt = Document("rebuilt.docx")
            texts_and_styles = {p.text: p.style.name for p in rebuilt.paragraphs}
            assert texts_and_styles.get("Chapter One") == "Heading 1"
            assert texts_and_styles.get("Section A") == "Heading 2"

    def test_missing_custom_style_falls_back_to_normal(self):
        """When a custom style doesn't exist in the target document, fall back to
        Normal and emit a UserWarning so the user can spot silent fidelity loss.
        """
        runner = CliRunner()
        with runner.isolated_filesystem():
            # Create a document with a custom style.
            doc = Document()
            doc.styles.add_style("Legal Citation", 1)  # WD_STYLE_TYPE.PARAGRAPH = 1
            doc.add_paragraph("Case law reference", style="Legal Citation")
            doc.save("original.docx")

            extract_result = runner.invoke(main, ["extract", "original.docx"])
            assert extract_result.exit_code == 0, extract_result.output

            # Build should succeed (Normal fallback) AND emit a warning naming the style.
            # CliRunner swallows warnings emitted in the subprocess by default, so call
            # _apply_block_formatting directly to verify the warning fires.
            from sidedoc.reconstruct import _apply_block_formatting

            target_doc = Document()
            target_para = target_doc.add_paragraph("Case law reference")
            with warnings.catch_warnings(record=True) as caught:
                warnings.simplefilter("always")
                _apply_block_formatting(
                    target_para,
                    {"docx_style": "Legal Citation", "alignment": "left"},
                )
            assert any(
                "Legal Citation" in str(w.message) and issubclass(w.category, UserWarning)
                for w in caught
            ), f"Expected UserWarning naming 'Legal Citation'; got {[str(w.message) for w in caught]}"

            # End-to-end build still succeeds and preserves the paragraph text.
            result = runner.invoke(main, ["build", "original.sidedoc", "-o", "rebuilt.docx"])
            assert result.exit_code == 0
            rebuilt = Document("rebuilt.docx")
            texts = [p.text for p in rebuilt.paragraphs]
            assert "Case law reference" in texts


class TestParagraphFormatExtraction:
    """Test that paragraph format properties are extracted."""

    def _extract(self, tmp_path, doc):
        """Helper: save doc, extract blocks+styles, return styles list."""
        docx_path = str(tmp_path / "test.docx")
        doc.save(docx_path)
        blocks, _ = extract_blocks(docx_path)
        styles = extract_styles(docx_path, blocks)
        return styles

    def test_extract_left_indent(self, tmp_path):
        """Left indent should be extracted from paragraph format."""
        doc = Document()
        para = doc.add_paragraph("Indented text")
        para.paragraph_format.left_indent = Inches(0.5)

        styles = self._extract(tmp_path, doc)
        para_styles = [s for s in styles if s.docx_style != "Table"]
        assert len(para_styles) > 0
        assert para_styles[0].left_indent is not None
        assert para_styles[0].left_indent > 0

    def test_extract_spacing(self, tmp_path):
        """Space before and after should be extracted from paragraph format."""
        doc = Document()
        para = doc.add_paragraph("Spaced text")
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)

        styles = self._extract(tmp_path, doc)
        para_styles = [s for s in styles if s.docx_style != "Table"]
        assert len(para_styles) > 0
        assert para_styles[0].space_before is not None
        assert para_styles[0].space_after is not None

    def test_extract_keep_together(self, tmp_path):
        """keep_together should be extracted from paragraph format."""
        doc = Document()
        para = doc.add_paragraph("Keep together text")
        para.paragraph_format.keep_together = True

        styles = self._extract(tmp_path, doc)
        para_styles = [s for s in styles if s.docx_style != "Table"]
        assert len(para_styles) > 0
        assert para_styles[0].keep_together is True

    def test_extract_page_break_before(self, tmp_path):
        """page_break_before should be extracted from paragraph format."""
        doc = Document()
        para = doc.add_paragraph("Page break text")
        para.paragraph_format.page_break_before = True

        styles = self._extract(tmp_path, doc)
        para_styles = [s for s in styles if s.docx_style != "Table"]
        assert len(para_styles) > 0
        assert para_styles[0].page_break_before is True


class TestParagraphFormatRoundtrip:
    """Test that paragraph format properties survive round-trip."""

    def test_roundtrip_preserves_indentation(self):
        """Left indent should survive extract → build round-trip."""
        runner = CliRunner()
        with runner.isolated_filesystem():
            doc = Document()
            para = doc.add_paragraph("Indented text")
            para.paragraph_format.left_indent = Inches(0.5)
            doc.save("original.docx")

            extract_result = runner.invoke(main, ["extract", "original.docx"])
            assert extract_result.exit_code == 0, extract_result.output
            result = runner.invoke(main, ["build", "original.sidedoc", "-o", "rebuilt.docx"])
            assert result.exit_code == 0

            rebuilt = Document("rebuilt.docx")
            indented = [p for p in rebuilt.paragraphs if "Indented" in p.text]
            assert len(indented) > 0
            assert indented[0].paragraph_format.left_indent is not None
            assert indented[0].paragraph_format.left_indent == Inches(0.5)

    def test_roundtrip_preserves_spacing(self):
        """Space before/after should survive extract → build round-trip."""
        runner = CliRunner()
        with runner.isolated_filesystem():
            doc = Document()
            para = doc.add_paragraph("Spaced text")
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
            doc.save("original.docx")

            extract_result = runner.invoke(main, ["extract", "original.docx"])
            assert extract_result.exit_code == 0, extract_result.output
            result = runner.invoke(main, ["build", "original.sidedoc", "-o", "rebuilt.docx"])
            assert result.exit_code == 0

            rebuilt = Document("rebuilt.docx")
            spaced = [p for p in rebuilt.paragraphs if "Spaced" in p.text]
            assert len(spaced) > 0
            assert spaced[0].paragraph_format.space_before == Pt(12)
            assert spaced[0].paragraph_format.space_after == Pt(6)

    def test_roundtrip_preserves_keep_with_next(self):
        """keep_with_next should survive extract → build round-trip."""
        runner = CliRunner()
        with runner.isolated_filesystem():
            doc = Document()
            para = doc.add_paragraph("Keep with next text")
            para.paragraph_format.keep_with_next = True
            doc.add_paragraph("Following text")
            doc.save("original.docx")

            extract_result = runner.invoke(main, ["extract", "original.docx"])
            assert extract_result.exit_code == 0, extract_result.output
            result = runner.invoke(main, ["build", "original.sidedoc", "-o", "rebuilt.docx"])
            assert result.exit_code == 0

            rebuilt = Document("rebuilt.docx")
            kept = [p for p in rebuilt.paragraphs if "Keep with next" in p.text]
            assert len(kept) > 0
            assert kept[0].paragraph_format.keep_with_next is True

    def test_roundtrip_preserves_first_line_indent(self):
        """First line indent should survive extract → build round-trip."""
        runner = CliRunner()
        with runner.isolated_filesystem():
            doc = Document()
            para = doc.add_paragraph("First line indented text")
            para.paragraph_format.first_line_indent = Inches(0.25)
            doc.save("original.docx")

            extract_result = runner.invoke(main, ["extract", "original.docx"])
            assert extract_result.exit_code == 0, extract_result.output
            result = runner.invoke(main, ["build", "original.sidedoc", "-o", "rebuilt.docx"])
            assert result.exit_code == 0

            rebuilt = Document("rebuilt.docx")
            indented = [p for p in rebuilt.paragraphs if "First line" in p.text]
            assert len(indented) > 0
            assert indented[0].paragraph_format.first_line_indent == Inches(0.25)


class TestStyleWithDirectFormatting:
    """Test that direct formatting overrides style defaults."""

    def test_direct_formatting_overrides_style(self):
        """Direct formatting on a paragraph should override style defaults after round-trip."""
        runner = CliRunner()
        with runner.isolated_filesystem():
            doc = Document()
            para = doc.add_paragraph("Custom spaced heading", style="Heading 1")
            para.paragraph_format.space_before = Pt(24)
            doc.save("original.docx")

            extract_result = runner.invoke(main, ["extract", "original.docx"])
            assert extract_result.exit_code == 0, extract_result.output
            result = runner.invoke(main, ["build", "original.sidedoc", "-o", "rebuilt.docx"])
            assert result.exit_code == 0

            rebuilt = Document("rebuilt.docx")
            heading = [p for p in rebuilt.paragraphs if "Custom spaced" in p.text]
            assert len(heading) > 0
            assert heading[0].style.name == "Heading 1"
            assert heading[0].paragraph_format.space_before == Pt(24)


class TestClaudeReviewRegressions:
    """Regression tests for issues flagged in the Claude bot review on PR #65."""

    # Issue #1: boolean falsy-check dropped explicit False overrides.
    def test_explicit_false_keep_with_next_survives_roundtrip(self):
        """A paragraph that overrides Heading 1's default keep_with_next=True with an
        explicit False must round-trip with the False preserved, not silently dropped.
        """
        runner = CliRunner()
        with runner.isolated_filesystem():
            doc = Document()
            para = doc.add_paragraph("Heading override", style="Heading 1")
            para.paragraph_format.keep_with_next = False
            doc.save("original.docx")

            extract_result = runner.invoke(main, ["extract", "original.docx"])
            assert extract_result.exit_code == 0, extract_result.output
            result = runner.invoke(main, ["build", "original.sidedoc", "-o", "rebuilt.docx"])
            assert result.exit_code == 0

            # Verify styles.json captured the False, not None
            with open("original.sidedoc/styles.json") as f:
                styles_data = json.load(f)
            block_styles = list(styles_data["block_styles"].values())
            assert any(bs.get("keep_with_next") is False for bs in block_styles), \
                f"keep_with_next=False was not persisted to styles.json: {block_styles}"

            rebuilt = Document("rebuilt.docx")
            paras = [p for p in rebuilt.paragraphs if "Heading override" in p.text]
            assert len(paras) > 0
            assert paras[0].paragraph_format.keep_with_next is False

    # Issue #2: int() cast truncated proportional line_spacing (e.g. 1.5 → 1).
    def test_proportional_line_spacing_survives_roundtrip(self):
        """A paragraph with line_spacing=1.5 (proportional) must round-trip as 1.5,
        not be silently truncated to 1 by an int() cast.
        """
        runner = CliRunner()
        with runner.isolated_filesystem():
            doc = Document()
            para = doc.add_paragraph("Proportionally spaced text")
            para.paragraph_format.line_spacing = 1.5
            doc.save("original.docx")

            extract_result = runner.invoke(main, ["extract", "original.docx"])
            assert extract_result.exit_code == 0, extract_result.output
            result = runner.invoke(main, ["build", "original.sidedoc", "-o", "rebuilt.docx"])
            assert result.exit_code == 0

            with open("original.sidedoc/styles.json") as f:
                styles_data = json.load(f)
            line_spacings = [bs.get("line_spacing") for bs in styles_data["block_styles"].values()]
            assert any(ls == pytest.approx(1.5) for ls in line_spacings), \
                f"line_spacing 1.5 was truncated/lost: {line_spacings}"

            rebuilt = Document("rebuilt.docx")
            spaced = [p for p in rebuilt.paragraphs if "Proportionally spaced" in p.text]
            assert len(spaced) > 0
            assert spaced[0].paragraph_format.line_spacing == pytest.approx(1.5)

    # Issue #3: para.style.font.* mutation leaked across paragraphs sharing a style.
    def test_block_style_font_does_not_leak_via_shared_style(self):
        """Applying font_name to one paragraph must not mutate the shared style object
        and bleed into other paragraphs that share the same docx_style.
        """
        from sidedoc.reconstruct import _apply_block_formatting

        doc = Document()
        para_a = doc.add_paragraph("Paragraph A", style="Heading 1")
        para_b = doc.add_paragraph("Paragraph B", style="Heading 1")

        # Snapshot Heading 1's font.name before any mutation so we have a baseline.
        baseline_font = para_b.style.font.name

        _apply_block_formatting(
            para_a,
            {"docx_style": "Heading 1", "font_name": "Courier New", "alignment": "left"},
        )
        _apply_block_formatting(
            para_b,
            {"docx_style": "Heading 1", "alignment": "left"},
        )

        # Pre-fix: para_b.style.font.name is "Courier New" (shared style mutated by A).
        # Post-fix: para_b.style.font.name is unchanged from baseline.
        assert para_b.style.font.name != "Courier New", (
            f"Shared 'Heading 1' style was mutated: font.name leaked to {para_b.style.font.name!r}"
        )
        assert para_b.style.font.name == baseline_font, (
            f"Shared 'Heading 1' style was mutated: was {baseline_font!r}, now {para_b.style.font.name!r}"
        )

    # Issue #8: pre-existing archives without the new paragraph format fields
    # must still build successfully (forward compatibility).
    def test_old_styles_json_without_new_fields_builds(self):
        """A styles.json that predates the JON-91 fields (no left_indent, line_spacing,
        keep_with_next, etc.) must still build without errors.
        """
        runner = CliRunner()
        with runner.isolated_filesystem():
            sidedoc_dir = Path("legacy.sidedoc")
            sidedoc_dir.mkdir()
            (sidedoc_dir / "content.md").write_text("Hello world\n")
            # Minimal styles.json with only the pre-JON-91 fields.
            legacy_styles = {
                "block_styles": {
                    "block-0": {
                        "docx_style": "Normal",
                        "font_name": "Calibri",
                        "font_size": 11,
                        "alignment": "left",
                        "bold": None,
                        "italic": None,
                        "underline": None,
                        "table_formatting": None,
                    },
                },
                "document_defaults": {"font_name": "Calibri", "font_size": 11},
            }
            (sidedoc_dir / "styles.json").write_text(json.dumps(legacy_styles))

            result = runner.invoke(main, ["build", str(sidedoc_dir), "-o", "rebuilt.docx"])
            assert result.exit_code == 0, result.output
            rebuilt = Document("rebuilt.docx")
            texts = [p.text for p in rebuilt.paragraphs]
            assert "Hello world" in texts
