"""Tests for footnote and endnote support.

Test progression follows the track changes pattern:
- Fixtures: Verify fixture files have correct XML structure
- Extract: Verify footnotes extracted as [^N] syntax in content.md
- Metadata: Verify footnote metadata stored in structure.json
- Reconstruct: Verify build creates proper footnote XML from [^N] syntax
- Round-trip: Verify extract -> build preserves footnotes
- Sync: Verify footnote edits handled correctly
"""

import json
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from sidedoc.extract import extract_blocks, blocks_to_markdown
from sidedoc.reconstruct import (
    build_docx_from_sidedoc,
    parse_markdown_to_blocks,
    create_docx_from_blocks,
)
from sidedoc.package import create_sidedoc_directory, create_sidedoc_archive

FIXTURES_DIR = Path(__file__).parent / "fixtures"

FOOTNOTES_RT = (
    "http://schemas.openxmlformats.org/officeDocument"
    "/2006/relationships/footnotes"
)
ENDNOTES_RT = (
    "http://schemas.openxmlformats.org/officeDocument"
    "/2006/relationships/endnotes"
)


# ============================================================================
# Fixture Verification Tests
# ============================================================================


class TestFootnoteFixtures:
    """Verify fixture .docx files have correct XML structure."""

    def test_simple_fixture_has_footnote_references_in_body(self):
        """footnotes_simple.docx body paragraphs contain w:footnoteReference elements."""
        doc = Document(str(FIXTURES_DIR / "footnotes_simple.docx"))
        refs = []
        for para in doc.paragraphs:
            para_refs = para._element.findall(
                ".//" + qn("w:footnoteReference")
            )
            refs.extend(para_refs)
        assert len(refs) == 2
        assert refs[0].get(qn("w:id")) == "1"
        assert refs[1].get(qn("w:id")) == "2"

    def test_simple_fixture_has_footnotes_xml(self):
        """footnotes_simple.docx contains word/footnotes.xml with 2 user footnotes."""
        with zipfile.ZipFile(str(FIXTURES_DIR / "footnotes_simple.docx")) as z:
            assert "word/footnotes.xml" in z.namelist()

    def test_mixed_fixture_has_both_parts(self):
        """footnotes_endnotes_mixed.docx has both footnotes.xml and endnotes.xml."""
        with zipfile.ZipFile(
            str(FIXTURES_DIR / "footnotes_endnotes_mixed.docx")
        ) as z:
            names = z.namelist()
            assert "word/footnotes.xml" in names
            assert "word/endnotes.xml" in names

    def test_consecutive_fixture_has_three_refs(self):
        """footnotes_consecutive.docx has three consecutive footnote refs."""
        doc = Document(str(FIXTURES_DIR / "footnotes_consecutive.docx"))
        refs = doc.paragraphs[0]._element.findall(
            ".//" + qn("w:footnoteReference")
        )
        assert len(refs) == 3


# ============================================================================
# Extract Tests
# ============================================================================


class TestFootnoteExtraction:
    """Verify footnotes are extracted as [^N] syntax."""

    def test_extract_simple_footnotes_in_body(self):
        """Body text contains [^1] and [^2] markers."""
        blocks, _ = extract_blocks(
            str(FIXTURES_DIR / "footnotes_simple.docx")
        )
        md = blocks_to_markdown(blocks)
        assert "[^1]" in md
        assert "[^2]" in md

    def test_extract_simple_footnote_definitions(self):
        """Content.md has footnote definitions at the end."""
        blocks, _ = extract_blocks(
            str(FIXTURES_DIR / "footnotes_simple.docx")
        )
        md = blocks_to_markdown(blocks)
        assert '[^1]: Smith, J. (2024). "Document Processing." Journal of AI, 15(3), 42-58.' in md
        assert "[^2]: This is an explanatory footnote with more detail." in md

    def test_extract_footnote_reference_position(self):
        """[^1] appears right after 'citation' in body text."""
        blocks, _ = extract_blocks(
            str(FIXTURES_DIR / "footnotes_simple.docx")
        )
        md = blocks_to_markdown(blocks)
        # The first block should contain the body text with markers
        body_line = md.split("\n")[0]
        assert "citation[^1]" in body_line
        assert "point[^2]" in body_line

    def test_extract_mixed_footnotes_and_endnotes(self):
        """Both footnotes and endnotes extracted with [^N] syntax."""
        blocks, _ = extract_blocks(
            str(FIXTURES_DIR / "footnotes_endnotes_mixed.docx")
        )
        md = blocks_to_markdown(blocks)
        assert "[^1]" in md
        assert "[^2]" in md
        assert "[^3]" in md
        assert "[^1]: This is a footnote." in md
        assert "[^2]: This is an endnote." in md
        assert "[^3]: Second endnote content." in md

    def test_extract_consecutive_footnotes(self):
        """Consecutive footnote refs [^1][^2][^3] preserved."""
        blocks, _ = extract_blocks(
            str(FIXTURES_DIR / "footnotes_consecutive.docx")
        )
        md = blocks_to_markdown(blocks)
        assert "[^1][^2][^3]" in md

    def test_extract_footnote_in_heading(self):
        """Footnote in heading extracted correctly."""
        blocks, _ = extract_blocks(
            str(FIXTURES_DIR / "footnotes_in_heading.docx")
        )
        md = blocks_to_markdown(blocks)
        assert "# Chapter Title[^1]" in md
        assert "[^1]: Chapter footnote content." in md

    def test_extract_multiblock_footnotes_sequential_numbering(self):
        """Footnotes across blocks numbered sequentially."""
        blocks, _ = extract_blocks(
            str(FIXTURES_DIR / "footnotes_multiblock.docx")
        )
        md = blocks_to_markdown(blocks)
        assert "[^1]" in md
        assert "[^2]" in md
        assert "[^3]" in md
        assert "[^1]: Introduction source." in md
        assert "[^2]: Methods reference one." in md
        assert "[^3]: Methods reference two." in md

    def test_extract_formatted_footnote_content(self):
        """Footnotes with bold/italic formatting preserved."""
        blocks, _ = extract_blocks(
            str(FIXTURES_DIR / "footnotes_formatted.docx")
        )
        md = blocks_to_markdown(blocks)
        assert "[^1]: **This footnote has bold text.**" in md
        assert "[^2]: *This footnote has italic text.*" in md


# ============================================================================
# Footnote Metadata Tests
# ============================================================================


class TestFootnoteMetadata:
    """Verify footnote metadata stored on blocks."""

    def test_block_has_footnote_references(self):
        """Blocks with footnotes store footnote_references metadata."""
        blocks, _ = extract_blocks(
            str(FIXTURES_DIR / "footnotes_simple.docx")
        )
        # First block should have footnote references
        body_block = blocks[0]
        assert body_block.footnote_references is not None
        assert len(body_block.footnote_references) == 2

    def test_footnote_reference_has_id_and_type(self):
        """Each footnote reference has note_id and note_type."""
        blocks, _ = extract_blocks(
            str(FIXTURES_DIR / "footnotes_simple.docx")
        )
        refs = blocks[0].footnote_references
        assert refs[0]["note_id"] == 1
        assert refs[0]["note_type"] == "footnote"

    def test_endnote_reference_has_correct_type(self):
        """Endnote references have note_type='endnote'."""
        blocks, _ = extract_blocks(
            str(FIXTURES_DIR / "footnotes_endnotes_mixed.docx")
        )
        md = blocks_to_markdown(blocks)
        # Find the block with the endnote
        body_block = blocks[0]
        refs = body_block.footnote_references
        # Should have footnote (id=1) and endnote (id=2)
        types = {r["note_type"] for r in refs}
        assert "footnote" in types
        assert "endnote" in types


# ============================================================================
# Reconstruct Tests
# ============================================================================


class TestFootnoteReconstruction:
    """Verify build creates proper footnote XML from [^N] syntax."""

    def test_build_creates_footnote_references_in_body(self, tmp_path):
        """Built docx has w:footnoteReference in body paragraphs."""
        content_md = (
            "This is a claim[^1]. Another point[^2].\n\n"
            '[^1]: Smith, J. (2024). "Document Processing."\n'
            "[^2]: Explanatory footnote.\n"
        )
        blocks = parse_markdown_to_blocks(content_md)
        styles = {"block_styles": {}, "document_defaults": {"font_name": "Calibri", "font_size": 11}}
        doc = create_docx_from_blocks(blocks, styles, content_md=content_md)
        output = str(tmp_path / "output.docx")
        doc.save(output)

        # Verify the output has footnote references
        rebuilt = Document(output)
        refs = []
        for para in rebuilt.paragraphs:
            refs.extend(
                para._element.findall(".//" + qn("w:footnoteReference"))
            )
        assert len(refs) == 2

    def test_build_creates_footnotes_xml(self, tmp_path):
        """Built docx contains word/footnotes.xml with correct content."""
        content_md = (
            "Text with note[^1].\n\n"
            "[^1]: Footnote content here.\n"
        )
        blocks = parse_markdown_to_blocks(content_md)
        styles = {"block_styles": {}, "document_defaults": {"font_name": "Calibri", "font_size": 11}}
        doc = create_docx_from_blocks(blocks, styles, content_md=content_md)
        output = str(tmp_path / "output.docx")
        doc.save(output)

        with zipfile.ZipFile(output) as z:
            assert "word/footnotes.xml" in z.namelist()
            fn_xml = z.read("word/footnotes.xml").decode()
            assert "Footnote content here." in fn_xml

    def test_build_creates_endnotes_from_metadata(self, tmp_path):
        """Built docx creates endnotes.xml for notes marked as endnotes in metadata."""
        # Create sidedoc with endnote metadata
        content_md = (
            "Text with endnote[^1].\n\n"
            "[^1]: Endnote content.\n"
        )
        sidedoc_dir = tmp_path / "test.sidedoc"
        sidedoc_dir.mkdir()
        (sidedoc_dir / "content.md").write_text(content_md)
        structure = {
            "blocks": [{
                "id": "block-0",
                "type": "paragraph",
                "docx_paragraph_index": 0,
                "content_start": 0,
                "content_end": len("Text with endnote[^1]."),
                "content_hash": "abc",
                "level": None,
                "image_path": None,
                "inline_formatting": None,
                "table_metadata": None,
                "track_changes": None,
                "footnote_references": [{"note_id": 1, "note_type": "endnote", "marker": "[^1]"}],
            }],
            "footnotes": {
                "1": {
                    "content": "Endnote content.",
                    "note_type": "endnote",
                    "original_id": "1",
                }
            },
        }
        (sidedoc_dir / "structure.json").write_text(json.dumps(structure))
        (sidedoc_dir / "styles.json").write_text(json.dumps({
            "block_styles": {},
            "document_defaults": {"font_name": "Calibri", "font_size": 11},
        }))
        (sidedoc_dir / "manifest.json").write_text(json.dumps({
            "sidedoc_version": "1.0.0",
            "created_at": "2026-01-01T00:00:00+00:00",
            "modified_at": "2026-01-01T00:00:00+00:00",
            "source_file": "test.docx",
            "source_hash": "abc",
            "content_hash": "abc",
            "generator": "sidedoc-cli/0.1.0",
        }))

        output = str(tmp_path / "output.docx")
        build_docx_from_sidedoc(str(sidedoc_dir), output)

        with zipfile.ZipFile(output) as z:
            assert "word/endnotes.xml" in z.namelist()
            en_xml = z.read("word/endnotes.xml").decode()
            assert "Endnote content." in en_xml


# ============================================================================
# Round-Trip Tests
# ============================================================================


class TestFootnoteRoundTrip:
    """Verify extract -> build preserves footnotes."""

    def test_roundtrip_simple_footnotes(self, tmp_path):
        """Extract then build preserves footnote content and references."""
        blocks, images = extract_blocks(
            str(FIXTURES_DIR / "footnotes_simple.docx")
        )
        md = blocks_to_markdown(blocks)
        from sidedoc.extract import extract_styles
        styles = extract_styles(str(FIXTURES_DIR / "footnotes_simple.docx"), blocks)

        # Create sidedoc directory and build
        sidedoc_dir = tmp_path / "test.sidedoc"
        create_sidedoc_directory(
            str(sidedoc_dir), md, blocks, styles,
            str(FIXTURES_DIR / "footnotes_simple.docx"),
        )

        output = str(tmp_path / "output.docx")
        build_docx_from_sidedoc(str(sidedoc_dir), output)

        # Verify output has footnotes
        rebuilt = Document(output)
        refs = []
        for para in rebuilt.paragraphs:
            refs.extend(
                para._element.findall(".//" + qn("w:footnoteReference"))
            )
        assert len(refs) >= 2

        with zipfile.ZipFile(output) as z:
            assert "word/footnotes.xml" in z.namelist()

    def test_roundtrip_mixed_footnotes_endnotes(self, tmp_path):
        """Extract then build preserves both footnotes and endnotes."""
        blocks, images = extract_blocks(
            str(FIXTURES_DIR / "footnotes_endnotes_mixed.docx")
        )
        md = blocks_to_markdown(blocks)
        from sidedoc.extract import extract_styles
        styles = extract_styles(str(FIXTURES_DIR / "footnotes_endnotes_mixed.docx"), blocks)

        sidedoc_dir = tmp_path / "test.sidedoc"
        create_sidedoc_directory(
            str(sidedoc_dir), md, blocks, styles,
            str(FIXTURES_DIR / "footnotes_endnotes_mixed.docx"),
        )

        output = str(tmp_path / "output.docx")
        build_docx_from_sidedoc(str(sidedoc_dir), output)

        with zipfile.ZipFile(output) as z:
            names = z.namelist()
            assert "word/footnotes.xml" in names
            assert "word/endnotes.xml" in names

    def test_roundtrip_preserves_footnote_content(self, tmp_path):
        """Round-trip preserves actual footnote text content."""
        blocks, images = extract_blocks(
            str(FIXTURES_DIR / "footnotes_simple.docx")
        )
        md = blocks_to_markdown(blocks)
        from sidedoc.extract import extract_styles
        styles = extract_styles(str(FIXTURES_DIR / "footnotes_simple.docx"), blocks)

        sidedoc_dir = tmp_path / "test.sidedoc"
        create_sidedoc_directory(
            str(sidedoc_dir), md, blocks, styles,
            str(FIXTURES_DIR / "footnotes_simple.docx"),
        )

        output = str(tmp_path / "output.docx")
        build_docx_from_sidedoc(str(sidedoc_dir), output)

        # Re-extract and verify content preserved
        blocks2, _ = extract_blocks(output)
        md2 = blocks_to_markdown(blocks2)
        assert "[^1]" in md2
        assert "[^2]" in md2
        assert "Document Processing" in md2

    def test_roundtrip_preserves_footnote_formatting(self, tmp_path):
        """Round-trip preserves bold/italic formatting in footnote text."""
        blocks, images = extract_blocks(
            str(FIXTURES_DIR / "footnotes_formatted.docx")
        )
        md = blocks_to_markdown(blocks)
        from sidedoc.extract import extract_styles
        styles = extract_styles(str(FIXTURES_DIR / "footnotes_formatted.docx"), blocks)

        sidedoc_dir = tmp_path / "test.sidedoc"
        create_sidedoc_directory(
            str(sidedoc_dir), md, blocks, styles,
            str(FIXTURES_DIR / "footnotes_formatted.docx"),
        )

        output = str(tmp_path / "output.docx")
        build_docx_from_sidedoc(str(sidedoc_dir), output)

        # Re-extract and verify bold/italic markers survived
        blocks2, _ = extract_blocks(output)
        md2 = blocks_to_markdown(blocks2)
        # The formatted fixture has bold and/or italic in footnote text
        assert "[^1]" in md2
        # Check that formatting markers exist in footnote definitions
        footnote_defs = [line for line in md2.split("\n") if line.startswith("[^")]
        assert len(footnote_defs) > 0
        # At least one definition should contain formatting markers
        has_formatting = any("**" in d or "*" in d for d in footnote_defs)
        assert has_formatting, f"Expected formatting in footnote defs, got: {footnote_defs}"


# ============================================================================
# Sync Tests
# ============================================================================


class TestFootnoteSync:
    """Verify sync handles footnote edits."""

    def test_sync_modified_footnote_content(self, tmp_path):
        """Editing a footnote definition in content.md updates the docx."""
        from sidedoc.sync import sync_sidedoc_to_docx
        from tests.helpers import create_sidedoc_dir

        content_md = (
            "Text with note[^1].\n\n"
            "[^1]: Original footnote text.\n"
        )
        sidedoc_dir = tmp_path / "test.sidedoc"
        structure = {
            "blocks": [{
                "id": "block-0",
                "type": "paragraph",
                "docx_paragraph_index": 0,
                "content_start": 0,
                "content_end": 22,
                "content_hash": "abc",
                "level": None,
                "image_path": None,
                "inline_formatting": None,
                "table_metadata": None,
                "track_changes": None,
                "footnote_references": [{"note_id": 1, "note_type": "footnote", "marker": "[^1]"}],
            }],
            "footnotes": {
                "1": {
                    "content": "Original footnote text.",
                    "note_type": "footnote",
                    "original_id": "1",
                }
            },
        }
        create_sidedoc_dir(sidedoc_dir, content_md, structure)

        # Edit: change the footnote text
        edited_md = (
            "Text with note[^1].\n\n"
            "[^1]: Updated footnote text.\n"
        )
        (sidedoc_dir / "content.md").write_text(edited_md)

        output = str(tmp_path / "output.docx")
        sync_sidedoc_to_docx(str(sidedoc_dir), output)

        with zipfile.ZipFile(output) as z:
            fn_xml = z.read("word/footnotes.xml").decode()
            assert "Updated footnote text." in fn_xml

    def test_sync_added_footnote(self, tmp_path):
        """Adding a new footnote in content.md creates it in docx."""
        from sidedoc.sync import sync_sidedoc_to_docx
        from tests.helpers import create_sidedoc_dir

        content_md = "Text without notes.\n"
        sidedoc_dir = tmp_path / "test.sidedoc"
        structure = {
            "blocks": [{
                "id": "block-0",
                "type": "paragraph",
                "docx_paragraph_index": 0,
                "content_start": 0,
                "content_end": 19,
                "content_hash": "abc",
                "level": None,
                "image_path": None,
                "inline_formatting": None,
                "table_metadata": None,
                "track_changes": None,
                "footnote_references": None,
            }],
        }
        create_sidedoc_dir(sidedoc_dir, content_md, structure)

        # Edit: add a footnote
        edited_md = (
            "Text with new note[^1].\n\n"
            "[^1]: Brand new footnote.\n"
        )
        (sidedoc_dir / "content.md").write_text(edited_md)

        output = str(tmp_path / "output.docx")
        sync_sidedoc_to_docx(str(sidedoc_dir), output)

        with zipfile.ZipFile(output) as z:
            assert "word/footnotes.xml" in z.namelist()
            fn_xml = z.read("word/footnotes.xml").decode()
            assert "Brand new footnote." in fn_xml

    def test_sync_removed_footnote(self, tmp_path):
        """Removing a footnote from content.md removes it from docx."""
        from sidedoc.sync import sync_sidedoc_to_docx
        from tests.helpers import create_sidedoc_dir

        content_md = (
            "Text with note[^1].\n\n"
            "[^1]: Will be removed.\n"
        )
        sidedoc_dir = tmp_path / "test.sidedoc"
        structure = {
            "blocks": [{
                "id": "block-0",
                "type": "paragraph",
                "docx_paragraph_index": 0,
                "content_start": 0,
                "content_end": 22,
                "content_hash": "abc",
                "level": None,
                "image_path": None,
                "inline_formatting": None,
                "table_metadata": None,
                "track_changes": None,
                "footnote_references": [{"note_id": 1, "note_type": "footnote", "marker": "[^1]"}],
            }],
            "footnotes": {
                "1": {
                    "content": "Will be removed.",
                    "note_type": "footnote",
                    "original_id": "1",
                }
            },
        }
        create_sidedoc_dir(sidedoc_dir, content_md, structure)

        # Edit: remove the footnote
        edited_md = "Text without note.\n"
        (sidedoc_dir / "content.md").write_text(edited_md)

        output = str(tmp_path / "output.docx")
        sync_sidedoc_to_docx(str(sidedoc_dir), output)

        # Body should not have footnote references
        rebuilt = Document(output)
        refs = []
        for para in rebuilt.paragraphs:
            refs.extend(
                para._element.findall(".//" + qn("w:footnoteReference"))
            )
        assert len(refs) == 0

        # Verify footnotes.xml has no user footnotes (only separator elements if present)
        footnotes_rt = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
        has_footnotes_part = False
        for rel in rebuilt.part.rels.values():
            if rel.reltype == footnotes_rt:
                has_footnotes_part = True
                from lxml import etree as et
                root = et.fromstring(rel.target_part.blob)
                ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                user_notes = [
                    fn for fn in root.findall("w:footnote", ns)
                    if fn.get(et.QName(ns["w"], "type")) is None  # separator notes have a type attr
                ]
                assert len(user_notes) == 0, f"Expected no user footnotes, found {len(user_notes)}"
                break

        if not has_footnotes_part:
            # Part was removed entirely — also a valid outcome
            pass


# ============================================================================
# No-Regression Tests
# ============================================================================


class TestFootnoteNoRegression:
    """Verify footnote changes don't break existing extraction."""

    def test_simple_docx_still_extracts(self):
        """Documents without footnotes still extract correctly."""
        blocks, _ = extract_blocks(str(FIXTURES_DIR / "simple.docx"))
        md = blocks_to_markdown(blocks)
        assert "Main Title" in md
        assert "[^" not in md

    def test_hyperlinks_docx_still_extracts(self):
        """Documents with hyperlinks still extract correctly."""
        blocks, _ = extract_blocks(str(FIXTURES_DIR / "hyperlinks.docx"))
        md = blocks_to_markdown(blocks)
        assert "[" in md  # hyperlinks present


# ============================================================================
# Bug Regression Tests
# ============================================================================


class TestFootnoteBugFixes:
    """Tests targeting specific bugs found in code review."""

    def test_multiple_footnotes_all_present_in_xml(self, tmp_path):
        """All footnotes survive when multiple are added to a new part (bug #2).

        Previously, _add_footnote_to_part re-parsed part.blob each call,
        so only the last footnote's content appeared in the serialized XML.
        """
        content_md = (
            "First claim[^1]. Second claim[^2]. Third claim[^3].\n\n"
            "[^1]: First source.\n"
            "[^2]: Second source.\n"
            "[^3]: Third source.\n"
        )
        blocks = parse_markdown_to_blocks(content_md)
        styles = {"block_styles": {}, "document_defaults": {"font_name": "Calibri", "font_size": 11}}
        doc = create_docx_from_blocks(blocks, styles, content_md=content_md)
        output = str(tmp_path / "output.docx")
        doc.save(output)

        import zipfile
        with zipfile.ZipFile(output) as z:
            fn_xml = z.read("word/footnotes.xml").decode()
            assert "First source." in fn_xml
            assert "Second source." in fn_xml
            assert "Third source." in fn_xml

    def test_endnote_type_preserved_on_roundtrip(self, tmp_path):
        """Endnotes retain note_type='endnote' after extract -> build (bug #4).

        Previously, structure.json had no top-level 'footnotes' key, so
        endnotes defaulted to footnotes on round-trip.
        """
        blocks, images = extract_blocks(
            str(FIXTURES_DIR / "footnotes_endnotes_mixed.docx")
        )
        md = blocks_to_markdown(blocks)
        from sidedoc.extract import extract_styles
        styles = extract_styles(
            str(FIXTURES_DIR / "footnotes_endnotes_mixed.docx"), blocks
        )

        sidedoc_dir = tmp_path / "test.sidedoc"
        create_sidedoc_directory(
            str(sidedoc_dir), md, blocks, styles,
            str(FIXTURES_DIR / "footnotes_endnotes_mixed.docx"),
        )

        # Verify structure.json has top-level footnotes key with endnote types
        structure = json.loads(
            (sidedoc_dir / "structure.json").read_text(encoding="utf-8")
        )
        assert "footnotes" in structure
        endnote_types = [
            v["note_type"] for v in structure["footnotes"].values()
        ]
        assert "endnote" in endnote_types

        # Build and verify endnotes.xml is present (not just footnotes.xml)
        output = str(tmp_path / "output.docx")
        build_docx_from_sidedoc(str(sidedoc_dir), output)

        import zipfile
        with zipfile.ZipFile(output) as z:
            assert "word/endnotes.xml" in z.namelist()

    def test_continuation_separator_present_in_footnotes_xml(self, tmp_path):
        """Built footnotes.xml has w:continuationSeparator element (bug #3)."""
        content_md = (
            "Text[^1].\n\n"
            "[^1]: A footnote.\n"
        )
        blocks = parse_markdown_to_blocks(content_md)
        styles = {"block_styles": {}, "document_defaults": {"font_name": "Calibri", "font_size": 11}}
        doc = create_docx_from_blocks(blocks, styles, content_md=content_md)
        output = str(tmp_path / "output.docx")
        doc.save(output)

        import zipfile
        with zipfile.ZipFile(output) as z:
            fn_xml = z.read("word/footnotes.xml").decode()
            assert "continuationSeparator" in fn_xml

    def test_multiblock_footnotes_correct_paragraph_assignment(self):
        """Footnotes across blocks are assigned to the correct paragraphs (#10)."""
        blocks, _ = extract_blocks(
            str(FIXTURES_DIR / "footnotes_multiblock.docx")
        )
        # Collect which blocks own which footnote IDs
        block_footnote_map: dict[int, set[int]] = {}
        for i, block in enumerate(blocks):
            if block.footnote_references:
                block_footnote_map[i] = {r["note_id"] for r in block.footnote_references}

        # There should be at least 2 different blocks with footnotes
        assert len(block_footnote_map) >= 2, (
            f"Expected footnotes spread across blocks, got: {block_footnote_map}"
        )

        # All three footnotes should be present across blocks
        all_ids = set()
        for ids in block_footnote_map.values():
            all_ids |= ids
        assert {1, 2, 3} == all_ids

        # [^1] and [^2]/[^3] should be in different blocks
        block_with_1 = [i for i, ids in block_footnote_map.items() if 1 in ids][0]
        blocks_with_2_or_3 = [i for i, ids in block_footnote_map.items() if ids & {2, 3}]
        assert all(b != block_with_1 for b in blocks_with_2_or_3)
