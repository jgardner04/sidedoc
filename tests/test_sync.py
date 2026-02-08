"""Tests for sync module - block matching algorithm."""

import hashlib
import json
import tempfile
import zipfile
from pathlib import Path
from docx import Document
from sidedoc.models import Block
from sidedoc.sync import match_blocks, generate_updated_docx, update_sidedoc_metadata, remap_styles


def compute_hash(content: str) -> str:
    """Helper to compute SHA256 hash of content."""
    return hashlib.sha256(content.encode()).hexdigest()


def test_match_unchanged_blocks():
    """Test that unchanged blocks are matched by content hash."""
    old_blocks = [
        Block(
            id="block-1",
            type="paragraph",
            content="This is unchanged.",
            docx_paragraph_index=0,
            content_start=0,
            content_end=19,
            content_hash=compute_hash("This is unchanged."),
        ),
    ]

    new_blocks = [
        Block(
            id="block-new-1",
            type="paragraph",
            content="This is unchanged.",
            docx_paragraph_index=0,
            content_start=0,
            content_end=19,
            content_hash=compute_hash("This is unchanged."),
        ),
    ]

    matches = match_blocks(old_blocks, new_blocks)

    assert "block-1" in matches
    assert matches["block-1"] == new_blocks[0]
    assert matches["block-1"].content == "This is unchanged."


def test_match_edited_blocks_by_position():
    """Test that edited blocks are matched by type, position, and similarity."""
    old_blocks = [
        Block(
            id="block-1",
            type="paragraph",
            content="Original text for testing.",
            docx_paragraph_index=0,
            content_start=0,
            content_end=26,
            content_hash=compute_hash("Original text for testing."),
        ),
    ]

    new_blocks = [
        Block(
            id="block-new-1",
            type="paragraph",
            content="Original text for testing edits.",
            docx_paragraph_index=0,
            content_start=0,
            content_end=32,
            content_hash=compute_hash("Original text for testing edits."),
        ),
    ]

    matches = match_blocks(old_blocks, new_blocks)

    # Should match by position when content has high similarity (edit)
    assert "block-1" in matches
    assert matches["block-1"] == new_blocks[0]
    assert matches["block-1"].content == "Original text for testing edits."


def test_identify_new_blocks():
    """Test that new blocks have no corresponding old block."""
    old_blocks = [
        Block(
            id="block-1",
            type="paragraph",
            content="First paragraph.",
            docx_paragraph_index=0,
            content_start=0,
            content_end=16,
            content_hash=compute_hash("First paragraph."),
        ),
    ]

    new_blocks = [
        Block(
            id="block-new-1",
            type="paragraph",
            content="First paragraph.",
            docx_paragraph_index=0,
            content_start=0,
            content_end=16,
            content_hash=compute_hash("First paragraph."),
        ),
        Block(
            id="block-new-2",
            type="paragraph",
            content="New paragraph.",
            docx_paragraph_index=1,
            content_start=17,
            content_end=31,
            content_hash=compute_hash("New paragraph."),
        ),
    ]

    matches = match_blocks(old_blocks, new_blocks)

    # First block should match
    assert "block-1" in matches
    # New block should not be in matches (it's truly new)
    # The matches dict only contains old block IDs
    assert len(matches) == 1


def test_identify_deleted_blocks():
    """Test that deleted blocks are identified when old blocks have no match."""
    old_blocks = [
        Block(
            id="block-1",
            type="paragraph",
            content="Keep this.",
            docx_paragraph_index=0,
            content_start=0,
            content_end=10,
            content_hash=compute_hash("Keep this."),
        ),
        Block(
            id="block-2",
            type="paragraph",
            content="Delete this.",
            docx_paragraph_index=1,
            content_start=11,
            content_end=23,
            content_hash=compute_hash("Delete this."),
        ),
    ]

    new_blocks = [
        Block(
            id="block-new-1",
            type="paragraph",
            content="Keep this.",
            docx_paragraph_index=0,
            content_start=0,
            content_end=10,
            content_hash=compute_hash("Keep this."),
        ),
    ]

    matches = match_blocks(old_blocks, new_blocks)

    # Only first block should be matched
    assert "block-1" in matches
    assert "block-2" not in matches
    assert len(matches) == 1


def test_match_multiple_blocks_complex():
    """Test complex scenario with unchanged, edited, new, and deleted blocks."""
    old_blocks = [
        Block(
            id="block-1",
            type="heading",
            content="# Title",
            docx_paragraph_index=0,
            content_start=0,
            content_end=7,
            content_hash=compute_hash("# Title"),
            level=1,
        ),
        Block(
            id="block-2",
            type="paragraph",
            content="First para.",
            docx_paragraph_index=1,
            content_start=8,
            content_end=19,
            content_hash=compute_hash("First para."),
        ),
        Block(
            id="block-3",
            type="paragraph",
            content="Delete me.",
            docx_paragraph_index=2,
            content_start=20,
            content_end=30,
            content_hash=compute_hash("Delete me."),
        ),
    ]

    new_blocks = [
        Block(
            id="block-new-1",
            type="heading",
            content="# Title",
            docx_paragraph_index=0,
            content_start=0,
            content_end=7,
            content_hash=compute_hash("# Title"),
            level=1,
        ),
        Block(
            id="block-new-2",
            type="paragraph",
            content="First para modified.",
            docx_paragraph_index=1,
            content_start=8,
            content_end=28,
            content_hash=compute_hash("First para modified."),
        ),
        Block(
            id="block-new-3",
            type="paragraph",
            content="Brand new para.",
            docx_paragraph_index=2,
            content_start=29,
            content_end=44,
            content_hash=compute_hash("Brand new para."),
        ),
    ]

    matches = match_blocks(old_blocks, new_blocks)

    # block-1 unchanged (matched by hash)
    assert "block-1" in matches
    assert matches["block-1"].content == "# Title"

    # block-2 edited with high similarity (matched by position and similarity)
    assert "block-2" in matches
    assert matches["block-2"].content == "First para modified."

    # block-3 NOT matched - low similarity means it's treated as delete + add
    # This is the improvement from issue #11 - the algorithm now uses similarity
    # to distinguish "delete + add" from "edit"
    assert "block-3" not in matches

    # Total matches = 2 (block-1 by hash, block-2 by position+similarity)
    # block-3 was deleted, block-new-3 is a new addition
    assert len(matches) == 2


def test_true_deletion_when_blocks_decrease():
    """Test that blocks are truly deleted when new content has fewer blocks."""
    old_blocks = [
        Block(
            id="block-1",
            type="paragraph",
            content="Keep this.",
            docx_paragraph_index=0,
            content_start=0,
            content_end=10,
            content_hash=compute_hash("Keep this."),
        ),
        Block(
            id="block-2",
            type="paragraph",
            content="Also keep.",
            docx_paragraph_index=1,
            content_start=11,
            content_end=21,
            content_hash=compute_hash("Also keep."),
        ),
        Block(
            id="block-3",
            type="paragraph",
            content="Delete this.",
            docx_paragraph_index=2,
            content_start=22,
            content_end=34,
            content_hash=compute_hash("Delete this."),
        ),
    ]

    new_blocks = [
        Block(
            id="block-new-1",
            type="paragraph",
            content="Keep this.",
            docx_paragraph_index=0,
            content_start=0,
            content_end=10,
            content_hash=compute_hash("Keep this."),
        ),
        Block(
            id="block-new-2",
            type="paragraph",
            content="Also keep.",
            docx_paragraph_index=1,
            content_start=11,
            content_end=21,
            content_hash=compute_hash("Also keep."),
        ),
    ]

    matches = match_blocks(old_blocks, new_blocks)

    # First two blocks match by hash
    assert "block-1" in matches
    assert "block-2" in matches
    # Third block is truly deleted (no block at position 2)
    assert "block-3" not in matches
    assert len(matches) == 2


def test_match_respects_type_when_matching_by_position():
    """Test that blocks with different types don't match even at same position."""
    old_blocks = [
        Block(
            id="block-1",
            type="paragraph",
            content="Text.",
            docx_paragraph_index=0,
            content_start=0,
            content_end=5,
            content_hash=compute_hash("Text."),
        ),
    ]

    new_blocks = [
        Block(
            id="block-new-1",
            type="heading",
            content="# Text.",
            docx_paragraph_index=0,
            content_start=0,
            content_end=7,
            content_hash=compute_hash("# Text."),
            level=1,
        ),
    ]

    matches = match_blocks(old_blocks, new_blocks)

    # Should not match because types differ
    assert len(matches) == 0


# Tests for generate_updated_docx


def test_generate_updated_docx_with_unchanged_blocks() -> None:
    """Test that unchanged blocks preserve their content."""
    new_blocks = [
        Block(
            id="block-new-1",
            type="paragraph",
            content="Unchanged paragraph.",
            docx_paragraph_index=0,
            content_start=0,
            content_end=20,
            content_hash=compute_hash("Unchanged paragraph."),
        ),
    ]

    styles = {
        "block_styles": {
            "block-1": {
                "font_name": "Arial",
                "font_size": 12,
                "alignment": "left",
            }
        }
    }

    matches = {
        "block-1": new_blocks[0],
    }

    with tempfile.TemporaryDirectory() as temp_dir:
        output_path = Path(temp_dir) / "test.docx"
        generate_updated_docx(new_blocks, matches, styles, str(output_path))

        # Verify docx was created
        assert output_path.exists()

        # Verify content
        doc = Document(str(output_path))
        assert len(doc.paragraphs) == 1
        assert doc.paragraphs[0].text == "Unchanged paragraph."


def test_generate_updated_docx_with_edited_blocks() -> None:
    """Test that edited blocks get updated content but preserve formatting."""
    old_block_id = "block-1"
    new_blocks = [
        Block(
            id="block-new-1",
            type="paragraph",
            content="Modified paragraph.",
            docx_paragraph_index=0,
            content_start=0,
            content_end=19,
            content_hash=compute_hash("Modified paragraph."),
        ),
    ]

    styles = {
        "block_styles": {
            old_block_id: {
                "font_name": "Times New Roman",
                "font_size": 14,
                "alignment": "center",
            }
        }
    }

    matches = {
        old_block_id: new_blocks[0],
    }

    with tempfile.TemporaryDirectory() as temp_dir:
        output_path = Path(temp_dir) / "test.docx"
        generate_updated_docx(new_blocks, matches, styles, str(output_path))

        # Verify docx was created
        assert output_path.exists()

        # Verify content was updated
        doc = Document(str(output_path))
        assert len(doc.paragraphs) == 1
        assert doc.paragraphs[0].text == "Modified paragraph."


def test_generate_updated_docx_with_new_blocks() -> None:
    """Test that new blocks receive default formatting."""
    new_blocks = [
        Block(
            id="block-new-1",
            type="paragraph",
            content="Existing paragraph.",
            docx_paragraph_index=0,
            content_start=0,
            content_end=19,
            content_hash=compute_hash("Existing paragraph."),
        ),
        Block(
            id="block-new-2",
            type="paragraph",
            content="Brand new paragraph.",
            docx_paragraph_index=1,
            content_start=20,
            content_end=40,
            content_hash=compute_hash("Brand new paragraph."),
        ),
    ]

    styles = {
        "block_styles": {
            "block-1": {
                "font_name": "Arial",
                "font_size": 12,
                "alignment": "left",
            }
        }
    }

    matches = {
        "block-1": new_blocks[0],
    }

    with tempfile.TemporaryDirectory() as temp_dir:
        output_path = Path(temp_dir) / "test.docx"
        generate_updated_docx(new_blocks, matches, styles, str(output_path))

        # Verify docx was created
        assert output_path.exists()

        # Verify both blocks present
        doc = Document(str(output_path))
        assert len(doc.paragraphs) == 2
        assert doc.paragraphs[0].text == "Existing paragraph."
        assert doc.paragraphs[1].text == "Brand new paragraph."


def test_generate_updated_docx_deleted_blocks_omitted() -> None:
    """Test that deleted blocks are not included in output."""
    # Old blocks: block-1, block-2
    # New blocks: only block-new-1 (which matches block-1)
    # Result: block-2 should be deleted (not in output)

    new_blocks = [
        Block(
            id="block-new-1",
            type="paragraph",
            content="Keep this.",
            docx_paragraph_index=0,
            content_start=0,
            content_end=10,
            content_hash=compute_hash("Keep this."),
        ),
    ]

    styles = {
        "block_styles": {
            "block-1": {"font_name": "Arial", "font_size": 12, "alignment": "left"},
            "block-2": {"font_name": "Arial", "font_size": 12, "alignment": "left"},
        }
    }

    matches = {
        "block-1": new_blocks[0],
        # block-2 not in matches = deleted
    }

    with tempfile.TemporaryDirectory() as temp_dir:
        output_path = Path(temp_dir) / "test.docx"
        generate_updated_docx(new_blocks, matches, styles, str(output_path))

        # Verify only one paragraph (deleted block omitted)
        doc = Document(str(output_path))
        assert len(doc.paragraphs) == 1
        assert doc.paragraphs[0].text == "Keep this."


def test_generate_updated_docx_with_inline_formatting() -> None:
    """Test that inline formatting from markdown is applied."""
    new_blocks = [
        Block(
            id="block-new-1",
            type="paragraph",
            content="This is **bold** and *italic* text.",
            docx_paragraph_index=0,
            content_start=0,
            content_end=36,
            content_hash=compute_hash("This is **bold** and *italic* text."),
        ),
    ]

    styles = {"block_styles": {}}
    matches = {}

    with tempfile.TemporaryDirectory() as temp_dir:
        output_path = Path(temp_dir) / "test.docx"
        generate_updated_docx(new_blocks, matches, styles, str(output_path))

        # Verify docx was created and has content
        assert output_path.exists()
        doc = Document(str(output_path))
        assert len(doc.paragraphs) == 1
        # Note: Detailed inline formatting verification would require
        # checking runs, but basic content check is sufficient for now


# Tests for update_sidedoc_metadata


def test_update_sidedoc_metadata_regenerates_structure() -> None:
    """Test that structure.json is regenerated with new block info."""
    with tempfile.TemporaryDirectory() as temp_dir:
        sidedoc_path = Path(temp_dir) / "test.sidedoc"

        # Create initial sidedoc archive
        old_content = "# Old Title\n\nOld paragraph."
        old_structure = {
            "blocks": [
                {
                    "id": "block-1",
                    "type": "heading",
                    "docx_paragraph_index": 0,
                    "content_start": 0,
                    "content_end": 11,
                    "content_hash": compute_hash("# Old Title"),
                    "level": 1,
                    "image_path": None,
                    "inline_formatting": None,
                }
            ]
        }
        old_styles = {"block_styles": {}, "document_defaults": {"font_name": "Arial", "font_size": 11}}
        old_manifest = {
            "sidedoc_version": "1.0.0",
            "created_at": "2024-01-01T00:00:00+00:00",
            "modified_at": "2024-01-01T00:00:00+00:00",
            "source_file": "test.docx",
            "source_hash": "abc123",
            "content_hash": "old_hash",
            "generator": "sidedoc-cli/0.1.0",
        }

        with zipfile.ZipFile(sidedoc_path, "w") as zip_file:
            zip_file.writestr("content.md", old_content)
            zip_file.writestr("structure.json", json.dumps(old_structure))
            zip_file.writestr("styles.json", json.dumps(old_styles))
            zip_file.writestr("manifest.json", json.dumps(old_manifest))

        # New content with edited blocks
        new_content = "# New Title\n\nNew paragraph."
        new_blocks = [
            Block(
                id="block-new-1",
                type="heading",
                content="# New Title",
                docx_paragraph_index=0,
                content_start=0,
                content_end=11,
                content_hash=compute_hash("# New Title"),
                level=1,
            ),
            Block(
                id="block-new-2",
                type="paragraph",
                content="New paragraph.",
                docx_paragraph_index=1,
                content_start=13,
                content_end=27,
                content_hash=compute_hash("New paragraph."),
            ),
        ]

        # Update metadata
        update_sidedoc_metadata(str(sidedoc_path), new_blocks, new_content)

        # Verify structure.json was updated
        with zipfile.ZipFile(sidedoc_path, "r") as zip_file:
            structure_data = json.loads(zip_file.read("structure.json"))
            assert len(structure_data["blocks"]) == 2
            assert structure_data["blocks"][0]["content_hash"] == compute_hash("# New Title")
            assert structure_data["blocks"][1]["content_hash"] == compute_hash("New paragraph.")


def test_update_sidedoc_metadata_updates_manifest_timestamp() -> None:
    """Test that manifest.json modified_at is updated."""
    with tempfile.TemporaryDirectory() as temp_dir:
        sidedoc_path = Path(temp_dir) / "test.sidedoc"

        # Create initial sidedoc
        old_content = "Old content"
        old_manifest = {
            "sidedoc_version": "1.0.0",
            "created_at": "2024-01-01T00:00:00+00:00",
            "modified_at": "2024-01-01T00:00:00+00:00",
            "source_file": "test.docx",
            "source_hash": "abc123",
            "content_hash": "old_hash",
            "generator": "sidedoc-cli/0.1.0",
        }

        with zipfile.ZipFile(sidedoc_path, "w") as zip_file:
            zip_file.writestr("content.md", old_content)
            zip_file.writestr("structure.json", json.dumps({"blocks": []}))
            zip_file.writestr("styles.json", json.dumps({"block_styles": {}}))
            zip_file.writestr("manifest.json", json.dumps(old_manifest))

        # Update with new content
        new_content = "New content"
        new_blocks = [
            Block(
                id="block-1",
                type="paragraph",
                content="New content",
                docx_paragraph_index=0,
                content_start=0,
                content_end=11,
                content_hash=compute_hash("New content"),
            )
        ]

        update_sidedoc_metadata(str(sidedoc_path), new_blocks, new_content)

        # Verify manifest was updated
        with zipfile.ZipFile(sidedoc_path, "r") as zip_file:
            manifest_data = json.loads(zip_file.read("manifest.json"))
            # modified_at should be different from original
            assert manifest_data["modified_at"] != "2024-01-01T00:00:00+00:00"
            # created_at should remain unchanged
            assert manifest_data["created_at"] == "2024-01-01T00:00:00+00:00"


def test_update_sidedoc_metadata_updates_content_hash() -> None:
    """Test that manifest.json content_hash is updated."""
    with tempfile.TemporaryDirectory() as temp_dir:
        sidedoc_path = Path(temp_dir) / "test.sidedoc"

        old_content = "Old content"
        old_manifest = {
            "sidedoc_version": "1.0.0",
            "created_at": "2024-01-01T00:00:00+00:00",
            "modified_at": "2024-01-01T00:00:00+00:00",
            "source_file": "test.docx",
            "source_hash": "abc123",
            "content_hash": compute_hash("Old content"),
            "generator": "sidedoc-cli/0.1.0",
        }

        with zipfile.ZipFile(sidedoc_path, "w") as zip_file:
            zip_file.writestr("content.md", old_content)
            zip_file.writestr("structure.json", json.dumps({"blocks": []}))
            zip_file.writestr("styles.json", json.dumps({"block_styles": {}}))
            zip_file.writestr("manifest.json", json.dumps(old_manifest))

        new_content = "New content here"
        new_blocks = [
            Block(
                id="block-1",
                type="paragraph",
                content="New content here",
                docx_paragraph_index=0,
                content_start=0,
                content_end=16,
                content_hash=compute_hash("New content here"),
            )
        ]

        update_sidedoc_metadata(str(sidedoc_path), new_blocks, new_content)

        # Verify content_hash was updated
        with zipfile.ZipFile(sidedoc_path, "r") as zip_file:
            manifest_data = json.loads(zip_file.read("manifest.json"))
            expected_hash = compute_hash("New content here")
            assert manifest_data["content_hash"] == expected_hash
            assert manifest_data["content_hash"] != compute_hash("Old content")


def test_update_sidedoc_metadata_preserves_styles() -> None:
    """Test that styles.json is preserved in repackaged archive."""
    with tempfile.TemporaryDirectory() as temp_dir:
        sidedoc_path = Path(temp_dir) / "test.sidedoc"

        old_styles = {
            "block_styles": {
                "block-1": {
                    "font_name": "Times New Roman",
                    "font_size": 14,
                    "alignment": "center",
                }
            },
            "document_defaults": {"font_name": "Arial", "font_size": 11},
        }

        with zipfile.ZipFile(sidedoc_path, "w") as zip_file:
            zip_file.writestr("content.md", "Old content")
            zip_file.writestr("structure.json", json.dumps({"blocks": []}))
            zip_file.writestr("styles.json", json.dumps(old_styles))
            zip_file.writestr("manifest.json", json.dumps({
                "sidedoc_version": "1.0.0",
                "created_at": "2024-01-01T00:00:00+00:00",
                "modified_at": "2024-01-01T00:00:00+00:00",
                "source_file": "test.docx",
                "source_hash": "abc123",
                "content_hash": "old_hash",
                "generator": "sidedoc-cli/0.1.0",
            }))

        new_blocks = [
            Block(
                id="block-1",
                type="paragraph",
                content="New content",
                docx_paragraph_index=0,
                content_start=0,
                content_end=11,
                content_hash=compute_hash("New content"),
            )
        ]

        update_sidedoc_metadata(str(sidedoc_path), new_blocks, "New content")

        # Verify styles were preserved
        with zipfile.ZipFile(sidedoc_path, "r") as zip_file:
            styles_data = json.loads(zip_file.read("styles.json"))
            assert styles_data == old_styles


# Tests for apply_inline_formatting edge cases (Issue #7)
from sidedoc.sync import apply_inline_formatting


def test_inline_formatting_nested_bold_italic() -> None:
    """Test that nested formatting like **bold *italic* text** works correctly.

    This is a regression test for Issue #7: the regex pattern [^*]+ rejects
    any asterisks inside bold text, causing nested formatting to fail.
    """
    with tempfile.TemporaryDirectory() as temp_dir:
        output_path = Path(temp_dir) / "test.docx"
        doc = Document()
        para = doc.add_paragraph()

        apply_inline_formatting(para, "**bold *italic* text**")

        # The paragraph should have runs with proper formatting
        # Expected: "bold " (bold), "italic" (bold+italic), " text" (bold)
        runs = list(para.runs)

        # Verify all text is present
        full_text = "".join(run.text for run in runs)
        assert full_text == "bold italic text", f"Got: {full_text}"

        # Verify bold is applied to outer content
        bold_runs = [run for run in runs if run.bold]
        assert len(bold_runs) >= 1, "Expected at least one bold run"

        # Verify italic is applied to inner content
        italic_runs = [run for run in runs if run.italic]
        assert len(italic_runs) >= 1, "Expected at least one italic run"


def test_inline_formatting_escaped_asterisks() -> None:
    """Test that escaped asterisks are preserved as literal asterisks.

    This is a regression test for Issue #7: the parser doesn't handle
    escaped asterisks like \\*literal\\*.
    """
    with tempfile.TemporaryDirectory() as temp_dir:
        output_path = Path(temp_dir) / "test.docx"
        doc = Document()
        para = doc.add_paragraph()

        # \* should be treated as a literal asterisk, not formatting
        apply_inline_formatting(para, r"Use \*asterisks\* for emphasis")

        runs = list(para.runs)
        full_text = "".join(run.text for run in runs)

        # Should contain literal asterisks, not be formatted as italic
        assert "*asterisks*" in full_text, f"Got: {full_text}"
        # Should not have italic formatting
        italic_runs = [run for run in runs if run.italic]
        assert len(italic_runs) == 0, "Escaped asterisks should not create italic"


def test_inline_formatting_malformed_bold_italic() -> None:
    """Test graceful handling of malformed markdown like **bold*italic**.

    This is a regression test for Issue #7: malformed markdown may cause
    incorrect formatting or corrupted output.
    """
    with tempfile.TemporaryDirectory() as temp_dir:
        doc = Document()
        para = doc.add_paragraph()

        # Malformed: mixing bold and italic markers incorrectly
        apply_inline_formatting(para, "**bold*italic**")

        runs = list(para.runs)
        full_text = "".join(run.text for run in runs)

        # The text should be preserved (graceful degradation)
        # Either render as-is or apply best-effort formatting
        assert "bold" in full_text.lower(), f"Content lost: {full_text}"
        assert "italic" in full_text.lower(), f"Content lost: {full_text}"


def test_inline_formatting_multiple_bold_sections() -> None:
    """Test that multiple separate bold sections work correctly."""
    with tempfile.TemporaryDirectory() as temp_dir:
        doc = Document()
        para = doc.add_paragraph()

        apply_inline_formatting(para, "**first** and **second** bold")

        runs = list(para.runs)
        full_text = "".join(run.text for run in runs)

        assert full_text == "first and second bold", f"Got: {full_text}"

        # Both "first" and "second" should be bold
        bold_runs = [run for run in runs if run.bold]
        bold_text = "".join(run.text for run in bold_runs)
        assert "first" in bold_text, "First bold section missing"
        assert "second" in bold_text, "Second bold section missing"


def test_inline_formatting_bold_with_asterisk_word() -> None:
    """Test bold text containing words that look like italic markers.

    This tests **bold *word* more** where *word* should be bold+italic.
    """
    with tempfile.TemporaryDirectory() as temp_dir:
        doc = Document()
        para = doc.add_paragraph()

        apply_inline_formatting(para, "This is **very *important* text**.")

        runs = list(para.runs)
        full_text = "".join(run.text for run in runs)

        # Should contain all the words
        assert "very" in full_text, f"Missing content: {full_text}"
        assert "important" in full_text, f"Missing content: {full_text}"
        assert "text" in full_text, f"Missing content: {full_text}"


def test_inline_formatting_unclosed_bold() -> None:
    """Test graceful handling of unclosed bold marker."""
    with tempfile.TemporaryDirectory() as temp_dir:
        doc = Document()
        para = doc.add_paragraph()

        # Unclosed bold - should degrade gracefully
        apply_inline_formatting(para, "Some **unclosed bold text")

        runs = list(para.runs)
        full_text = "".join(run.text for run in runs)

        # Text should be preserved even if formatting is incorrect
        assert "unclosed" in full_text or "**" in full_text, f"Content lost: {full_text}"


def test_inline_formatting_adjacent_formatting() -> None:
    """Test adjacent bold and italic without space."""
    with tempfile.TemporaryDirectory() as temp_dir:
        doc = Document()
        para = doc.add_paragraph()

        apply_inline_formatting(para, "**bold***italic*")

        runs = list(para.runs)
        full_text = "".join(run.text for run in runs)

        assert "bold" in full_text, f"Missing bold: {full_text}"
        assert "italic" in full_text, f"Missing italic: {full_text}"


# Tests for asset size validation (Issue #8)


def test_update_sidedoc_metadata_rejects_oversized_assets() -> None:
    """Test that oversized assets are rejected to prevent ZIP bomb attacks.

    This is a regression test for Issue #8: assets are read without size
    validation, creating a potential ZIP bomb vector.
    """
    import pytest
    from unittest.mock import patch

    with tempfile.TemporaryDirectory() as temp_dir:
        sidedoc_path = Path(temp_dir) / "test.sidedoc"

        # Create a sidedoc with an asset
        old_manifest = {
            "sidedoc_version": "1.0.0",
            "created_at": "2024-01-01T00:00:00+00:00",
            "modified_at": "2024-01-01T00:00:00+00:00",
            "source_file": "test.docx",
            "source_hash": "abc123",
            "content_hash": "old_hash",
            "generator": "sidedoc-cli/0.1.0",
        }

        # Create a modest-sized asset for testing (1KB)
        asset_data = b"X" * 1024

        with zipfile.ZipFile(sidedoc_path, "w", zipfile.ZIP_DEFLATED) as zip_file:
            zip_file.writestr("content.md", "Test content")
            zip_file.writestr("structure.json", json.dumps({"blocks": []}))
            zip_file.writestr("styles.json", json.dumps({"block_styles": {}}))
            zip_file.writestr("manifest.json", json.dumps(old_manifest))
            # Write asset
            zip_file.writestr("assets/huge_image.png", asset_data)

        # Attempt to update metadata with a low MAX_ASSET_SIZE limit
        new_blocks = [
            Block(
                id="block-1",
                type="paragraph",
                content="Test",
                docx_paragraph_index=0,
                content_start=0,
                content_end=4,
                content_hash=compute_hash("Test"),
            )
        ]

        # Patch MAX_ASSET_SIZE to be smaller than our test asset
        with patch("sidedoc.sync.MAX_ASSET_SIZE", 100):
            with pytest.raises(ValueError) as exc_info:
                update_sidedoc_metadata(str(sidedoc_path), new_blocks, "Test")

        # Verify error message mentions the asset and ZIP bomb
        error_msg = str(exc_info.value)
        assert "assets/huge_image.png" in error_msg
        assert "exceeds maximum size" in error_msg
        assert "ZIP bomb" in error_msg


# Tests for tempfile cleanup (Issue #17)


def test_update_sidedoc_metadata_successful_no_tempfile_left() -> None:
    """Test that temp file does not exist after successful update.

    This is a regression test for Issue #17: the finally block attempts
    to delete a temp file that was already moved via .replace().
    """
    with tempfile.TemporaryDirectory() as temp_dir:
        sidedoc_path = Path(temp_dir) / "test.sidedoc"

        # Create initial sidedoc
        old_manifest = {
            "sidedoc_version": "1.0.0",
            "created_at": "2024-01-01T00:00:00+00:00",
            "modified_at": "2024-01-01T00:00:00+00:00",
            "source_file": "test.docx",
            "source_hash": "abc123",
            "content_hash": "old_hash",
            "generator": "sidedoc-cli/0.1.0",
        }

        with zipfile.ZipFile(sidedoc_path, "w") as zip_file:
            zip_file.writestr("content.md", "Old content")
            zip_file.writestr("structure.json", json.dumps({"blocks": []}))
            zip_file.writestr("styles.json", json.dumps({"block_styles": {}}))
            zip_file.writestr("manifest.json", json.dumps(old_manifest))

        # Update with new content
        new_blocks = [
            Block(
                id="block-1",
                type="paragraph",
                content="New content",
                docx_paragraph_index=0,
                content_start=0,
                content_end=11,
                content_hash=compute_hash("New content"),
            )
        ]

        # Track temp files before and after
        temp_files_before = set(Path(temp_dir).glob("*.sidedoc"))

        # This should succeed without leaving temp files
        update_sidedoc_metadata(str(sidedoc_path), new_blocks, "New content")

        temp_files_after = set(Path(temp_dir).glob("*.sidedoc"))

        # Should only have the original file, no temp files left behind
        assert temp_files_after == temp_files_before
        assert sidedoc_path.exists()


def test_update_sidedoc_metadata_cleanup_on_error() -> None:
    """Test that temp file is cleaned up when an error occurs.

    This is a regression test for Issue #17: temp files should be
    cleaned up only when an error prevents successful completion.
    """
    import pytest
    from unittest.mock import patch, MagicMock

    with tempfile.TemporaryDirectory() as temp_dir:
        sidedoc_path = Path(temp_dir) / "test.sidedoc"

        # Create initial sidedoc
        old_manifest = {
            "sidedoc_version": "1.0.0",
            "created_at": "2024-01-01T00:00:00+00:00",
            "modified_at": "2024-01-01T00:00:00+00:00",
            "source_file": "test.docx",
            "source_hash": "abc123",
            "content_hash": "old_hash",
            "generator": "sidedoc-cli/0.1.0",
        }

        with zipfile.ZipFile(sidedoc_path, "w") as zip_file:
            zip_file.writestr("content.md", "Old content")
            zip_file.writestr("structure.json", json.dumps({"blocks": []}))
            zip_file.writestr("styles.json", json.dumps({"block_styles": {}}))
            zip_file.writestr("manifest.json", json.dumps(old_manifest))

        new_blocks = [
            Block(
                id="block-1",
                type="paragraph",
                content="New content",
                docx_paragraph_index=0,
                content_start=0,
                content_end=11,
                content_hash=compute_hash("New content"),
            )
        ]

        # Patch Path.replace to raise an error during the replace operation
        with patch("pathlib.Path.replace") as mock_replace:
            mock_replace.side_effect = OSError("Simulated replace error")

            # Should raise the error
            with pytest.raises(OSError, match="Simulated replace error"):
                update_sidedoc_metadata(str(sidedoc_path), new_blocks, "New content")

        # After error, no temp files should remain in the directory
        temp_files = [f for f in Path(temp_dir).glob("*.sidedoc") if f.name != "test.sidedoc"]
        assert len(temp_files) == 0, f"Found leftover temp files: {temp_files}"


def test_delete_and_add_with_low_similarity_not_matched():
    """Test that deleting a paragraph and adding a completely different one is NOT matched.

    This test addresses issue #11 - when content at the same position is completely different
    (low similarity), it should be treated as delete + add, not as an edit.
    """
    old_blocks = [
        Block(
            id="block-1",
            type="paragraph",
            content="The quick brown fox jumps over the lazy dog.",
            docx_paragraph_index=0,
            content_start=0,
            content_end=45,
            content_hash=compute_hash("The quick brown fox jumps over the lazy dog."),
        ),
    ]

    new_blocks = [
        Block(
            id="block-new-1",
            type="paragraph",
            content="Python is a great programming language.",
            docx_paragraph_index=0,
            content_start=0,
            content_end=40,
            content_hash=compute_hash("Python is a great programming language."),
        ),
    ]

    matches = match_blocks(old_blocks, new_blocks)

    # With low similarity, block should NOT be matched
    # This means block-1 was deleted and block-new-1 was added
    assert "block-1" not in matches


def test_edit_with_high_similarity_is_matched():
    """Test that editing a paragraph with high similarity IS matched.

    This test addresses issue #11 - when content at the same position has high similarity
    (e.g., minor edits, additions), it should be treated as an edit and preserve formatting.
    """
    old_blocks = [
        Block(
            id="block-1",
            type="paragraph",
            content="The quick brown fox jumps over the lazy dog.",
            docx_paragraph_index=0,
            content_start=0,
            content_end=45,
            content_hash=compute_hash("The quick brown fox jumps over the lazy dog."),
        ),
    ]

    new_blocks = [
        Block(
            id="block-new-1",
            type="paragraph",
            content="The quick brown fox jumps over the sleepy dog.",
            docx_paragraph_index=0,
            content_start=0,
            content_end=47,
            content_hash=compute_hash("The quick brown fox jumps over the sleepy dog."),
        ),
    ]

    matches = match_blocks(old_blocks, new_blocks)

    # With high similarity, block SHOULD be matched (treated as edit)
    assert "block-1" in matches
    assert matches["block-1"].content == "The quick brown fox jumps over the sleepy dog."


def test_similarity_threshold_boundary_cases():
    """Test similarity threshold boundary cases.

    Tests content that is at or near the similarity threshold (0.7 or 70%).
    """
    # Create a base paragraph
    base_content = "This is a paragraph with about fifty characters."

    old_blocks = [
        Block(
            id="block-1",
            type="paragraph",
            content=base_content,
            docx_paragraph_index=0,
            content_start=0,
            content_end=len(base_content),
            content_hash=compute_hash(base_content),
        ),
    ]

    # Test 1: Low similarity (well below threshold) - should NOT match
    # Completely different content
    new_blocks_low = [
        Block(
            id="block-new-1",
            type="paragraph",
            content="Python programming language features and syntax.",
            docx_paragraph_index=0,
            content_start=0,
            content_end=48,
            content_hash=compute_hash("Python programming language features and syntax."),
        ),
    ]

    matches_low = match_blocks(old_blocks, new_blocks_low)
    assert "block-1" not in matches_low, "Low similarity should not match"

    # Test 2: High similarity (above threshold) - SHOULD match
    # Only changed one word
    new_blocks_high = [
        Block(
            id="block-new-2",
            type="paragraph",
            content="This is a paragraph with about sixty characters.",
            docx_paragraph_index=0,
            content_start=0,
            content_end=49,
            content_hash=compute_hash("This is a paragraph with about sixty characters."),
        ),
    ]

    matches_high = match_blocks(old_blocks, new_blocks_high)
    assert "block-1" in matches_high, "High similarity should match"


def test_remap_styles_remaps_block_ids() -> None:
    """Test that remap_styles correctly remaps block IDs in styles_data."""
    styles_data = {
        "block_styles": {
            "block-0": {"font_name": "Arial", "font_size": 12, "alignment": "left"},
            "block-1": {"font_name": "Times New Roman", "font_size": 14, "alignment": "center"},
        },
        "document_defaults": {"font_name": "Calibri", "font_size": 11},
    }

    # Matches map old IDs to new blocks
    new_block_0 = Block(id="block-0", type="heading", content="# Title", docx_paragraph_index=0,
                        content_start=0, content_end=7, content_hash="h1", level=1)
    new_block_1 = Block(id="block-1", type="paragraph", content="Para", docx_paragraph_index=1,
                        content_start=8, content_end=12, content_hash="h2")
    matches = {"block-0": new_block_0, "block-1": new_block_1}

    result = remap_styles(styles_data, matches)

    # Styles should now use new block IDs (same in this case since IDs match)
    assert "block-0" in result["block_styles"]
    assert "block-1" in result["block_styles"]
    assert result["document_defaults"] == {"font_name": "Calibri", "font_size": 11}


def test_remap_styles_handles_shifted_ids() -> None:
    """Test that remap_styles handles ID shifts from block additions/removals."""
    styles_data = {
        "block_styles": {
            "block-0": {"font_name": "Arial", "font_size": 12, "alignment": "left"},
            "block-1": {"font_name": "Courier", "font_size": 10, "alignment": "left"},
        },
        "document_defaults": {"font_name": "Calibri", "font_size": 11},
    }

    # After editing, block-0 matched to block-0 (same), block-1 matched to block-2 (shifted)
    new_block_0 = Block(id="block-0", type="heading", content="# Title", docx_paragraph_index=0,
                        content_start=0, content_end=7, content_hash="h1", level=1)
    new_block_2 = Block(id="block-2", type="paragraph", content="Para", docx_paragraph_index=2,
                        content_start=20, content_end=24, content_hash="h2")
    matches = {"block-0": new_block_0, "block-1": new_block_2}

    result = remap_styles(styles_data, matches)

    # block-0 style stays at block-0 (matched to same ID)
    assert "block-0" in result["block_styles"]
    assert result["block_styles"]["block-0"]["font_name"] == "Arial"
    # block-1's style should now be at block-2 (shifted)
    assert "block-2" in result["block_styles"]
    assert result["block_styles"]["block-2"]["font_name"] == "Courier"
    # block-1 should not exist in remapped styles
    assert "block-1" not in result["block_styles"]


def test_remap_styles_drops_unmatched_blocks() -> None:
    """Test that styles for deleted (unmatched) blocks are dropped."""
    styles_data = {
        "block_styles": {
            "block-0": {"font_name": "Arial", "font_size": 12, "alignment": "left"},
            "block-1": {"font_name": "Courier", "font_size": 10, "alignment": "left"},
        },
        "document_defaults": {"font_name": "Calibri", "font_size": 11},
    }

    # Only block-0 matched; block-1 was deleted
    new_block_0 = Block(id="block-0", type="heading", content="# Title", docx_paragraph_index=0,
                        content_start=0, content_end=7, content_hash="h1", level=1)
    matches = {"block-0": new_block_0}

    result = remap_styles(styles_data, matches)

    assert "block-0" in result["block_styles"]
    assert "block-1" not in result["block_styles"]


def test_different_block_types_not_matched():
    """Test that blocks with different types are not matched even with similar content.

    This ensures type checking happens before similarity checking.
    """
    old_blocks = [
        Block(
            id="block-1",
            type="paragraph",
            content="Important Note",
            docx_paragraph_index=0,
            content_start=0,
            content_end=14,
            content_hash=compute_hash("Important Note"),
        ),
    ]

    new_blocks = [
        Block(
            id="block-new-1",
            type="heading",
            level=1,
            content="# Important Note",
            docx_paragraph_index=0,
            content_start=0,
            content_end=16,
            content_hash=compute_hash("# Important Note"),
        ),
    ]

    matches = match_blocks(old_blocks, new_blocks)

    # Different types should never match, regardless of content similarity
    assert "block-1" not in matches


# Tests for table block handling in sync


def test_sync_handles_table_blocks() -> None:
    """Verify sync creates actual tables, not raw GFM text.

    This is a critical test: table blocks must result in actual Table objects
    in the docx, not paragraphs containing raw GFM pipe syntax.
    """
    table_content = "| Name | Age |\n| --- | --- |\n| Alice | 30 |"
    new_blocks = [
        Block(
            id="block-new-1",
            type="table",
            content=table_content,
            docx_paragraph_index=0,
            content_start=0,
            content_end=len(table_content),
            content_hash=compute_hash(table_content),
        ),
    ]

    styles = {"block_styles": {}}
    matches = {}

    with tempfile.TemporaryDirectory() as temp_dir:
        output_path = Path(temp_dir) / "test.docx"
        generate_updated_docx(new_blocks, matches, styles, str(output_path))

        # Verify docx was created
        assert output_path.exists()

        # Verify it contains a table, not a paragraph with GFM
        doc = Document(str(output_path))

        # Should have a table
        assert len(doc.tables) == 1, f"Expected 1 table, got {len(doc.tables)}"

        # Verify table content
        table = doc.tables[0]
        assert len(table.rows) == 2  # header + 1 data row
        assert len(table.columns) == 2

        # Check cell content
        assert table.cell(0, 0).text == "Name"
        assert table.cell(0, 1).text == "Age"
        assert table.cell(1, 0).text == "Alice"
        assert table.cell(1, 1).text == "30"

        # Should NOT have paragraphs with GFM pipe syntax
        for para in doc.paragraphs:
            assert "|" not in para.text, f"Found raw GFM in paragraph: {para.text}"


def test_sync_preserves_table_metadata() -> None:
    """Test that update_sidedoc_metadata preserves table_metadata in structure.json.

    This is a regression test for PR #40 review: table_metadata was omitted from
    the block dict serialization, causing table structure (rows, cols, cells,
    alignments) to be lost after sync.
    """
    with tempfile.TemporaryDirectory() as temp_dir:
        sidedoc_path = Path(temp_dir) / "test.sidedoc"

        # Create initial sidedoc archive
        old_content = "| Name | Age |\n| --- | --- |\n| Alice | 30 |"
        old_structure = {
            "blocks": [
                {
                    "id": "block-0",
                    "type": "table",
                    "docx_paragraph_index": -1,
                    "content_start": 0,
                    "content_end": len(old_content),
                    "content_hash": compute_hash(old_content),
                    "level": None,
                    "image_path": None,
                    "inline_formatting": None,
                    "table_metadata": {
                        "rows": 2,
                        "cols": 2,
                        "cells": [[{"row": 0, "col": 0, "content_hash": "h1"},
                                    {"row": 0, "col": 1, "content_hash": "h2"}],
                                   [{"row": 1, "col": 0, "content_hash": "h3"},
                                    {"row": 1, "col": 1, "content_hash": "h4"}]],
                        "column_alignments": ["left", "left"],
                        "docx_table_index": 0
                    }
                }
            ]
        }
        old_styles = {"block_styles": {}}
        old_manifest = {
            "sidedoc_version": "1.0.0",
            "created_at": "2024-01-01T00:00:00+00:00",
            "modified_at": "2024-01-01T00:00:00+00:00",
            "source_file": "test.docx",
            "source_hash": "abc123",
            "content_hash": compute_hash(old_content),
            "generator": "sidedoc-cli/0.1.0",
        }

        with zipfile.ZipFile(sidedoc_path, "w") as zip_file:
            zip_file.writestr("content.md", old_content)
            zip_file.writestr("structure.json", json.dumps(old_structure))
            zip_file.writestr("styles.json", json.dumps(old_styles))
            zip_file.writestr("manifest.json", json.dumps(old_manifest))

        # Create new blocks with table_metadata
        table_metadata = {
            "rows": 2,
            "cols": 2,
            "cells": [[{"row": 0, "col": 0, "content_hash": "h1"},
                        {"row": 0, "col": 1, "content_hash": "h2"}],
                       [{"row": 1, "col": 0, "content_hash": "h3"},
                        {"row": 1, "col": 1, "content_hash": "h4"}]],
            "column_alignments": ["left", "left"],
            "docx_table_index": 0
        }
        new_blocks = [
            Block(
                id="block-0",
                type="table",
                content=old_content,
                docx_paragraph_index=-1,
                content_start=0,
                content_end=len(old_content),
                content_hash=compute_hash(old_content),
                table_metadata=table_metadata,
            ),
        ]

        # Update metadata
        update_sidedoc_metadata(str(sidedoc_path), new_blocks, old_content)

        # Verify table_metadata is preserved in structure.json
        with zipfile.ZipFile(sidedoc_path, "r") as zip_file:
            structure_data = json.loads(zip_file.read("structure.json"))
            block_data = structure_data["blocks"][0]
            assert "table_metadata" in block_data, \
                "table_metadata should be serialized in structure.json"
            assert block_data["table_metadata"]["rows"] == 2
            assert block_data["table_metadata"]["cols"] == 2
            assert len(block_data["table_metadata"]["cells"]) == 2
            assert block_data["table_metadata"]["column_alignments"] == ["left", "left"]
            assert block_data["table_metadata"]["docx_table_index"] == 0


def test_sync_table_roundtrip_preserves_structure_and_data() -> None:
    """E2E: extract docx with table → modify GFM → sync → verify docx table.

    This integration test exercises the complete table sync workflow:
    1. Extract tables_simple.docx to blocks
    2. Modify a cell value in the GFM content
    3. Call generate_updated_docx with modified blocks
    4. Verify the output docx has a Table object with correct data and alignment
    """
    from sidedoc.extract import extract_blocks, extract_styles
    from sidedoc.reconstruct import parse_markdown_to_blocks

    fixture_path = Path("tests/fixtures/tables_simple.docx")
    assert fixture_path.exists(), "tables_simple.docx fixture required"

    # Step 1: Extract blocks and styles from the fixture
    original_blocks, _ = extract_blocks(str(fixture_path))
    original_styles = extract_styles(str(fixture_path), original_blocks)

    # Build styles dict in the format generate_updated_docx expects
    styles_dict = {"block_styles": {}}
    for style in original_styles:
        style_data = {}
        if style.font_name:
            style_data["font_name"] = style.font_name
        if style.font_size:
            style_data["font_size"] = style.font_size
        if style.alignment:
            style_data["alignment"] = style.alignment
        if style.table_formatting:
            style_data["table_formatting"] = style.table_formatting
        styles_dict["block_styles"][style.block_id] = style_data

    # Find the table block and get its GFM content
    table_blocks = [b for b in original_blocks if b.type == "table"]
    assert len(table_blocks) >= 1, "Fixture should have at least one table"

    # Step 2: Modify a cell value — change "Alice" to "Alice Smith"
    original_table = table_blocks[0]
    modified_content = original_table.content.replace("Alice", "Alice Smith")
    assert "Alice Smith" in modified_content, "Modification should be applied"

    # Re-parse the modified GFM to get new blocks
    # Build modified markdown by replacing the table content in full markdown
    from sidedoc.extract import blocks_to_markdown
    original_md = blocks_to_markdown(original_blocks)
    modified_md = original_md.replace(original_table.content, modified_content)
    new_blocks = parse_markdown_to_blocks(modified_md)

    # Build matches: match old block IDs to new blocks by position/type
    matches = match_blocks(original_blocks, new_blocks)

    # Step 3: Generate updated docx
    with tempfile.TemporaryDirectory() as temp_dir:
        output_path = Path(temp_dir) / "roundtrip.docx"
        generate_updated_docx(new_blocks, matches, styles_dict, str(output_path))

        assert output_path.exists(), "Output docx should be created"

        # Step 4: Verify the output docx
        doc = Document(str(output_path))

        # Should contain a Table object (not paragraph text)
        assert len(doc.tables) >= 1, (
            f"Expected at least 1 table, got {len(doc.tables)}. "
            "Table blocks should create Table objects, not paragraphs."
        )

        table = doc.tables[0]

        # Verify the modified cell content survived the roundtrip
        found_alice_smith = False
        for row in table.rows:
            for cell in row.cells:
                if "Alice Smith" in cell.text:
                    found_alice_smith = True
        assert found_alice_smith, (
            "Modified cell 'Alice Smith' should survive the sync roundtrip. "
            f"Table cells: {[[c.text for c in r.cells] for r in table.rows]}"
        )

        # Verify header row is preserved
        header_texts = [cell.text for cell in table.rows[0].cells]
        assert "Name" in header_texts, f"Header should contain 'Name', got {header_texts}"
        assert "Role" in header_texts, f"Header should contain 'Role', got {header_texts}"

        # Verify no raw GFM pipe syntax leaked into paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                assert not (para.text.strip().startswith("|") and para.text.strip().endswith("|")), (
                    f"Raw GFM leaked into paragraph: {para.text}"
                )


def test_sync_handles_mixed_blocks_with_table() -> None:
    """Verify sync handles documents with paragraphs and tables together."""
    table_content = "| Col1 | Col2 |\n| --- | --- |\n| A | B |"
    new_blocks = [
        Block(
            id="block-1",
            type="paragraph",
            content="Before the table.",
            docx_paragraph_index=0,
            content_start=0,
            content_end=17,
            content_hash=compute_hash("Before the table."),
        ),
        Block(
            id="block-2",
            type="table",
            content=table_content,
            docx_paragraph_index=1,
            content_start=18,
            content_end=18 + len(table_content),
            content_hash=compute_hash(table_content),
        ),
        Block(
            id="block-3",
            type="paragraph",
            content="After the table.",
            docx_paragraph_index=2,
            content_start=100,
            content_end=116,
            content_hash=compute_hash("After the table."),
        ),
    ]

    styles = {"block_styles": {}}
    matches = {}

    with tempfile.TemporaryDirectory() as temp_dir:
        output_path = Path(temp_dir) / "test.docx"
        generate_updated_docx(new_blocks, matches, styles, str(output_path))

        doc = Document(str(output_path))

        # Should have 2 paragraphs and 1 table
        assert len(doc.tables) == 1
        # Paragraphs include those before and after table
        para_texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "Before the table." in para_texts
        assert "After the table." in para_texts


# Tests for sync_sidedoc_to_docx table handling (CriticMarkup sync path)


def test_sync_sidedoc_to_docx_creates_table_objects() -> None:
    """Test that sync_sidedoc_to_docx creates actual Table objects for table blocks.

    The CriticMarkup sync path was treating ALL non-heading blocks as paragraphs,
    rendering table blocks as raw GFM pipe text instead of actual Table objects.
    """
    from sidedoc.sync import sync_sidedoc_to_docx

    table_content = "| Name | Age |\n| --- | --- |\n| Alice | 30 |"
    content_md = f"# Title\n\n{table_content}\n\nA paragraph."

    structure = {
        "blocks": [
            {"id": "block-0", "type": "heading", "docx_paragraph_index": 0,
             "content_start": 0, "content_end": 7, "content_hash": "", "level": 1,
             "image_path": None, "inline_formatting": None, "table_metadata": None},
            {"id": "block-1", "type": "table", "docx_paragraph_index": -1,
             "content_start": 9, "content_end": 9 + len(table_content), "content_hash": "",
             "level": None, "image_path": None, "inline_formatting": None,
             "table_metadata": {"rows": 2, "cols": 2, "cells": [],
                                "column_alignments": ["left", "left"],
                                "docx_table_index": 0}},
            {"id": "block-2", "type": "paragraph", "docx_paragraph_index": 1,
             "content_start": 100, "content_end": 112, "content_hash": "", "level": None,
             "image_path": None, "inline_formatting": None, "table_metadata": None},
        ]
    }
    styles = {"block_styles": {}}
    manifest = {
        "sidedoc_version": "1.0.0",
        "created_at": "2024-01-01T00:00:00+00:00",
        "modified_at": "2024-01-01T00:00:00+00:00",
        "source_file": "test.docx",
        "source_hash": "abc123",
        "content_hash": "",
        "generator": "sidedoc-cli/0.1.0",
    }

    with tempfile.TemporaryDirectory() as temp_dir:
        sidedoc_path = Path(temp_dir) / "test.sidedoc"
        output_path = Path(temp_dir) / "output.docx"

        with zipfile.ZipFile(str(sidedoc_path), "w") as zf:
            zf.writestr("content.md", content_md)
            zf.writestr("structure.json", json.dumps(structure))
            zf.writestr("styles.json", json.dumps(styles))
            zf.writestr("manifest.json", json.dumps(manifest))

        sync_sidedoc_to_docx(str(sidedoc_path), str(output_path))

        doc = Document(str(output_path))

        # Should have at least 1 table
        assert len(doc.tables) >= 1, (
            f"Expected at least 1 table, got {len(doc.tables)}. "
            "Table blocks should create Table objects, not paragraphs."
        )

        # Verify table content
        table = doc.tables[0]
        assert table.cell(0, 0).text == "Name"
        assert table.cell(1, 0).text == "Alice"

        # No raw GFM pipe syntax in paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                assert not (text.startswith("|") and text.endswith("|")), \
                    f"Raw GFM leaked into paragraph: {text}"
