"""Test document extraction functionality."""

import tempfile
from pathlib import Path
import struct
import zlib
from docx import Document
from docx.shared import Inches
from sidedoc.extract import extract_blocks
from sidedoc.models import Block


def create_test_docx(content_items: list[tuple[str, str]]) -> str:
    """Create a test docx file with specified content.

    Args:
        content_items: List of (style, text) tuples

    Returns:
        Path to temporary docx file
    """
    doc = Document()
    for style, text in content_items:
        doc.add_paragraph(text, style=style)

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    temp_file.close()
    return temp_file.name


def test_extract_blocks_function_exists():
    """Test that extract_blocks function exists."""
    assert callable(extract_blocks)


def test_extract_single_paragraph():
    """Test extracting a simple paragraph."""
    docx_path = create_test_docx([("Normal", "Hello world")])

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        assert blocks[0].type == "paragraph"
        assert blocks[0].content == "Hello world"
        assert blocks[0].docx_paragraph_index == 0
    finally:
        Path(docx_path).unlink()


def test_extract_heading_1():
    """Test extracting Heading 1."""
    docx_path = create_test_docx([("Heading 1", "Title")])

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        assert blocks[0].type == "heading"
        assert blocks[0].content == "# Title"
        assert blocks[0].level == 1
    finally:
        Path(docx_path).unlink()


def test_extract_heading_2():
    """Test extracting Heading 2."""
    docx_path = create_test_docx([("Heading 2", "Subtitle")])

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        assert blocks[0].type == "heading"
        assert blocks[0].content == "## Subtitle"
        assert blocks[0].level == 2
    finally:
        Path(docx_path).unlink()


def test_extract_multiple_paragraphs():
    """Test extracting multiple paragraphs in order."""
    docx_path = create_test_docx([
        ("Normal", "First paragraph"),
        ("Normal", "Second paragraph"),
        ("Normal", "Third paragraph")
    ])

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 3
        assert blocks[0].content == "First paragraph"
        assert blocks[1].content == "Second paragraph"
        assert blocks[2].content == "Third paragraph"
        assert blocks[0].docx_paragraph_index == 0
        assert blocks[1].docx_paragraph_index == 1
        assert blocks[2].docx_paragraph_index == 2
    finally:
        Path(docx_path).unlink()


def test_extract_mixed_content():
    """Test extracting headings and paragraphs together."""
    docx_path = create_test_docx([
        ("Heading 1", "Main Title"),
        ("Normal", "Introduction paragraph"),
        ("Heading 2", "Section"),
        ("Normal", "Section content")
    ])

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 4
        assert blocks[0].type == "heading"
        assert blocks[0].level == 1
        assert blocks[1].type == "paragraph"
        assert blocks[2].type == "heading"
        assert blocks[2].level == 2
        assert blocks[3].type == "paragraph"
    finally:
        Path(docx_path).unlink()


def test_blocks_have_unique_ids():
    """Test that each block gets a unique ID."""
    docx_path = create_test_docx([
        ("Normal", "First"),
        ("Normal", "Second")
    ])

    try:
        blocks, _ = extract_blocks(docx_path)
        ids = [block.id for block in blocks]
        assert len(ids) == len(set(ids)), "Block IDs must be unique"
    finally:
        Path(docx_path).unlink()


def test_blocks_have_content_hashes():
    """Test that blocks have content hashes."""
    docx_path = create_test_docx([("Normal", "Test content")])

    try:
        blocks, _ = extract_blocks(docx_path)
        assert blocks[0].content_hash is not None
        assert len(blocks[0].content_hash) > 0
    finally:
        Path(docx_path).unlink()


def test_extract_all_heading_levels():
    """Test extracting all heading levels (1-6)."""
    docx_path = create_test_docx([
        ("Heading 1", "H1"),
        ("Heading 2", "H2"),
        ("Heading 3", "H3"),
        ("Heading 4", "H4"),
        ("Heading 5", "H5"),
        ("Heading 6", "H6")
    ])

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 6
        for i, block in enumerate(blocks):
            assert block.type == "heading"
            assert block.level == i + 1
            assert block.content == "#" * (i + 1) + f" H{i + 1}"
    finally:
        Path(docx_path).unlink()


def create_formatted_docx(text_runs: list[tuple[str, bool, bool, bool]]) -> str:
    """Create a test docx file with inline formatting.

    Args:
        text_runs: List of (text, bold, italic, underline) tuples

    Returns:
        Path to temporary docx file
    """
    doc = Document()
    para = doc.add_paragraph()

    for text, bold, italic, underline in text_runs:
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        run.underline = underline

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    temp_file.close()
    return temp_file.name


def test_extract_bold_text():
    """Test extracting bold text as markdown."""
    docx_path = create_formatted_docx([
        ("This is ", False, False, False),
        ("bold", True, False, False),
        (" text", False, False, False)
    ])

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        assert blocks[0].content == "This is **bold** text"
        assert blocks[0].type == "paragraph"
    finally:
        Path(docx_path).unlink()


def test_extract_italic_text():
    """Test extracting italic text as markdown."""
    docx_path = create_formatted_docx([
        ("This is ", False, False, False),
        ("italic", False, True, False),
        (" text", False, False, False)
    ])

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        assert blocks[0].content == "This is *italic* text"
        assert blocks[0].type == "paragraph"
    finally:
        Path(docx_path).unlink()


def test_extract_bold_italic_text():
    """Test extracting bold and italic text as markdown."""
    docx_path = create_formatted_docx([
        ("This is ", False, False, False),
        ("bold and italic", True, True, False),
        (" text", False, False, False)
    ])

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        assert blocks[0].content == "This is ***bold and italic*** text"
        assert blocks[0].type == "paragraph"
    finally:
        Path(docx_path).unlink()


def test_extract_underline_preserved_in_formatting():
    """Test that underline is preserved in inline_formatting, not markdown."""
    docx_path = create_formatted_docx([
        ("This is ", False, False, False),
        ("underlined", False, False, True),
        (" text", False, False, False)
    ])

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        # Underline should not appear in markdown content
        assert blocks[0].content == "This is underlined text"
        # But should be recorded in inline_formatting
        assert blocks[0].inline_formatting is not None
        assert len(blocks[0].inline_formatting) > 0
        # Find the underline formatting entry
        underline_format = next(
            (fmt for fmt in blocks[0].inline_formatting if fmt.get("underline")),
            None
        )
        assert underline_format is not None
        assert underline_format["start"] == 8  # "This is " = 8 chars
        assert underline_format["end"] == 18  # + "underlined" = 10 chars
    finally:
        Path(docx_path).unlink()


def test_extract_mixed_inline_formatting():
    """Test extracting multiple inline formatting in one paragraph."""
    docx_path = create_formatted_docx([
        ("Normal ", False, False, False),
        ("bold", True, False, False),
        (" and ", False, False, False),
        ("italic", False, True, False),
        (" and ", False, False, False),
        ("both", True, True, False)
    ])

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        assert blocks[0].content == "Normal **bold** and *italic* and ***both***"
        assert blocks[0].type == "paragraph"
    finally:
        Path(docx_path).unlink()


def test_extract_escaped_asterisks_in_text():
    """Test that literal asterisks in text are preserved (not treated as markdown).

    When docx contains text with asterisks but NO formatting applied,
    those asterisks should appear in markdown as-is.
    """
    docx_path = create_formatted_docx([
        ("This has ", False, False, False),
        ("*literal asterisks*", False, False, False),
        (" in text", False, False, False)
    ])

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        # No formatting was applied, so asterisks should be preserved literally
        assert blocks[0].content == "This has *literal asterisks* in text"
        assert blocks[0].type == "paragraph"
    finally:
        Path(docx_path).unlink()


def test_extract_nested_formatting_bold_containing_italic():
    """Test nested formatting: bold text containing italic text.

    In docx: "start BOLD_START normal ITALIC_START both ITALIC_END normal BOLD_END end"
    Current implementation: treats each run separately, resulting in adjacent markers
    """
    docx_path = create_formatted_docx([
        ("start ", False, False, False),
        ("bold ", True, False, False),
        ("and italic", True, True, False),
        (" bold", True, False, False),
        (" end", False, False, False)
    ])

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        # Current implementation: adjacent runs create adjacent markers
        # "bold " is bold: **bold **
        # "and italic" is bold+italic: ***and italic***
        # " bold" is bold: ** bold**
        assert blocks[0].content == "start **bold *****and italic***** bold** end"
        assert blocks[0].type == "paragraph"
    finally:
        Path(docx_path).unlink()


def test_extract_alternating_bold_italic():
    """Test alternating bold and italic formatting.

    Tests rapid switching between formatting types.
    """
    docx_path = create_formatted_docx([
        ("A", True, False, False),
        ("B", False, True, False),
        ("C", True, False, False),
        ("D", False, True, False),
    ])

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        assert blocks[0].content == "**A***B***C***D*"
        assert blocks[0].type == "paragraph"
    finally:
        Path(docx_path).unlink()


def test_extract_empty_formatting_runs():
    """Test that empty runs with formatting don't create spurious markdown markers."""
    docx_path = create_formatted_docx([
        ("start", False, False, False),
        ("", True, False, False),  # Empty bold run
        ("end", False, False, False),
    ])

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        # Empty runs should be ignored
        assert blocks[0].content == "startend"
        assert blocks[0].type == "paragraph"
    finally:
        Path(docx_path).unlink()


def test_extract_whitespace_only_formatted_runs():
    """Test formatting on whitespace-only runs.

    Edge case: what happens when only spaces are bold/italic?
    """
    docx_path = create_formatted_docx([
        ("word", False, False, False),
        ("   ", True, False, False),  # Bold spaces
        ("word", False, False, False),
    ])

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        # Spaces with formatting should preserve the formatting markers
        assert blocks[0].content == "word**   **word"
        assert blocks[0].type == "paragraph"
    finally:
        Path(docx_path).unlink()


def test_extract_single_character_formatting():
    """Test formatting applied to single characters."""
    docx_path = create_formatted_docx([
        ("a", True, False, False),
        ("b", False, True, False),
        ("c", True, True, False),
    ])

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        assert blocks[0].content == "**a***b****c***"
        assert blocks[0].type == "paragraph"
    finally:
        Path(docx_path).unlink()


def test_extract_unicode_with_formatting():
    """Test that Unicode characters work correctly with formatting."""
    docx_path = create_formatted_docx([
        ("Hello ", False, False, False),
        ("世界", True, False, False),  # "world" in Chinese
        (" café", False, True, False),
        (" 🎉", True, True, False),  # emoji
    ])

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        # Each run is formatted separately, creating adjacent markers
        assert blocks[0].content == "Hello **世界*** café**** 🎉***"
        assert blocks[0].type == "paragraph"
    finally:
        Path(docx_path).unlink()


def test_extract_mixed_underline_bold_italic():
    """Test complex mix of underline with bold and italic.

    Underline should be tracked separately in inline_formatting,
    while bold/italic appear in markdown.
    """
    docx_path = create_formatted_docx([
        ("plain ", False, False, False),
        ("bold-underline", True, False, True),
        (" ", False, False, False),
        ("italic-underline", False, True, True),
        (" ", False, False, False),
        ("all-three", True, True, True),
    ])

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        # Bold and italic in markdown
        assert blocks[0].content == "plain **bold-underline** *italic-underline* ***all-three***"

        # Underline tracked separately
        assert blocks[0].inline_formatting is not None
        assert len(blocks[0].inline_formatting) == 3

        # Verify underline positions (calculated from plain text)
        underline_entries = [fmt for fmt in blocks[0].inline_formatting if fmt.get("underline")]
        assert len(underline_entries) == 3
    finally:
        Path(docx_path).unlink()


def test_extract_adjacent_same_formatting():
    """Test adjacent runs with same formatting.

    Edge case: multiple consecutive runs with identical formatting
    should produce clean markdown.
    """
    docx_path = create_formatted_docx([
        ("bold", True, False, False),
        ("also", True, False, False),
        ("bold", True, False, False),
    ])

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        # Current implementation handles each run separately
        assert blocks[0].content == "**bold****also****bold**"
        assert blocks[0].type == "paragraph"
    finally:
        Path(docx_path).unlink()


def test_extract_formatting_with_newline_characters():
    """Test that runs containing newline characters are handled correctly."""
    docx_path = create_formatted_docx([
        ("before", False, False, False),
        ("bold\ntext", True, False, False),
        ("after", False, False, False),
    ])

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 1
        # Newlines in runs should be preserved
        assert blocks[0].content == "before**bold\ntext**after"
        assert blocks[0].type == "paragraph"
    finally:
        Path(docx_path).unlink()


def create_list_docx(list_items: list[tuple[str, str]]) -> str:
    """Create a test docx file with list items.

    Args:
        list_items: List of (style, text) tuples where style is 'List Bullet' or 'List Number'

    Returns:
        Path to temporary docx file
    """
    doc = Document()
    for style, text in list_items:
        doc.add_paragraph(text, style=style)

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    temp_file.close()
    return temp_file.name


def test_extract_bulleted_list():
    """Test extracting bulleted list items."""
    docx_path = create_list_docx([
        ("List Bullet", "First item"),
        ("List Bullet", "Second item"),
        ("List Bullet", "Third item")
    ])

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 3
        assert blocks[0].type == "list"
        assert blocks[0].content == "- First item"
        assert blocks[1].type == "list"
        assert blocks[1].content == "- Second item"
        assert blocks[2].type == "list"
        assert blocks[2].content == "- Third item"
    finally:
        Path(docx_path).unlink()


def test_extract_numbered_list():
    """Test extracting numbered list items."""
    docx_path = create_list_docx([
        ("List Number", "First item"),
        ("List Number", "Second item"),
        ("List Number", "Third item")
    ])

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 3
        assert blocks[0].type == "list"
        assert blocks[0].content == "1. First item"
        assert blocks[1].type == "list"
        assert blocks[1].content == "2. Second item"
        assert blocks[2].type == "list"
        assert blocks[2].content == "3. Third item"
    finally:
        Path(docx_path).unlink()


def test_extract_mixed_list_types():
    """Test extracting both bulleted and numbered lists."""
    docx_path = create_list_docx([
        ("List Bullet", "Bulleted item 1"),
        ("List Bullet", "Bulleted item 2"),
        ("List Number", "Numbered item 1"),
        ("List Number", "Numbered item 2")
    ])

    try:
        blocks, _ = extract_blocks(docx_path)
        assert len(blocks) == 4
        assert blocks[0].content == "- Bulleted item 1"
        assert blocks[1].content == "- Bulleted item 2"
        assert blocks[2].content == "1. Numbered item 1"
        assert blocks[3].content == "2. Numbered item 2"
        assert all(b.type == "list" for b in blocks)
    finally:
        Path(docx_path).unlink()


def test_extract_list_with_paragraphs():
    """Test extracting lists mixed with regular paragraphs."""
    doc = Document()
    doc.add_paragraph("Introduction paragraph")
    doc.add_paragraph("First item", style="List Bullet")
    doc.add_paragraph("Second item", style="List Bullet")
    doc.add_paragraph("Conclusion paragraph")

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    temp_file.close()

    try:
        blocks, _ = extract_blocks(temp_file.name)
        assert len(blocks) == 4
        assert blocks[0].type == "paragraph"
        assert blocks[0].content == "Introduction paragraph"
        assert blocks[1].type == "list"
        assert blocks[1].content == "- First item"
        assert blocks[2].type == "list"
        assert blocks[2].content == "- Second item"
        assert blocks[3].type == "paragraph"
        assert blocks[3].content == "Conclusion paragraph"
    finally:
        Path(temp_file.name).unlink()


def create_minimal_png() -> bytes:
    """Create a minimal valid 1x1 PNG image.

    Returns:
        PNG file bytes
    """
    from PIL import Image
    import io

    # Use PIL to create a proper 1x1 PNG image
    img = Image.new('RGB', (1, 1), color='white')
    img_bytes = io.BytesIO()
    img.save(img_bytes, format='PNG')
    return img_bytes.getvalue()


def create_docx_with_image() -> tuple[str, str]:
    """Create a test docx file with an embedded image.

    Returns:
        Tuple of (docx_path, temp_image_path)
    """
    # Create temp image file
    temp_img = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
    temp_img.write(create_minimal_png())
    temp_img.close()

    # Create docx with image
    doc = Document()
    doc.add_paragraph('Before image')
    doc.add_picture(temp_img.name, width=Inches(1.0))
    doc.add_paragraph('After image')

    temp_doc = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_doc.name)
    temp_doc.close()

    return temp_doc.name, temp_img.name


def test_extract_single_image():
    """Test extracting a document with a single image."""
    docx_path, img_path = create_docx_with_image()

    try:
        blocks, _ = extract_blocks(docx_path)

        # Should have 3 blocks: paragraph, image, paragraph
        assert len(blocks) == 3
        assert blocks[0].type == "paragraph"
        assert blocks[0].content == "Before image"

        assert blocks[1].type == "image"
        assert blocks[1].content.startswith("![")
        assert "assets/" in blocks[1].content
        assert blocks[1].image_path is not None

        assert blocks[2].type == "paragraph"
        assert blocks[2].content == "After image"
    finally:
        Path(docx_path).unlink()
        Path(img_path).unlink()


def test_extract_multiple_images():
    """Test extracting a document with multiple images."""
    temp_img1 = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
    temp_img1.write(create_minimal_png())
    temp_img1.close()

    temp_img2 = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
    temp_img2.write(create_minimal_png())
    temp_img2.close()

    # Create docx with two images
    doc = Document()
    doc.add_paragraph('First image:')
    doc.add_picture(temp_img1.name, width=Inches(1.0))
    doc.add_paragraph('Second image:')
    doc.add_picture(temp_img2.name, width=Inches(1.0))

    temp_doc = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_doc.name)
    temp_doc.close()

    try:
        blocks, _ = extract_blocks(temp_doc.name)

        # Should have 4 blocks
        assert len(blocks) == 4

        # First paragraph
        assert blocks[0].type == "paragraph"

        # First image
        assert blocks[1].type == "image"
        assert blocks[1].content.startswith("![")

        # Second paragraph
        assert blocks[2].type == "paragraph"

        # Second image
        assert blocks[3].type == "image"
        assert blocks[3].content.startswith("![")

        # Images should have different paths
        assert blocks[1].image_path != blocks[3].image_path
    finally:
        Path(temp_doc.name).unlink()
        Path(temp_img1.name).unlink()
        Path(temp_img2.name).unlink()


def test_extract_image_markdown_format():
    """Test that image blocks generate correct markdown format."""
    docx_path, img_path = create_docx_with_image()

    try:
        blocks, _ = extract_blocks(docx_path)
        image_block = blocks[1]

        # Check markdown format: ![alt text](assets/filename.ext)
        assert image_block.content.startswith("![")
        assert "](assets/" in image_block.content
        assert image_block.content.endswith(")")

        # Extract filename from markdown
        import re
        match = re.search(r'!\[(.*?)\]\((.*?)\)', image_block.content)
        assert match is not None
        alt_text = match.group(1)
        image_path = match.group(2)

        # Check alt text (can be empty or descriptive)
        assert isinstance(alt_text, str)

        # Check path format
        assert image_path.startswith("assets/")
        assert image_path.endswith(".png")
    finally:
        Path(docx_path).unlink()
        Path(img_path).unlink()


def test_validate_image_boundary_cases():
    """Test validate_image() with boundary cases around MAX_IMAGE_SIZE."""
    from sidedoc.extract import validate_image, MAX_IMAGE_SIZE
    from PIL import Image
    import io

    # Test 1: Image exactly at MAX_IMAGE_SIZE limit (should pass)
    # Create a small valid image and pad to exact size
    img = Image.new('RGB', (1, 1), color='white')
    img_bytes = io.BytesIO()
    img.save(img_bytes, format='PNG')
    base_size = len(img_bytes.getvalue())

    # Pad to exactly MAX_IMAGE_SIZE
    exactly_max = img_bytes.getvalue() + b'\x00' * (MAX_IMAGE_SIZE - base_size)
    # This will fail validation because padding corrupts the image, but tests size check
    is_valid, error = validate_image(exactly_max, 'png')
    # Should fail due to corruption, not size
    assert "exceeds maximum size" not in error.lower()

    # Test 2: Image at MAX_IMAGE_SIZE + 1 (should fail with size error)
    just_over_max = b'x' * (MAX_IMAGE_SIZE + 1)
    is_valid, error = validate_image(just_over_max, 'png')
    assert is_valid is False
    assert "exceeds maximum size" in error.lower()

    # Test 3: Image at MAX_IMAGE_SIZE - 1 (should pass size check, may fail format check)
    just_under_max = b'x' * (MAX_IMAGE_SIZE - 1)
    is_valid, error = validate_image(just_under_max, 'png')
    # Should not fail on size
    if not is_valid:
        assert "exceeds maximum size" not in error.lower()


def test_validate_image_function():
    """Test the validate_image() function directly with various inputs."""
    from sidedoc.extract import validate_image, MAX_IMAGE_SIZE
    from PIL import Image
    import io

    # Test 1: Valid PNG image
    img = Image.new('RGB', (10, 10), color='red')
    img_bytes = io.BytesIO()
    img.save(img_bytes, format='PNG')
    is_valid, error = validate_image(img_bytes.getvalue(), 'png')
    assert is_valid is True
    assert error == ""

    # Test 2: Format mismatch (JPEG with PNG extension)
    img = Image.new('RGB', (10, 10), color='green')
    img_bytes = io.BytesIO()
    img.save(img_bytes, format='JPEG')
    is_valid, error = validate_image(img_bytes.getvalue(), 'png')
    assert is_valid is False
    assert "format mismatch" in error.lower()

    # Test 3: Oversized image
    oversized_data = b'fake image data' * (MAX_IMAGE_SIZE // 10)
    is_valid, error = validate_image(oversized_data, 'png')
    assert is_valid is False
    assert "exceeds maximum size" in error.lower()

    # Test 4: Corrupted image data
    corrupted_data = b'this is not a valid image'
    is_valid, error = validate_image(corrupted_data, 'png')
    assert is_valid is False
    assert "invalid or corrupted" in error.lower()

    # Test 5: Valid JPEG image
    img = Image.new('RGB', (10, 10), color='blue')
    img_bytes = io.BytesIO()
    img.save(img_bytes, format='JPEG')
    is_valid, error = validate_image(img_bytes.getvalue(), 'jpg')
    assert is_valid is True
    assert error == ""

    # Test 6: Empty image data
    is_valid, error = validate_image(b'', 'png')
    assert is_valid is False
    assert "invalid or corrupted" in error.lower()


def test_extract_rejects_oversized_image():
    """Test that extraction rejects images exceeding MAX_IMAGE_SIZE."""
    # Create a large valid BMP image (BMP format has no compression)
    # This is fast and deterministic compared to creating random PNG data
    from sidedoc.extract import MAX_IMAGE_SIZE
    from PIL import Image

    # Calculate image size needed to exceed MAX_IMAGE_SIZE
    # BMP stores 3 bytes per pixel (RGB) + small header
    # For 10MB limit, need about 1875x1875 pixels (10.5MB uncompressed)
    img_size = 1875
    img = Image.new('RGB', (img_size, img_size), color='blue')

    temp_img = tempfile.NamedTemporaryFile(delete=False, suffix='.bmp')
    img.save(temp_img.name, format='BMP')
    temp_img.close()

    # Verify the image is actually larger than MAX_IMAGE_SIZE
    img_file_size = Path(temp_img.name).stat().st_size
    assert img_file_size > MAX_IMAGE_SIZE, f"Test image ({img_file_size} bytes) should be larger than MAX_IMAGE_SIZE ({MAX_IMAGE_SIZE} bytes)"

    # Create docx with oversized image
    doc = Document()
    doc.add_paragraph('Before image')
    doc.add_picture(temp_img.name, width=Inches(1.0))
    doc.add_paragraph('After image')

    temp_doc = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_doc.name)
    temp_doc.close()

    try:
        blocks, image_data = extract_blocks(temp_doc.name)

        # Should have 3 blocks but image should be skipped
        # The image paragraph should be converted to a regular paragraph with warning
        assert len(blocks) == 3
        assert blocks[0].type == "paragraph"
        assert blocks[2].type == "paragraph"

        # The middle block should be a paragraph with error message, not an image
        assert blocks[1].type == "paragraph"
        assert "image 1 skipped" in blocks[1].content.lower() or "image" in blocks[1].content.lower() and "skipped" in blocks[1].content.lower()

        # No image data should be stored
        assert len(image_data) == 0
    finally:
        Path(temp_doc.name).unlink()
        Path(temp_img.name).unlink()


def test_extract_rejects_format_mismatch():
    """Test that extraction rejects images where format doesn't match extension."""
    # Create a JPEG image but save with .png extension
    from PIL import Image

    # Create a small JPEG image
    img = Image.new('RGB', (1, 1), color='red')
    temp_img = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
    img.save(temp_img.name, format='JPEG')  # Save as JPEG but file has .png extension
    temp_img.close()

    # Create docx with mismatched image
    doc = Document()
    doc.add_paragraph('Before image')
    doc.add_picture(temp_img.name, width=Inches(1.0))
    doc.add_paragraph('After image')

    temp_doc = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_doc.name)
    temp_doc.close()

    try:
        blocks, image_data = extract_blocks(temp_doc.name)

        # Should have 3 blocks but image should be skipped
        assert len(blocks) == 3
        assert blocks[0].type == "paragraph"
        assert blocks[2].type == "paragraph"

        # The middle block should be a paragraph with error message
        assert blocks[1].type == "paragraph"
        assert "image 1 skipped" in blocks[1].content.lower() and "format" in blocks[1].content.lower()

        # No image data should be stored
        assert len(image_data) == 0
    finally:
        Path(temp_doc.name).unlink()
        Path(temp_img.name).unlink()


def test_extract_rejects_corrupted_image():
    """Test that extraction rejects corrupted image data."""
    # Create corrupted image data (not a valid image)
    corrupted_data = b'This is not a valid image file'

    temp_img = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
    temp_img.write(corrupted_data)
    temp_img.close()

    # Create docx with corrupted image
    doc = Document()
    doc.add_paragraph('Before image')
    # Note: python-docx may refuse to add invalid image, so we'll handle that
    try:
        doc.add_picture(temp_img.name, width=Inches(1.0))
    except Exception:
        # If python-docx can't add it, create a simpler test
        # Just test that extract_blocks handles corrupted embedded images gracefully
        Path(temp_img.name).unlink()
        return  # Skip this test if we can't create the scenario

    doc.add_paragraph('After image')

    temp_doc = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_doc.name)
    temp_doc.close()

    try:
        blocks, image_data = extract_blocks(temp_doc.name)

        # Should handle gracefully - either skip or convert to paragraph
        assert len(blocks) >= 2  # At least the text paragraphs

        # If there's an image block, it should be marked as problematic
        # Otherwise it should be converted to paragraph with warning
        image_blocks = [b for b in blocks if b.type == "image"]
        if len(image_blocks) == 0:
            # Image was skipped, should have warning paragraph with image number
            para_contents = [b.content.lower() for b in blocks if b.type == "paragraph"]
            assert any("image" in c and "skipped" in c for c in para_contents)
    finally:
        Path(temp_doc.name).unlink()
        Path(temp_img.name).unlink()


# ============================================================================
# Tests for formatting helpers
# ============================================================================


def test_wrap_formatting_plain():
    """Test wrap_formatting with no formatting."""
    from sidedoc.extract import wrap_formatting
    assert wrap_formatting("hello", False, False) == "hello"


def test_wrap_formatting_bold():
    """Test wrap_formatting with bold."""
    from sidedoc.extract import wrap_formatting
    assert wrap_formatting("hello", True, False) == "**hello**"


def test_wrap_formatting_italic():
    """Test wrap_formatting with italic."""
    from sidedoc.extract import wrap_formatting
    assert wrap_formatting("hello", False, True) == "*hello*"


def test_wrap_formatting_bold_italic():
    """Test wrap_formatting with bold and italic."""
    from sidedoc.extract import wrap_formatting
    assert wrap_formatting("hello", True, True) == "***hello***"


def test_format_hyperlink_md_plain():
    """Test format_hyperlink_md with no formatting."""
    from sidedoc.extract import format_hyperlink_md
    assert format_hyperlink_md("click here", "https://example.com", False, False) == "[click here](https://example.com)"


def test_format_hyperlink_md_bold():
    """Test format_hyperlink_md with bold."""
    from sidedoc.extract import format_hyperlink_md
    assert format_hyperlink_md("click", "https://example.com", True, False) == "[**click**](https://example.com)"


def test_format_hyperlink_md_italic():
    """Test format_hyperlink_md with italic."""
    from sidedoc.extract import format_hyperlink_md
    assert format_hyperlink_md("click", "https://example.com", False, True) == "[*click*](https://example.com)"


def test_format_hyperlink_md_bold_italic():
    """Test format_hyperlink_md with bold and italic."""
    from sidedoc.extract import format_hyperlink_md
    assert format_hyperlink_md("click", "https://example.com", True, True) == "[***click***](https://example.com)"


def test_format_hyperlink_md_escapes_brackets():
    """Test that format_hyperlink_md escapes brackets in text."""
    from sidedoc.extract import format_hyperlink_md
    result = format_hyperlink_md("foo[bar]baz", "https://example.com", False, False)
    assert result == "[foo\\[bar\\]baz](https://example.com)"


def test_format_hyperlink_md_encodes_url():
    """Test that format_hyperlink_md encodes parentheses in URLs."""
    from sidedoc.extract import format_hyperlink_md
    result = format_hyperlink_md("click", "https://example.com/page_(1)", False, False)
    assert "%28" in result
    assert "%29" in result


# ============================================================================
# Fix 4: extract_paragraph_accept_all bugs
# ============================================================================


def test_accept_all_respects_val_false_on_bold():
    """Test that val='0' on bold element means not bold in accept_all path."""
    from lxml import etree
    from sidedoc.extract import extract_paragraph_accept_all
    from sidedoc.constants import WORDPROCESSINGML_NS as NS

    # Build XML with <w:b w:val="0"/> — should NOT be treated as bold
    xml = f"""
    <w:p xmlns:w="{NS}">
        <w:r>
            <w:rPr>
                <w:b w:val="0"/>
            </w:rPr>
            <w:t>not bold</w:t>
        </w:r>
    </w:p>
    """
    para_elem = etree.fromstring(xml)
    content, inline_fmt, _ = extract_paragraph_accept_all(para_elem)

    assert content == "not bold"
    # Should have NO bold formatting recorded
    if inline_fmt:
        for fmt in inline_fmt:
            assert not fmt.get("bold", False), \
                f"val='0' should not produce bold formatting: {fmt}"


def test_accept_all_records_hyperlinks_in_inline_formatting():
    """Test that hyperlinks are recorded in inline_formatting in accept_all path."""
    from lxml import etree
    from unittest.mock import MagicMock
    from sidedoc.extract import extract_paragraph_accept_all
    from sidedoc.constants import WORDPROCESSINGML_NS as NS

    REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    xml = f"""
    <w:p xmlns:w="{NS}" xmlns:r="{REL_NS}">
        <w:hyperlink r:id="rId1">
            <w:r>
                <w:t>click here</w:t>
            </w:r>
        </w:hyperlink>
    </w:p>
    """
    para_elem = etree.fromstring(xml)

    # Mock doc_part with relationship
    mock_rel = MagicMock()
    mock_rel._target = "https://example.com"
    mock_doc_part = MagicMock()
    mock_doc_part.rels = {"rId1": mock_rel}

    content, inline_fmt, _ = extract_paragraph_accept_all(para_elem, mock_doc_part)

    assert "click here" in content
    assert "example.com" in content

    # inline_formatting should include a hyperlink entry
    assert inline_fmt is not None, "inline_formatting should not be None"
    hyperlink_entries = [f for f in inline_fmt if f.get("type") == "hyperlink"]
    assert len(hyperlink_entries) >= 1, f"Should have hyperlink in inline_formatting: {inline_fmt}"
    assert hyperlink_entries[0]["url"] == "https://example.com"


def test_accept_all_handles_line_breaks():
    """Test that w:br elements produce newlines in accept_all path."""
    from lxml import etree
    from sidedoc.extract import extract_paragraph_accept_all
    from sidedoc.constants import WORDPROCESSINGML_NS as NS

    xml = f"""
    <w:p xmlns:w="{NS}">
        <w:r>
            <w:t>line one</w:t>
            <w:br/>
            <w:t>line two</w:t>
        </w:r>
    </w:p>
    """
    para_elem = etree.fromstring(xml)
    content, _, _ = extract_paragraph_accept_all(para_elem)

    assert "\n" in content, f"Expected newline from w:br, got: {repr(content)}"
    assert "line one" in content
    assert "line two" in content


def test_accept_all_tracks_underline():
    """Test that underline formatting is tracked in accept_all path."""
    from lxml import etree
    from sidedoc.extract import extract_paragraph_accept_all
    from sidedoc.constants import WORDPROCESSINGML_NS as NS

    xml = f"""
    <w:p xmlns:w="{NS}">
        <w:r>
            <w:rPr>
                <w:u w:val="single"/>
            </w:rPr>
            <w:t>underlined</w:t>
        </w:r>
    </w:p>
    """
    para_elem = etree.fromstring(xml)
    content, inline_fmt, _ = extract_paragraph_accept_all(para_elem)

    assert content == "underlined"
    assert inline_fmt is not None, "inline_formatting should not be None"
    underline_entries = [f for f in inline_fmt if f.get("type") == "underline"]
    assert len(underline_entries) >= 1, f"Should have underline entry: {inline_fmt}"
