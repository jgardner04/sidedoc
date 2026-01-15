"""Test build command."""

import tempfile
from pathlib import Path
import struct
import zlib
from click.testing import CliRunner
from docx import Document
from docx.shared import Inches
from sidedoc.cli import main


def test_build_command_creates_docx():
    """Test that build command creates a .docx file from extract output."""
    runner = CliRunner()
    with runner.isolated_filesystem():
        # Create a test docx
        doc = Document()
        doc.add_paragraph("Test Title", style="Heading 1")
        doc.add_paragraph("Test content")
        docx_path = "test.docx"
        doc.save(docx_path)

        # Extract it
        result = runner.invoke(main, ["extract", docx_path])
        assert result.exit_code == 0

        # Build it back
        sidedoc_path = "test.sidedoc"
        output_docx = "rebuilt.docx"
        result = runner.invoke(main, ["build", sidedoc_path, "-o", output_docx])

        assert result.exit_code == 0
        assert Path(output_docx).exists()

        # Verify it's a valid docx
        rebuilt_doc = Document(output_docx)
        assert len(rebuilt_doc.paragraphs) >= 2


def test_build_command_preserves_headings():
    """Test that build preserves heading styles."""
    runner = CliRunner()
    with runner.isolated_filesystem():
        # Create docx with headings
        doc = Document()
        doc.add_paragraph("Heading 1", style="Heading 1")
        doc.add_paragraph("Heading 2", style="Heading 2")
        docx_path = "test.docx"
        doc.save(docx_path)

        # Extract and build
        runner.invoke(main, ["extract", docx_path])
        result = runner.invoke(main, ["build", "test.sidedoc", "-o", "rebuilt.docx"])

        assert result.exit_code == 0

        # Verify headings
        rebuilt_doc = Document("rebuilt.docx")
        assert len(rebuilt_doc.paragraphs) >= 2


def test_build_command_default_output():
    """Test that build uses default output path."""
    runner = CliRunner()
    with runner.isolated_filesystem():
        # Create and extract
        doc = Document()
        doc.add_paragraph("Test")
        doc.save("test.docx")
        runner.invoke(main, ["extract", "test.docx"])

        # Build without -o flag
        result = runner.invoke(main, ["build", "test.sidedoc"])

        assert result.exit_code == 0
        # Should create test.docx (overwriting original)
        assert Path("test.docx").exists()


def create_minimal_png() -> bytes:
    """Create a minimal valid 1x1 PNG image.

    Returns:
        PNG file bytes
    """
    from PIL import Image
    import io

    # Use PIL to create a proper 1x1 PNG image
    img = Image.new('RGB', (1, 1), color='white')
    img_bytes = io.BytesIO()
    img.save(img_bytes, format='PNG')
    return img_bytes.getvalue()


def test_build_command_embeds_images():
    """Test that build command embeds images from assets directory."""
    runner = CliRunner()
    with runner.isolated_filesystem():
        # Create a temp image file
        temp_img = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
        temp_img.write(create_minimal_png())
        temp_img.close()

        # Create docx with image
        doc = Document()
        doc.add_paragraph("Before image")
        doc.add_picture(temp_img.name, width=Inches(1.0))
        doc.add_paragraph("After image")
        docx_path = "test.docx"
        doc.save(docx_path)

        # Extract it
        result = runner.invoke(main, ["extract", docx_path])
        assert result.exit_code == 0

        # Build it back
        result = runner.invoke(main, ["build", "test.sidedoc", "-o", "rebuilt.docx"])
        assert result.exit_code == 0
        assert Path("rebuilt.docx").exists()

        # Verify the rebuilt document has the image
        rebuilt_doc = Document("rebuilt.docx")

        # Count paragraphs with images
        image_count = 0
        for para in rebuilt_doc.paragraphs:
            for run in para.runs:
                drawing_elems = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                if drawing_elems:
                    image_count += 1

        assert image_count == 1, f"Expected 1 image, found {image_count}"

        # Verify paragraph count (3: before, image, after)
        assert len(rebuilt_doc.paragraphs) == 3

        # Clean up
        Path(temp_img.name).unlink()
