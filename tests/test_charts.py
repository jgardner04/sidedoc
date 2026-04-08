"""Tests for chart extraction and reconstruction."""

import tempfile
from pathlib import Path

from docx import Document

from sidedoc.constants import WORDPROCESSINGML_NS
from sidedoc.extract import extract_blocks, extract_chart_from_paragraph, validate_image


FIXTURES_DIR = Path(__file__).parent / "fixtures"


# --- Step 2: Chart detection ---


def test_extract_chart_from_paragraph_detects_chart():
    """Chart paragraph returns chart info with cached image bytes and rel ID."""
    doc = Document(str(FIXTURES_DIR / "charts.docx"))
    chart_para = doc.paragraphs[1]  # The chart paragraph

    result = extract_chart_from_paragraph(chart_para, doc.part, image_counter=1)

    assert result is not None
    image_filename, extension, image_bytes, error_message, chart_rel_id = result
    assert image_filename == "chart1.png"
    assert extension == "png"
    assert len(image_bytes) > 0
    assert error_message == ""
    assert chart_rel_id == "rId5"


def test_extract_chart_from_paragraph_returns_none_for_regular_paragraph():
    """Non-chart paragraph returns None."""
    doc = Document(str(FIXTURES_DIR / "charts.docx"))
    text_para = doc.paragraphs[0]  # "Before chart"

    result = extract_chart_from_paragraph(text_para, doc.part, image_counter=1)

    assert result is None


def test_extract_chart_from_paragraph_returns_none_for_regular_image():
    """Paragraph with a regular image (no chart) returns None."""
    doc = Document(str(FIXTURES_DIR / "images.docx"))
    # Paragraph index 2 is the first image per create_fixtures.py:
    # [0] heading, [1] "First image:", [2] image, [3] "Second image:", [4] image, [5] text
    image_para = doc.paragraphs[2]

    # Precondition: confirm this paragraph actually contains a drawing (image)
    drawings = image_para._element.findall(f'.//{{{WORDPROCESSINGML_NS}}}drawing')
    assert len(drawings) > 0, "Expected paragraph[2] to contain a drawing element"

    result = extract_chart_from_paragraph(image_para, doc.part, image_counter=1)
    assert result is None


# --- Step 3: EMF validation pass-through ---


def test_validate_image_emf_passes_without_pil():
    """EMF bytes pass validation without PIL (pass-through format)."""
    # EMF files start with a specific header but PIL can't open them.
    # Use arbitrary bytes — EMF validation should skip PIL entirely.
    fake_emf_bytes = b"\x01\x00\x00\x00" * 10
    is_valid, error_message = validate_image(fake_emf_bytes, "emf")
    assert is_valid is True
    assert error_message == ""


def test_validate_image_wmf_passes_without_pil():
    """WMF bytes pass validation without PIL (pass-through format)."""
    fake_wmf_bytes = b"\xd7\xcd\xc6\x9a" * 10
    is_valid, error_message = validate_image(fake_wmf_bytes, "wmf")
    assert is_valid is True
    assert error_message == ""


def test_validate_image_emf_respects_size_limit():
    """EMF pass-through still enforces the max size limit."""
    from sidedoc.constants import MAX_IMAGE_SIZE

    oversized_bytes = b"\x00" * (MAX_IMAGE_SIZE + 1)
    is_valid, error_message = validate_image(oversized_bytes, "emf")
    assert is_valid is False
    assert "exceeds maximum size" in error_message


def test_validate_image_png_unchanged():
    """Regular PNG validation is unaffected by EMF changes."""
    from PIL import Image
    import io

    img = Image.new("RGB", (1, 1), color="red")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    png_bytes = buf.getvalue()

    is_valid, error_message = validate_image(png_bytes, "png")
    assert is_valid is True
    assert error_message == ""


# --- Step 4: Chart block creation in extract_blocks() ---


def test_extract_blocks_chart_produces_chart_block():
    """Chart docx produces a Block with type='chart' and correct content."""
    blocks, image_data = extract_blocks(str(FIXTURES_DIR / "charts.docx"))

    chart_blocks = [b for b in blocks if b.type == "chart"]
    assert len(chart_blocks) == 1

    chart = chart_blocks[0]
    assert chart.image_path is not None
    assert chart.image_path.startswith("assets/chart")
    assert chart.content.startswith("![Chart")
    assert chart.chart_metadata == {"chart_rel_id": "rId5"}


def test_extract_blocks_chart_image_in_image_data():
    """Chart cached image bytes are included in image_data dict."""
    blocks, image_data = extract_blocks(str(FIXTURES_DIR / "charts.docx"))

    chart_blocks = [b for b in blocks if b.type == "chart"]
    assert len(chart_blocks) == 1

    # image_path is "assets/chartN.ext", image_data key is just "chartN.ext"
    image_filename = chart_blocks[0].image_path.split("/")[-1]
    assert image_filename in image_data
    assert len(image_data[image_filename]) > 0


def test_extract_blocks_regular_images_not_affected():
    """Regular images still produce type='image' blocks (no regression)."""
    blocks, image_data = extract_blocks(str(FIXTURES_DIR / "images.docx"))

    image_blocks = [b for b in blocks if b.type == "image"]
    chart_blocks = [b for b in blocks if b.type == "chart"]

    assert len(image_blocks) > 0
    assert len(chart_blocks) == 0


def test_extract_blocks_chart_text_paragraphs_preserved():
    """Text paragraphs before and after chart are preserved."""
    blocks, _ = extract_blocks(str(FIXTURES_DIR / "charts.docx"))

    assert blocks[0].type == "paragraph"
    assert blocks[0].content == "Before chart"
    assert blocks[-1].type == "paragraph"
    assert blocks[-1].content == "After chart"


# --- Step 5: Missing cached image fallback ---


def test_extract_blocks_chart_no_fallback_produces_placeholder():
    """Chart with no cached image produces a paragraph placeholder."""
    blocks, image_data = extract_blocks(str(FIXTURES_DIR / "charts_no_fallback.docx"))

    # Should not produce a chart block (no image to extract)
    chart_blocks = [b for b in blocks if b.type == "chart"]
    assert len(chart_blocks) == 0

    # Should produce a placeholder paragraph
    placeholder_blocks = [b for b in blocks if "no preview available" in b.content]
    assert len(placeholder_blocks) == 1
    assert placeholder_blocks[0].type == "paragraph"
    assert "[Chart: no preview available]" in placeholder_blocks[0].content
    # chart_metadata preserves the rel ID even on degraded blocks (for JON-107)
    assert placeholder_blocks[0].chart_metadata is not None


# --- Step 6: Chart markdown parsing ---


def test_parse_markdown_chart_produces_chart_block():
    """Markdown ![Chart: Revenue](assets/chart1.png) parsed as type='chart'."""
    from sidedoc.reconstruct import parse_markdown_to_blocks

    blocks = parse_markdown_to_blocks("![Chart: Revenue](assets/chart1.png)")
    assert len(blocks) == 1
    assert blocks[0].type == "chart"
    assert blocks[0].image_path == "assets/chart1.png"


def test_parse_markdown_chart_no_title_produces_chart_block():
    """Markdown ![Chart](assets/chart1.png) parsed as type='chart'."""
    from sidedoc.reconstruct import parse_markdown_to_blocks

    blocks = parse_markdown_to_blocks("![Chart](assets/chart1.png)")
    assert len(blocks) == 1
    assert blocks[0].type == "chart"


def test_parse_markdown_regular_image_unchanged():
    """Markdown ![Screenshot](assets/image1.png) still parsed as type='image'."""
    from sidedoc.reconstruct import parse_markdown_to_blocks

    blocks = parse_markdown_to_blocks("![Screenshot](assets/image1.png)")
    assert len(blocks) == 1
    assert blocks[0].type == "image"


# --- Step 7: Chart reconstruction ---


def _make_styles_dict() -> dict:
    """Create a minimal styles dict for testing reconstruction."""
    return {
        "block_styles": {},
        "document_defaults": {"font_name": "Arial", "font_size": 11},
    }


def test_reconstruct_chart_block_with_image():
    """Chart block with image_path produces a docx with the image embedded."""
    from sidedoc.reconstruct import create_docx_from_blocks
    from sidedoc.models import Block

    chart_block = Block(
        id="block-0",
        type="chart",
        content="![Chart](assets/chart1.png)",
        docx_paragraph_index=0,
        content_start=0,
        content_end=27,
        content_hash="abc",
        image_path="assets/chart1.png",
    )

    from PIL import Image

    with tempfile.TemporaryDirectory() as tmpdir:
        assets_dir = Path(tmpdir) / "assets"
        assets_dir.mkdir()
        img = Image.new("RGB", (10, 10), color="green")
        img.save(assets_dir / "chart1.png", format="PNG")

        doc = create_docx_from_blocks(
            [chart_block], _make_styles_dict(), assets_dir=assets_dir
        )

        # Verify the document has an embedded image (not just text)
        # Check for inline shapes (pictures) in the document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        image_rels = [
            rel for rel in doc.part.rels.values()
            if "image" in rel.reltype
        ]
        assert len(image_rels) > 0, "Chart should be embedded as an image"


def test_reconstruct_chart_block_without_image():
    """Chart block with no image_path produces placeholder text."""
    from sidedoc.reconstruct import create_docx_from_blocks
    from sidedoc.models import Block

    chart_block = Block(
        id="block-0",
        type="chart",
        content="[Chart: no preview available]",
        docx_paragraph_index=0,
        content_start=0,
        content_end=29,
        content_hash="abc",
    )

    doc = create_docx_from_blocks([chart_block], _make_styles_dict())
    texts = [p.text for p in doc.paragraphs]
    assert any("[Chart" in t for t in texts)


# --- Step 8: Integration tests ---


def test_extract_build_roundtrip_preserves_chart():
    """Extract → build round-trip preserves chart as image in rebuilt docx."""
    from click.testing import CliRunner
    from sidedoc.cli import main

    runner = CliRunner()
    with runner.isolated_filesystem():
        # Copy chart fixture to working directory
        import shutil
        shutil.copy(FIXTURES_DIR / "charts.docx", "charts.docx")

        # Extract
        result = runner.invoke(main, ["extract", "charts.docx", "--force"])
        assert result.exit_code == 0, f"Extract failed: {result.output}"

        # Verify content.md has chart reference
        content = Path("charts.sidedoc/content.md").read_text()
        assert "![Chart]" in content
        assert "assets/chart" in content

        # Verify chart image in assets
        assets = list(Path("charts.sidedoc/assets").glob("chart*"))
        assert len(assets) > 0

        # Build
        result = runner.invoke(main, ["build", "charts.sidedoc", "-o", "rebuilt.docx"])
        assert result.exit_code == 0, f"Build failed: {result.output}"

        # Verify rebuilt docx exists and has content
        doc = Document("rebuilt.docx")
        assert len(doc.paragraphs) > 0


def test_cli_extract_chart_succeeds():
    """sidedoc extract on chart docx succeeds without errors."""
    from click.testing import CliRunner
    from sidedoc.cli import main

    runner = CliRunner()
    with runner.isolated_filesystem():
        import shutil
        shutil.copy(FIXTURES_DIR / "charts.docx", "charts.docx")

        result = runner.invoke(main, ["extract", "charts.docx", "--force"])
        assert result.exit_code == 0
        assert "Error" not in result.output


# --- PR #58 review feedback: additional tests ---


def test_block_to_structure_dict_includes_chart_metadata():
    """chart_metadata must be serialized into structure.json dict."""
    from sidedoc.models import Block
    from sidedoc.package import block_to_structure_dict

    block = Block(
        id="block-0",
        type="chart",
        content="![Chart](assets/chart1.png)",
        docx_paragraph_index=0,
        content_start=0,
        content_end=27,
        content_hash="abc",
        image_path="assets/chart1.png",
        chart_metadata={"chart_rel_id": "rId5"},
    )
    d = block_to_structure_dict(block)
    assert "chart_metadata" in d
    assert d["chart_metadata"] == {"chart_rel_id": "rId5"}


def test_extract_blocks_chart_validation_error_produces_placeholder():
    """Chart with cached image that fails validation produces a skipped placeholder."""
    from unittest.mock import patch

    # Patch validate_image to simulate a validation failure
    with patch("sidedoc.extract.validate_image", return_value=(False, "exceeds maximum size")):
        blocks, image_data = extract_blocks(str(FIXTURES_DIR / "charts.docx"))

    # The chart should degrade to a paragraph with a "skipped" message
    skipped = [b for b in blocks if "skipped" in b.content]
    assert len(skipped) == 1
    assert skipped[0].type == "paragraph"
    assert "exceeds maximum size" in skipped[0].content


def test_chart_metadata_preserved_on_degraded_blocks():
    """chart_metadata with chart_rel_id is set even when chart degrades to paragraph."""
    blocks, _ = extract_blocks(str(FIXTURES_DIR / "charts_no_fallback.docx"))

    # The no-fallback chart has a chart_rel_id but no cached image,
    # so it degrades to a paragraph. chart_metadata should still be set.
    placeholder = [b for b in blocks if "no preview available" in b.content]
    assert len(placeholder) == 1
    assert placeholder[0].chart_metadata is not None
    assert "chart_rel_id" in placeholder[0].chart_metadata


def test_parse_markdown_lowercase_chart_is_image():
    """Lowercase ![chart](...) is classified as image, not chart (case-sensitive)."""
    from sidedoc.reconstruct import parse_markdown_to_blocks

    blocks = parse_markdown_to_blocks("![chart](assets/chart1.png)")
    assert len(blocks) == 1
    assert blocks[0].type == "image"


# --- PR #58 review feedback: security hardening tests ---


def test_extract_blip_image_skips_external_relationship():
    """_extract_blip_image returns None for external relationships."""
    from unittest.mock import MagicMock, PropertyMock
    from sidedoc.extract import _extract_blip_image

    blip = MagicMock()
    blip.get.return_value = "rId1"

    # Mock an external relationship — target_part raises if accessed
    rel = MagicMock()
    rel.is_external = True
    type(rel).target_part = PropertyMock(
        side_effect=AssertionError("target_part should not be accessed")
    )
    doc_part = MagicMock()
    doc_part.rels = {"rId1": rel}

    result = _extract_blip_image(blip, doc_part, "image", 1)
    assert result is None


def test_validate_image_emf_rejects_wrong_magic_bytes():
    """EMF extension with wrong magic bytes is rejected."""
    bad_emf = b"\xFF\xFF\xFF\xFF" * 10
    is_valid, error_message = validate_image(bad_emf, "emf")
    assert is_valid is False
    assert "magic bytes" in error_message.lower() or "EMF" in error_message


def test_validate_image_wmf_rejects_wrong_magic_bytes():
    """WMF extension with wrong magic bytes is rejected."""
    bad_wmf = b"\xFF\xFF\xFF\xFF" * 10
    is_valid, error_message = validate_image(bad_wmf, "wmf")
    assert is_valid is False
    assert "magic bytes" in error_message.lower() or "WMF" in error_message


def test_extract_blip_image_sanitizes_extension():
    """_extract_blip_image sanitizes pathological extensions."""
    from unittest.mock import MagicMock, PropertyMock
    from sidedoc.extract import _extract_blip_image

    blip = MagicMock()
    blip.get.return_value = "rId1"

    # Mock a part with a pathological partname (no dot)
    rel = MagicMock()
    rel.is_external = False
    image_part = MagicMock()
    type(image_part).partname = PropertyMock(return_value="/word/media/nodotfile")
    image_part.blob = b"\x89PNG\r\n\x1a\n" + b"\x00" * 100  # minimal bytes

    rel.target_part = image_part
    doc_part = MagicMock()
    doc_part.rels = {"rId1": rel}

    result = _extract_blip_image(blip, doc_part, "image", 1)
    assert result is not None
    filename, extension, _, _ = result
    # Extension should be sanitized — no path separators, fallback to "bin" if empty
    assert "/" not in extension
    assert "\\" not in extension
    assert len(extension) <= 10
    assert extension == "bin"


def test_validate_image_wmf_disk_variant_passes():
    """WMF disk metafile variant (0x0002) is accepted as valid."""
    # MetafileType=0x0002 (DISKMETAFILE), HeaderSize=0x0009 per MS-WMF spec
    disk_wmf = b"\x02\x00\x09\x00" * 10
    is_valid, error_message = validate_image(disk_wmf, "wmf")
    assert is_valid is True
    assert error_message == ""
