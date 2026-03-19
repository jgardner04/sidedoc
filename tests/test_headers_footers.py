"""Test headers and footers extraction, storage, and reconstruction."""

import io
import json
import tempfile
import zipfile
from pathlib import Path

from click.testing import CliRunner
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from PIL import Image as PILImage

from sidedoc.cli import main
from sidedoc.extract import extract_blocks, extract_sections


def _create_docx_with_headers_footers() -> str:
    """Create a test docx with default headers and footers."""
    doc = Document()
    section = doc.sections[0]

    header = section.header
    header.is_linked_to_previous = False
    header.paragraphs[0].text = "Default Header"

    footer = section.footer
    footer.is_linked_to_previous = False
    footer.paragraphs[0].text = "Default Footer"

    doc.add_paragraph("Body paragraph one.")
    doc.add_paragraph("Body paragraph two.")

    temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp.name)
    temp.close()
    return temp.name


def _create_docx_with_first_page_variant() -> str:
    """Create a test docx with different first page headers/footers."""
    doc = Document()
    section = doc.sections[0]

    section.different_first_page_header_footer = True

    header = section.header
    header.is_linked_to_previous = False
    header.paragraphs[0].text = "Default Header"

    footer = section.footer
    footer.is_linked_to_previous = False
    footer.paragraphs[0].text = "Default Footer"

    first_header = section.first_page_header
    first_header.paragraphs[0].text = "First Page Header"

    first_footer = section.first_page_footer
    first_footer.paragraphs[0].text = "First Page Footer"

    doc.add_paragraph("Body content.")

    temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp.name)
    temp.close()
    return temp.name


def _create_docx_with_multi_section() -> str:
    """Create a test docx with multiple sections having different headers."""
    doc = Document()

    section1 = doc.sections[0]
    header1 = section1.header
    header1.is_linked_to_previous = False
    header1.paragraphs[0].text = "Section 1 Header"

    footer1 = section1.footer
    footer1.is_linked_to_previous = False
    footer1.paragraphs[0].text = "Section 1 Footer"

    doc.add_paragraph("Section 1 body.")

    doc.add_section()
    section2 = doc.sections[1]
    header2 = section2.header
    header2.is_linked_to_previous = False
    header2.paragraphs[0].text = "Section 2 Header"

    footer2 = section2.footer
    footer2.is_linked_to_previous = False
    footer2.paragraphs[0].text = "Section 2 Footer"

    doc.add_paragraph("Section 2 body.")

    temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp.name)
    temp.close()
    return temp.name


def _create_docx_with_section_properties() -> str:
    """Create a test docx with custom section properties (margins, orientation)."""
    doc = Document()
    section = doc.sections[0]

    section.orientation = WD_ORIENT.LANDSCAPE
    old_width = section.page_width
    old_height = section.page_height
    section.page_width = old_height
    section.page_height = old_width

    section.top_margin = Inches(1.5)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    header = section.header
    header.is_linked_to_previous = False
    header.paragraphs[0].text = "Landscape Header"

    doc.add_paragraph("Landscape body.")

    temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp.name)
    temp.close()
    return temp.name


# ============================================================
# Extraction tests
# ============================================================


class TestExtractSections:
    """Test section metadata extraction from docx files."""

    def test_extract_sections_returns_list(self):
        """extract_sections returns a list of section dicts."""
        docx_path = _create_docx_with_headers_footers()
        try:
            sections, _ = extract_sections(docx_path)
            assert isinstance(sections, list)
            assert len(sections) >= 1
        finally:
            Path(docx_path).unlink()

    def test_extract_default_header(self):
        """Default header text is extracted."""
        docx_path = _create_docx_with_headers_footers()
        try:
            sections, _ = extract_sections(docx_path)
            section = sections[0]
            assert "header_default" in section
            paragraphs = section["header_default"]
            assert len(paragraphs) >= 1
            assert paragraphs[0]["content"] == "Default Header"
        finally:
            Path(docx_path).unlink()

    def test_extract_default_footer(self):
        """Default footer text is extracted."""
        docx_path = _create_docx_with_headers_footers()
        try:
            sections, _ = extract_sections(docx_path)
            section = sections[0]
            assert "footer_default" in section
            paragraphs = section["footer_default"]
            assert len(paragraphs) >= 1
            assert paragraphs[0]["content"] == "Default Footer"
        finally:
            Path(docx_path).unlink()

    def test_extract_first_page_header(self):
        """First page header variant is extracted."""
        docx_path = _create_docx_with_first_page_variant()
        try:
            sections, _ = extract_sections(docx_path)
            section = sections[0]
            assert "header_first" in section
            assert section["header_first"][0]["content"] == "First Page Header"
        finally:
            Path(docx_path).unlink()

    def test_extract_first_page_footer(self):
        """First page footer variant is extracted."""
        docx_path = _create_docx_with_first_page_variant()
        try:
            sections, _ = extract_sections(docx_path)
            section = sections[0]
            assert "footer_first" in section
            assert section["footer_first"][0]["content"] == "First Page Footer"
        finally:
            Path(docx_path).unlink()

    def test_extract_multi_section(self):
        """Multiple sections with different headers are extracted."""
        docx_path = _create_docx_with_multi_section()
        try:
            sections, _ = extract_sections(docx_path)
            assert len(sections) == 2
            assert sections[0]["header_default"][0]["content"] == "Section 1 Header"
            assert sections[1]["header_default"][0]["content"] == "Section 2 Header"
            assert sections[0]["footer_default"][0]["content"] == "Section 1 Footer"
            assert sections[1]["footer_default"][0]["content"] == "Section 2 Footer"
        finally:
            Path(docx_path).unlink()

    def test_extract_section_properties(self):
        """Section properties (margins, orientation) are extracted."""
        docx_path = _create_docx_with_section_properties()
        try:
            sections, _ = extract_sections(docx_path)
            section = sections[0]
            assert "page_setup" in section
            setup = section["page_setup"]
            assert setup["orientation"] == "landscape"
            assert setup["top_margin"] == Inches(1.5)
            assert setup["bottom_margin"] == Inches(1.0)
            assert setup["left_margin"] == Inches(0.75)
            assert setup["right_margin"] == Inches(0.75)
        finally:
            Path(docx_path).unlink()

    def test_extract_empty_header_not_included(self):
        """Sections with no real header content should have empty lists."""
        doc = Document()
        doc.add_paragraph("Just body.")
        temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp.name)
        temp.close()
        try:
            sections, _ = extract_sections(temp.name)
            section = sections[0]
            assert section["header_default"] == []
            assert section["footer_default"] == []
        finally:
            Path(temp.name).unlink()

    def test_extract_blocks_unchanged(self):
        """extract_blocks should still only return body blocks (no regression)."""
        docx_path = _create_docx_with_headers_footers()
        try:
            blocks, _ = extract_blocks(docx_path)
            contents = [b.content for b in blocks]
            assert "Body paragraph one." in contents
            assert "Body paragraph two." in contents
            assert "Default Header" not in contents
            assert "Default Footer" not in contents
        finally:
            Path(docx_path).unlink()

    def test_header_with_multiple_paragraphs(self):
        """Headers with multiple paragraphs are all extracted."""
        doc = Document()
        section = doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False
        header.paragraphs[0].text = "Line 1"
        header.add_paragraph("Line 2")

        doc.add_paragraph("Body.")

        temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp.name)
        temp.close()
        try:
            sections, _ = extract_sections(temp.name)
            paragraphs = sections[0]["header_default"]
            assert len(paragraphs) == 2
            assert paragraphs[0]["content"] == "Line 1"
            assert paragraphs[1]["content"] == "Line 2"
        finally:
            Path(temp.name).unlink()

    def test_header_paragraph_type(self):
        """Header paragraphs include type field."""
        docx_path = _create_docx_with_headers_footers()
        try:
            sections, _ = extract_sections(docx_path)
            para = sections[0]["header_default"][0]
            assert para["type"] == "paragraph"
        finally:
            Path(docx_path).unlink()


# ============================================================
# Packaging tests
# ============================================================


class TestSectionMetadataPackaging:
    """Test that section metadata is stored in structure.json."""

    def test_sections_in_structure_json(self):
        """Section metadata appears in structure.json."""
        runner = CliRunner()
        with runner.isolated_filesystem():
            docx_path = _create_docx_with_headers_footers()
            Path(docx_path).rename("test.docx")

            result = runner.invoke(main, ["extract", "test.docx"])
            assert result.exit_code == 0

            structure = json.loads(Path("test.sidedoc/structure.json").read_text())
            assert "sections" in structure
            sections = structure["sections"]
            assert len(sections) >= 1
            assert "header_default" in sections[0]

    def test_sections_in_sdoc_zip(self):
        """Section metadata in structure.json within .sdoc ZIP."""
        runner = CliRunner()
        with runner.isolated_filesystem():
            docx_path = _create_docx_with_headers_footers()
            Path(docx_path).rename("test.docx")

            result = runner.invoke(main, ["extract", "test.docx", "--pack"])
            assert result.exit_code == 0

            with zipfile.ZipFile("test.sdoc", "r") as zf:
                structure = json.loads(zf.read("structure.json"))
            assert "sections" in structure


# ============================================================
# Reconstruction tests
# ============================================================


class TestHeaderFooterReconstruction:
    """Test build command reconstructs headers and footers."""

    def test_roundtrip_default_header_footer(self):
        """Extract -> build preserves default header and footer."""
        runner = CliRunner()
        with runner.isolated_filesystem():
            docx_path = _create_docx_with_headers_footers()
            Path(docx_path).rename("test.docx")

            extract_result = runner.invoke(main, ["extract", "test.docx"])
            assert extract_result.exit_code == 0
            result = runner.invoke(main, ["build", "test.sidedoc", "-o", "rebuilt.docx"])
            assert result.exit_code == 0

            rebuilt = Document("rebuilt.docx")
            section = rebuilt.sections[0]
            header_text = section.header.paragraphs[0].text
            footer_text = section.footer.paragraphs[0].text
            assert header_text == "Default Header"
            assert footer_text == "Default Footer"

    def test_roundtrip_first_page_variant(self):
        """Extract -> build preserves first page header/footer variants."""
        runner = CliRunner()
        with runner.isolated_filesystem():
            docx_path = _create_docx_with_first_page_variant()
            Path(docx_path).rename("test.docx")

            extract_result = runner.invoke(main, ["extract", "test.docx"])
            assert extract_result.exit_code == 0
            result = runner.invoke(main, ["build", "test.sidedoc", "-o", "rebuilt.docx"])
            assert result.exit_code == 0

            rebuilt = Document("rebuilt.docx")
            section = rebuilt.sections[0]
            assert section.different_first_page_header_footer is True
            assert section.first_page_header.paragraphs[0].text == "First Page Header"
            assert section.first_page_footer.paragraphs[0].text == "First Page Footer"
            assert section.header.paragraphs[0].text == "Default Header"
            assert section.footer.paragraphs[0].text == "Default Footer"

    def test_roundtrip_multi_section(self):
        """Extract -> build preserves headers across multiple sections."""
        runner = CliRunner()
        with runner.isolated_filesystem():
            docx_path = _create_docx_with_multi_section()
            Path(docx_path).rename("test.docx")

            extract_result = runner.invoke(main, ["extract", "test.docx"])
            assert extract_result.exit_code == 0
            result = runner.invoke(main, ["build", "test.sidedoc", "-o", "rebuilt.docx"])
            assert result.exit_code == 0

            rebuilt = Document("rebuilt.docx")
            assert len(rebuilt.sections) >= 2
            assert rebuilt.sections[0].header.paragraphs[0].text == "Section 1 Header"
            assert rebuilt.sections[1].header.paragraphs[0].text == "Section 2 Header"

    def test_roundtrip_section_properties(self):
        """Extract -> build preserves section properties."""
        runner = CliRunner()
        with runner.isolated_filesystem():
            docx_path = _create_docx_with_section_properties()
            Path(docx_path).rename("test.docx")

            extract_result = runner.invoke(main, ["extract", "test.docx"])
            assert extract_result.exit_code == 0
            result = runner.invoke(main, ["build", "test.sidedoc", "-o", "rebuilt.docx"])
            assert result.exit_code == 0

            rebuilt = Document("rebuilt.docx")
            section = rebuilt.sections[0]
            assert section.orientation == WD_ORIENT.LANDSCAPE
            assert section.top_margin == Inches(1.5)
            assert section.bottom_margin == Inches(1.0)
            assert section.left_margin == Inches(0.75)
            assert section.right_margin == Inches(0.75)

    def test_roundtrip_body_content_preserved(self):
        """Extract -> build preserves body content alongside headers/footers."""
        runner = CliRunner()
        with runner.isolated_filesystem():
            docx_path = _create_docx_with_headers_footers()
            Path(docx_path).rename("test.docx")

            extract_result = runner.invoke(main, ["extract", "test.docx"])
            assert extract_result.exit_code == 0
            result = runner.invoke(main, ["build", "test.sidedoc", "-o", "rebuilt.docx"])
            assert result.exit_code == 0

            rebuilt = Document("rebuilt.docx")
            body_texts = [p.text for p in rebuilt.paragraphs]
            assert "Body paragraph one." in body_texts
            assert "Body paragraph two." in body_texts


# ============================================================
# Header/footer image tests
# ============================================================


def _create_docx_with_header_image() -> str:
    """Create docx with an image in the header (logo)."""
    doc = Document()
    section = doc.sections[0]

    img = PILImage.new("RGB", (50, 50), color="red")
    img_bytes = io.BytesIO()
    img.save(img_bytes, format="PNG")
    img_bytes.seek(0)

    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(img_bytes, width=Inches(0.5))

    doc.add_paragraph("Body with header image.")

    temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp.name)
    temp.close()
    return temp.name


class TestHeaderFooterImages:
    """Test images in headers/footers."""

    def test_header_image_extracted(self):
        """Images in headers are extracted to assets."""
        docx_path = _create_docx_with_header_image()
        try:
            sections, _ = extract_sections(docx_path)
            section = sections[0]
            header_paras = section["header_default"]
            has_image = any(
                p.get("image_path") is not None
                for p in header_paras
            )
            assert has_image
        finally:
            Path(docx_path).unlink()

    def test_header_image_content_is_empty_string(self):
        """Image paragraphs in headers should have empty string content, not file extension."""
        docx_path = _create_docx_with_header_image()
        try:
            sections, _ = extract_sections(docx_path)
            section = sections[0]
            header_paras = section["header_default"]
            image_paras = [p for p in header_paras if p.get("type") == "image"]
            assert len(image_paras) >= 1
            assert image_paras[0]["content"] == ""
        finally:
            Path(docx_path).unlink()

    def test_extract_sections_returns_image_data(self):
        """extract_sections returns a (list, dict) tuple with image data."""
        docx_path = _create_docx_with_header_image()
        try:
            sections, image_data = extract_sections(docx_path)
            assert isinstance(sections, list)
            assert isinstance(image_data, dict)
            assert len(image_data) > 0
        finally:
            Path(docx_path).unlink()

    def test_header_image_has_hf_prefix(self):
        """Header/footer image filenames start with hf_ prefix."""
        docx_path = _create_docx_with_header_image()
        try:
            sections, image_data = extract_sections(docx_path)
            header_paras = sections[0]["header_default"]
            image_paras = [p for p in header_paras if p.get("type") == "image"]
            assert len(image_paras) >= 1
            image_path = image_paras[0]["image_path"]
            filename = image_path.split("/")[-1]
            assert filename.startswith("hf_"), f"Expected hf_ prefix, got {filename}"
            # Also check image_data keys
            for key in image_data:
                assert key.startswith("hf_"), f"Expected hf_ prefix in image_data key, got {key}"
        finally:
            Path(docx_path).unlink()

    def test_no_image_key_collision_with_body_images(self):
        """Header/footer and body images cannot have colliding filenames."""
        docx_path = _create_docx_with_header_image()
        try:
            _, body_image_data = extract_blocks(docx_path)
            sections, hf_image_data = extract_sections(docx_path)
            # Verify no key overlap
            overlap = set(body_image_data.keys()) & set(hf_image_data.keys())
            assert len(overlap) == 0, f"Image key collision: {overlap}"
        finally:
            Path(docx_path).unlink()

    def test_roundtrip_header_image(self):
        """Extract -> build preserves images in headers."""
        runner = CliRunner()
        with runner.isolated_filesystem():
            docx_path = _create_docx_with_header_image()
            Path(docx_path).rename("test.docx")

            extract_result = runner.invoke(main, ["extract", "test.docx"])
            assert extract_result.exit_code == 0
            result = runner.invoke(main, ["build", "test.sidedoc", "-o", "rebuilt.docx"])
            assert result.exit_code == 0

            rebuilt = Document("rebuilt.docx")
            header = rebuilt.sections[0].header
            from docx.oxml.ns import qn
            drawings = header._element.findall(f".//{qn('w:drawing')}")
            assert len(drawings) >= 1


# ============================================================
# Even-page header/footer tests
# ============================================================


def _create_docx_with_even_page_variant() -> str:
    """Create a test docx with even-page headers/footers."""
    doc = Document()
    section = doc.sections[0]

    doc.settings.odd_and_even_pages_header_footer = True

    header = section.header
    header.is_linked_to_previous = False
    header.paragraphs[0].text = "Default Header"

    footer = section.footer
    footer.is_linked_to_previous = False
    footer.paragraphs[0].text = "Default Footer"

    even_header = section.even_page_header
    even_header.is_linked_to_previous = False
    even_header.paragraphs[0].text = "Even Page Header"

    even_footer = section.even_page_footer
    even_footer.is_linked_to_previous = False
    even_footer.paragraphs[0].text = "Even Page Footer"

    doc.add_paragraph("Body content.")

    temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp.name)
    temp.close()
    return temp.name


class TestEvenPageHeaderFooter:
    """Test even-page header/footer extraction and reconstruction."""

    def test_extract_even_page_header(self):
        """Even page header variant is extracted."""
        docx_path = _create_docx_with_even_page_variant()
        try:
            sections, _ = extract_sections(docx_path)
            section = sections[0]
            assert "header_even" in section
            assert section["header_even"][0]["content"] == "Even Page Header"
        finally:
            Path(docx_path).unlink()

    def test_extract_even_page_footer(self):
        """Even page footer variant is extracted."""
        docx_path = _create_docx_with_even_page_variant()
        try:
            sections, _ = extract_sections(docx_path)
            section = sections[0]
            assert "footer_even" in section
            assert section["footer_even"][0]["content"] == "Even Page Footer"
        finally:
            Path(docx_path).unlink()

    def test_roundtrip_even_page_variant(self):
        """Extract -> build preserves even page header/footer variants."""
        runner = CliRunner()
        with runner.isolated_filesystem():
            docx_path = _create_docx_with_even_page_variant()
            Path(docx_path).rename("test.docx")

            extract_result = runner.invoke(main, ["extract", "test.docx"])
            assert extract_result.exit_code == 0
            result = runner.invoke(main, ["build", "test.sidedoc", "-o", "rebuilt.docx"])
            assert result.exit_code == 0

            rebuilt = Document("rebuilt.docx")
            section = rebuilt.sections[0]
            assert section.even_page_header.paragraphs[0].text == "Even Page Header"
            assert section.even_page_footer.paragraphs[0].text == "Even Page Footer"
            assert section.header.paragraphs[0].text == "Default Header"
            assert section.footer.paragraphs[0].text == "Default Footer"

    def test_odd_and_even_pages_stored_in_page_setup(self):
        """extract_sections stores odd_and_even_pages in page_setup."""
        docx_path = _create_docx_with_even_page_variant()
        try:
            sections, _ = extract_sections(docx_path)
            setup = sections[0]["page_setup"]
            assert "odd_and_even_pages" in setup
            assert setup["odd_and_even_pages"] is True
        finally:
            Path(docx_path).unlink()

    def test_roundtrip_preserves_odd_and_even_pages_setting(self):
        """Extract -> build preserves the odd_and_even_pages document setting."""
        runner = CliRunner()
        with runner.isolated_filesystem():
            docx_path = _create_docx_with_even_page_variant()
            Path(docx_path).rename("test.docx")

            runner.invoke(main, ["extract", "test.docx"])
            runner.invoke(main, ["build", "test.sidedoc", "-o", "rebuilt.docx"])

            rebuilt = Document("rebuilt.docx")
            assert rebuilt.settings.odd_and_even_pages_header_footer is True

    def test_even_page_not_extracted_when_setting_disabled(self):
        """Even-page content is NOT extracted when odd_and_even_pages is False."""
        doc = Document()
        section = doc.sections[0]
        # Do NOT set odd_and_even_pages_header_footer
        header = section.header
        header.is_linked_to_previous = False
        header.paragraphs[0].text = "Default Header"
        even_header = section.even_page_header
        even_header.is_linked_to_previous = False
        even_header.paragraphs[0].text = "Even Page Header"
        doc.add_paragraph("Body.")
        temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp.name)
        temp.close()
        try:
            sections, _ = extract_sections(temp.name)
            section_dict = sections[0]
            assert "header_even" not in section_dict
        finally:
            Path(temp.name).unlink()


# ============================================================
# Conditional setdefault tests
# ============================================================


class TestConditionalSetdefault:
    """Test that only relevant header/footer keys are defaulted."""

    def test_simple_doc_no_first_or_even_keys(self):
        """Simple doc (no first-page, no even-page) omits those keys."""
        doc = Document()
        section = doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False
        header.paragraphs[0].text = "Header"
        footer = section.footer
        footer.is_linked_to_previous = False
        footer.paragraphs[0].text = "Footer"
        doc.add_paragraph("Body.")
        temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp.name)
        temp.close()
        try:
            sections, _ = extract_sections(temp.name)
            section_dict = sections[0]
            assert "header_default" in section_dict
            assert "footer_default" in section_dict
            assert "header_first" not in section_dict
            assert "footer_first" not in section_dict
            assert "header_even" not in section_dict
            assert "footer_even" not in section_dict
        finally:
            Path(temp.name).unlink()

    def test_first_page_doc_has_first_keys(self):
        """Doc with different_first_page has first-page keys defaulted."""
        docx_path = _create_docx_with_first_page_variant()
        try:
            sections, _ = extract_sections(docx_path)
            section_dict = sections[0]
            assert "header_first" in section_dict
            assert "footer_first" in section_dict
        finally:
            Path(docx_path).unlink()

    def test_even_page_doc_has_even_keys(self):
        """Doc with odd_and_even_pages has even-page keys defaulted."""
        docx_path = _create_docx_with_even_page_variant()
        try:
            sections, _ = extract_sections(docx_path)
            section_dict = sections[0]
            assert "header_even" in section_dict
            assert "footer_even" in section_dict
        finally:
            Path(docx_path).unlink()
