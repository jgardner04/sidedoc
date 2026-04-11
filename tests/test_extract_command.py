"""Test extract command integration."""

import json
import shutil
import tempfile
import zipfile
from pathlib import Path
from click.testing import CliRunner
from docx import Document
from sidedoc.cli import main


def create_test_docx() -> str:
    """Create a test docx file."""
    doc = Document()
    doc.add_paragraph("Test Title", style="Heading 1")
    doc.add_paragraph("This is a test paragraph.")

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    temp_file.close()
    return temp_file.name


def test_extract_command_creates_sidedoc_directory():
    """Test that extract command creates a .sidedoc directory."""
    runner = CliRunner()
    with runner.isolated_filesystem():
        docx_path = create_test_docx()

        try:
            result = runner.invoke(main, ["extract", docx_path])

            assert result.exit_code == 0

            sidedoc_path = Path(docx_path).with_suffix(".sidedoc")
            assert sidedoc_path.exists()
            assert sidedoc_path.is_dir()
        finally:
            Path(docx_path).unlink(missing_ok=True)


def test_extract_command_with_custom_output():
    """Test extract command with -o flag."""
    runner = CliRunner()
    with runner.isolated_filesystem():
        docx_path = create_test_docx()
        output_path = "custom_output.sidedoc"

        try:
            result = runner.invoke(main, ["extract", docx_path, "-o", output_path])

            assert result.exit_code == 0
            assert Path(output_path).exists()
            assert Path(output_path).is_dir()
        finally:
            Path(docx_path).unlink(missing_ok=True)
            shutil.rmtree(output_path, ignore_errors=True)


def test_extract_creates_valid_directory_structure():
    """Test that extract creates proper directory structure."""
    runner = CliRunner()
    with runner.isolated_filesystem():
        docx_path = create_test_docx()

        try:
            result = runner.invoke(main, ["extract", docx_path])
            assert result.exit_code == 0

            sidedoc_path = Path(docx_path).with_suffix(".sidedoc")
            assert (sidedoc_path / "content.md").exists()
            assert (sidedoc_path / "structure.json").exists()
            assert (sidedoc_path / "styles.json").exists()
            assert (sidedoc_path / "manifest.json").exists()
        finally:
            Path(docx_path).unlink(missing_ok=True)


def test_extract_creates_valid_manifest():
    """Test that manifest.json has required fields."""
    runner = CliRunner()
    with runner.isolated_filesystem():
        docx_path = create_test_docx()

        try:
            result = runner.invoke(main, ["extract", docx_path])
            assert result.exit_code == 0

            sidedoc_path = Path(docx_path).with_suffix(".sidedoc")
            manifest_data = json.loads((sidedoc_path / "manifest.json").read_text())

            assert "sidedoc_version" in manifest_data
            assert "created_at" in manifest_data
            assert "modified_at" in manifest_data
            assert "source_file" in manifest_data
            assert "source_hash" in manifest_data
            assert "content_hash" in manifest_data
            assert "generator" in manifest_data
        finally:
            Path(docx_path).unlink(missing_ok=True)


def test_manifest_content_hash_matches_content_md():
    """manifest.json content_hash should be SHA256 of content.md, not the source file."""
    import hashlib

    runner = CliRunner()
    with runner.isolated_filesystem():
        docx_path = create_test_docx()

        try:
            result = runner.invoke(main, ["extract", docx_path])
            assert result.exit_code == 0

            sidedoc_path = Path(docx_path).with_suffix(".sidedoc")
            content_md = (sidedoc_path / "content.md").read_text()
            manifest = json.loads((sidedoc_path / "manifest.json").read_text())

            expected_hash = hashlib.sha256(content_md.encode()).hexdigest()
            assert manifest["content_hash"] == expected_hash, (
                f"content_hash should be SHA256 of content.md, "
                f"got {manifest['content_hash']}, expected {expected_hash}"
            )
            assert manifest["content_hash"] != manifest["source_hash"], (
                "content_hash and source_hash should differ "
                "(content.md is markdown, source is binary docx)"
            )
        finally:
            Path(docx_path).unlink(missing_ok=True)


def test_extract_error_on_missing_file():
    """Test that extract returns error code 2 for missing file."""
    runner = CliRunner()
    result = runner.invoke(main, ["extract", "nonexistent.docx"])

    # Click handles missing files with exit code 2
    assert result.exit_code == 2


def test_extract_force_overwrites_existing():
    """Test that --force allows overwriting an existing directory."""
    runner = CliRunner()
    with runner.isolated_filesystem():
        docx_path = create_test_docx()

        try:
            # First extract
            result1 = runner.invoke(main, ["extract", docx_path])
            assert result1.exit_code == 0

            # Second extract without --force should fail
            result2 = runner.invoke(main, ["extract", docx_path])
            assert result2.exit_code != 0
            assert "already exists" in result2.output

            # With --force should succeed
            result3 = runner.invoke(main, ["extract", docx_path, "--force"])
            assert result3.exit_code == 0
        finally:
            Path(docx_path).unlink(missing_ok=True)


def test_extract_force_rejects_symlink(tmp_path: Path):
    """Test that --force rejects symlinked output paths."""
    runner = CliRunner()
    docx_path = create_test_docx()

    try:
        # Create a symlink where the output would go
        target_dir = tmp_path / "real_target"
        target_dir.mkdir()
        symlink_path = Path(docx_path).with_suffix(".sidedoc")
        symlink_path.symlink_to(target_dir)

        result = runner.invoke(main, ["extract", docx_path, "--force"])

        assert result.exit_code != 0
        assert "symlink" in result.output.lower()
    finally:
        Path(docx_path).unlink(missing_ok=True)
        if symlink_path.is_symlink():
            symlink_path.unlink()


def test_extract_pack_creates_zip():
    """Test that --pack creates a .sdoc ZIP archive."""
    runner = CliRunner()
    with runner.isolated_filesystem():
        docx_path = create_test_docx()

        try:
            result = runner.invoke(main, ["extract", docx_path, "--pack"])
            assert result.exit_code == 0

            sdoc_path = Path(docx_path).with_suffix(".sdoc")
            assert sdoc_path.exists()
            assert zipfile.is_zipfile(sdoc_path)

            with zipfile.ZipFile(sdoc_path, "r") as zf:
                names = zf.namelist()
                assert "content.md" in names
                assert "styles.json" in names
        finally:
            Path(docx_path).unlink(missing_ok=True)
