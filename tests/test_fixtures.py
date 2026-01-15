"""Test to create and verify comprehensive test fixtures."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io


def create_minimal_png() -> bytes:
    """Create a minimal valid PNG file for testing."""
    from PIL import Image

    # Use PIL to create a proper 1x1 PNG image
    img = Image.new('RGB', (1, 1), color='red')
    img_bytes = io.BytesIO()
    img.save(img_bytes, format='PNG')
    return img_bytes.getvalue()


def test_create_simple_fixture():
    """Create simple.docx with headings and paragraphs only."""
    fixtures_dir = Path("tests/fixtures")
    fixtures_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Add various heading levels
    doc.add_heading("Main Title", level=1)
    doc.add_paragraph("This is a simple document with headings and paragraphs.")

    doc.add_heading("Section 1", level=2)
    doc.add_paragraph("This is a paragraph under section 1.")
    doc.add_paragraph("This is another paragraph.")

    doc.add_heading("Section 2", level=2)
    doc.add_paragraph("This is a paragraph under section 2.")

    doc.add_heading("Subsection 2.1", level=3)
    doc.add_paragraph("This is a paragraph under subsection 2.1.")

    output_path = fixtures_dir / "simple.docx"
    doc.save(str(output_path))

    # Verify file was created and can be opened
    assert output_path.exists()
    verify_doc = Document(str(output_path))
    assert len(verify_doc.paragraphs) > 0


def test_create_lists_fixture():
    """Create lists.docx with bulleted and numbered lists."""
    fixtures_dir = Path("tests/fixtures")
    fixtures_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()

    doc.add_heading("Lists Example", level=1)

    # Bulleted list
    doc.add_heading("Bulleted List", level=2)
    doc.add_paragraph("First bullet item", style="List Bullet")
    doc.add_paragraph("Second bullet item", style="List Bullet")
    doc.add_paragraph("Third bullet item", style="List Bullet")

    # Regular paragraph between lists
    doc.add_paragraph("This is a regular paragraph between lists.")

    # Numbered list
    doc.add_heading("Numbered List", level=2)
    doc.add_paragraph("First numbered item", style="List Number")
    doc.add_paragraph("Second numbered item", style="List Number")
    doc.add_paragraph("Third numbered item", style="List Number")

    output_path = fixtures_dir / "lists.docx"
    doc.save(str(output_path))

    # Verify file was created and can be opened
    assert output_path.exists()
    verify_doc = Document(str(output_path))
    assert len(verify_doc.paragraphs) > 0


def test_create_formatted_fixture():
    """Create formatted.docx with bold, italic, and mixed inline formatting."""
    fixtures_dir = Path("tests/fixtures")
    fixtures_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()

    doc.add_heading("Formatting Examples", level=1)

    # Bold text
    para1 = doc.add_paragraph()
    run1 = para1.add_run("This is ")
    run2 = para1.add_run("bold text")
    run2.bold = True
    run3 = para1.add_run(" in a sentence.")

    # Italic text
    para2 = doc.add_paragraph()
    run1 = para2.add_run("This is ")
    run2 = para2.add_run("italic text")
    run2.italic = True
    run3 = para2.add_run(" in a sentence.")

    # Bold and italic
    para3 = doc.add_paragraph()
    run1 = para3.add_run("This is ")
    run2 = para3.add_run("bold and italic")
    run2.bold = True
    run2.italic = True
    run3 = para3.add_run(" in a sentence.")

    # Underlined text
    para4 = doc.add_paragraph()
    run1 = para4.add_run("This is ")
    run2 = para4.add_run("underlined text")
    run2.underline = True
    run3 = para4.add_run(" in a sentence.")

    # Multiple formatting in same paragraph
    para5 = doc.add_paragraph()
    run1 = para5.add_run("This paragraph has ")
    run2 = para5.add_run("bold")
    run2.bold = True
    run3 = para5.add_run(", ")
    run4 = para5.add_run("italic")
    run4.italic = True
    run5 = para5.add_run(", and ")
    run6 = para5.add_run("underlined")
    run6.underline = True
    run7 = para5.add_run(" text.")

    output_path = fixtures_dir / "formatted.docx"
    doc.save(str(output_path))

    # Verify file was created and can be opened
    assert output_path.exists()
    verify_doc = Document(str(output_path))
    assert len(verify_doc.paragraphs) > 0


def test_create_images_fixture():
    """Create images.docx with multiple images."""
    fixtures_dir = Path("tests/fixtures")
    fixtures_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()

    doc.add_heading("Images Example", level=1)

    doc.add_paragraph("This document contains multiple images.")

    # Add first image
    png_data = create_minimal_png()
    image_stream1 = io.BytesIO(png_data)
    doc.add_picture(image_stream1, width=Inches(2.0))

    doc.add_paragraph("Text between images.")

    # Add second image
    image_stream2 = io.BytesIO(png_data)
    doc.add_picture(image_stream2, width=Inches(2.5))

    doc.add_paragraph("Text after images.")

    # Add third image
    image_stream3 = io.BytesIO(png_data)
    doc.add_picture(image_stream3, width=Inches(1.5))

    output_path = fixtures_dir / "images.docx"
    doc.save(str(output_path))

    # Verify file was created and can be opened
    assert output_path.exists()
    verify_doc = Document(str(output_path))
    assert len(verify_doc.paragraphs) > 0


def test_create_complex_fixture():
    """Create complex.docx with all supported elements combined."""
    fixtures_dir = Path("tests/fixtures")
    fixtures_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Title
    doc.add_heading("Comprehensive Document", level=1)

    # Introduction with formatting
    para1 = doc.add_paragraph()
    run1 = para1.add_run("This is a ")
    run2 = para1.add_run("comprehensive document")
    run2.bold = True
    run3 = para1.add_run(" that includes ")
    run4 = para1.add_run("all supported elements")
    run4.italic = True
    run5 = para1.add_run(".")

    # Section with lists
    doc.add_heading("Features", level=2)
    doc.add_paragraph("Key features:", style="List Bullet")
    doc.add_paragraph("Multiple heading levels", style="List Bullet")
    doc.add_paragraph("Inline formatting (bold, italic, underline)", style="List Bullet")
    doc.add_paragraph("Lists (bulleted and numbered)", style="List Bullet")
    doc.add_paragraph("Images", style="List Bullet")

    # Numbered list
    doc.add_heading("Steps", level=2)
    doc.add_paragraph("First step", style="List Number")
    doc.add_paragraph("Second step", style="List Number")
    doc.add_paragraph("Third step", style="List Number")

    # Image section
    doc.add_heading("Visual Content", level=2)
    doc.add_paragraph("Below is an embedded image:")

    png_data = create_minimal_png()
    image_stream = io.BytesIO(png_data)
    doc.add_picture(image_stream, width=Inches(2.0))

    # Conclusion with mixed formatting
    doc.add_heading("Conclusion", level=2)
    para2 = doc.add_paragraph()
    run1 = para2.add_run("This document demonstrates ")
    run2 = para2.add_run("all")
    run2.bold = True
    run2.underline = True
    run3 = para2.add_run(" the capabilities of the ")
    run4 = para2.add_run("sidedoc")
    run4.italic = True
    run5 = para2.add_run(" format.")

    output_path = fixtures_dir / "complex.docx"
    doc.save(str(output_path))

    # Verify file was created and can be opened
    assert output_path.exists()
    verify_doc = Document(str(output_path))
    assert len(verify_doc.paragraphs) > 0


def test_all_fixtures_exist_and_valid():
    """Verify all fixture files exist and can be opened successfully."""
    fixtures_dir = Path("tests/fixtures")

    expected_fixtures = [
        "simple.docx",
        "lists.docx",
        "formatted.docx",
        "images.docx",
        "complex.docx"
    ]

    for fixture_name in expected_fixtures:
        fixture_path = fixtures_dir / fixture_name
        assert fixture_path.exists(), f"{fixture_name} should exist"

        # Verify it can be opened
        doc = Document(str(fixture_path))
        assert len(doc.paragraphs) > 0, f"{fixture_name} should have paragraphs"
