"""Fidelity scorer utility (US-013 to US-016).

This module provides format fidelity scoring for comparing
original and rebuilt documents.
"""

import subprocess
import tempfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn


_SOFFICE_CANDIDATES = [
    "soffice",
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    "/usr/bin/soffice",
    "/usr/bin/libreoffice",
]


def _find_soffice() -> str | None:
    """Find a working LibreOffice soffice binary.

    Returns:
        Path to soffice binary, or None if not found.
    """
    for path in _SOFFICE_CANDIDATES:
        try:
            result = subprocess.run(
                [path, "--version"],
                capture_output=True,
                check=True,
            )
            if "LibreOffice" in result.stdout.decode("utf-8", errors="ignore"):
                return path
        except (subprocess.CalledProcessError, FileNotFoundError):
            continue
    return None


class FidelityScorer:
    """Scorer for measuring format fidelity between original and rebuilt documents.

    Supports scoring:
    - Structural fidelity (heading levels, paragraph, list, table counts)
    - Formatting fidelity (per-run bold, italic, underline, font name, font size)
    - Table fidelity (structure, merged cells, backgrounds, style)
    - Hyperlink fidelity (text + URL pairs)
    - Track changes fidelity (insertion/deletion counts, authors)
    - Visual fidelity (perceptual hash comparison)
    - Combined weighted score
    """

    def _validate_path(self, path: Path) -> Path:
        """Validate a path is safe for subprocess use.

        Args:
            path: Path to validate.

        Returns:
            Resolved absolute path.

        Raises:
            ValueError: If path does not exist, is not a file, or is outside allowed directories.
        """
        resolved = path.resolve()
        if not resolved.exists():
            raise ValueError(f"Path does not exist: {resolved}")
        if not resolved.is_file():
            raise ValueError(f"Path is not a file: {resolved}")

        # Whitelist allowed directories to prevent path traversal
        benchmarks_dir = Path(__file__).parent.parent.resolve()
        project_root = benchmarks_dir.parent
        allowed_parents = [
            benchmarks_dir,                          # benchmarks/
            project_root / "tests",                  # tests/ (for fixtures)
            Path(tempfile.gettempdir()).resolve(),   # temp files
        ]
        if not any(self._is_subpath(resolved, parent) for parent in allowed_parents):
            raise ValueError(f"Path outside allowed directories: {resolved}")

        return resolved

    def _is_subpath(self, path: Path, parent: Path) -> bool:
        """Check if path is under parent directory.

        Args:
            path: Path to check.
            parent: Parent directory to check against.

        Returns:
            True if path is under parent directory.
        """
        try:
            path.relative_to(parent)
            return True
        except ValueError:
            return False

    # ── Helpers ──────────────────────────────────────────────────────────

    def _extract_heading_levels(self, docx_path: Path) -> list[int]:
        """Ordered heading levels (e.g., [1, 2, 2, 1]).

        Args:
            docx_path: Path to the document.

        Returns:
            List of heading level integers in document order.
        """
        doc = Document(str(docx_path))
        levels: list[int] = []
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.replace("Heading", "").strip())
                    levels.append(level)
                except ValueError:
                    pass
        return levels

    def _extract_all_run_formatting(self, docx_path: Path) -> list[list[dict]]:
        """Per-paragraph, per-run formatting.

        Each dict has: bold, italic, underline, font_name, font_size.

        Args:
            docx_path: Path to the document.

        Returns:
            List of lists. Outer list = paragraphs, inner list = runs in that paragraph.
        """
        doc = Document(str(docx_path))
        result: list[list[dict]] = []
        for para in doc.paragraphs:
            runs_fmt: list[dict] = []
            for run in para.runs:
                runs_fmt.append({
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                    "font_name": run.font.name,
                    "font_size": run.font.size,
                })
            result.append(runs_fmt)
        return result

    def _extract_tables_metadata(self, docx_path: Path) -> list[dict]:
        """List of table dicts with: rows, cols, merged_cells, cell_backgrounds, table_style.

        Args:
            docx_path: Path to the document.

        Returns:
            List of table metadata dicts.
        """
        doc = Document(str(docx_path))
        tables_meta: list[dict] = []
        for table in doc.tables:
            tbl = table._tbl
            rows = tbl.findall(qn("w:tr"))
            row_count = len(rows)

            # Determine col count from the first row
            col_count = 0
            if rows:
                first_row_cells = rows[0].findall(qn("w:tc"))
                col_count = len(first_row_cells)

            merged_cells: list[dict] = []
            cell_backgrounds: list[str | None] = []

            for tr in rows:
                for tc in tr.findall(qn("w:tc")):
                    tcPr = tc.find(qn("w:tcPr"))
                    # Check gridSpan (horizontal merge)
                    if tcPr is not None:
                        gridSpan = tcPr.find(qn("w:gridSpan"))
                        if gridSpan is not None:
                            span_val = gridSpan.get(qn("w:val"))
                            if span_val and int(span_val) > 1:
                                merged_cells.append({"type": "gridSpan", "span": int(span_val)})

                        vMerge = tcPr.find(qn("w:vMerge"))
                        if vMerge is not None:
                            merge_val = vMerge.get(qn("w:val"), "continue")
                            merged_cells.append({"type": "vMerge", "val": merge_val})

                        shd = tcPr.find(qn("w:shd"))
                        if shd is not None:
                            fill = shd.get(qn("w:fill"))
                            cell_backgrounds.append(fill)
                        else:
                            cell_backgrounds.append(None)
                    else:
                        cell_backgrounds.append(None)

            # Table style
            tblPr = tbl.find(qn("w:tblPr"))
            table_style = None
            if tblPr is not None:
                tblStyle = tblPr.find(qn("w:tblStyle"))
                if tblStyle is not None:
                    table_style = tblStyle.get(qn("w:val"))

            tables_meta.append({
                "rows": row_count,
                "cols": col_count,
                "merged_cells": merged_cells,
                "cell_backgrounds": cell_backgrounds,
                "table_style": table_style,
            })

        return tables_meta

    def _extract_hyperlinks(self, docx_path: Path) -> list[tuple[str, str]]:
        """(text, url) pairs from w:hyperlink elements.

        Args:
            docx_path: Path to the document.

        Returns:
            List of (text, url) tuples.
        """
        doc = Document(str(docx_path))
        rels = doc.part.rels
        body = doc.element.body
        links: list[tuple[str, str]] = []

        for para in body.findall(qn("w:p")):
            for hyperlink in para.findall(qn("w:hyperlink")):
                r_id = hyperlink.get(qn("r:id"))
                # Collect text from all runs within the hyperlink
                text_parts: list[str] = []
                for run in hyperlink.findall(qn("w:r")):
                    for t in run.findall(qn("w:t")):
                        if t.text:
                            text_parts.append(t.text)
                text = "".join(text_parts)
                url = ""
                if r_id and r_id in rels:
                    url = rels[r_id].target_ref
                if text or url:
                    links.append((text, url))

        return links

    def _extract_track_changes(self, docx_path: Path) -> dict:
        """Extract track change counts and authors.

        Args:
            docx_path: Path to the document.

        Returns:
            Dict with 'insertions' (int), 'deletions' (int), 'authors' (set[str]).
        """
        doc = Document(str(docx_path))
        body = doc.element.body

        insertions = 0
        deletions = 0
        authors: set[str] = set()

        for ins in body.iter(qn("w:ins")):
            insertions += 1
            author = ins.get(qn("w:author"))
            if author:
                authors.add(author)

        for del_elem in body.iter(qn("w:del")):
            deletions += 1
            author = del_elem.get(qn("w:author"))
            if author:
                authors.add(author)

        return {
            "insertions": insertions,
            "deletions": deletions,
            "authors": authors,
        }

    def _is_list_item(self, para: Any) -> bool:
        """Check if a paragraph is a list item by examining its XML.

        Args:
            para: A python-docx paragraph object.

        Returns:
            True if the paragraph is a list item.
        """
        pPr = para._p.find(qn("w:pPr"))
        if pPr is not None:
            numPr = pPr.find(qn("w:numPr"))
            return numPr is not None
        return False

    def _count_structure(self, docx_path: Path) -> dict[str, int]:
        """Count structural elements in a document.

        Args:
            docx_path: Path to the document.

        Returns:
            Dict with counts for headings, paragraphs, list_items, and tables.
        """
        doc = Document(str(docx_path))

        heading_count = 0
        paragraph_count = 0
        list_item_count = 0

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""

            if style_name.startswith("Heading"):
                heading_count += 1
            elif style_name == "List Paragraph" or self._is_list_item(para):
                list_item_count += 1
            else:
                paragraph_count += 1

        table_count = len(doc.tables)

        return {
            "headings": heading_count,
            "paragraphs": paragraph_count,
            "list_items": list_item_count,
            "tables": table_count,
        }

    # ── Scoring Methods ─────────────────────────────────────────────────

    def _count_similarity(self, a: int, b: int) -> float:
        """Compute similarity between two counts as a 0-100 score.

        Args:
            a: First count.
            b: Second count.

        Returns:
            Similarity score from 0 to 100.
        """
        return (1 - abs(a - b) / max(a, b, 1)) * 100

    def score_structure(
        self, original_docx: str | Path, rebuilt_docx: str | Path
    ) -> float:
        """Score structural fidelity between two documents.

        Compares heading levels at positions, paragraph count, list item count,
        and table count. Returns 0-100.

        Weights: 30% heading level match, 30% paragraph count, 20% list items, 20% table count.

        Args:
            original_docx: Path to the original document.
            rebuilt_docx: Path to the rebuilt document.

        Returns:
            Score from 0 to 100.
        """
        original_path = Path(original_docx)
        rebuilt_path = Path(rebuilt_docx)

        # Heading level comparison (30%)
        orig_levels = self._extract_heading_levels(original_path)
        rebuilt_levels = self._extract_heading_levels(rebuilt_path)

        if orig_levels or rebuilt_levels:
            max_len = max(len(orig_levels), len(rebuilt_levels))
            matches = 0
            for i in range(max_len):
                orig_level = orig_levels[i] if i < len(orig_levels) else None
                rebuilt_level = rebuilt_levels[i] if i < len(rebuilt_levels) else None
                if orig_level == rebuilt_level:
                    matches += 1
            heading_score = (matches / max_len) * 100 if max_len > 0 else 100.0
        else:
            heading_score = 100.0

        # Count-based comparisons
        original_counts = self._count_structure(original_path)
        rebuilt_counts = self._count_structure(rebuilt_path)

        # Paragraph count comparison (30%)
        para_score = self._count_similarity(
            original_counts["paragraphs"], rebuilt_counts["paragraphs"]
        )

        # List items comparison (20%)
        list_score = self._count_similarity(
            original_counts["list_items"], rebuilt_counts["list_items"]
        )

        # Table count comparison (20%)
        table_score = self._count_similarity(
            original_counts["tables"], rebuilt_counts["tables"]
        )

        total = (
            0.30 * heading_score
            + 0.30 * para_score
            + 0.20 * list_score
            + 0.20 * table_score
        )

        return max(0.0, min(100.0, total))

    def score_formatting(
        self, original_docx: str | Path, rebuilt_docx: str | Path
    ) -> float:
        """Score formatting fidelity between two documents.

        Full run-level comparison across ALL paragraphs. Compares bold, italic,
        underline (emphasis) and font_name, font_size (typography) for each run.

        Args:
            original_docx: Path to the original document.
            rebuilt_docx: Path to the rebuilt document.

        Returns:
            Score from 0 to 100. 100 if no paragraphs with runs.
        """
        original_path = Path(original_docx)
        rebuilt_path = Path(rebuilt_docx)

        orig_fmt = self._extract_all_run_formatting(original_path)
        rebuilt_fmt = self._extract_all_run_formatting(rebuilt_path)

        total_attributes = 0
        matching_attributes = 0

        # Compare paragraph by paragraph
        max_paras = max(len(orig_fmt), len(rebuilt_fmt))
        for i in range(max_paras):
            orig_runs = orig_fmt[i] if i < len(orig_fmt) else []
            rebuilt_runs = rebuilt_fmt[i] if i < len(rebuilt_fmt) else []

            max_runs = max(len(orig_runs), len(rebuilt_runs))
            for j in range(max_runs):
                orig_run = orig_runs[j] if j < len(orig_runs) else {}
                rebuilt_run = rebuilt_runs[j] if j < len(rebuilt_runs) else {}

                # Compare 5 attributes: bold, italic, underline, font_name, font_size
                attrs = ["bold", "italic", "underline", "font_name", "font_size"]
                for attr in attrs:
                    total_attributes += 1
                    if orig_run.get(attr) == rebuilt_run.get(attr):
                        matching_attributes += 1

        if total_attributes == 0:
            return 100.0

        return (matching_attributes / total_attributes) * 100

    def score_styles(
        self, original_docx: str | Path, rebuilt_docx: str | Path
    ) -> float:
        """Score style fidelity between two documents (deprecated).

        This method is deprecated. Use score_formatting() instead.

        Args:
            original_docx: Path to the original document.
            rebuilt_docx: Path to the rebuilt document.

        Returns:
            Score from 0 to 100.
        """
        return self.score_formatting(original_docx, rebuilt_docx)

    def score_tables(
        self, original_docx: str | Path, rebuilt_docx: str | Path
    ) -> float | None:
        """Score table fidelity between two documents.

        Compares table structure (row/col count), merged cells, cell backgrounds,
        and table style.

        Args:
            original_docx: Path to the original document.
            rebuilt_docx: Path to the rebuilt document.

        Returns:
            Score from 0 to 100, or None if no tables in original.
        """
        original_path = Path(original_docx)
        rebuilt_path = Path(rebuilt_docx)

        orig_tables = self._extract_tables_metadata(original_path)
        rebuilt_tables = self._extract_tables_metadata(rebuilt_path)

        if not orig_tables:
            return None

        total_checks = 0
        matching_checks = 0

        max_tables = max(len(orig_tables), len(rebuilt_tables))

        # Table count match
        total_checks += 1
        if len(orig_tables) == len(rebuilt_tables):
            matching_checks += 1

        # Per-table comparison
        for i in range(max_tables):
            orig = orig_tables[i] if i < len(orig_tables) else None
            rebuilt = rebuilt_tables[i] if i < len(rebuilt_tables) else None

            if orig is None or rebuilt is None:
                # Missing table in one doc => penalize all checks for this table
                total_checks += 4  # rows, cols, merged, style
                continue

            # Row count
            total_checks += 1
            if orig["rows"] == rebuilt["rows"]:
                matching_checks += 1

            # Col count
            total_checks += 1
            if orig["cols"] == rebuilt["cols"]:
                matching_checks += 1

            # Merged cells
            total_checks += 1
            if orig["merged_cells"] == rebuilt["merged_cells"]:
                matching_checks += 1

            # Cell backgrounds
            total_checks += 1
            if orig["cell_backgrounds"] == rebuilt["cell_backgrounds"]:
                matching_checks += 1

            # Table style
            total_checks += 1
            if orig["table_style"] == rebuilt["table_style"]:
                matching_checks += 1

        if total_checks == 0:
            return 100.0

        return (matching_checks / total_checks) * 100

    def score_hyperlinks(
        self, original_docx: str | Path, rebuilt_docx: str | Path
    ) -> float | None:
        """Score hyperlink fidelity between two documents.

        Compares (text, url) pairs in order.

        Args:
            original_docx: Path to the original document.
            rebuilt_docx: Path to the rebuilt document.

        Returns:
            Score from 0 to 100, or None if no hyperlinks in original.
        """
        original_path = Path(original_docx)
        rebuilt_path = Path(rebuilt_docx)

        orig_links = self._extract_hyperlinks(original_path)
        rebuilt_links = self._extract_hyperlinks(rebuilt_path)

        if not orig_links:
            return None

        max_links = max(len(orig_links), len(rebuilt_links))

        matches = 0
        for i in range(max_links):
            orig = orig_links[i] if i < len(orig_links) else (None, None)
            rebuilt = rebuilt_links[i] if i < len(rebuilt_links) else (None, None)
            if orig == rebuilt:
                matches += 1

        return (matches / max_links) * 100

    def score_track_changes(
        self, original_docx: str | Path, rebuilt_docx: str | Path
    ) -> float | None:
        """Score track change preservation between two documents.

        Compares insertion count, deletion count, and author sets.

        Args:
            original_docx: Path to the original document.
            rebuilt_docx: Path to the rebuilt document.

        Returns:
            Score from 0 to 100, or None if no track changes in original.
        """
        original_path = Path(original_docx)
        rebuilt_path = Path(rebuilt_docx)

        orig_tc = self._extract_track_changes(original_path)
        rebuilt_tc = self._extract_track_changes(rebuilt_path)

        # No track changes in original
        if orig_tc["insertions"] == 0 and orig_tc["deletions"] == 0:
            return None

        total_checks = 0
        matching_checks = 0

        # Insertion count
        total_checks += 1
        max_ins = max(orig_tc["insertions"], rebuilt_tc["insertions"], 1)
        ins_similarity = 1 - abs(orig_tc["insertions"] - rebuilt_tc["insertions"]) / max_ins
        matching_checks += ins_similarity

        # Deletion count
        total_checks += 1
        max_del = max(orig_tc["deletions"], rebuilt_tc["deletions"], 1)
        del_similarity = 1 - abs(orig_tc["deletions"] - rebuilt_tc["deletions"]) / max_del
        matching_checks += del_similarity

        # Author sets
        total_checks += 1
        if orig_tc["authors"] and rebuilt_tc["authors"]:
            union = orig_tc["authors"] | rebuilt_tc["authors"]
            intersection = orig_tc["authors"] & rebuilt_tc["authors"]
            matching_checks += len(intersection) / len(union) if union else 1.0
        elif not orig_tc["authors"] and not rebuilt_tc["authors"]:
            matching_checks += 1.0
        # else: one has authors, the other doesn't => 0 credit

        if total_checks == 0:
            return 100.0

        return (matching_checks / total_checks) * 100

    def score_visual(
        self, original_docx: str | Path, rebuilt_docx: str | Path
    ) -> float:
        """Score visual fidelity between two documents.

        Renders first page of each docx to PNG and computes perceptual hash
        difference using imagehash.

        Args:
            original_docx: Path to the original document.
            rebuilt_docx: Path to the rebuilt document.

        Returns:
            Score from 0 to 100 (100 = visually identical).

        Raises:
            RuntimeError: If LibreOffice or Poppler are not available.
        """
        import imagehash
        from pdf2image import convert_from_path
        from PIL import Image

        original_path = Path(original_docx)
        rebuilt_path = Path(rebuilt_docx)

        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)

            # Convert docx to PDF using LibreOffice
            original_pdf = self._convert_docx_to_pdf(original_path, tmp_path)
            rebuilt_pdf = self._convert_docx_to_pdf(rebuilt_path, tmp_path)

            # Convert first page of each PDF to image
            original_images = convert_from_path(str(original_pdf), first_page=1, last_page=1)
            rebuilt_images = convert_from_path(str(rebuilt_pdf), first_page=1, last_page=1)

            if not original_images or not rebuilt_images:
                return 0.0

            original_img = original_images[0]
            rebuilt_img = rebuilt_images[0]

            # Compute perceptual hashes
            original_hash = imagehash.phash(original_img)
            rebuilt_hash = imagehash.phash(rebuilt_img)

            # Hash difference: 0 = identical, higher = more different
            # Max possible difference for 64-bit hash is 64
            hash_diff = original_hash - rebuilt_hash

            # Convert to 0-100 score (0 diff = 100 score)
            score = max(0.0, 100 - (hash_diff / 64) * 100)

            return score

    def _convert_docx_to_pdf(self, docx_path: Path, output_dir: Path) -> Path:
        """Convert a docx file to PDF using LibreOffice.

        Args:
            docx_path: Path to the docx file.
            output_dir: Directory for output PDF.

        Returns:
            Path to the generated PDF.

        Raises:
            RuntimeError: If LibreOffice is not available.
            ValueError: If docx_path is invalid.
        """
        # Validate input path before subprocess call
        validated_docx = self._validate_path(docx_path)
        validated_outdir = output_dir.resolve()

        soffice_cmd = _find_soffice()
        if soffice_cmd is None:
            raise RuntimeError(
                "LibreOffice (soffice) not found. Install LibreOffice for visual comparison."
            )

        # Convert docx to PDF using validated paths
        subprocess.run(
            [
                soffice_cmd,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(validated_outdir),
                str(validated_docx),
            ],
            capture_output=True,
            check=True,
        )

        # Return path to generated PDF
        pdf_path = validated_outdir / f"{validated_docx.stem}.pdf"
        return pdf_path

    def score_total(
        self, original_docx: str | Path, rebuilt_docx: str | Path
    ) -> dict[str, float | None]:
        """Compute combined fidelity score across all dimensions.

        Computes structure, formatting, tables, hyperlinks, and track_changes.
        Total is the mean of non-None dimension scores.
        Visual is reported separately and NOT included in total.

        Args:
            original_docx: Path to the original document.
            rebuilt_docx: Path to the rebuilt document.

        Returns:
            Dict with keys: structure, formatting, tables, hyperlinks,
            track_changes, visual, total.
        """
        structure = self.score_structure(original_docx, rebuilt_docx)
        formatting = self.score_formatting(original_docx, rebuilt_docx)
        tables = self.score_tables(original_docx, rebuilt_docx)
        hyperlinks = self.score_hyperlinks(original_docx, rebuilt_docx)
        track_changes = self.score_track_changes(original_docx, rebuilt_docx)

        visual: float | None = None
        try:
            visual = self.score_visual(original_docx, rebuilt_docx)
        except Exception:
            # Visual scoring not available (missing LibreOffice/Poppler)
            pass

        # Total = mean of non-None dimensions (excluding visual)
        scores = [structure, formatting] + [s for s in (tables, hyperlinks, track_changes) if s is not None]

        total = sum(scores) / len(scores) if scores else 0.0

        return {
            "structure": structure,
            "formatting": formatting,
            "tables": tables,
            "hyperlinks": hyperlinks,
            "track_changes": track_changes,
            "visual": visual,
            "total": total,
        }
