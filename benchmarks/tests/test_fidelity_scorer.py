"""Tests for fidelity scorer utility (US-013 to US-016)."""

import shutil
import tempfile
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt

BENCHMARKS_DIR = Path(__file__).parent.parent
SYNTHETIC_DIR = BENCHMARKS_DIR / "corpus" / "synthetic"


class TestFidelityStructuralScorer:
    """Test the structural fidelity scorer (US-013)."""

    def test_module_exists(self, benchmarks_dir: Path) -> None:
        """Test that fidelity_scorer.py exists."""
        module_path = benchmarks_dir / "metrics" / "fidelity_scorer.py"
        assert module_path.exists(), "benchmarks/metrics/fidelity_scorer.py does not exist"

    def test_fidelity_scorer_is_importable(self) -> None:
        """Test that FidelityScorer can be imported."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        assert FidelityScorer is not None

    def test_score_structure_returns_numeric(self) -> None:
        """Test that score_structure returns a 0-100 score."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        simple_docx = SYNTHETIC_DIR / "simple.docx"

        # Comparing a document to itself should return 100
        score = scorer.score_structure(simple_docx, simple_docx)

        assert isinstance(score, (int, float))
        assert 0 <= score <= 100

    def test_identical_documents_score_100(self) -> None:
        """Test that identical documents get a perfect structural score."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        simple_docx = SYNTHETIC_DIR / "simple.docx"

        score = scorer.score_structure(simple_docx, simple_docx)

        assert score == 100

    def test_different_heading_counts_reduce_score(self) -> None:
        """Test that different heading counts reduce the score."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()

        # Create two temporary documents with different heading counts
        with tempfile.TemporaryDirectory() as tmp_dir:
            original_path = Path(tmp_dir) / "original.docx"
            rebuilt_path = Path(tmp_dir) / "rebuilt.docx"

            # Original: 2 headings
            doc1 = Document()
            doc1.add_heading("Heading 1", level=1)
            doc1.add_paragraph("Content 1")
            doc1.add_heading("Heading 2", level=2)
            doc1.add_paragraph("Content 2")
            doc1.save(str(original_path))

            # Rebuilt: 1 heading (missing one)
            doc2 = Document()
            doc2.add_heading("Heading 1", level=1)
            doc2.add_paragraph("Content 1")
            doc2.add_paragraph("Content 2")  # No heading here
            doc2.save(str(rebuilt_path))

            score = scorer.score_structure(original_path, rebuilt_path)

            assert score < 100

    def test_different_paragraph_counts_reduce_score(self) -> None:
        """Test that different paragraph counts reduce the score."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            original_path = Path(tmp_dir) / "original.docx"
            rebuilt_path = Path(tmp_dir) / "rebuilt.docx"

            # Original: 3 paragraphs
            doc1 = Document()
            doc1.add_paragraph("Paragraph 1")
            doc1.add_paragraph("Paragraph 2")
            doc1.add_paragraph("Paragraph 3")
            doc1.save(str(original_path))

            # Rebuilt: 2 paragraphs
            doc2 = Document()
            doc2.add_paragraph("Paragraph 1")
            doc2.add_paragraph("Paragraph 2")
            doc2.save(str(rebuilt_path))

            score = scorer.score_structure(original_path, rebuilt_path)

            assert score < 100

    def test_different_list_item_counts_reduce_score(self) -> None:
        """Test that different list item counts reduce the score."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        lists_docx = SYNTHETIC_DIR / "lists.docx"

        with tempfile.TemporaryDirectory() as tmp_dir:
            rebuilt_path = Path(tmp_dir) / "rebuilt.docx"

            # Create a document with fewer list items
            doc = Document()
            doc.add_paragraph("Just a paragraph, no lists")
            doc.save(str(rebuilt_path))

            score = scorer.score_structure(lists_docx, rebuilt_path)

            # Should be less than 100 due to missing list items
            assert score < 100

    def test_score_never_negative(self) -> None:
        """Test that the score is never negative."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            original_path = Path(tmp_dir) / "original.docx"
            rebuilt_path = Path(tmp_dir) / "rebuilt.docx"

            # Original: lots of structure
            doc1 = Document()
            for i in range(10):
                doc1.add_heading(f"Heading {i}", level=1)
                doc1.add_paragraph(f"Paragraph {i}")
            doc1.save(str(original_path))

            # Rebuilt: minimal structure
            doc2 = Document()
            doc2.add_paragraph("Minimal content")
            doc2.save(str(rebuilt_path))

            score = scorer.score_structure(original_path, rebuilt_path)

            assert score >= 0

    def test_accepts_path_strings(self) -> None:
        """Test that score_structure accepts string paths."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        simple_docx = str(SYNTHETIC_DIR / "simple.docx")

        score = scorer.score_structure(simple_docx, simple_docx)

        assert score == 100


class TestFidelityStructureEnhanced:
    """Test enhanced structural scoring with heading level comparison and table count."""

    def test_heading_level_mismatch_reduces_score(self) -> None:
        """Test that mismatched heading levels reduce the score."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            original_path = Path(tmp_dir) / "original.docx"
            rebuilt_path = Path(tmp_dir) / "rebuilt.docx"

            # Original: H1 then H2
            doc1 = Document()
            doc1.add_heading("Title", level=1)
            doc1.add_heading("Subtitle", level=2)
            doc1.add_paragraph("Content")
            doc1.save(str(original_path))

            # Rebuilt: H1 then H3 (wrong level)
            doc2 = Document()
            doc2.add_heading("Title", level=1)
            doc2.add_heading("Subtitle", level=3)
            doc2.add_paragraph("Content")
            doc2.save(str(rebuilt_path))

            score = scorer.score_structure(original_path, rebuilt_path)

            assert score < 100

    def test_extract_heading_levels(self) -> None:
        """Test _extract_heading_levels returns correct ordered levels."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            doc_path = Path(tmp_dir) / "test.docx"

            doc = Document()
            doc.add_heading("H1", level=1)
            doc.add_paragraph("Content")
            doc.add_heading("H2", level=2)
            doc.add_heading("H2b", level=2)
            doc.add_heading("H1b", level=1)
            doc.save(str(doc_path))

            levels = scorer._extract_heading_levels(doc_path)

            assert levels == [1, 2, 2, 1]

    def test_table_count_difference_reduces_score(self) -> None:
        """Test that different table counts reduce the structural score."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            original_path = Path(tmp_dir) / "original.docx"
            rebuilt_path = Path(tmp_dir) / "rebuilt.docx"

            # Original: 1 table
            doc1 = Document()
            doc1.add_paragraph("Content")
            table = doc1.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "A"
            table.cell(0, 1).text = "B"
            doc1.save(str(original_path))

            # Rebuilt: no tables
            doc2 = Document()
            doc2.add_paragraph("Content")
            doc2.save(str(rebuilt_path))

            score = scorer.score_structure(original_path, rebuilt_path)

            assert score < 100

    def test_identical_heading_levels_contribute_full_score(self) -> None:
        """Test that identical heading levels at all positions give full heading credit."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            doc_path = Path(tmp_dir) / "test.docx"

            doc = Document()
            doc.add_heading("H1", level=1)
            doc.add_heading("H2", level=2)
            doc.add_paragraph("Content")
            doc.save(str(doc_path))

            score = scorer.score_structure(doc_path, doc_path)

            assert score == 100.0


class TestFidelityFormatting:
    """Test the formatting fidelity scorer (replaces TestFidelityStyleScorer)."""

    def test_score_formatting_method_exists(self) -> None:
        """Test that score_formatting method exists on FidelityScorer."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        assert hasattr(scorer, "score_formatting")
        assert callable(getattr(scorer, "score_formatting"))

    def test_score_formatting_returns_numeric(self) -> None:
        """Test that score_formatting returns a 0-100 score."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        formatted_docx = SYNTHETIC_DIR / "formatted.docx"

        score = scorer.score_formatting(formatted_docx, formatted_docx)

        assert isinstance(score, (int, float))
        assert 0 <= score <= 100

    def test_identical_documents_score_100(self) -> None:
        """Test that identical documents get a perfect formatting score."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        formatted_docx = SYNTHETIC_DIR / "formatted.docx"

        score = scorer.score_formatting(formatted_docx, formatted_docx)

        assert score == 100.0

    def test_different_bold_reduces_score(self) -> None:
        """Test that different bold settings reduce the formatting score."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            original_path = Path(tmp_dir) / "original.docx"
            rebuilt_path = Path(tmp_dir) / "rebuilt.docx"

            # Original: bold text
            doc1 = Document()
            para = doc1.add_paragraph()
            run = para.add_run("Test content")
            run.bold = True
            doc1.save(str(original_path))

            # Rebuilt: no bold
            doc2 = Document()
            para = doc2.add_paragraph()
            run = para.add_run("Test content")
            run.bold = False
            doc2.save(str(rebuilt_path))

            score = scorer.score_formatting(original_path, rebuilt_path)

            assert score < 100

    def test_different_italic_reduces_score(self) -> None:
        """Test that different italic settings reduce the formatting score."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            original_path = Path(tmp_dir) / "original.docx"
            rebuilt_path = Path(tmp_dir) / "rebuilt.docx"

            doc1 = Document()
            para = doc1.add_paragraph()
            run = para.add_run("Test content")
            run.italic = True
            doc1.save(str(original_path))

            doc2 = Document()
            para = doc2.add_paragraph()
            run = para.add_run("Test content")
            run.italic = False
            doc2.save(str(rebuilt_path))

            score = scorer.score_formatting(original_path, rebuilt_path)

            assert score < 100

    def test_different_underline_reduces_score(self) -> None:
        """Test that different underline settings reduce the formatting score."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            original_path = Path(tmp_dir) / "original.docx"
            rebuilt_path = Path(tmp_dir) / "rebuilt.docx"

            doc1 = Document()
            para = doc1.add_paragraph()
            run = para.add_run("Test content")
            run.underline = True
            doc1.save(str(original_path))

            doc2 = Document()
            para = doc2.add_paragraph()
            run = para.add_run("Test content")
            run.underline = False
            doc2.save(str(rebuilt_path))

            score = scorer.score_formatting(original_path, rebuilt_path)

            assert score < 100

    def test_different_font_name_reduces_score(self) -> None:
        """Test that different font names reduce the formatting score."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            original_path = Path(tmp_dir) / "original.docx"
            rebuilt_path = Path(tmp_dir) / "rebuilt.docx"

            doc1 = Document()
            para = doc1.add_paragraph()
            run = para.add_run("Test content")
            run.font.name = "Arial"
            doc1.save(str(original_path))

            doc2 = Document()
            para = doc2.add_paragraph()
            run = para.add_run("Test content")
            run.font.name = "Times New Roman"
            doc2.save(str(rebuilt_path))

            score = scorer.score_formatting(original_path, rebuilt_path)

            assert score < 100

    def test_different_font_size_reduces_score(self) -> None:
        """Test that different font sizes reduce the formatting score."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            original_path = Path(tmp_dir) / "original.docx"
            rebuilt_path = Path(tmp_dir) / "rebuilt.docx"

            doc1 = Document()
            para = doc1.add_paragraph()
            run = para.add_run("Test content")
            run.font.size = Pt(12)
            doc1.save(str(original_path))

            doc2 = Document()
            para = doc2.add_paragraph()
            run = para.add_run("Test content")
            run.font.size = Pt(24)
            doc2.save(str(rebuilt_path))

            score = scorer.score_formatting(original_path, rebuilt_path)

            assert score < 100

    def test_no_paragraphs_with_runs_returns_100(self) -> None:
        """Test that documents with no runs return 100."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            doc_path = Path(tmp_dir) / "empty.docx"

            doc = Document()
            doc.add_paragraph()  # empty paragraph, no runs
            doc.save(str(doc_path))

            score = scorer.score_formatting(doc_path, doc_path)

            assert score == 100.0

    def test_multiple_runs_per_paragraph_compared(self) -> None:
        """Test that all runs in a paragraph are compared, not just the first."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            original_path = Path(tmp_dir) / "original.docx"
            rebuilt_path = Path(tmp_dir) / "rebuilt.docx"

            # Original: two runs, second is bold
            doc1 = Document()
            para = doc1.add_paragraph()
            para.add_run("Normal ")
            run2 = para.add_run("Bold")
            run2.bold = True
            doc1.save(str(original_path))

            # Rebuilt: two runs, second is NOT bold
            doc2 = Document()
            para = doc2.add_paragraph()
            para.add_run("Normal ")
            run2 = para.add_run("Bold")
            run2.bold = False
            doc2.save(str(rebuilt_path))

            score = scorer.score_formatting(original_path, rebuilt_path)

            assert score < 100

    def test_score_styles_deprecated_wrapper(self) -> None:
        """Test that score_styles still works as a deprecated wrapper."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        formatted_docx = SYNTHETIC_DIR / "formatted.docx"

        score = scorer.score_styles(formatted_docx, formatted_docx)

        assert score == 100.0

    def test_extract_all_run_formatting(self) -> None:
        """Test _extract_all_run_formatting returns per-paragraph, per-run formatting."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            doc_path = Path(tmp_dir) / "test.docx"

            doc = Document()
            para = doc.add_paragraph()
            run = para.add_run("Bold text")
            run.bold = True
            run.italic = False
            run.underline = False
            run.font.name = "Arial"
            run.font.size = Pt(12)
            doc.save(str(doc_path))

            result = scorer._extract_all_run_formatting(doc_path)

            # Should have at least one paragraph with one run
            assert len(result) >= 1
            # Find the paragraph with our run
            found = False
            for para_runs in result:
                for run_fmt in para_runs:
                    if run_fmt.get("font_name") == "Arial":
                        assert run_fmt["bold"] is True
                        assert run_fmt["italic"] is False
                        assert run_fmt["underline"] is False
                        assert run_fmt["font_size"] == Pt(12)
                        found = True
            assert found, "Expected to find a run with Arial font"


class TestFidelityTables:
    """Test the table fidelity scorer."""

    def test_score_tables_method_exists(self) -> None:
        """Test that score_tables method exists on FidelityScorer."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        assert hasattr(scorer, "score_tables")
        assert callable(getattr(scorer, "score_tables"))

    def test_identical_tables_score_100(self) -> None:
        """Test that identical table documents get 100."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        tables_docx = SYNTHETIC_DIR / "tables_simple.docx"

        score = scorer.score_tables(tables_docx, tables_docx)

        assert score == 100.0

    def test_no_tables_returns_none(self) -> None:
        """Test that documents without tables return None."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        simple_docx = SYNTHETIC_DIR / "simple.docx"

        score = scorer.score_tables(simple_docx, simple_docx)

        assert score is None

    def test_different_table_structure_reduces_score(self) -> None:
        """Test that different table row/col counts reduce the score."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            original_path = Path(tmp_dir) / "original.docx"
            rebuilt_path = Path(tmp_dir) / "rebuilt.docx"

            # Original: 3x3 table
            doc1 = Document()
            doc1.add_table(rows=3, cols=3)
            doc1.save(str(original_path))

            # Rebuilt: 2x2 table
            doc2 = Document()
            doc2.add_table(rows=2, cols=2)
            doc2.save(str(rebuilt_path))

            score = scorer.score_tables(original_path, rebuilt_path)

            assert score is not None
            assert score < 100

    def test_different_table_count_reduces_score(self) -> None:
        """Test that different number of tables reduces score."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            original_path = Path(tmp_dir) / "original.docx"
            rebuilt_path = Path(tmp_dir) / "rebuilt.docx"

            # Original: 2 tables
            doc1 = Document()
            doc1.add_table(rows=2, cols=2)
            doc1.add_table(rows=2, cols=2)
            doc1.save(str(original_path))

            # Rebuilt: 1 table
            doc2 = Document()
            doc2.add_table(rows=2, cols=2)
            doc2.save(str(rebuilt_path))

            score = scorer.score_tables(original_path, rebuilt_path)

            assert score is not None
            assert score < 100

    def test_score_tables_returns_0_to_100(self) -> None:
        """Test that score_tables returns a value between 0 and 100."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            doc_path = Path(tmp_dir) / "test.docx"

            doc = Document()
            doc.add_table(rows=2, cols=2)
            doc.save(str(doc_path))

            score = scorer.score_tables(doc_path, doc_path)

            assert score is not None
            assert 0 <= score <= 100

    def test_extract_tables_metadata(self) -> None:
        """Test _extract_tables_metadata returns correct table info."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            doc_path = Path(tmp_dir) / "test.docx"

            doc = Document()
            doc.add_table(rows=3, cols=4)
            doc.save(str(doc_path))

            metadata = scorer._extract_tables_metadata(doc_path)

            assert len(metadata) == 1
            assert metadata[0]["rows"] == 3
            assert metadata[0]["cols"] == 4
            assert "merged_cells" in metadata[0]
            assert "cell_backgrounds" in metadata[0]
            assert "table_style" in metadata[0]


class TestFidelityHyperlinks:
    """Test the hyperlink fidelity scorer."""

    def test_score_hyperlinks_method_exists(self) -> None:
        """Test that score_hyperlinks method exists on FidelityScorer."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        assert hasattr(scorer, "score_hyperlinks")
        assert callable(getattr(scorer, "score_hyperlinks"))

    def test_identical_hyperlinks_score_100(self) -> None:
        """Test that identical hyperlink documents get 100."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        hyperlinks_docx = SYNTHETIC_DIR / "hyperlinks.docx"

        score = scorer.score_hyperlinks(hyperlinks_docx, hyperlinks_docx)

        assert score == 100.0

    def test_no_hyperlinks_returns_none(self) -> None:
        """Test that documents without hyperlinks return None."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        simple_docx = SYNTHETIC_DIR / "simple.docx"

        score = scorer.score_hyperlinks(simple_docx, simple_docx)

        assert score is None

    def test_different_hyperlinks_reduce_score(self) -> None:
        """Test that different hyperlinks reduce the score."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            original_path = Path(tmp_dir) / "original.docx"
            rebuilt_path = Path(tmp_dir) / "rebuilt.docx"

            # Original: one hyperlink
            doc1 = Document()
            para1 = doc1.add_paragraph()
            _add_hyperlink(doc1, para1, "https://example.com", "Example")
            doc1.save(str(original_path))

            # Rebuilt: different hyperlink URL
            doc2 = Document()
            para2 = doc2.add_paragraph()
            _add_hyperlink(doc2, para2, "https://different.com", "Example")
            doc2.save(str(rebuilt_path))

            score = scorer.score_hyperlinks(original_path, rebuilt_path)

            assert score is not None
            assert score < 100

    def test_extract_hyperlinks(self) -> None:
        """Test _extract_hyperlinks returns (text, url) pairs."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            doc_path = Path(tmp_dir) / "test.docx"

            doc = Document()
            para = doc.add_paragraph()
            _add_hyperlink(doc, para, "https://example.com", "Click here")
            doc.save(str(doc_path))

            links = scorer._extract_hyperlinks(doc_path)

            assert len(links) >= 1
            texts = [t for t, u in links]
            urls = [u for t, u in links]
            assert "Click here" in texts
            assert "https://example.com" in urls


class TestFidelityTrackChanges:
    """Test the track changes fidelity scorer."""

    def test_score_track_changes_method_exists(self) -> None:
        """Test that score_track_changes method exists on FidelityScorer."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        assert hasattr(scorer, "score_track_changes")
        assert callable(getattr(scorer, "score_track_changes"))

    def test_identical_track_changes_score_100(self) -> None:
        """Test that identical track change documents get 100."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        tc_docx = SYNTHETIC_DIR / "track_changes_simple.docx"

        score = scorer.score_track_changes(tc_docx, tc_docx)

        assert score == 100.0

    def test_no_track_changes_returns_none(self) -> None:
        """Test that documents without track changes return None."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        simple_docx = SYNTHETIC_DIR / "simple.docx"

        score = scorer.score_track_changes(simple_docx, simple_docx)

        assert score is None

    def test_extract_track_changes(self) -> None:
        """Test _extract_track_changes returns insertions, deletions, and authors."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        tc_docx = SYNTHETIC_DIR / "track_changes_simple.docx"

        result = scorer._extract_track_changes(tc_docx)

        assert "insertions" in result
        assert "deletions" in result
        assert "authors" in result
        assert isinstance(result["insertions"], int)
        assert isinstance(result["deletions"], int)
        assert isinstance(result["authors"], set)
        # track_changes_simple.docx should have at least some changes
        assert result["insertions"] + result["deletions"] > 0

    def test_different_track_change_counts_reduce_score(self) -> None:
        """Test that different track change counts reduce the score."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            # Use a track changes doc as original, and simple doc as rebuilt
            tc_docx = SYNTHETIC_DIR / "track_changes_simple.docx"
            no_tc_path = Path(tmp_dir) / "no_tc.docx"

            doc = Document()
            doc.add_paragraph("No track changes here")
            doc.save(str(no_tc_path))

            score = scorer.score_track_changes(tc_docx, no_tc_path)

            # Original has track changes, rebuilt doesn't => should be less than 100
            assert score is not None
            assert score < 100


class TestFidelityTotalUpdated:
    """Test the updated score_total with new dimensions."""

    def test_score_total_returns_all_dimension_keys(self) -> None:
        """Test that score_total returns dict with all 5 dimensions + visual + total."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        simple_docx = SYNTHETIC_DIR / "simple.docx"

        result = scorer.score_total(simple_docx, simple_docx)

        assert isinstance(result, dict)
        assert "structure" in result
        assert "formatting" in result
        assert "tables" in result
        assert "hyperlinks" in result
        assert "track_changes" in result
        assert "visual" in result
        assert "total" in result

    def test_score_total_excludes_none_from_average(self) -> None:
        """Test that None dimensions are excluded from the total average."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        # simple.docx has no tables, hyperlinks, or track changes
        simple_docx = SYNTHETIC_DIR / "simple.docx"

        result = scorer.score_total(simple_docx, simple_docx)

        # tables, hyperlinks, track_changes should be None
        assert result["tables"] is None
        assert result["hyperlinks"] is None
        assert result["track_changes"] is None
        # Total should be mean of only structure + formatting
        expected = (result["structure"] + result["formatting"]) / 2
        assert abs(result["total"] - expected) < 0.01

    def test_score_total_with_tables(self) -> None:
        """Test that total includes table score when tables are present."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        tables_docx = SYNTHETIC_DIR / "tables_simple.docx"

        result = scorer.score_total(tables_docx, tables_docx)

        assert result["tables"] is not None
        assert result["tables"] == 100.0
        # Total should include tables in the average
        assert result["total"] >= 90

    def test_score_total_with_hyperlinks(self) -> None:
        """Test that total includes hyperlink score when hyperlinks are present."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        hyperlinks_docx = SYNTHETIC_DIR / "hyperlinks.docx"

        result = scorer.score_total(hyperlinks_docx, hyperlinks_docx)

        assert result["hyperlinks"] is not None
        assert result["hyperlinks"] == 100.0

    def test_score_total_with_track_changes(self) -> None:
        """Test that total includes track change score when track changes are present."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        tc_docx = SYNTHETIC_DIR / "track_changes_simple.docx"

        result = scorer.score_total(tc_docx, tc_docx)

        assert result["track_changes"] is not None
        assert result["track_changes"] == 100.0

    def test_score_total_visual_not_in_total_calculation(self) -> None:
        """Test that visual is reported but not included in total calculation."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        simple_docx = SYNTHETIC_DIR / "simple.docx"

        result = scorer.score_total(simple_docx, simple_docx)

        # Total should be mean of non-None dimensions, NOT including visual
        non_null = []
        for key in ["structure", "formatting", "tables", "hyperlinks", "track_changes"]:
            if result[key] is not None:
                non_null.append(result[key])
        expected_total = sum(non_null) / len(non_null) if non_null else 0
        assert abs(result["total"] - expected_total) < 0.01

    def test_score_total_identical_documents_score_high(self) -> None:
        """Test that identical documents get high combined score."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        simple_docx = SYNTHETIC_DIR / "simple.docx"

        result = scorer.score_total(simple_docx, simple_docx)

        assert result["structure"] == 100
        assert result["formatting"] == 100
        assert result["total"] >= 90


def visual_scorer_available() -> bool:
    """Check if visual scoring dependencies (LibreOffice, Poppler) are available."""
    from benchmarks.metrics.fidelity_scorer import _SOFFICE_CANDIDATES

    return any(shutil.which(p) for p in _SOFFICE_CANDIDATES) and shutil.which("pdfinfo") is not None


visual_scorer_skip = pytest.mark.skipif(
    not visual_scorer_available(),
    reason="Visual scoring requires LibreOffice and Poppler (pdfinfo)"
)


class TestFidelityVisualScorer:
    """Test the visual fidelity scorer (US-015)."""

    def test_score_visual_method_exists(self) -> None:
        """Test that score_visual method exists on FidelityScorer."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        assert hasattr(scorer, "score_visual")
        assert callable(getattr(scorer, "score_visual"))

    @visual_scorer_skip
    def test_score_visual_returns_numeric(self) -> None:
        """Test that score_visual returns a 0-100 score."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        simple_docx = SYNTHETIC_DIR / "simple.docx"

        score = scorer.score_visual(simple_docx, simple_docx)

        assert isinstance(score, (int, float))
        assert 0 <= score <= 100

    @visual_scorer_skip
    def test_identical_documents_score_high(self) -> None:
        """Test that identical documents get a high visual score."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        simple_docx = SYNTHETIC_DIR / "simple.docx"

        score = scorer.score_visual(simple_docx, simple_docx)

        assert score >= 95

    @visual_scorer_skip
    def test_different_documents_score_lower(self) -> None:
        """Test that visually different documents get a lower score."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            doc1_path = Path(tmp_dir) / "doc1.docx"
            doc2_path = Path(tmp_dir) / "doc2.docx"

            doc1 = Document()
            doc1.add_heading("Title One", level=1)
            doc1.add_paragraph("This is the first document with specific content.")
            doc1.save(str(doc1_path))

            doc2 = Document()
            doc2.add_heading("Different Title", level=1)
            doc2.add_paragraph("Completely different text that looks nothing alike.")
            doc2.add_paragraph("Even more content that makes them different.")
            doc2.save(str(doc2_path))

            score = scorer.score_visual(doc1_path, doc2_path)

            assert score < 100

    @visual_scorer_skip
    def test_score_visual_accepts_path_strings(self) -> None:
        """Test that score_visual accepts string paths."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        simple_docx = str(SYNTHETIC_DIR / "simple.docx")

        score = scorer.score_visual(simple_docx, simple_docx)

        assert score >= 95

    def test_score_visual_handles_missing_dependencies(self) -> None:
        """Test that score_visual raises clear error for missing dependencies."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        simple_docx = SYNTHETIC_DIR / "simple.docx"

        try:
            score = scorer.score_visual(simple_docx, simple_docx)
            assert 0 <= score <= 100
        except Exception as e:
            error_msg = str(e).lower()
            assert (
                "poppler" in error_msg
                or "libreoffice" in error_msg
                or "soffice" in error_msg
                or "pdfinfo" in error_msg
            )


class TestFidelityCombinedScorer:
    """Test the combined fidelity scorer (US-016) - backward compatibility."""

    def test_score_total_method_exists(self) -> None:
        """Test that score_total method exists on FidelityScorer."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        assert hasattr(scorer, "score_total")
        assert callable(getattr(scorer, "score_total"))

    def test_score_total_returns_dict(self) -> None:
        """Test that score_total returns a dict with expected keys."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        simple_docx = SYNTHETIC_DIR / "simple.docx"

        result = scorer.score_total(simple_docx, simple_docx)

        assert isinstance(result, dict)
        assert "structure" in result
        assert "formatting" in result
        assert "visual" in result
        assert "total" in result

    def test_score_total_returns_numeric_scores(self) -> None:
        """Test that non-None scores are numeric 0-100."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        simple_docx = SYNTHETIC_DIR / "simple.docx"

        result = scorer.score_total(simple_docx, simple_docx)

        for key in ["structure", "formatting", "total"]:
            assert isinstance(result[key], (int, float))
            assert 0 <= result[key] <= 100

        if result["visual"] is not None:
            assert isinstance(result["visual"], (int, float))
            assert 0 <= result["visual"] <= 100

    def test_score_total_accepts_path_strings(self) -> None:
        """Test that score_total accepts string paths."""
        from benchmarks.metrics.fidelity_scorer import FidelityScorer

        scorer = FidelityScorer()
        simple_docx = str(SYNTHETIC_DIR / "simple.docx")

        result = scorer.score_total(simple_docx, simple_docx)

        assert "total" in result
        assert result["total"] >= 90


def _add_hyperlink(doc: Document, paragraph, url: str, text: str) -> None:
    """Helper to add a hyperlink to a paragraph in a docx document."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Create the relationship
    part = doc.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the hyperlink element
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create a run with the text
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
