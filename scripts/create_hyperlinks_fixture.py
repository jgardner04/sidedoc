#!/usr/bin/env python3
"""Create hyperlinks.docx test fixture with various hyperlink scenarios."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


def add_hyperlink(paragraph, text: str, url: str, bold: bool = False, italic: bool = False):
    """Add a hyperlink to a paragraph.

    Args:
        paragraph: The paragraph to add the hyperlink to
        text: The display text for the hyperlink
        url: The URL the hyperlink points to
        bold: Whether the hyperlink text should be bold
        italic: Whether the hyperlink text should be italic
    """
    # Get the document part to access relationships
    part = paragraph.part

    # Create the relationship
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a run for the text
    run = OxmlElement('w:r')

    # Add run properties (for formatting)
    rPr = OxmlElement('w:rPr')

    # Add hyperlink style (blue color and underline)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')  # Standard hyperlink blue
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    # Add bold if requested
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)

    # Add italic if requested
    if italic:
        i = OxmlElement('w:i')
        rPr.append(i)

    run.append(rPr)

    # Add the text
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    run.append(text_elem)

    # Add run to hyperlink
    hyperlink.append(run)

    # Add hyperlink to paragraph
    paragraph._p.append(hyperlink)

    return hyperlink


def create_hyperlinks_docx():
    """Create hyperlinks.docx with various hyperlink scenarios."""
    doc = Document()

    # 1. Simple hyperlink
    para = doc.add_paragraph("Visit ")
    add_hyperlink(para, "Google", "https://www.google.com")
    run = para.add_run(" for search.")

    # 2. Multiple hyperlinks in one paragraph
    para = doc.add_paragraph("Check out ")
    add_hyperlink(para, "GitHub", "https://github.com")
    run = para.add_run(" and ")
    add_hyperlink(para, "Stack Overflow", "https://stackoverflow.com")
    run = para.add_run(" for coding help.")

    # 3. Bold hyperlink
    para = doc.add_paragraph("This is a ")
    add_hyperlink(para, "bold link", "https://example.com/bold", bold=True)
    run = para.add_run(" in the text.")

    # 4. Italic hyperlink
    para = doc.add_paragraph("This is an ")
    add_hyperlink(para, "italic link", "https://example.com/italic", italic=True)
    run = para.add_run(" in the text.")

    # 5. Bold and italic hyperlink
    para = doc.add_paragraph("This is a ")
    add_hyperlink(para, "bold italic link", "https://example.com/bolditalic", bold=True, italic=True)
    run = para.add_run(" in the text.")

    # 6. Hyperlink in heading
    heading = doc.add_paragraph("", style="Heading 1")
    heading.add_run("Heading with ")
    add_hyperlink(heading, "link", "https://example.com/heading")

    # 7. Hyperlink in Heading 2
    heading2 = doc.add_paragraph("", style="Heading 2")
    heading2.add_run("Subheading with ")
    add_hyperlink(heading2, "another link", "https://example.com/heading2")

    # 8. Hyperlink in bulleted list
    bullet1 = doc.add_paragraph("First item with ", style="List Bullet")
    add_hyperlink(bullet1, "link one", "https://example.com/list1")

    bullet2 = doc.add_paragraph("Second item with ", style="List Bullet")
    add_hyperlink(bullet2, "link two", "https://example.com/list2")

    # 9. Hyperlink in numbered list
    num1 = doc.add_paragraph("Numbered item with ", style="List Number")
    add_hyperlink(num1, "link three", "https://example.com/list3")

    num2 = doc.add_paragraph("Numbered item with ", style="List Number")
    add_hyperlink(num2, "link four", "https://example.com/list4")

    # 10. Long URL
    para = doc.add_paragraph("Here's a ")
    add_hyperlink(para, "long URL link", "https://example.com/very/long/path/with/many/segments/and/query?param1=value1&param2=value2&param3=value3")
    run = para.add_run(".")

    # 11. URL with special characters (parentheses, spaces encoded)
    para = doc.add_paragraph("Wikipedia article: ")
    add_hyperlink(para, "Article with (parentheses)", "https://en.wikipedia.org/wiki/Python_(programming_language)")

    # 12. Link text with special markdown characters
    para = doc.add_paragraph("Link with special chars: ")
    add_hyperlink(para, "text with *asterisks* and [brackets]", "https://example.com/special")

    # 13. URL with percent-encoded characters
    para = doc.add_paragraph("Link with encoded URL: ")
    add_hyperlink(para, "encoded spaces", "https://example.com/path%20with%20spaces")

    # 14. Email link (mailto)
    para = doc.add_paragraph("Contact us at ")
    add_hyperlink(para, "support@example.com", "mailto:support@example.com")
    run = para.add_run(".")

    # 15. Plain paragraph after links (for context)
    doc.add_paragraph("This is a plain paragraph with no links.")

    # Save the document
    output_path = Path(__file__).parent.parent / "tests" / "fixtures" / "hyperlinks.docx"
    doc.save(str(output_path))
    print(f"Created {output_path}")


if __name__ == "__main__":
    create_hyperlinks_docx()
